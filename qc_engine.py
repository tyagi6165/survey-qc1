"""
================================================================
  qc_engine.py — v8.1 Complete QC Engine
================================================================
8 checks in Phase 1:
  1. Termination rules (AI-powered, batched)
  2. Missing words/keywords
  3. Question text match (doc vs live)
  4. Options text match
  5. Mandatory markers
  6. Piping markers (raw unresolved)
  7. Answer codes sequence
  8. Question order

Works with ANY language. URL required for checks 2-4.
"""

import re
import json
import numpy as np
from difflib import SequenceMatcher
from docx import Document
from docx.oxml.ns import qn

from rule_engine import RuleEngine
from ai_judge import AIJudge as _AIJudge, make_judge as _make_judge
_rule_engine = RuleEngine()

try:
    from embeddings import (
        semantic_similarity, batch_embeddings, get_embedding, clear_cache,
        MATCH as SEM_MATCH, LIKELY_MATCH as SEM_LIKELY, MISMATCH as SEM_MISMATCH,
    )
    _SEMANTIC_AVAILABLE = True
except Exception:
    _SEMANTIC_AVAILABLE = False
    SEM_MATCH = 0.85
    SEM_LIKELY = 0.70
    SEM_MISMATCH = 0.60

    def semantic_similarity(a, b):
        return SequenceMatcher(None, _normalize(a)[:300], _normalize(b)[:300]).ratio()

    def batch_embeddings(texts):
        return []

    def get_embedding(text):
        return None

    def clear_cache():
        pass


# ============================================================
# CONSTANTS
# ============================================================
QID_PATTERN = re.compile(
    r'\[?\s*([RSQ]\d+(?:bis|ter|Info|info|Ex)?)\s*[\.\-–—\s\]]',
    re.IGNORECASE
)

PIPING_PATTERNS = [
    r'\[PIPE[^\]]*\]',
    r'\{\{[^}]+\}\}',
    r'<pipe[^>]*>',
    r'\[PIPED[^\]]*\]',
    r'\[Q\d+\.[A-Z]\]',
    r'\[Q\d+\]',
]
PIPING_RE = re.compile('|'.join(PIPING_PATTERNS), re.IGNORECASE)

MANDATORY_RE = re.compile(
    r'\*\s*$|\bmandatory\b|\bobligation\b|\bobbligatorio\b|\bobligatoire\b|\berforderlich\b',
    re.IGNORECASE
)

TERMINATION_KEYWORDS = [
    'close', 'terminate', 'thank', 'thanks', 'screen out',
    'grazie', 'chiudi', 'gracias', 'cierre', 'merci', 'fermer',
    'vielen dank', 'beenden', 'qualify', 'continue',
]

SKIP_PATTERNS = re.compile(
    r'^\s*(PROGRAMMING\s+TABLE|\[NEXT SCREEN\]|\[SAME SCREEN|---+|\|\s*---|\|\s*$)',
    re.IGNORECASE
)

STOPWORDS = {
    # English
    'the','and','or','but','is','are','was','were','be','been','have','has',
    'had','do','does','did','will','would','could','a','an','of','in','on',
    'at','to','for','with','by','from','this','that','these','those','it',
    'its','i','you','he','she','we','they','me','him','her','us','them',
    'my','your','his','their','our','all','any','some','not','so','if','as',
    # French
    'le','la','les','un','une','des','de','du','au','aux','et','ou','mais',
    'donc','car','ni','que','qui','quoi','dont','ou','je','tu','il','elle',
    'nous','vous','ils','elles','ce','se','sa','son','ses','mon','ton','ma',
    # Italian
    'il','lo','gli','un','una','di','da','del','della','dei','delle','degli',
    'al','alla','ai','alle','con','su','per','tra','fra','che','chi','come',
    'quando','dove','non','anche','ancora','sempre','molto','poco',
    # Spanish
    'el','los','las','una','unos','unas','del','al','con','por','para',
    'sin','sobre','entre','hacia','desde','hasta','durante','mediante',
}


# ============================================================
# DOCUMENT PARSER
# ============================================================
def _extract_table_options(question, rows):
    """Extract answer options from a doc table row list into a question dict."""
    for r in rows:
        if len(r) < 2:
            continue
        # Format 1: [code, text, marker?] — first cell is a purely numeric code
        if re.match(r'^\d+$', r[0]):
            code, text = r[0], r[1]
        # Format 2: [text, code] — exactly two cells, last is numeric, first is not
        elif len(r) == 2 and re.match(r'^\d+$', r[-1]) and not re.match(r'^\d+$', r[0]):
            code, text = r[-1], r[0]
        else:
            continue
        if text and code not in question["raw_codes"]:
            question["options"].append({"code": code, "text": text})
            question["raw_codes"].append(code)


def parse_document(doc_path):
    """
    Parse document → extract questions, options, logic tables.
    Returns: dict
    """
    doc = Document(doc_path)
    questions = {}
    logic_tables = []
    current_qid = None
    qid_order = []
    _pending_opts = None  # last option table seen, for PROG TABLE look-back

    # Walk body in order
    for child in doc.element.body.iterchildren():
        if child.tag == qn('w:p'):
            for para in doc.paragraphs:
                if para._element is child:
                    text = para.text.strip()
                    if not text or SKIP_PATTERNS.search(text):
                        break
                    m = QID_PATTERN.match(text)
                    if m:
                        qid = m.group(1)
                        current_qid = qid
                        if qid not in questions:
                            questions[qid] = {
                                "text": "", "options": [],
                                "is_mandatory": False, "has_piping": False,
                                "piping_found": [], "raw_codes": [],
                                "question_type": "",
                            }
                            qid_order.append(qid)
                        rest = text[m.end():].strip()
                        rest = re.sub(r'^[\-–—\s]+', '', rest).strip()
                        if rest:
                            questions[qid]["text"] += " " + rest
                    elif current_qid:
                        opt = re.match(r'^(\d+)[\.\)]\s+(.+)', text)
                        if opt:
                            questions[current_qid]["options"].append({
                                "code": opt.group(1),
                                "text": opt.group(2).strip()
                            })
                            questions[current_qid]["raw_codes"].append(opt.group(1))
                        else:
                            questions[current_qid]["text"] += " " + text
                            # Check mandatory
                            if MANDATORY_RE.search(text):
                                questions[current_qid]["is_mandatory"] = True
                            # Check piping
                            pipes = PIPING_RE.findall(text)
                            if pipes:
                                questions[current_qid]["has_piping"] = True
                                questions[current_qid]["piping_found"].extend(pipes)
                    break

        elif child.tag == qn('w:tbl'):
            for tbl in doc.tables:
                if tbl._element is child:
                    rows = []
                    for row in tbl.rows:
                        row_text = []
                        for cell in row.cells:
                            ct = ' '.join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                            if ct:
                                row_text.append(ct)
                        if row_text:
                            rows.append(row_text)

                    flat = ' | '.join(' / '.join(r) for r in rows)

                    # Is this a LOGIC table with termination?
                    has_logic = bool(re.search(r'\b(LOGIC|ROUTING|ROUTINE)\b', flat, re.I))
                    has_term = any(kw in flat.lower() for kw in TERMINATION_KEYWORDS)

                    if has_logic and has_term:
                        logic_tables.append({
                            'host_qid': current_qid,
                            'flat_text': flat,
                            'rows': rows
                        })
                        _pending_opts = None
                    elif rows:
                        first_cell = rows[0][0] if rows[0] else ''
                        is_prog = bool(re.match(r'PROGRAMM?ING\s+TABLE', first_cell, re.I))

                        if is_prog:
                            # Extract QID from second cell of first row
                            prog_qid = None
                            if len(rows[0]) >= 2:
                                m = re.match(
                                    r'([RSQ]\d+(?:bis|ter|Info|info|Ex)?)',
                                    rows[0][1].strip(), re.I
                                )
                                if m:
                                    prog_qid = m.group(1)
                            if prog_qid:
                                if prog_qid not in questions:
                                    questions[prog_qid] = {
                                        "text": "", "options": [],
                                        "is_mandatory": False, "has_piping": False,
                                        "piping_found": [], "raw_codes": [],
                                        "question_type": "",
                                    }
                                    qid_order.append(prog_qid)
                                current_qid = prog_qid
                                # Check TYPE row for mandatory marker and question type
                                for r in rows[1:5]:
                                    if r and r[0].strip().upper() == 'TYPE' and len(r) >= 2:
                                        if MANDATORY_RE.search(r[1]):
                                            questions[prog_qid]["is_mandatory"] = True
                                        questions[prog_qid]["question_type"] = r[1].strip().upper()
                                # Assign buffered options if this question has none yet
                                if _pending_opts and not questions[prog_qid]["options"]:
                                    _extract_table_options(questions[prog_qid], _pending_opts)
                            _pending_opts = None
                        else:
                            # Potential answer-option table
                            if current_qid and current_qid in questions:
                                # Only assign if the question has no options yet — prevents
                                # a following question's option table being appended here
                                # when current_qid hasn't advanced past its own PROG TABLE yet.
                                if not questions[current_qid]["options"]:
                                    _extract_table_options(questions[current_qid], rows)
                            # Always buffer for the next PROG TABLE (handles the case where
                            # the option table appears before the PROG TABLE that names the QID)
                            _pending_opts = rows
                    break

    # Clean up texts
    for qid in questions:
        questions[qid]["text"] = re.sub(r'\s+', ' ', questions[qid]["text"]).strip()

    return {
        "questions": questions,
        "qid_order": qid_order,
        "logic_tables": logic_tables,
        "stats": {
            "total_questions": len(questions),
            "with_options": sum(1 for q in questions.values() if q["options"]),
            "with_piping": sum(1 for q in questions.values() if q["has_piping"]),
            "with_mandatory": sum(1 for q in questions.values() if q["is_mandatory"]),
            "logic_tables": len(logic_tables),
        }
    }


# ============================================================
# CHECK 1: TERMINATION RULES (AI-powered, BATCHED)
# ============================================================
def check_termination(logic_tables, gemini_model):
    """
    Extract ALL termination rules using ONE batched AI call.
    Avoids quota issues by sending all tables at once.
    """
    if not logic_tables:
        return [], {"status": "no_logic_tables"}

    # Build batched prompt
    tables_text = ""
    for i, t in enumerate(logic_tables, 1):
        tables_text += f"\n--- TABLE {i} (QID: {t['host_qid']}) ---\n{t['flat_text'][:500]}\n"

    prompt = f"""You are a survey QC expert. Extract ALL termination rules from these {len(logic_tables)} LOGIC tables.

{tables_text}

RULES:
- Language can be ANY (English, French, Italian, Spanish, etc.)
- "thanks and close", "thank and close", "merci et fermer", "grazie e chiudi" = TERMINATE
- If multiple codes close in ONE table, create SEPARATE rule for EACH code
- "If code 1: continue, If code 2: close" → only code 2 terminates
- Compound = involves NOT/AND/OR or cross-question reference

Return ONLY this JSON (no markdown):
{{
  "rules": [
    {{
      "table_num": 1,
      "test_qid": "R0",
      "answer_code": "1",
      "action": "terminate",
      "complexity": "simple",
      "reason": "code 1 → thanks and close"
    }}
  ]
}}"""

    try:
        response = gemini_model.generate_content(prompt)
        text = response.text.strip()
        # Clean any markdown or thinking tags
        text = re.sub(r'```json|```', '', text).strip()
        text = re.sub(r'<[^>]+>.*?</[^>]+>', '', text, flags=re.DOTALL).strip()
        # Extract JSON
        m = re.search(r'\{.*\}', text, re.DOTALL)
        if m:
            data = json.loads(m.group(0))
            rules = data.get("rules", [])
            return rules, {"status": "ok", "total": len(rules)}
        return [], {"status": "json_parse_error", "raw": text[:200]}
    except Exception as e:
        return [], {"status": "error", "message": str(e)}


# ============================================================
# CHECK 2: MISSING WORDS (Doc vs Live)
# ============================================================
def check_missing_words(doc_questions, live_questions):
    """
    For each question, find words in doc that are missing from live.
    Semantic gate: if the texts are semantically equivalent (translations,
    synonyms, paraphrases) the word differences are not bugs.
    """
    issues = []

    for qid, doc_q in doc_questions.items():
        if qid not in live_questions:
            continue

        doc_text = doc_q["text"]
        live_text = live_questions[qid].get("text", "")

        # Semantic gate — skip word-diff noise when meaning is preserved
        try:
            sem_score = semantic_similarity(doc_text[:400], live_text[:400])
            if sem_score >= SEM_MATCH:
                continue
        except Exception:
            sem_score = None

        doc_words = _tokenize(doc_text)
        live_words = set(_tokenize(live_text))

        missing = [w for w in doc_words if w not in live_words]
        missing = list(dict.fromkeys(missing))

        if missing:
            score_note = f" | sem={sem_score:.2f}" if sem_score is not None else ""
            issues.append({
                "qid": qid,
                "check": "MISSING_WORDS",
                "severity": "HIGH" if len(missing) >= 3 else "MEDIUM",
                "details": f"Words in doc but not in live: {missing[:10]}{score_note}",
                "doc_snippet": doc_text[:100],
                "live_snippet": live_text[:100],
                "similarity_score": round(sem_score, 3) if sem_score is not None else None,
            })

    return issues


# ============================================================
# CHECK 3: QUESTION TEXT MATCH
# ============================================================
def _find_base_match(doc_qid, live_qids):
    """
    Startswith fallback: doc='R2', live='R2new' → match; doc='R2', live='R20' → no match.
    Suffix must be non-empty and start with a non-digit character.
    Dot-to-x: doc='R2.2' also tries 'R2x2' against live QIDs.
    """
    doc_lower = doc_qid.lower()
    for live_qid in live_qids:
        if live_qid.lower().startswith(doc_lower):
            suffix = live_qid[len(doc_qid):]
            if suffix and re.match(r'^[^0-9]', suffix, re.IGNORECASE):
                return live_qid

    if '.' in doc_qid:
        dot_as_x = doc_qid.replace('.', 'x')
        for live_qid in live_qids:
            if live_qid.lower() == dot_as_x.lower():
                return live_qid

    return None


def check_question_text(doc_questions, live_questions, threshold=0.65):
    """
    Semantic match question text between doc and live (multilingual).

    Buckets:
      score >= MATCH (0.85)          → same meaning, no issue
      MISMATCH (0.60) <= score < MATCH → POSSIBLE — needs review (MEDIUM)
      score < MISMATCH (0.60)         → confirmed mismatch (HIGH)

    threshold parameter kept for API compatibility but no longer used;
    the semantic thresholds from embeddings.py are authoritative.
    """
    issues = []

    for qid, doc_q in doc_questions.items():
        if qid not in live_questions:
            matched = _find_base_match(qid, live_questions)
            if matched:
                issues.append({
                    "qid": qid,
                    "check": "QID_NAMING_VARIANT",
                    "severity": "MEDIUM",
                    "details": f"Doc QID '{qid}' matched live as '{matched}' (naming variant)",
                    "live_qid": matched,
                })
                live_key = matched
            else:
                issues.append({
                    "qid": qid,
                    "check": "QUESTION_MISSING_IN_LIVE",
                    "severity": "HIGH",
                    "details": f"Question {qid} found in doc but NOT in live survey",
                })
                continue
        else:
            live_key = qid

        doc_text = doc_q["text"]
        live_text = live_questions[live_key].get("text", "")

        if not doc_text or len(doc_text) < 10:
            continue

        try:
            score = semantic_similarity(doc_text[:500], live_text[:500])
        except Exception:
            # Fallback to fuzzy if embeddings unavailable
            score = SequenceMatcher(
                None, _normalize(doc_text)[:300], _normalize(live_text)[:300]
            ).ratio()

        if score >= SEM_MATCH:
            continue  # Same meaning — not a bug

        if score >= SEM_MISMATCH:
            check_label = "TEXT_POSSIBLE_MISMATCH"
            severity = "MEDIUM"
            verdict = "POSSIBLE - needs review"
        else:
            check_label = "TEXT_MISMATCH"
            severity = "HIGH"
            verdict = "confirmed mismatch"

        issues.append({
            "qid": qid,
            "check": check_label,
            "severity": severity,
            "details": f"Semantic similarity: {score:.0%} — {verdict}",
            "doc_snippet": doc_text[:120],
            "live_snippet": live_text[:120],
            "similarity_score": round(score, 3),
        })

    return issues


# ============================================================
# CHECK 4: OPTIONS MATCH
# ============================================================
def check_options(doc_questions, live_questions):
    """
    Check if all answer options from doc appear in live survey.

    Uses batch semantic similarity so cross-lingual option lists
    (French doc / English live, etc.) are matched correctly.
    An option is considered found if any live option scores >= LIKELY_MATCH (0.70).
    Similarity scores are included in reported evidence.
    """
    issues = []

    for qid, doc_q in doc_questions.items():
        if "NUMERIC" in doc_q.get("question_type", "").upper():
            continue
        doc_opts = doc_q.get("options", [])
        if not doc_opts:
            continue
        if qid not in live_questions:
            continue

        live_opts = live_questions[qid].get("options", [])
        live_raw = [o.get("text", "") for o in live_opts]

        # Pre-encode all live option texts for this QID in one batch call
        try:
            if live_raw:
                live_vecs = np.array(batch_embeddings(live_raw))  # (N, D)
            else:
                live_vecs = None
        except Exception:
            live_vecs = None

        missing_opts = []
        for opt in doc_opts:
            opt_text = opt.get("text", "")
            if len(_normalize(opt_text)) < 3:
                continue

            best_score = 0.0
            best_match = ""

            if live_vecs is not None and live_vecs.size:
                try:
                    opt_vec = get_embedding(opt_text)  # (D,) normalised
                    scores = live_vecs @ opt_vec        # cosine sim for each live option
                    idx = int(np.argmax(scores))
                    best_score = float(scores[idx])
                    best_match = live_raw[idx][:50]
                except Exception:
                    # Fallback to fuzzy for this option
                    live_norms = [_normalize(t) for t in live_raw]
                    opt_norm = _normalize(opt_text)
                    fuzz = [SequenceMatcher(None, opt_norm, lt).ratio() for lt in live_norms]
                    best_score = max(fuzz) if fuzz else 0.0
                    best_match = live_raw[int(np.argmax(fuzz))][:50] if fuzz else ""
            elif live_raw:
                live_norms = [_normalize(t) for t in live_raw]
                opt_norm = _normalize(opt_text)
                fuzz = [SequenceMatcher(None, opt_norm, lt).ratio() for lt in live_norms]
                best_score = max(fuzz) if fuzz else 0.0

            if best_score < SEM_LIKELY:
                score_str = f"{best_score:.0%}"
                label = f"[{opt['code']}] {opt_text[:50]} (best={score_str}"
                if best_match:
                    label += f", closest: '{best_match}'"
                label += ")"
                missing_opts.append(label)

        if missing_opts:
            issues.append({
                "qid": qid,
                "check": "OPTIONS_MISSING",
                "severity": "HIGH",
                "details": f"Options in doc but not in live: {missing_opts[:5]}",
            })

    return issues


# ============================================================
# CHECK 5: MANDATORY MARKERS
# ============================================================
def check_mandatory(doc_questions, live_questions=None):
    """
    Check mandatory questions have * marker in live.
    If no live data, report mandatory questions found in doc.
    """
    issues = []

    for qid, doc_q in doc_questions.items():
        if not doc_q.get("is_mandatory"):
            continue

        if live_questions is None:
            # Doc-only mode: just report what's mandatory
            issues.append({
                "qid": qid,
                "check": "MANDATORY_CHECK",
                "severity": "INFO",
                "details": f"Marked mandatory in doc — verify * marker in live survey",
            })
        else:
            live_q = live_questions.get(qid, {})
            if not live_q.get("has_mandatory_marker"):
                issues.append({
                    "qid": qid,
                    "check": "MANDATORY_MARKER_MISSING",
                    "severity": "MEDIUM",
                    "details": "Doc says mandatory but live survey has no * marker",
                })

    return issues


# ============================================================
# CHECK 6: PIPING MARKERS
# ============================================================
def check_piping(doc_questions, live_questions=None):
    """
    Check for raw unresolved piping markers in live survey.
    Doc piping = expected. Live piping = BUG.
    """
    issues = []

    # Check doc for piping expectations
    doc_piping_qids = {
        qid for qid, q in doc_questions.items()
        if q.get("has_piping")
    }

    if live_questions:
        for qid, live_q in live_questions.items():
            if live_q.get("has_raw_piping"):
                raw = live_q.get("raw_piping_found", [])
                issues.append({
                    "qid": qid,
                    "check": "PIPING_NOT_RESOLVED",
                    "severity": "HIGH",
                    "details": f"Raw piping markers visible in live survey: {raw[:5]}",
                })
    else:
        # Doc-only: check if doc has piping expectations
        for qid in doc_piping_qids:
            pipes = doc_questions[qid].get("piping_found", [])
            issues.append({
                "qid": qid,
                "check": "PIPING_EXPECTED",
                "severity": "INFO",
                "details": f"Doc has piping markers {pipes[:3]} — verify resolved in live",
            })

    return issues


# ============================================================
# CHECK 7: ANSWER CODES SEQUENCE
# ============================================================
def check_answer_codes(doc_questions):
    """
    Verify answer codes are sequential (1,2,3,4...) without gaps or duplicates.
    """
    issues = []

    for qid, q in doc_questions.items():
        codes = q.get("raw_codes", [])
        if len(codes) < 2:
            continue

        try:
            int_codes = [int(c) for c in codes]
        except ValueError:
            continue

        # Check sequential
        expected = list(range(int_codes[0], int_codes[0] + len(int_codes)))
        if int_codes != expected:
            # Find gaps or duplicates
            gaps = [e for e in expected if e not in int_codes]
            dupes = [c for c in int_codes if int_codes.count(c) > 1]

            detail = []
            if gaps:
                detail.append(f"Missing codes: {gaps}")
            if dupes:
                detail.append(f"Duplicate codes: {list(set(dupes))}")

            if detail:
                issues.append({
                    "qid": qid,
                    "check": "CODE_SEQUENCE_ERROR",
                    "severity": "MEDIUM",
                    "details": " | ".join(detail),
                    "found_codes": int_codes,
                })

    return issues


# ============================================================
# CHECK 8: QUESTION ORDER
# ============================================================
def check_question_order(doc_qid_order, live_questions=None):
    """
    Verify questions appear in same order in doc and live.
    """
    issues = []

    if live_questions is None:
        # Doc-only: check doc order is logical (R before S, numbered sequence)
        r_questions = [q for q in doc_qid_order if q.startswith('R')]
        s_questions = [q for q in doc_qid_order if q.startswith('S')]

        # Extract numbers
        def get_num(qid):
            m = re.search(r'\d+', qid)
            return int(m.group()) if m else 0

        r_nums = [get_num(q) for q in r_questions]
        s_nums = [get_num(q) for q in s_questions]

        # Check R questions are in order
        for i in range(1, len(r_nums)):
            if r_nums[i] < r_nums[i-1]:
                issues.append({
                    "qid": r_questions[i],
                    "check": "ORDER_ISSUE",
                    "severity": "MEDIUM",
                    "details": f"{r_questions[i]} appears after {r_questions[i-1]} — unexpected order",
                })

        # Check S questions are in order
        for i in range(1, len(s_nums)):
            if s_nums[i] < s_nums[i-1]:
                issues.append({
                    "qid": s_questions[i],
                    "check": "ORDER_ISSUE",
                    "severity": "MEDIUM",
                    "details": f"{s_questions[i]} appears after {s_questions[i-1]} — unexpected order",
                })
    else:
        # Compare doc order vs live order
        live_order = list(live_questions.keys())
        doc_in_live = [q for q in doc_qid_order if q in live_questions]

        for i, qid in enumerate(doc_in_live):
            if qid in live_order:
                live_pos = live_order.index(qid)
                doc_pos = i
                # Check if relative order is maintained
                if i > 0:
                    prev_qid = doc_in_live[i-1]
                    if prev_qid in live_order:
                        prev_live_pos = live_order.index(prev_qid)
                        if live_pos < prev_live_pos:
                            issues.append({
                                "qid": qid,
                                "check": "ORDER_MISMATCH",
                                "severity": "MEDIUM",
                                "details": f"Doc order: {prev_qid}→{qid}, but live order is reversed",
                            })

    return issues


# ============================================================
# LAYER 1+2 BRIDGE: deterministic rules → semantic tiebreak
# ============================================================
def compare_question_pair(doc_q, live_q, xml_q=None):
    """
    Run Layer 1 (rule engine) then Layer 2 (semantic) for a single QID pair.

    WARN results are escalated to semantic check:
      - texts semantically similar (≥ MATCH) → PASS_SEMANTIC (suppressed)
      - texts semantically different          → issue confirmed
    FAIL results bypass semantic and are always reported.

    Returns list of issue dicts ready for the results structure.
    """
    issues = []
    rule_results = _rule_engine.run_all_checks(doc_q, live_q, xml_q)

    for ri in rule_results:
        if ri["result"] == "FAIL":
            issues.append(ri)

        elif ri["result"] == "WARN":
            doc_text = (doc_q or {}).get("text", "")
            live_text = (live_q or {}).get("text", "")

            if doc_text and live_text:
                try:
                    score = semantic_similarity(doc_text[:500], live_text[:500])
                    ri["similarity_score"] = round(score, 3)
                    if score >= SEM_MATCH:
                        ri["result"] = "PASS_SEMANTIC"
                        ri["confidence"] = max(ri["confidence"] - 20, 30)
                        # suppressed — don't append
                        continue
                except Exception:
                    pass  # fall through to append

            issues.append(ri)

    return issues


def _rule_issue_to_standard(qid, ri):
    """Convert a rule engine result dict to the standard qc_engine issue format."""
    severity_map = {"FAIL": "HIGH", "WARN": "MEDIUM"}
    return {
        "qid": qid,
        "check": ri.get("type", "RULE_CHECK"),
        "severity": severity_map.get(ri.get("result", "WARN"), "MEDIUM"),
        "details": ri.get("reason", ""),
        "rule": ri.get("rule", ""),
        "confidence": ri.get("confidence", 0),
        "similarity_score": ri.get("similarity_score"),
    }


# ============================================================
# MASTER RUN — ALL 8 CHECKS
# ============================================================
def generate_test_cases(doc_data: dict) -> tuple:
    """
    Generate test cases from survey logic/routing rules.
    Returns (test_cases: list[dict], summary: dict).
    Safe to call with any doc_data — returns empty on error.
    """
    try:
        from test_generator import generate_test_cases as _gen
        return _gen(doc_data)
    except Exception:
        return [], {"total": 0, "auto_runnable": 0, "manual_only": 0,
                    "by_type": {}, "by_priority": {}}


def run_playwright_tests(
    test_cases: list,
    survey_url: str,
    screenshot_dir: str,
    max_tests: int = 20,
    timeout_ms: int = 30_000,
) -> dict:
    """
    Run auto_runnable test cases against the live survey using Playwright.

    Safe to call whether or not Playwright is installed — returns gracefully
    with error key set if the dependency is missing or the runner crashes.

    Returns:
        {"results": [...], "summary": {...}, "error": None | str}
    """
    try:
        from test_runner import run_playwright_tests as _run
        return _run(
            test_cases,
            survey_url,
            screenshot_dir,
            max_tests=max_tests,
            timeout_ms=timeout_ms,
        )
    except Exception as e:
        return {
            "results": [],
            "summary": {"total": 0, "passed": 0, "failed": 0,
                        "errors": 0, "skipped": 0, "pass_rate": "N/A"},
            "error": f"test_runner import failed: {str(e)[:200]}",
        }


def check_export_schema(
    doc_questions: dict,
    export_input: str,
    xml_questions: list | None = None,
) -> list[dict]:
    """
    Validate data export schema against survey definition.
    Wraps export_validator.run_export_validation() for direct use from app.py.

    Args:
        doc_questions:  from parse_document()["questions"]
        export_input:   CSV header line, TSV, or whitespace-separated variable names
        xml_questions:  optional parsed XML questions list

    Returns list of issue dicts with rule R031–R038.
    """
    try:
        from export_validator import run_export_validation
        result = run_export_validation(doc_questions, export_input,
                                       xml_questions=xml_questions)
        return result.get("issues", [])
    except Exception:
        return []


def run_all_checks(doc_data, live_questions=None, gemini_model=None,
                   threshold=0.65, export_schema: str | None = None):
    """
    Run all QC checks. Returns structured results.

    Args:
        doc_data:       output from parse_document()
        live_questions: dict of live survey data (optional, from crawler)
        gemini_model:   Gemini model instance (for termination check)
        threshold:      kept for API compatibility
        export_schema:  CSV header string for data export validation (optional)

    Returns:
        dict with all check results + summary
    """
    # Drop embedding cache so stale vectors from previous runs don't persist
    clear_cache()

    results = {
        "rule_engine": [],          # Layer 1+2 deterministic results
        "termination": {"rules": [], "meta": {}},
        "missing_words": [],
        "text_match": [],
        "options_match": [],
        "mandatory": [],
        "piping": [],
        "answer_codes": [],
        "question_order": [],
        "export_validation": [],    # R031-R038 data export checks
        "test_cases": [],           # auto-generated test cases
        "test_cases_summary": {},
    }

    doc_questions = doc_data["questions"]
    qid_order = doc_data["qid_order"]
    logic_tables = doc_data["logic_tables"]

    # ── LAYER 1+2: Deterministic rule engine (runs before AI, always) ─────────
    # Per-question-pair checks (R001–R006)
    if live_questions:
        re_issues = []
        all_qids = set(doc_questions) | set(live_questions)
        for qid in all_qids:
            doc_q  = doc_questions.get(qid)
            live_q = live_questions.get(qid)
            for ri in compare_question_pair(doc_q, live_q):
                re_issues.append(_rule_issue_to_standard(qid, ri))
        results["rule_engine"] = re_issues

    # Survey-level deterministic checks: R007 termination, R008 order
    re_survey = []
    for lt in logic_tables:
        ri = _rule_engine.check_termination(lt, live_tested=False)
        if ri["result"] in ("FAIL", "WARN"):
            re_survey.append(_rule_issue_to_standard(lt.get("host_qid", "?"), ri))
    if live_questions:
        live_order = list(live_questions.keys())
        ri = _rule_engine.check_question_order(qid_order, live_order)
        if ri["result"] in ("FAIL", "WARN"):
            re_survey.append(_rule_issue_to_standard("ORDER", ri))
    results["rule_engine"].extend(re_survey)
    # ─────────────────────────────────────────────────────────────────────────

    # CHECK 1: Termination (AI, batched)
    if gemini_model and logic_tables:
        rules, meta = check_termination(logic_tables, gemini_model)
        results["termination"] = {"rules": rules, "meta": meta}

    # CHECK 2: Missing words (needs live)
    if live_questions:
        results["missing_words"] = check_missing_words(doc_questions, live_questions)

    # CHECK 3: Text match (needs live)
    if live_questions:
        results["text_match"] = check_question_text(doc_questions, live_questions, threshold)

    # CHECK 4: Options match (needs live)
    if live_questions:
        results["options_match"] = check_options(doc_questions, live_questions)

    # CHECK 5: Mandatory (doc-only or with live)
    results["mandatory"] = check_mandatory(doc_questions, live_questions)

    # CHECK 6: Piping (doc-only or with live)
    results["piping"] = check_piping(doc_questions, live_questions)

    # CHECK 7: Answer codes (doc-only)
    results["answer_codes"] = check_answer_codes(doc_questions)

    # CHECK 8: Question order (doc-only or with live)
    results["question_order"] = check_question_order(qid_order, live_questions)

    # ── DATA EXPORT VALIDATION (R031-R038) ────────────────────────────────────
    if export_schema:
        results["export_validation"] = check_export_schema(
            doc_questions,
            export_schema,
            xml_questions=doc_data.get("xml_questions"),
        )
    # ─────────────────────────────────────────────────────────────────────────

    # ── TEST CASE GENERATION ──────────────────────────────────────────────────
    _tc_list, _tc_summary = generate_test_cases(doc_data)
    results["test_cases"] = _tc_list
    results["test_cases_summary"] = _tc_summary
    # ─────────────────────────────────────────────────────────────────────────

    # Flatten all issues for Layer 3 review
    all_issues = (
        results["rule_engine"] +
        results["missing_words"] +
        results["text_match"] +
        results["options_match"] +
        results["mandatory"] +
        results["piping"] +
        results["answer_codes"] +
        results["question_order"] +
        results["export_validation"]
    )

    # ── LAYER 3: AI Judge — reviews only uncertain issues ─────────────────────
    # Rules with confidence ≥ 85 are NEVER overridden (rules are king).
    # If Gemini is unavailable the report is returned unchanged.
    judge = _make_judge(gemini_model)
    all_issues, _aj_stats = judge.review_all(all_issues, doc_questions, live_questions)
    results["ai_judge"] = _aj_stats
    # ─────────────────────────────────────────────────────────────────────────

    term_rules = results["termination"]["rules"]
    terminate_rules = [r for r in term_rules if r.get("action") == "terminate"]
    simple_rules = [r for r in terminate_rules if r.get("complexity") == "simple"]
    compound_rules = [r for r in terminate_rules if r.get("complexity") == "compound"]

    sev = {"HIGH": 0, "MEDIUM": 0, "INFO": 0}
    for issue in all_issues:
        s = issue.get("severity", "INFO")
        sev[s] = sev.get(s, 0) + 1

    results["summary"] = {
        "total_issues": len(all_issues),
        "severity": sev,
        "termination_rules_found": len(term_rules),
        "terminate_simple": len(simple_rules),
        "terminate_compound": len(compound_rules),
        "checks_run": sum([
            1,                            # rule engine (always)
            1 if gemini_model else 0,     # termination (AI)
            1 if live_questions else 0,   # missing words
            1 if live_questions else 0,   # text match
            1 if live_questions else 0,   # options
            1,  # mandatory
            1,  # piping
            1,  # codes
            1,  # order
        ]),
        "mode": "full" if live_questions else "doc_only",
        "ai_judge": _aj_stats,
        "test_cases_generated": _tc_summary.get("total", 0),
        "test_cases_auto": _tc_summary.get("auto_runnable", 0),
    }

    # Verdict
    if sev["HIGH"] == 0 and len(compound_rules) == 0:
        results["summary"]["verdict"] = "PASS"
        results["summary"]["verdict_msg"] = "✅ No critical issues found"
    elif sev["HIGH"] > 0:
        results["summary"]["verdict"] = "FAIL"
        results["summary"]["verdict_msg"] = f"❌ {sev['HIGH']} HIGH severity issues found"
    else:
        results["summary"]["verdict"] = "REVIEW"
        results["summary"]["verdict_msg"] = f"⚠️ {sev['MEDIUM']} issues need review"

    return results


# ============================================================
# HELPERS
# ============================================================
def _normalize(text):
    if not text:
        return ""
    text = re.sub(r'\s+', ' ', text).strip().lower()
    text = re.sub(r"[''`\"\"']", "'", text)
    return text


def _tokenize(text):
    if not text:
        return []
    text = re.sub(r'\[[^\]]{1,30}\]|\{\{[^}]+\}\}|<[^>]+>|https?://\S+', ' ', text)
    text = text.lower()
    words = re.findall(r"[a-zàèéìòùáíóúüâêîôûñçäöüßœæ']+", text)
    return [w for w in words if len(w) >= 3 and w not in STOPWORDS]
