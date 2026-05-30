"""
normalizer.py — Survey document normalizer v1.0
Converts any survey doc/text into a structured JSON list using Gemini.
Works for any language, any platform, any question type.

Public API:
    normalize_document(raw_text, model=None)           -> list[dict]
    normalize_document_from_path(doc_path, model=None) -> list[dict]
    normalize_chunk(batch_text, model=None)            -> list[dict]
    chunk_document(raw_text)                           -> list[str]
    extract_raw_text(doc_path)                         -> str
"""

import re
import json
import os
import logging
import time

from docx import Document
from docx.oxml.ns import qn

log = logging.getLogger(__name__)

# ── Constants ────────────────────────────────────────────────────────────────

BATCH_SIZE = 12        # QID chunks per Gemini call
_MAX_BATCH_RETRIES = 2  # attempts after the first before giving up
_RETRY_DELAY_S = 3      # base sleep (seconds); multiplied by attempt number

# Matches a QID at the START of a line (question header), e.g. "R2.", "[S1bis]"
_QID_LINE_RE = re.compile(
    r'^\s*\[?\s*([RSQ]\d+(?:bis|ter|info|Info|Ex)?)\s*[\.\-–—\s\]]',
    re.IGNORECASE,
)

# Matches any QID token anywhere in text (for unique-count validation)
_QID_TOKEN_RE = re.compile(
    r'\b([RSQ]\d+(?:bis|ter|info|Info|Ex)?)\b',
    re.IGNORECASE,
)

# ── Prompt ────────────────────────────────────────────────────────────────────

_SCHEMA_DOC = """
Output schema (one object per question):
{
  "qid":            "<exactly as written in doc>",
  "qid_normalized": "<uppercase, dots→x, no spaces>",
  "text":           "<question text, original language>",
  "type":           "<UNIQUE|MULTIPLE|OPEN|MATRIX|GRID|UNKNOWN>",
  "options": [
    {"code": "01", "text": "<option text>", "marker": "<m/q/e/s/ne/null>"}
  ],
  "routing":     "<display/skip logic verbatim, or null>",
  "termination": "<screen-out/close condition verbatim, or null>",
  "is_matrix_group": <true if QID is a row/column inside a battery>
}

Type rules:
  UNIQUE   — single-select (radio)
  MULTIPLE — multi-select (checkbox)
  OPEN     — free text / numeric entry
  MATRIX   — grid where rows and columns both have QIDs
  GRID     — rating scale grid
  UNKNOWN  — cannot determine from context

Option rules:
  - code: the numeric/alpha code preceding the option text, null if absent
  - marker: any label AFTER the text such as m, q, e, s, ne — null if absent
  - Never skip an option even if code is missing
  - Include codes from LOGIC/ROUTING tables as references only — do NOT
    create phantom options for codes mentioned only in logic tables

Termination: extract the full condition ("If S2 < 2 or S2 > 40 → close"),
  not just "close". Include the triggering condition.

Return ONLY a valid JSON array. No markdown fences, no prose.
"""

_SYSTEM_PROMPT = "You are a survey specification parser.\n" + _SCHEMA_DOC


# ── Text extraction ───────────────────────────────────────────────────────────

def extract_raw_text(doc_path: str) -> str:
    """
    Extract all content from a .docx in document order.
    Tables are wrapped with [TABLE_START] / [TABLE_END] sentinels so
    chunk_document can attach them to the preceding QID.
    """
    doc = Document(doc_path)
    parts = []

    for child in doc.element.body.iterchildren():
        if child.tag == qn('w:p'):
            for para in doc.paragraphs:
                if para._element is child:
                    t = para.text.strip()
                    if t:
                        parts.append(t)
                    break

        elif child.tag == qn('w:tbl'):
            for tbl in doc.tables:
                if tbl._element is child:
                    parts.append('[TABLE_START]')
                    for row in tbl.rows:
                        cells = []
                        for cell in row.cells:
                            ct = ' '.join(
                                p.text.strip()
                                for p in cell.paragraphs
                                if p.text.strip()
                            )
                            if ct:
                                cells.append(ct)
                        if cells:
                            parts.append(' | '.join(cells))
                    parts.append('[TABLE_END]')
                    break

    return '\n'.join(parts)


# ── Chunking ──────────────────────────────────────────────────────────────────

def chunk_document(raw_text: str) -> list:
    """
    Split raw_text into one chunk per QID.
    Each chunk = QID header line + option lines + any immediately following
    TABLE block (LOGIC/PROG tables belong to the preceding QID).
    Returns list[str], one string per question cluster.
    """
    lines = raw_text.splitlines()
    chunks = []
    current = []

    for line in lines:
        if _QID_LINE_RE.match(line):
            if current:
                chunks.append('\n'.join(current))
            current = [line]
        else:
            current.append(line)

    if current:
        chunks.append('\n'.join(current))

    # Drop leading preamble blocks that contain no QID at all
    return [c for c in chunks if _QID_LINE_RE.search(c)]


# ── Model factory ─────────────────────────────────────────────────────────────

class _GeminiAdapter:
    """
    Wraps google.genai.Client to expose the same generate_content(prompt).text
    interface as google.generativeai.GenerativeModel, so normalize_chunk works
    with either an injected legacy model or one we create here.
    """
    def __init__(self, client, model_name, config):
        self._client = client
        self._model_name = model_name
        self._config = config

    def generate_content(self, prompt):
        return self._client.models.generate_content(
            model=self._model_name,
            contents=prompt,
            config=self._config,
        )


def _make_model():
    """
    Create a Gemini model from GEMINI_API_KEY env var using google.genai.
    Sets temperature=0 and requests JSON output via response_mime_type.
    Returns a _GeminiAdapter with a .generate_content(prompt).text interface.
    """
    from google import genai
    from google.genai import types

    key = os.environ.get('GEMINI_API_KEY', '').strip()
    if not key:
        raise RuntimeError(
            'GEMINI_API_KEY not set — export it or pass a model explicitly'
        )
    client = genai.Client(api_key=key)
    config = types.GenerateContentConfig(
        temperature=0,
        response_mime_type='application/json',
    )
    return _GeminiAdapter(client, 'gemini-2.5-flash', config)


# ── Gemini call ───────────────────────────────────────────────────────────────

def normalize_chunk(batch_text: str, model=None) -> list:
    """
    Normalize one batch of raw survey text via a single Gemini call.
    Returns list of question dicts. Returns [] on any parse/API failure
    (caller handles retry).
    """
    if model is None:
        model = _make_model()

    full_prompt = _SYSTEM_PROMPT + '\n\n--- SURVEY TEXT ---\n' + batch_text

    try:
        response = model.generate_content(full_prompt)
        raw = response.text.strip()

        # Strip markdown fences in case model ignores response_mime_type
        raw = re.sub(r'^```(?:json)?\s*', '', raw, flags=re.IGNORECASE)
        raw = re.sub(r'\s*```\s*$', '', raw)

        data = json.loads(raw)

        if isinstance(data, list):
            return data
        # Some models wrap in {"questions": [...]}
        if isinstance(data, dict):
            for key in ('questions', 'items', 'results'):
                if key in data and isinstance(data[key], list):
                    return data[key]

        log.warning('normalize_chunk: unexpected response shape (%s) — skipping', type(data))
        return []

    except json.JSONDecodeError as exc:
        log.error('normalize_chunk: JSON parse error — %s', exc)
        log.debug('normalize_chunk: raw response was: %s', raw[:300] if 'raw' in dir() else '(no response)')
        return []
    except Exception as exc:
        log.error('normalize_chunk: Gemini error — %s', exc)
        return []


# ── Validation helpers ────────────────────────────────────────────────────────

def _unique_qids_in_text(text: str) -> set:
    """Return the set of unique QID strings that appear as question headers."""
    found = set()
    for line in text.splitlines():
        m = _QID_LINE_RE.match(line)
        if m:
            found.add(m.group(1).upper())
    return found


# ── Orchestration ─────────────────────────────────────────────────────────────

def normalize_document(raw_text: str, model=None) -> list:
    """
    Primary entry point. Pure function: text in → list of question dicts out.
    Works for doc text or live-survey HTML text.

    Steps:
      1. chunk_document → one chunk per QID
      2. Group into batches of BATCH_SIZE
      3. normalize_chunk per batch; retry once on <90% yield
      4. Deduplicate by qid_normalized (first occurrence wins)
      5. Validate coverage: warn only if expected doc-header QIDs are MISSING.
         Extra sub-item QIDs (matrix rows, battery items) are valid — no warning.
    """
    if model is None:
        model = _make_model()

    chunks = chunk_document(raw_text)
    if not chunks:
        log.warning('normalize_document: no QID chunks found — returning empty list')
        return []

    expected_qids = _unique_qids_in_text(raw_text)
    log.info(
        'normalize_document: %d chunks from %d expected QIDs',
        len(chunks), len(expected_qids),
    )

    results = []

    for batch_idx in range(0, len(chunks), BATCH_SIZE):
        batch = chunks[batch_idx : batch_idx + BATCH_SIZE]
        batch_num = batch_idx // BATCH_SIZE + 1
        batch_text = '\n\n'.join(batch)

        log.info('  batch %d: %d chunks', batch_num, len(batch))
        batch_results = normalize_chunk(batch_text, model)

        # Retry up to _MAX_BATCH_RETRIES times if yield is suspiciously low (< 90%)
        for attempt in range(1, _MAX_BATCH_RETRIES + 1):
            if len(batch_results) >= len(batch) * 0.9:
                break
            log.warning(
                '  batch %d returned %d questions for %d chunks — retry %d/%d',
                batch_num, len(batch_results), len(batch), attempt, _MAX_BATCH_RETRIES,
            )
            time.sleep(_RETRY_DELAY_S * attempt)
            retry_results = normalize_chunk(batch_text, model)
            if len(retry_results) >= len(batch_results):
                batch_results = retry_results
                log.info('  batch %d retry %d: got %d questions', batch_num, attempt, len(batch_results))

        results.extend(batch_results)

    # Deduplicate by qid_normalized, preserving order
    seen: set = set()
    merged = []
    for q in results:
        key = (q.get('qid_normalized') or q.get('qid') or '').upper().strip()
        if key and key not in seen:
            seen.add(key)
            merged.append(q)

    # Final coverage validation: are all doc-header QIDs present in output?
    # Extra sub-item QIDs (matrix rows, battery items) are expected and valid.
    returned_qid_set = {
        (q.get('qid_normalized') or q.get('qid') or '').upper()
        for q in merged
    }
    missing = expected_qids - returned_qid_set
    extra_count = len(returned_qid_set - expected_qids)

    if missing:
        log.warning(
            'normalize_document: %d doc-header QID(s) missing from output: %s',
            len(missing), sorted(missing)[:10],
        )
    else:
        log.info(
            'normalize_document: all %d doc-header QIDs covered. '
            '%d extra sub-item QIDs also returned.',
            len(expected_qids), extra_count,
        )

    return merged


def normalize_document_from_path(doc_path: str, model=None) -> list:
    """Convenience wrapper — extracts text from .docx then normalizes."""
    raw_text = extract_raw_text(doc_path)
    return normalize_document(raw_text, model=model)
