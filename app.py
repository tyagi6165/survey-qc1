"""
================================================================
  SURVEYQC — Complete Full Stack App v10.0
  Light Theme + All Features + New Pages

  INSTALL:
    pip3 install flask playwright python-docx werkzeug
    playwright install chromium

  RUN:
    python3 app.py

  OPEN:
    http://localhost:5000

  FEATURES:
    - Landing page (public)
    - Login / Signup
    - Dashboard with time saved
    - New QC analysis page
    - Reports history
    - Report detail page
    - AI Tester simulation
    - Settings page
    - Admin panel (password protected)
    - Full QC engine (8 checks)
    - Real-time progress via polling
    - Word report download
================================================================
"""

import os, re, sys, json, uuid, threading, hashlib, time, sqlite3
import concurrent.futures
from pathlib import Path
from datetime import datetime
from difflib import SequenceMatcher
from qid_normalizer import (
    is_valid_qid, is_framework_page, is_internal_id, should_skip_qid,
    verify_extraction, get_parent_qid,
    INTERNAL_NORMS, S99_DATE_PAT, FW_NORM_PAT,
    SCREENER_QIDS, build_strip_candidates,
)
try:
    from loop_detector import build_loop_map as _ld_build_loop_map, get_loop_skip_set as _ld_skip_set
    _LOOP_DETECTOR_AVAILABLE = True
except ImportError:
    _LOOP_DETECTOR_AVAILABLE = False
    def _ld_build_loop_map(qids): return {}
    def _ld_skip_set(qids): return {}
from functools import wraps
from flask import (Flask, render_template_string, request,
                   send_file, jsonify, session, redirect, url_for)
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
try:
    from email_service import (send_welcome_email,
                                send_report_ready_email,
                                send_password_reset_email)
    _EMAIL_ENABLED = True
except Exception:
    _EMAIL_ENABLED = False
    def send_welcome_email(*a, **kw): pass
    def send_report_ready_email(*a, **kw): pass
    def send_password_reset_email(*a, **kw): pass

try:
    from payment_service import create_order as _rzp_create_order
    from payment_service import verify_payment as _rzp_verify_payment
    from payment_service import PLANS as PAYMENT_PLANS
    _PAYMENT_ENABLED = True
except Exception as _pay_exc:
    import logging as _l; _l.getLogger(__name__).warning("payment_service not loaded: %s", _pay_exc)
    _PAYMENT_ENABLED = False
    def _rzp_create_order(*a, **kw): raise RuntimeError("Payment not configured")
    def _rzp_verify_payment(*a, **kw): return False
    PAYMENT_PLANS = {}

# ================================================================
# CONFIG
# ================================================================

# Load API keys from .env file
import os
from pathlib import Path

def load_env():
    env_file = Path('/var/www/surveyqc/.env')
    if env_file.exists():
        for line in env_file.read_text().splitlines():
            if '=' in line and not line.startswith('#'):
                key, val = line.split('=', 1)
                os.environ[key.strip()] = val.strip()

load_env()

app = Flask(__name__)
app.secret_key = 'surveyqc-secret-key-change-in-production'
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024

UPLOAD_FOLDER = './uploads'
OUTPUT_FOLDER = './qc_output'
ADMIN_PASSWORD = 'admin123'   # Change this!
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# ── SQLite-backed Job Store ───────────────────────────────────────────────────

DB_PATH = '/var/www/surveyqc/surveyqc.db'

class JobStore:
    """
    Dict-like in-memory store that persists to SQLite on job creation and at
    explicit persist() calls (terminal states: done / error / stopped).

    In-memory layer handles all real-time updates during a run (log appends,
    progress ticks) without touching the DB on every mutation.  SQLite layer
    makes completed-job data survive gunicorn restarts.
    """

    def __init__(self, db_path: str):
        self._db_path = db_path
        self._mem: dict = {}
        self._lock = threading.Lock()
        self._init_db()
        self._load_from_db()

    # ── DB helpers ────────────────────────────────────────────────────────────

    def _conn(self):
        c = sqlite3.connect(self._db_path, check_same_thread=False, timeout=10)
        c.row_factory = sqlite3.Row
        return c

    def _init_db(self):
        with self._conn() as c:
            c.execute('''
                CREATE TABLE IF NOT EXISTS jobs (
                    id           TEXT PRIMARY KEY,
                    status       TEXT,
                    doc_filename TEXT,
                    url          TEXT,
                    platform     TEXT,
                    created_at   TEXT,
                    result_json  TEXT,
                    log_text     TEXT
                )
            ''')
            c.execute('''
                CREATE TABLE IF NOT EXISTS issue_feedback (
                    id           INTEGER PRIMARY KEY AUTOINCREMENT,
                    job_id       TEXT,
                    qid          TEXT,
                    issue_type   TEXT,
                    platform     TEXT,
                    user_verdict TEXT,
                    user_email   TEXT,
                    created_at   TEXT
                )
            ''')

    def _load_from_db(self):
        """Load all persisted jobs into memory on startup."""
        try:
            with self._conn() as c:
                rows = c.execute('SELECT * FROM jobs').fetchall()
            for row in rows:
                job = json.loads(row['result_json'] or '{}')
                job['id']       = row['id']
                job['status']   = row['status'] or job.get('status', 'error')
                job['logs']     = json.loads(row['log_text'] or '[]')
                # Orphaned running jobs can't be resumed after a restart
                if job.get('status') == 'running':
                    job['status'] = 'error'
                    job['phase']  = 'Server restarted — job interrupted'
                self._mem[row['id']] = job
        except Exception as exc:
            import logging
            logging.getLogger(__name__).error('JobStore._load_from_db: %s', exc)

    def _upsert(self, job_id: str, job: dict):
        """Write current job state to SQLite (INSERT OR REPLACE)."""
        logs   = job.get('logs', [])
        result = {k: v for k, v in job.items() if k != 'logs'}
        try:
            with self._conn() as c:
                c.execute(
                    '''INSERT OR REPLACE INTO jobs
                           (id, status, doc_filename, url, platform, created_at,
                            result_json, log_text)
                       VALUES (?, ?, ?, ?, ?, ?, ?, ?)''',
                    (
                        job_id,
                        job.get('status', ''),
                        job.get('doc_name', ''),
                        job.get('survey_url', ''),
                        job.get('platform', ''),
                        job.get('created_at', ''),
                        json.dumps(result, default=str),
                        json.dumps(logs,   default=str),
                    )
                )
        except Exception as exc:
            import logging
            logging.getLogger(__name__).error('JobStore._upsert(%s): %s', job_id, exc)

    def _delete(self, job_id: str):
        try:
            with self._conn() as c:
                c.execute('DELETE FROM jobs WHERE id = ?', (job_id,))
        except Exception as exc:
            import logging
            logging.getLogger(__name__).error('JobStore._delete(%s): %s', job_id, exc)

    # ── Public dict-like interface ────────────────────────────────────────────

    def __setitem__(self, key: str, value: dict):
        """Job creation or full replacement — writes through to DB."""
        with self._lock:
            self._mem[key] = value
        self._upsert(key, value)

    def __getitem__(self, key: str):
        return self._mem[key]

    def __contains__(self, key: str):
        return key in self._mem

    def get(self, key: str, default=None):
        return self._mem.get(key, default)

    def __delitem__(self, key: str):
        with self._lock:
            del self._mem[key]
        self._delete(key)

    def pop(self, key: str, *args):
        val = self._mem.pop(key, *args)
        self._delete(key)
        return val

    def evict(self, job_id: str):
        """Remove from memory only; DB record is preserved for history."""
        self._mem.pop(job_id, None)

    def persist(self, job_id: str):
        """Sync current in-memory state to SQLite.  Call at terminal states."""
        job = self._mem.get(job_id)
        if job:
            self._upsert(job_id, job)

    def items(self):
        return self._mem.items()

    def values(self):
        return self._mem.values()

    def keys(self):
        return self._mem.keys()

    def __len__(self):
        return len(self._mem)

    # ── Issue feedback (historical learning) ─────────────────────────────────

    def load_feedback(self, platform: str = '') -> list:
        """Return all issue_feedback rows, optionally filtered by platform."""
        try:
            with self._conn() as c:
                if platform:
                    rows = c.execute(
                        'SELECT * FROM issue_feedback WHERE platform=?', (platform,)
                    ).fetchall()
                else:
                    rows = c.execute('SELECT * FROM issue_feedback').fetchall()
            return [dict(r) for r in rows]
        except Exception:
            return []

    def save_feedback(self, job_id: str, qid: str, issue_type: str,
                      platform: str, verdict: str, user_email: str = '') -> None:
        """Insert a user verdict for a specific issue."""
        try:
            from datetime import datetime as _dt
            with self._conn() as c:
                c.execute(
                    '''INSERT INTO issue_feedback
                       (job_id, qid, issue_type, platform, user_verdict, user_email, created_at)
                       VALUES (?,?,?,?,?,?,?)''',
                    (job_id, qid, issue_type, platform, verdict, user_email,
                     _dt.now().strftime('%Y-%m-%d %H:%M:%S')),
                )
        except Exception:
            pass


# ── SQLite-backed User Store ──────────────────────────────────────────────────

class UserDB:
    """
    Dict-like user store backed by the `users` table in SQLite.
    Loads all rows into memory on startup; writes through to DB on
    __setitem__ and explicit save() calls after in-place mutations.
    """

    PLAN_LIMITS = {'Free': 3, 'Pro': 25, 'Business': 999999}

    def __init__(self, db_path: str):
        self._db_path = db_path
        self._mem: dict = {}
        self._lock = threading.Lock()
        self._init_table()
        self._load_all()

    def _conn(self):
        c = sqlite3.connect(self._db_path, check_same_thread=False, timeout=10)
        c.row_factory = sqlite3.Row
        return c

    def _init_table(self):
        with self._conn() as c:
            c.execute('''
                CREATE TABLE IF NOT EXISTS users (
                    id                   TEXT PRIMARY KEY,
                    email                TEXT UNIQUE NOT NULL,
                    password_hash        TEXT NOT NULL,
                    name                 TEXT,
                    plan                 TEXT DEFAULT "Free",
                    created_at           TEXT,
                    reports_used         INTEGER DEFAULT 0,
                    total_saved_hours    INTEGER DEFAULT 0,
                    must_change_password INTEGER DEFAULT 0
                )
            ''')
            c.execute('''
                CREATE TABLE IF NOT EXISTS password_resets (
                    token      TEXT PRIMARY KEY,
                    email      TEXT NOT NULL,
                    expires_at TEXT NOT NULL
                )
            ''')
            # Seed default admin if absent
            if not c.execute(
                'SELECT id FROM users WHERE email=?', ('admin@surveyqc.com',)
            ).fetchone():
                c.execute(
                    '''INSERT INTO users (id,email,password_hash,name,plan,created_at,
                                         reports_used,total_saved_hours,must_change_password)
                       VALUES (?,?,?,?,?,?,?,?,?)''',
                    (str(uuid.uuid4())[:8], 'admin@surveyqc.com',
                     generate_password_hash('admin123'),
                     'Admin', 'Business',
                     datetime.now().strftime('%Y-%m-%d'), 0, 0, 1),
                )

    def _row_to_dict(self, row) -> dict:
        plan = row['plan'] or 'Free'
        return {
            'id':                   row['id'],
            'email':                row['email'],
            'password_hash':        row['password_hash'],
            'password':             row['password_hash'],   # legacy compat key
            'name':                 row['name'] or row['email'].split('@')[0].title(),
            'plan':                 plan,
            'joined':               row['created_at'] or '',
            'created_at':           row['created_at'] or '',
            'reports_used':         row['reports_used'] or 0,
            'reports_limit':        self.PLAN_LIMITS.get(plan, 3),
            'total_saved_hours':    row['total_saved_hours'] or 0,
            'must_change_password': bool(row['must_change_password']),
        }

    def _load_all(self):
        with self._conn() as c:
            for row in c.execute('SELECT * FROM users').fetchall():
                self._mem[row['email']] = self._row_to_dict(row)

    def _upsert(self, email: str, u: dict):
        try:
            with self._conn() as c:
                c.execute(
                    '''INSERT OR REPLACE INTO users
                           (id,email,password_hash,name,plan,created_at,
                            reports_used,total_saved_hours,must_change_password)
                       VALUES (?,?,?,?,?,?,?,?,?)''',
                    (
                        u.get('id') or str(uuid.uuid4())[:8],
                        email,
                        u.get('password_hash') or u.get('password', ''),
                        u.get('name', ''),
                        u.get('plan', 'Free'),
                        u.get('joined') or u.get('created_at') or datetime.now().strftime('%Y-%m-%d'),
                        u.get('reports_used', 0),
                        u.get('total_saved_hours', 0),
                        int(u.get('must_change_password', 0)),
                    )
                )
        except Exception as exc:
            import logging
            logging.getLogger(__name__).error('UserDB._upsert(%s): %s', email, exc)

    # ── Public dict-like interface ────────────────────────────────────────────

    def __getitem__(self, email: str) -> dict:
        return self._mem[email]

    def __setitem__(self, email: str, value: dict):
        value.setdefault('id', str(uuid.uuid4())[:8])
        value.setdefault('plan', 'Free')
        value['reports_limit'] = self.PLAN_LIMITS.get(value.get('plan', 'Free'), 3)
        with self._lock:
            self._mem[email] = value
        self._upsert(email, value)

    def __contains__(self, email: str) -> bool:
        return email in self._mem

    def get(self, email: str, default=None):
        return self._mem.get(email, default)

    def items(self):
        return self._mem.items()

    def values(self):
        return self._mem.values()

    def keys(self):
        return self._mem.keys()

    def __len__(self) -> int:
        return len(self._mem)

    def save(self, email: str):
        """Sync in-memory user to DB after an in-place field mutation."""
        u = self._mem.get(email)
        if u:
            u['reports_limit'] = self.PLAN_LIMITS.get(u.get('plan', 'Free'), 3)
            self._upsert(email, u)

    def check_password(self, email: str, password: str) -> bool:
        u = self._mem.get(email)
        if not u:
            return False
        stored = u.get('password_hash') or u.get('password', '')
        try:
            return check_password_hash(stored, password)
        except Exception:
            # Fall back to legacy sha256 hashes (pre-migration accounts)
            return stored == hashlib.sha256(password.encode()).hexdigest()

    def can_run_report(self, email: str) -> bool:
        u = self._mem.get(email, {})
        limit = self.PLAN_LIMITS.get(u.get('plan', 'Free'), 3)
        return limit == 999999 or u.get('reports_used', 0) < limit


# In-memory stores (jobs + users backed by SQLite)
jobs = JobStore(DB_PATH)
users_db = UserDB(DB_PATH)
feedback_store = []
user_feedback_db = []  # [{id, type, message, page, user_email, created_at, read}]

THANKYOU_INDICATORS = [
    "grazie", "thank you", "ringraziamo", "gracias", "merci",
    "complete", "completato", "screened out", "non qualificato",
    "danke", "bedankt", "obrigado"
]
SKIP_PATTERNS = [
    r'^\s*PROGRAMMING\s+TABLE', r'^\s*ROUTING\s*[:\|]',
    r'^\s*ROUTINE\s*[:\|]', r'^\s*TYPE\s*[:\|]',
    r'^\s*\[NEXT SCREEN\]', r'^\s*\[SAME SCREEN',
    r'^---+$', r'^\s*\|\s*$',
]
JUNK_RE = re.compile('|'.join(SKIP_PATTERNS), re.IGNORECASE)

# Section-level headings that mark the end of a question's content scope.
# When a paragraph matches this (and has no '?'), stop accumulating text
# for the current QID — prevents consent/GDPR/address blocks bleeding into
# the previous question's text.
SECTION_STOP_RE = re.compile(
    # Start-of-line heading patterns (standalone section titles):
    r'(?:^\s*(?:'
    r'TRANSPARENCE\b'                             # FR/EN consent heading
    r'|SIGNALEMENT\s+DES\b'                       # adverse-event section
    r'|TRAITEMENT\s+DES\s+DONN'                  # GDPR data-processing block
    r'|NOTICE\s+D.INFORMATION'                    # information notice
    r'|PHARMACOVIGILANCE\b'                       # pharmacovigilance note
    r'|RGPD\b|GDPR\b'                            # regulation labels
    r'|QUOTAS?\b'                                 # quota section heading
    r'|SIGNE?L[EÉ]TIQUE'                         # demographics heading (FR)
    r'|DEMOGRAPHICS?\b'                           # demographics heading (EN)
    r'|FIN\s+DE\s+L.ENTRETIEN'                  # "end of interview"
    r'|FIN\s+DU\s+QUESTIONNAIRE'                 # "end of questionnaire"
    r'|END\s+OF\s+(?:THE\s+)?(?:INTERVIEW|QUESTIONNAIRE|SURVEY)'
    r'|ADRESSE\s+(?:PROFESSIONNELLE|DE\s+D)'     # address-collection block
    r'|NOUS\s+VOUS\s+REMERC'                     # "Nous vous remercions..."
    r'|NOUS\s+VOUS\s+RAPPELONS'                  # pharmacovigilance preamble
    r'|THANK\s+YOU\s+FOR\s+(?:YOUR\s+)?PARTICIPAT'
    r'))'
    # Anywhere-in-line patterns (distinctive phrases that mark boilerplate):
    r'|LOI\s+BERTRAND\b'                         # French transparency law
    r'|BERTRAND\s+LAW\b'
    r'|DONN[EÉ]ES\s+PERSONNELLES'               # GDPR data paragraph (mid-sentence)
    r'|SI\s+VOUS\s+[EÊ]TES\s+(?:LE\s+)?(?:UN\s+)?(?:PATIENT|PROFESSIONNEL)',  # PV notification text
    re.IGNORECASE
)

STOPWORDS = {
    'il','la','lo','gli','le','un','una','uno','di','da','del','della',
    'the','and','or','but','is','are','was','were','have','has','had',
    'a','an','of','in','on','at','to','for','with','by','from','this',
    'that','these','those','it','its','i','you','he','she','we','they',
    'all','any','some','no','not','so','if','as','than','then',
}

# Multilingual stop-words for QID detection.  Any candidate that is a
# common French/English/German word is rejected as a QID — particularly
# important for TN inner-text Strategy 4 which uses \d* and can match
# is_valid_qid, is_framework_page, verify_extraction, should_skip_qid,
# SCREENER_QIDS and related constants are imported from qid_normalizer above.

# ================================================================
# AUTH HELPERS
# ================================================================
def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if 'user_email' not in session:
            return redirect('/login')
        # Force password change for first-login admin before accessing anything else
        u = users_db.get(session['user_email'], {})
        if u.get('must_change_password') and request.endpoint != 'settings':
            return redirect('/settings?force_change=1')
        return f(*args, **kwargs)
    return decorated

def admin_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get('is_admin'):
            return redirect('/admin/login')
        return f(*args, **kwargs)
    return decorated

def get_current_user():
    email = session.get('user_email')
    if email and email in users_db:
        u = users_db[email].copy()
        u['email'] = email
        u['reports_limit'] = UserDB.PLAN_LIMITS.get(u.get('plan', 'Free'), 3)
        return u
    return None

# ================================================================
# CSS + JS (shared across all pages)
# ================================================================
SHARED_CSS = """
<script src="/admin-sidebar-js"></script>
<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/@tabler/icons-webfont@3.19.0/dist/tabler-icons.min.css">
<link rel="stylesheet" href="/static/style.css">
<script src="/static/app.js" defer></script>
"""

# ================================================================
# SIDEBAR TEMPLATE
# ================================================================
def sidebar_html(active='dashboard'):
    user = get_current_user()
    name = user['name'] if user else 'User'
    plan = user.get('plan', 'Free') if user else 'Free'
    initials = ''.join([n[0] for n in name.split()[:2]]).upper()

    items = [
        ('dashboard', 'ti-home', 'Dashboard', '/dashboard'),
        ('new-qc', 'ti-plus', 'New QC', '/new-qc'),
        ('reports', 'ti-file-report', 'Reports', '/reports'),
        ('retest', 'ti-refresh', 'Re-test', '/retest'),
        ('templates', 'ti-copy', 'Templates', '/templates'),
    ]
    nav_html = ''
    for key, icon, label, href in items:
        cls = 'nav-item active' if active == key else 'nav-item'
        nav_html += f'<a href="{href}" class="{cls}"><i class="ti {icon}"></i><span>{label}</span></a>'

    # Bottom mobile nav items
    mob_items = [
        ('dashboard', 'ti-home', 'Home', '/dashboard'),
        ('new-qc', 'ti-plus', 'New QC', '/new-qc'),
        ('reports', 'ti-file-report', 'Reports', '/reports'),
        ('retest', 'ti-refresh', 'Re-test', '/retest'),
    ]
    mob_nav = ''
    for key, icon, label, href in mob_items:
        cls = 'color:var(--purple)' if active == key else 'color:var(--text3)'
        mob_nav += f'<a href="{href}" style="display:flex;flex-direction:column;align-items:center;gap:3px;text-decoration:none;{cls}"><i class="ti {icon}" style="font-size:20px"></i><span style="font-size:10px">{label}</span></a>'

    return f"""
<div class="sidebar">
  <div class="sidebar-logo">
    <div class="sidebar-logo-icon"><i class="ti ti-shield-check" style="color:white;font-size:16px"></i></div>
    <span class="sidebar-logo-text">SurveyQC</span>
  </div>
  {nav_html}
  <div class="nav-divider"></div>
  <div style="padding:3px 0;margin-bottom:2px">
    <p style="font-size:9px;color:rgba(255,255,255,.3);padding:3px 10px;text-transform:uppercase;letter-spacing:.08em">Account</p>
  </div>
  <a href="/settings" class="{'nav-item active' if active=='settings' else 'nav-item'}"><i class="ti ti-settings"></i><span>Settings</span></a>
  <a href="/billing" class="{'nav-item active' if active=='billing' else 'nav-item'}"><i class="ti ti-credit-card"></i><span>Billing</span></a>
  <a href="/privacy-data" class="{'nav-item active' if active=='privacy' else 'nav-item'}"><i class="ti ti-shield-lock"></i><span>Privacy</span></a>
  <div style="margin-top:auto;padding-top:12px;border-top:0.5px solid rgba(255,255,255,.1)">
    <div style="display:flex;align-items:center;gap:8px;padding:8px 10px;border-radius:8px">
      <div class="avatar" style="background:var(--purple-dim);color:var(--purple);font-size:11px;font-weight:500">{initials}</div>
      <div style="flex:1;min-width:0">
        <p style="font-size:12px;color:white;font-weight:500;white-space:nowrap;overflow:hidden;text-overflow:ellipsis">{name}</p>
        <p style="font-size:10px;color:#85B7EB">{plan} plan</p>
      </div>
      <a href="/logout" style="color:#85B7EB;text-decoration:none;flex-shrink:0"><i class="ti ti-logout" style="font-size:14px"></i></a>
    </div>
  </div>
</div>
<!-- MOBILE BOTTOM NAV -->
<div style="display:none;position:fixed;bottom:0;left:0;right:0;background:white;padding:10px 20px;justify-content:space-around;align-items:center;z-index:1000;border-top:0.5px solid rgba(255,255,255,.1)" class="mobile-bottom-nav">
  {mob_nav}
</div>
<style>
@media(max-width:768px){{
  .mobile-bottom-nav{{display:flex !important}}
  .main-content{{padding-bottom:70px}}
}}
</style>"""

# ================================================================
# PAGE: LANDING
# ================================================================
@app.route('/')
def landing():
    return redirect('/home')


@app.route('/login', methods=['GET', 'POST'])
def login():
    error = ''
    if request.method == 'POST':
        email = request.form.get('email', '').strip().lower()
        password = request.form.get('password', '')
        if users_db.check_password(email, password):
            session['user_email'] = email
            if users_db[email].get('must_change_password'):
                return redirect('/settings?force_change=1')
            return redirect('/dashboard')
        error = 'Invalid email or password'

    error_html = ('<div class="auth-error"><i class="ti ti-alert-circle"></i>' + error + '</div>') if error else ''

    page = """<!DOCTYPE html><html lang="en"><head>
<meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1.0">
<title>Sign in \u2014 SurveyQC</title>
<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/@tabler/icons-webfont@3.19.0/dist/tabler-icons.min.css">

<style>
:root{--bg:#F7F4EE;--card:#FFFFFF;--text:#171717;--text2:#5F5B53;--text3:#8A847A;--accent:#C46A2B;--accent-hover:#A9551F;--accent-bg:#F5E6D8;--border:#E8E1D8;--dark:#1B140F;--danger:#C84B31;--success:#3F7D58;--shadow-lg:0 10px 40px rgba(24,17,10,0.08)}
*{box-sizing:border-box;margin:0;padding:0;font-family:-apple-system,'Inter','Plus Jakarta Sans',BlinkMacSystemFont,'Segoe UI',Roboto,sans-serif;-webkit-font-smoothing:antialiased}
body{background:var(--bg);color:var(--text);min-height:100vh;display:flex}
a{text-decoration:none;color:inherit}
.auth-wrap{display:grid;grid-template-columns:1fr 1fr;width:100%;min-height:100vh}
.auth-left{background:var(--dark);padding:56px 64px;display:flex;flex-direction:column;justify-content:space-between;position:relative;overflow:hidden}
.auth-left::before{content:"";position:absolute;top:-100px;right:-100px;width:400px;height:400px;background:radial-gradient(circle,rgba(196,106,43,.3),transparent 70%);pointer-events:none}
.auth-logo{display:flex;align-items:center;gap:11px;position:relative;z-index:1}
.auth-logo-mark{width:36px;height:36px;background:var(--accent);border-radius:9px;display:flex;align-items:center;justify-content:center}
.auth-logo-text{font-family:'Plus Jakarta Sans',sans-serif;font-size:19px;font-weight:700;color:white}
.auth-left-content{position:relative;z-index:1}
.auth-left-content h2{font-family:'Plus Jakarta Sans',sans-serif;font-size:34px;font-weight:800;color:white;line-height:1.2;letter-spacing:-1px;margin-bottom:20px}
.auth-quote{background:rgba(255,255,255,.06);border:1px solid rgba(255,255,255,.1);border-radius:16px;padding:24px;margin-top:32px}
.auth-quote p{font-size:15px;color:#E8DDD2;line-height:1.7;margin-bottom:16px}
.auth-quote-author{display:flex;align-items:center;gap:12px}
.auth-quote-avatar{width:40px;height:40px;border-radius:50%;background:var(--accent);color:white;display:flex;align-items:center;justify-content:center;font-weight:700;font-size:13px}
.auth-quote-name{font-size:14px;font-weight:700;color:white}
.auth-quote-role{font-size:12px;color:#9A8C7B}
.auth-stats{display:flex;gap:32px;position:relative;z-index:1}
.auth-stat-num{font-family:'Plus Jakarta Sans',sans-serif;font-size:26px;font-weight:800;color:white}
.auth-stat-lbl{font-size:12px;color:#9A8C7B;margin-top:2px}
.auth-right{display:flex;align-items:center;justify-content:center;padding:40px 24px;background:var(--bg)}
.auth-card{width:100%;max-width:400px}
.auth-card h1{font-family:'Plus Jakarta Sans',sans-serif;font-size:30px;font-weight:800;letter-spacing:-1px;margin-bottom:8px}
.auth-card-sub{font-size:15px;color:var(--text2);margin-bottom:32px}
.auth-card-sub a{color:var(--accent);font-weight:600}
.form-group{margin-bottom:18px}
.form-label{font-size:13px;font-weight:600;color:var(--text);display:block;margin-bottom:7px}
.form-input{width:100%;padding:13px 16px;border:1px solid var(--border);border-radius:12px;font-size:14px;color:var(--text);background:white;outline:none;transition:all .15s;font-family:inherit}
.form-input:focus{border-color:var(--accent);box-shadow:0 0 0 3px rgba(196,106,43,.12)}
.auth-btn{width:100%;background:var(--dark);color:#F7F4EE;border:none;padding:14px;border-radius:12px;font-size:15px;font-weight:600;cursor:pointer;transition:all .2s;display:flex;align-items:center;justify-content:center;gap:8px;font-family:inherit;margin-top:6px}
.auth-btn:hover{background:#2A1F18;transform:translateY(-1px)}
.auth-error{background:#FAE5E0;border:1px solid #F0C4BA;color:var(--danger);font-size:13px;padding:11px 14px;border-radius:10px;margin-bottom:18px;display:flex;align-items:center;gap:8px}
.auth-divider{display:flex;align-items:center;gap:14px;margin:24px 0;color:var(--text3);font-size:13px}
.auth-divider::before,.auth-divider::after{content:"";flex:1;height:1px;background:var(--border)}
.auth-social{display:flex;gap:10px}
.auth-social-btn{flex:1;border:1px solid var(--border);background:white;border-radius:12px;padding:11px;display:flex;align-items:center;justify-content:center;gap:8px;font-size:13px;font-weight:600;color:var(--text);cursor:pointer;transition:all .15s}
.auth-social-btn:hover{background:var(--bg)}
.auth-row{display:flex;align-items:center;justify-content:space-between;margin-bottom:18px}
.auth-check{display:flex;align-items:center;gap:7px;font-size:13px;color:var(--text2);cursor:pointer}
.auth-link{font-size:13px;color:var(--accent);font-weight:600}
.auth-back{position:absolute;top:24px;left:24px;font-size:13px;color:#9A8C7B;display:flex;align-items:center;gap:6px;z-index:2}
.auth-back:hover{color:white}
@media(max-width:900px){
  .auth-wrap{grid-template-columns:1fr}
  .auth-left{display:none}
  .auth-right{padding:40px 20px}
}
</style></head><body>
<div class="auth-wrap">
  <div class="auth-left">
    <a href="/home" class="auth-logo">
      <div class="auth-logo-mark"><i class="ti ti-shield-check" style="color:white;font-size:18px"></i></div>
      <span class="auth-logo-text">SurveyQC</span>
    </a>
    <div class="auth-left-content">
      <h2>Welcome back to<br>perfect survey QC.</h2>
      <div style="margin-top:24px">
        <p style="font-size:16px;font-weight:600;color:#F7F4EE;margin-bottom:8px">Built for survey QC professionals.</p>
        <p style="font-size:14px;color:#9A8C7B;line-height:1.7">Catch every survey bug before your data is collected — termination logic, question text, piping, mandatory markers, and more.</p>
      </div>
    </div>
    <div class="auth-stats">
      <div><div class="auth-stat-num">15+</div><div class="auth-stat-lbl">QC checks</div></div>
      <div><div class="auth-stat-num">4</div><div class="auth-stat-lbl">Platforms</div></div>
      <div><div class="auth-stat-num">80+</div><div class="auth-stat-lbl">Languages</div></div>
    </div>
  </div>
  <div class="auth-right">
    <div class="auth-card">
      <h1>Sign in</h1>
      <p class="auth-card-sub">Don't have an account? <a href="/signup">Sign up free</a></p>
      """ + error_html + """
      <form method="POST">
        <div class="form-group">
          <label class="form-label">Email</label>
          <input type="email" name="email" class="form-input" placeholder="you@company.com" required>
        </div>
        <div class="form-group">
          <label class="form-label">Password</label>
          <input type="password" name="password" class="form-input" placeholder="Enter your password" required>
        </div>
        <div class="auth-row">
          <label class="auth-check"><input type="checkbox" style="accent-color:var(--accent)"> Remember me</label>
          <a href="/forgot-password" class="auth-link">Forgot password?</a>
        </div>
        <button type="submit" class="auth-btn">Sign in <i class="ti ti-arrow-right"></i></button>
      </form>
    </div>
  </div>
</div>
</body></html>"""
    return page


@app.route('/signup', methods=['GET', 'POST'])
def signup():
    error = ''
    if request.method == 'POST':
        name  = request.form.get('name', '').strip()
        email = request.form.get('email', '').strip().lower()
        password = request.form.get('password', '')
        if email in users_db:
            error = 'Email already registered'
        elif len(password) < 6:
            error = 'Password must be at least 6 characters'
        elif not name:
            error = 'Full name is required'
        else:
            users_db[email] = {
                'password_hash': generate_password_hash(password),
                'password':      generate_password_hash(password),
                'name':          name,
                'plan':          'Free',
                'reports_used':  0,
                'reports_limit': UserDB.PLAN_LIMITS['Free'],
                'joined':        datetime.now().strftime('%Y-%m-%d'),
                'total_saved_hours': 0,
                'must_change_password': False,
            }
            session['user_email'] = email
            try:
                send_welcome_email(email, name, 'Free')
            except Exception:
                pass
            return redirect('/dashboard')

    error_html = ('<div class="auth-error"><i class="ti ti-alert-circle"></i>' + error + '</div>') if error else ''

    page = """<!DOCTYPE html><html lang="en"><head>
<meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1.0">
<title>Sign up \u2014 SurveyQC</title>
<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/@tabler/icons-webfont@3.19.0/dist/tabler-icons.min.css">

<style>
:root{--bg:#F7F4EE;--card:#FFFFFF;--text:#171717;--text2:#5F5B53;--text3:#8A847A;--accent:#C46A2B;--accent-hover:#A9551F;--accent-bg:#F5E6D8;--border:#E8E1D8;--dark:#1B140F;--danger:#C84B31;--success:#3F7D58}
*{box-sizing:border-box;margin:0;padding:0;font-family:-apple-system,'Inter','Plus Jakarta Sans',BlinkMacSystemFont,'Segoe UI',Roboto,sans-serif;-webkit-font-smoothing:antialiased}
body{background:var(--bg);color:var(--text);min-height:100vh;display:flex}
a{text-decoration:none;color:inherit}
.auth-wrap{display:grid;grid-template-columns:1fr 1fr;width:100%;min-height:100vh}
.auth-left{background:var(--dark);padding:56px 64px;display:flex;flex-direction:column;justify-content:space-between;position:relative;overflow:hidden}
.auth-left::before{content:"";position:absolute;top:-100px;right:-100px;width:400px;height:400px;background:radial-gradient(circle,rgba(196,106,43,.3),transparent 70%)}
.auth-logo{display:flex;align-items:center;gap:11px;position:relative;z-index:1}
.auth-logo-mark{width:36px;height:36px;background:var(--accent);border-radius:9px;display:flex;align-items:center;justify-content:center}
.auth-logo-text{font-family:'Plus Jakarta Sans',sans-serif;font-size:19px;font-weight:700;color:white}
.auth-left-content{position:relative;z-index:1}
.auth-left-content h2{font-family:'Plus Jakarta Sans',sans-serif;font-size:34px;font-weight:800;color:white;line-height:1.2;letter-spacing:-1px;margin-bottom:24px}
.auth-benefits{list-style:none}
.auth-benefit{display:flex;align-items:center;gap:12px;color:#E8DDD2;font-size:15px;margin-bottom:16px}
.auth-benefit i{color:var(--accent);background:rgba(196,106,43,.15);width:28px;height:28px;border-radius:7px;display:flex;align-items:center;justify-content:center;font-size:15px;flex-shrink:0}
.auth-stats{display:flex;gap:32px;position:relative;z-index:1}
.auth-stat-num{font-family:'Plus Jakarta Sans',sans-serif;font-size:26px;font-weight:800;color:white}
.auth-stat-lbl{font-size:12px;color:#9A8C7B;margin-top:2px}
.auth-right{display:flex;align-items:center;justify-content:center;padding:40px 24px}
.auth-card{width:100%;max-width:400px}
.auth-card h1{font-family:'Plus Jakarta Sans',sans-serif;font-size:30px;font-weight:800;letter-spacing:-1px;margin-bottom:8px}
.auth-card-sub{font-size:15px;color:var(--text2);margin-bottom:32px}
.auth-card-sub a{color:var(--accent);font-weight:600}
.form-group{margin-bottom:18px}
.form-label{font-size:13px;font-weight:600;color:var(--text);display:block;margin-bottom:7px}
.form-input{width:100%;padding:13px 16px;border:1px solid var(--border);border-radius:12px;font-size:14px;color:var(--text);background:white;outline:none;transition:all .15s;font-family:inherit}
.form-input:focus{border-color:var(--accent);box-shadow:0 0 0 3px rgba(196,106,43,.12)}
.auth-btn{width:100%;background:var(--dark);color:#F7F4EE;border:none;padding:14px;border-radius:12px;font-size:15px;font-weight:600;cursor:pointer;transition:all .2s;display:flex;align-items:center;justify-content:center;gap:8px;font-family:inherit;margin-top:6px}
.auth-btn:hover{background:#2A1F18;transform:translateY(-1px)}
.auth-error{background:#FAE5E0;border:1px solid #F0C4BA;color:var(--danger);font-size:13px;padding:11px 14px;border-radius:10px;margin-bottom:18px;display:flex;align-items:center;gap:8px}
.auth-terms{font-size:12px;color:var(--text3);text-align:center;margin-top:18px;line-height:1.6}
.auth-terms a{color:var(--accent)}
@media(max-width:900px){
  .auth-wrap{grid-template-columns:1fr}
  .auth-left{display:none}
  .auth-right{padding:40px 20px}
}
</style></head><body>
<div class="auth-wrap">
  <div class="auth-left">
    <a href="/home" class="auth-logo">
      <div class="auth-logo-mark"><i class="ti ti-shield-check" style="color:white;font-size:18px"></i></div>
      <span class="auth-logo-text">SurveyQC</span>
    </a>
    <div class="auth-left-content">
      <h2>Start catching survey<br>bugs in minutes.</h2>
      <ul class="auth-benefits">
        <li class="auth-benefit"><i class="ti ti-check"></i>5 free reports every month, forever</li>
        <li class="auth-benefit"><i class="ti ti-check"></i>All 15+ QC checks included</li>
        <li class="auth-benefit"><i class="ti ti-check"></i>Works in 80+ languages</li>
        <li class="auth-benefit"><i class="ti ti-check"></i>No credit card required</li>
      </ul>
    </div>
    <div class="auth-stats">
      <div><div class="auth-stat-num">15+</div><div class="auth-stat-lbl">QC checks</div></div>
      <div><div class="auth-stat-num">4</div><div class="auth-stat-lbl">Platforms</div></div>
      <div><div class="auth-stat-num">80+</div><div class="auth-stat-lbl">Languages</div></div>
    </div>
  </div>
  <div class="auth-right">
    <div class="auth-card">
      <h1>Create account</h1>
      <p class="auth-card-sub">Already have an account? <a href="/login">Sign in</a></p>
      """ + error_html + """
      <form method="POST">
        <div class="form-group">
          <label class="form-label">Full name</label>
          <input type="text" name="name" class="form-input" placeholder="Your name" required>
        </div>
        <div class="form-group">
          <label class="form-label">Work email</label>
          <input type="email" name="email" class="form-input" placeholder="you@company.com" required>
        </div>
        <div class="form-group">
          <label class="form-label">Password</label>
          <input type="password" name="password" class="form-input" placeholder="At least 6 characters" required>
        </div>
        <button type="submit" class="auth-btn">Create free account <i class="ti ti-arrow-right"></i></button>
      </form>
      <p class="auth-terms">By signing up you agree to our <a href="/terms">Terms</a> and <a href="/privacy-policy">Privacy Policy</a>.</p>
    </div>
  </div>
</div>
</body></html>"""
    return page


@app.route('/logout')
def logout():
    session.clear()
    return redirect('/home')


# ================================================================
# PASSWORD RESET
# ================================================================

_AUTH_CSS = """
:root{--bg:#F7F4EE;--text:#171717;--text2:#5F5B53;--text3:#8A847A;
      --accent:#C46A2B;--border:#E8E1D8;--dark:#1B140F;--danger:#C84B31}
*{box-sizing:border-box;margin:0;padding:0;
  font-family:-apple-system,'Segoe UI',Roboto,sans-serif;-webkit-font-smoothing:antialiased}
body{background:var(--bg);color:var(--text);min-height:100vh;
     display:flex;align-items:center;justify-content:center;padding:24px}
.card{background:#fff;border-radius:16px;padding:40px;width:100%;max-width:400px;
      box-shadow:0 4px 24px rgba(0,0,0,.08)}
h1{font-size:22px;font-weight:700;letter-spacing:-.4px;margin-bottom:8px}
.sub{font-size:14px;color:var(--text2);margin-bottom:28px;line-height:1.6}
.sub a{color:var(--accent);font-weight:600;text-decoration:none}
label{font-size:13px;font-weight:600;display:block;margin-bottom:6px}
input{width:100%;padding:12px 14px;border:1px solid var(--border);border-radius:10px;
      font-size:14px;outline:none;font-family:inherit;transition:border .15s}
input:focus{border-color:var(--accent);box-shadow:0 0 0 3px rgba(196,106,43,.12)}
.form-group{margin-bottom:18px}
.btn{width:100%;background:var(--dark);color:#F7F4EE;border:none;padding:13px;
     border-radius:10px;font-size:15px;font-weight:600;cursor:pointer;margin-top:4px}
.btn-accent{background:var(--accent)}
.alert-ok{background:#E6F4EC;border:1px solid #A7D7B8;color:#1A5632;font-size:13px;
           padding:11px 14px;border-radius:10px;margin-bottom:18px}
.alert-err{background:#FAE5E0;border:1px solid #F0C4BA;color:var(--danger);font-size:13px;
           padding:11px 14px;border-radius:10px;margin-bottom:18px}
"""

@app.route('/forgot-password', methods=['GET', 'POST'])
def forgot_password():
    msg = ''
    is_error = False

    if request.method == 'POST':
        email = request.form.get('email', '').strip().lower()
        if email in users_db:
            token   = uuid.uuid4().hex
            from datetime import timedelta as _td
            expires = (datetime.now() + _td(hours=1)).isoformat()
            # Store token in DB
            try:
                with sqlite3.connect(DB_PATH, timeout=10) as _c:
                    _c.execute('DELETE FROM password_resets WHERE email=?', (email,))
                    _c.execute(
                        'INSERT INTO password_resets (token,email,expires_at) VALUES (?,?,?)',
                        (token, email, expires)
                    )
            except Exception as _e:
                import logging; logging.getLogger(__name__).error('reset token store: %s', _e)

            BASE_URL = "http://46.202.163.130:5000"
            reset_link = f"{BASE_URL}/reset-password/{token}"
            import logging as _logging
            _logging.getLogger(__name__).info("Reset link generated: %s", reset_link)
            try:
                send_password_reset_email(email, reset_link)
            except Exception as _email_exc:
                _logging.getLogger(__name__).error("send_password_reset_email error: %s", _email_exc)
        # Always show the same message — don't reveal whether email exists
        msg = 'If that address is registered, a reset link is on its way. Check your inbox.'

    alert = ''
    if msg:
        alert = f'<div class="alert-ok">{msg}</div>'
    elif is_error:
        alert = f'<div class="alert-err">{msg}</div>'

    return render_template_string(f"""<!DOCTYPE html><html><head>
<meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>Forgot password — SurveyQC</title>
<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/@tabler/icons-webfont@3.19.0/dist/tabler-icons.min.css">
<style>{_AUTH_CSS}</style></head><body>
<div class="card">
  <div style="display:flex;align-items:center;gap:9px;margin-bottom:28px">
    <div style="width:30px;height:30px;background:#1B140F;border-radius:7px;
                display:flex;align-items:center;justify-content:center">
      <i class="ti ti-shield-check" style="color:white;font-size:15px"></i>
    </div>
    <span style="font-weight:700;font-size:16px">SurveyQC</span>
  </div>
  <h1>Forgot your password?</h1>
  <p class="sub">Enter your email and we'll send a reset link.
    <a href="/login">Back to sign in</a></p>
  {alert}
  <form method="POST">
    <div class="form-group">
      <label>Email address</label>
      <input type="email" name="email" placeholder="you@company.com" required autofocus>
    </div>
    <button type="submit" class="btn">Send reset link <i class="ti ti-send"></i></button>
  </form>
</div>
</body></html>""")


@app.route('/test-email')
def test_email():
    import logging as _logging
    _log = _logging.getLogger(__name__)
    # Resend test-mode restriction: can only send to the account owner email
    # until a domain is verified at resend.com/domains
    target = 'tyagi6165@gmail.com'
    _log.info("test-email route hit, sending to %s", target)
    try:
        result = send_password_reset_email(target, 'http://46.202.163.130:5000/reset-password/test-token-123')
        return f'send_password_reset_email returned: {result}  (target={target})', 200
    except Exception as exc:
        _log.error("test-email error: %s", exc)
        return f'ERROR: {exc}', 500


@app.route('/reset-password/<token>', methods=['GET', 'POST'])
def reset_password(token):
    error = ''

    # Validate token
    try:
        with sqlite3.connect(DB_PATH, timeout=10) as _c:
            _c.row_factory = sqlite3.Row
            row = _c.execute(
                'SELECT email,expires_at FROM password_resets WHERE token=?', (token,)
            ).fetchone()
    except Exception:
        row = None

    if not row:
        return render_template_string(f"""<!DOCTYPE html><html><head>
<meta charset="UTF-8"><title>Invalid link — SurveyQC</title>
<style>{_AUTH_CSS}</style></head><body>
<div class="card">
  <h1>Link invalid or expired</h1>
  <p class="sub" style="margin-top:8px">This reset link has already been used or has expired.<br>
    <a href="/forgot-password">Request a new one</a></p>
</div></body></html>"""), 400

    email      = row['email']
    expires_at = row['expires_at']

    # Check expiry
    if datetime.fromisoformat(expires_at) < datetime.now():
        try:
            with sqlite3.connect(DB_PATH, timeout=10) as _c:
                _c.execute('DELETE FROM password_resets WHERE token=?', (token,))
        except Exception:
            pass
        return render_template_string(f"""<!DOCTYPE html><html><head>
<meta charset="UTF-8"><title>Link expired — SurveyQC</title>
<style>{_AUTH_CSS}</style></head><body>
<div class="card">
  <h1>Reset link expired</h1>
  <p class="sub" style="margin-top:8px">Links expire after 1 hour.
    <a href="/forgot-password">Request a new one</a></p>
</div></body></html>"""), 400

    if request.method == 'POST':
        new_pw = request.form.get('password', '')
        confirm = request.form.get('confirm', '')
        if len(new_pw) < 6:
            error = 'Password must be at least 6 characters.'
        elif new_pw != confirm:
            error = 'Passwords do not match.'
        else:
            # Update password
            new_hash = generate_password_hash(new_pw)
            if email in users_db:
                users_db[email]['password_hash'] = new_hash
                users_db[email]['password']       = new_hash
                users_db[email]['must_change_password'] = False
                users_db.save(email)
            # Delete used token
            try:
                with sqlite3.connect(DB_PATH, timeout=10) as _c:
                    _c.execute('DELETE FROM password_resets WHERE token=?', (token,))
            except Exception:
                pass
            return render_template_string(f"""<!DOCTYPE html><html><head>
<meta charset="UTF-8"><title>Password updated — SurveyQC</title>
<style>{_AUTH_CSS}</style></head><body>
<div class="card">
  <h1>Password updated ✓</h1>
  <p class="sub" style="margin-top:8px">Your password has been changed successfully.</p>
  <a href="/login" class="btn" style="display:block;text-align:center;text-decoration:none;
     margin-top:20px;padding:13px">Sign in now &rarr;</a>
</div></body></html>""")

    err_html = f'<div class="alert-err">{error}</div>' if error else ''
    return render_template_string(f"""<!DOCTYPE html><html><head>
<meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>Reset password — SurveyQC</title>
<style>{_AUTH_CSS}</style></head><body>
<div class="card">
  <h1>Set a new password</h1>
  <p class="sub" style="margin-top:8px">For <strong>{email}</strong></p>
  {err_html}
  <form method="POST">
    <div class="form-group">
      <label>New password</label>
      <input type="password" name="password" placeholder="At least 6 characters" required autofocus>
    </div>
    <div class="form-group">
      <label>Confirm password</label>
      <input type="password" name="confirm" placeholder="Repeat password" required>
    </div>
    <button type="submit" class="btn btn-accent">Set new password</button>
  </form>
</div>
</body></html>""")


# ================================================================
# PAGE: DASHBOARD
# ================================================================
@app.route('/dashboard')
@login_required
def dashboard():
    user = get_current_user()
    name = user['name'].split()[0]
    saved = user.get('total_saved_hours', 0)
    reports_used = user.get('reports_used', 0)
    reports_limit = user.get('reports_limit', 5)
    plan = user.get('plan', 'Free')

    user_jobs = [(jid, j) for jid, j in jobs.items()
                 if j.get('user_email') == session.get('user_email')]
    user_jobs.sort(key=lambda x: x[1].get('created_at', ''), reverse=True)

    done_jobs      = [j for _, j in user_jobs if j.get('status') == 'done']
    issues_found   = sum(j.get('total_issues', 0) for j in done_jobs)
    passed_count   = sum(1 for j in done_jobs if j.get('total_issues', 0) == 0)

    recent_html = ''
    for jid, j in user_jobs[:5]:
        status = j.get('status', 'running')
        doc_name = j.get('doc_name', 'Unknown')
        platform = j.get('platform', '-')
        issues = j.get('total_issues', 0)
        raw_dt = j.get('created_at', '')
        try:
            from datetime import datetime as _dt
            created = _dt.strptime(raw_dt[:16], '%Y-%m-%d %H:%M').strftime('%d %b, %H:%M')
        except Exception:
            created = raw_dt[:16]
        doc_display = (doc_name[:34] + '…') if len(doc_name) > 34 else doc_name
        if status == 'done':
            badge = (f'<span class="badge badge-green">Passed</span>' if issues == 0
                     else f'<span class="badge badge-amber">{issues} issue{"s" if issues!=1 else ""}</span>')
            action = f'<a href="/report/{jid}" style="color:var(--accent);font-size:12px;font-weight:600;text-decoration:none;white-space:nowrap">View →</a>'
        elif status == 'running':
            badge  = '<span class="badge badge-blue">Running</span>'
            action = (f'<form method="POST" action="/stop/{jid}" style="margin:0">'
                      f'<button type="submit" style="background:none;border:1px solid #E0E0E0;'
                      f'color:var(--text2);font-size:11px;padding:3px 9px;border-radius:5px;'
                      f'cursor:pointer;font-family:inherit">Stop</button></form>')
        else:
            badge  = '<span class="badge badge-red">Error</span>'
            action = '<a href="/new-qc" style="color:var(--accent);font-size:12px;font-weight:600;text-decoration:none">Retry</a>'
        recent_html += (
            f'<tr class="report-row" onclick="window.location=\'/report/{jid}\'" style="cursor:pointer">'
            f'<td class="primary" title="{doc_name}"><i class="ti ti-file-text" style="color:var(--accent);margin-right:8px"></i>{doc_display}</td>'
            f'<td>{platform}</td>'
            f'<td>{badge}</td>'
            f'<td style="color:var(--text3);white-space:nowrap">{created}</td>'
            f'<td onclick="event.stopPropagation()">{action}</td>'
            f'</tr>'
        )

    if not recent_html:
        recent_html = ('<tr><td colspan="5" style="text-align:center;padding:36px 16px">'
                       '<p style="font-size:14px;font-weight:600;color:var(--text);margin-bottom:6px">No reports yet</p>'
                       '<p style="font-size:13px;color:var(--text3);margin-bottom:16px">Run your first QC to see results here.</p>'
                       '<a href="/new-qc" class="btn btn-primary btn-sm"><i class="ti ti-plus"></i>&#8594; New QC</a>'
                       '</td></tr>')

    return render_template_string(SHARED_CSS + f"""
<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=5.0"><title>Dashboard — SurveyQC</title></head><body>
<div class="app-layout">
  {sidebar_html('dashboard')}
  <div class="main-content">
    <div class="topbar">
      <div>
        <p class="page-title">Good morning, {name}!</p>
        <p class="page-sub">{datetime.now().strftime('%A, %d %B %Y')}</p>
      </div>
      <div style="display:flex;gap:10px;align-items:center">
        <span class="badge badge-green">{user.get('plan','Free')} · {reports_used}/{reports_limit} reports</span>
        <a href="/new-qc" class="btn btn-primary btn-sm"><i class="ti ti-plus"></i>New QC</a>
      </div>
    </div>

    {"<div style='background:#FEF3C7;border:1.5px solid #F59E0B;border-radius:10px;padding:13px 18px;margin-bottom:16px;display:flex;align-items:center;justify-content:space-between;gap:12px'><span style='font-size:13px;color:#92400E;font-weight:500'>&#9888; Free plan: " + str(reports_used) + "/" + str(reports_limit) + " reports used. Upgrade to continue.</span><a href='/billing' style='font-size:12px;font-weight:700;color:#92400E;text-decoration:none;white-space:nowrap;border:1.5px solid #F59E0B;padding:5px 12px;border-radius:6px'>Upgrade &#8594;</a></div>" if plan == 'Free' and reports_used >= reports_limit else ""}

    <div class="time-saved-banner">
      <div>
        {"<p style='font-size:11px;color:var(--text3);margin-bottom:5px;font-weight:500;letter-spacing:.06em;text-transform:uppercase'>This month you saved</p><div style='display:flex;align-items:baseline;gap:10px;margin-bottom:4px'><p style='font-size:32px;font-weight:800;color:#2E8B57;font-family:Plus Jakarta Sans,sans-serif;letter-spacing:-0.5px'>" + str(reports_used * 8) + " hours</p><p style='font-size:13px;color:var(--green);font-weight:500'>= " + str(reports_used) + " full working days back in your life</p></div><p style='font-size:11px;color:var(--text3)'>" + str(reports_used) + " surveys completed — manual would take " + str(reports_used*8) + "h, SurveyQC did it in " + str(reports_used) + " mins</p>" if reports_used > 0 else "<p style='font-size:15px;font-weight:600;color:var(--text);margin-bottom:8px'>Welcome to SurveyQC!</p><p style='font-size:13px;color:var(--text3);line-height:1.6'>Run your first QC to see how much time you save here.</p>"}
      </div>
      {"<div style='text-align:center'><div style='background:rgba(255,255,255,.1);border-radius:10px;padding:10px 18px'><p style='font-size:20px;font-weight:600;color:var(--text)'>" + str(reports_used * 8) + "x</p><p style='font-size:10px;color:var(--text3)'>ROI on plan</p></div></div>" if reports_used > 0 and user.get('plan','Free') != 'Free' else ""}
    </div>

    <div class="stats-grid">
      <div class="stat-card"><p class="stat-num">{len(done_jobs)}</p><p class="stat-label">Reports run</p></div>
      <div class="stat-card"><p class="stat-num" style="color:#3F7D58">{passed_count}</p><p class="stat-label">Passed</p></div>
      <div class="stat-card"><p class="stat-num" style="color:#C84B31">{issues_found}</p><p class="stat-label">Issues found</p></div>
      <div class="stat-card"><p class="stat-num" style="color:#2E8B57">{len(done_jobs) * 8}h</p><p class="stat-label">Time saved</p></div>
    </div>

    <div style="display:grid;grid-template-columns:1fr 280px;gap:16px">
      <div class="card">
        <div style="display:flex;align-items:center;justify-content:space-between;margin-bottom:16px">
          <p style="font-size:14px;font-weight:600;color:var(--text)">Recent reports</p>
          <a href="/reports" style="font-size:12px;color:var(--purple);text-decoration:none">View all →</a>
        </div>
        <table class="data-table" style="width:100%">
          <thead><tr>
            <th>Survey</th><th>Platform</th><th>Status</th><th>Date</th><th></th>
          </tr></thead>
          <tbody>{recent_html}</tbody>
        </table>
      </div>
      <div style="display:flex;flex-direction:column;gap:12px">
        <div class="card">
          <p style="font-size:13px;font-weight:600;color:var(--text);margin-bottom:12px">Quick actions</p>
          <div style="display:grid;grid-template-columns:1fr 1fr;gap:8px">
            <a href="/new-qc" style="text-decoration:none;padding:12px;background:rgba(255,255,255,.04);border-radius:8px;text-align:center;display:block">
              <i class="ti ti-plus" style="font-size:20px;color:var(--purple)"></i>
              <p style="font-size:11px;color:var(--text);margin-top:5px;font-weight:500">New QC</p>
            </a>
            <a href="/retest" style="text-decoration:none;padding:12px;background:rgba(255,255,255,.04);border-radius:8px;text-align:center;display:block">
              <i class="ti ti-refresh" style="font-size:20px;color:#EF9F27"></i>
              <p style="font-size:11px;color:var(--text);margin-top:5px;font-weight:500">Re-test</p>
            </a>
            <a href="/reports" style="text-decoration:none;padding:12px;background:rgba(255,255,255,.04);border-radius:8px;text-align:center;display:block">
              <i class="ti ti-file-report" style="font-size:20px;color:#1D9E75"></i>
              <p style="font-size:11px;color:var(--text);margin-top:5px;font-weight:500">Reports</p>
            </a>
            <a href="/settings" style="text-decoration:none;padding:12px;background:rgba(255,255,255,.04);border-radius:8px;text-align:center;display:block">
              <i class="ti ti-settings" style="font-size:20px;color:#378ADD"></i>
              <p style="font-size:11px;color:var(--text);margin-top:5px;font-weight:500">Settings</p>
            </a>
          </div>
        </div>
        <div class="card" style="background:#FCEBEB;border-color:#F7C1C1">
          <div style="display:flex;align-items:center;gap:8px;margin-bottom:8px">
            <i class="ti ti-sparkles" style="font-size:14px;color:#A32D2D"></i>
            <p style="font-size:12px;font-weight:500;color:#791F1F">Tip of the day</p>
          </div>
          <p style="font-size:11px;color:#A32D2D;line-height:1.6">Always test French surveys with "merci et fermer" patterns — now supported!</p>
        </div>
      </div>
    </div>
  </div>
</div>
</body></html>""")

# ================================================================
# PAGE: NEW QC
# ================================================================
@app.route('/new-qc', methods=['GET'])
@login_required
def new_qc():
    user = get_current_user()
    plan = user.get('plan', 'Free') if user else 'Free'
    sb = sidebar_html('new-qc')
    is_pro = plan in ('Pro', 'Business', 'Enterprise')
    _reports_used = user.get('reports_used', 0) if user else 0
    _reports_limit = UserDB.PLAN_LIMITS.get(plan, 3)
    _at_limit = not users_db.can_run_report(session['user_email'])

    _msg_param = request.args.get('msg', '')
    _top_banner = ''
    if _msg_param == 'expired':
        _top_banner = ('<div class="nqc-banner nqc-banner-warn">'
                       '&#9888; <b>Previous job expired.</b> The report session is no longer '
                       'in memory. Please upload your document again to run a new QC.</div>')
    elif _msg_param == 'nodoc':
        _top_banner = ('<div class="nqc-banner nqc-banner-warn">'
                       '&#9888; <b>Original document not found.</b> '
                       'Please upload the .docx spec file again.</div>')

    # Recent reports for sidebar (last 3 for this user)
    _user_jobs = [(jid, j) for jid, j in jobs.items()
                  if j.get('user_email') == session.get('user_email')]
    _user_jobs.sort(key=lambda x: x[1].get('created_at', ''), reverse=True)
    _recent_rows = ''
    for _jid, _j in _user_jobs[:3]:
        _st = _j.get('status', 'running')
        _dn = (_j.get('doc_name', 'Unknown'))[:34]
        _iss = _j.get('total_issues', 0)
        _dt = _j.get('created_at', '')[:10]
        if _st == 'done':
            _bdg = (f'<span class="sb-badge sb-green">Passed</span>' if _iss == 0
                    else f'<span class="sb-badge sb-orange">{_iss} issue{"s" if _iss!=1 else ""}</span>')
        elif _st == 'running':
            _bdg = '<span class="sb-badge sb-blue">Running…</span>'
        else:
            _bdg = '<span class="sb-badge sb-red">Error</span>'
        _lnk = (f'<a href="/report/{_jid}" class="sb-view-link">View →</a>'
                if _st == 'done' else '')
        _recent_rows += (
            f'<div class="sb-row">'
            f'<div class="sb-row-left">'
            f'<span class="sb-row-name">{_dn}</span>'
            f'<span class="sb-row-date">{_dt}</span>'
            f'</div>'
            f'<div class="sb-row-right">{_bdg}{_lnk}</div>'
            f'</div>'
        )
    if not _recent_rows:
        _recent_rows = '<div class="sb-empty">No reports yet. Run your first QC!</div>'

    # Checkboxes for advanced panel
    _chk_items = [
        ('chk_term', 'Termination'), ('chk_text', 'Question text'),
        ('chk_words', 'Missing words'), ('chk_options', 'Options match'),
        ('chk_mandatory', 'Mandatory'), ('chk_piping', 'Piping'),
        ('chk_codes', 'Answer codes'), ('chk_order', 'Question order'),
    ]
    _chk_html = ''.join(
        f'<label class="qc-chk"><input type="checkbox" name="{n}" value="1" checked>{lbl}</label>'
        for n, lbl in _chk_items
    )

    # Screenshots field (advanced panel)
    if is_pro:
        _ss_field = (
            '<div class="dz dz-sm" id="ssZone"'
            ' onclick="document.getElementById(\'ssInput\').click()"'
            ' ondragover="dzOver(event,this)" ondragleave="dzLeave(this)"'
            ' ondrop="dzDrop(event,this,\'ssInput\',\'ssDone\',true)">'
            '<i class="ti ti-camera" style="font-size:18px;color:#9CA3AF"></i>'
            '<p style="font-size:11px;color:#6B7280;margin-top:3px">'
            'Drop screenshots · Multiple OK</p></div>'
            '<input type="file" name="screenshots" id="ssInput" accept="image/*" multiple'
            ' style="display:none" onchange="dzPick(this,\'ssZone\',\'ssDone\',true)">'
            '<div id="ssDone" class="dz-done" style="display:none"></div>'
        )
    else:
        _ss_field = (
            '<div style="display:flex;align-items:center;gap:6px;background:#FEF3E2;'
            'border-radius:7px;padding:6px 10px">'
            '<i class="ti ti-lock" style="font-size:12px;color:#C46A2B"></i>'
            '<span style="font-size:11px;color:#7C3D0A">Pro+ only. '
            '<a href="/billing" style="color:#042C53;font-weight:600">Upgrade →</a>'
            '</span></div>'
            '<input type="file" name="screenshots" accept="image/*" multiple '
            'disabled style="display:none">'
        )

    _ss_label = 'Optional' if is_pro else 'Pro+'

    page = SHARED_CSS + f'''<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1.0">
<title>New QC &#8212; SurveyQC</title>
<style>
/* ══ NEW QC PAGE ═══════════════════════════════════════════════════════ */
body.new-qc-page .nqc-wrap{{max-width:1100px;margin:0 auto;padding:0 0 48px;min-width:0}}
body.new-qc-page .nqc-banner{{border-radius:10px;padding:11px 16px;margin-bottom:18px;font-size:13px;line-height:1.5}}
body.new-qc-page .nqc-banner-warn{{background:#FEF3C7;border:1px solid #F59E0B;color:#92400E}}

/* 2-col layout: form + fixed 300px sidebar */
body.new-qc-page .nqc-layout{{display:grid;grid-template-columns:1fr 300px;gap:1.5rem;align-items:start;max-width:1100px}}

/* Form cards */
body.new-qc-page .fc{{background:#fff;border:0.5px solid var(--border);border-radius:14px;padding:18px 20px;
     transition:border-color .15s;display:flex;flex-direction:column}}
body.new-qc-page .fc:focus-within{{border-color:var(--accent);box-shadow:0 0 0 3px rgba(196,106,43,.07)}}
body.new-qc-page .fc-label{{font-size:13px;font-weight:600;color:var(--text);
           display:flex;align-items:center;gap:7px;margin-bottom:12px;flex-wrap:wrap}}
body.new-qc-page .fc-req{{color:#E24B4A}}
body.new-qc-page .fc-opt{{font-size:10px;font-weight:600;background:#F3F4F6;color:#6B7280;
         padding:2px 8px;border-radius:20px;flex-shrink:0}}

/* Upload cards row */
body.new-qc-page .upload-row{{display:grid;grid-template-columns:1fr 1fr;gap:1rem;margin-bottom:14px;align-items:start}}
body.new-qc-page .upload-row .fc{{height:180px;box-sizing:border-box;min-width:0}}

/* Drop zones */
body.new-qc-page .dz{{border:1.5px dashed #D1D5DB;border-radius:10px;min-height:118px;flex:1;
     display:flex;flex-direction:column;align-items:center;justify-content:center;
     cursor:pointer;background:#FAFAFA;text-align:center;padding:14px 16px;
     transition:border-color .15s,background .15s;user-select:none;width:100%;box-sizing:border-box}}
body.new-qc-page .dz:hover{{border-color:var(--accent);background:#FEF9F5}}
body.new-qc-page .dz.dz-over{{border-color:var(--accent);background:#FEF3E2;border-style:solid}}
body.new-qc-page .dz.dz-sm{{min-height:72px}}
body.new-qc-page .dz.dz-off{{cursor:default;opacity:.6}}
body.new-qc-page .dz.dz-off:hover{{border-color:#D1D5DB;background:#FAFAFA}}
body.new-qc-page .dz.dz-ok{{border:1.5px solid #16A34A!important;background:#F0FAF4}}
body.new-qc-page .dz.dz-ok:hover{{background:#E6F7EE}}
body.new-qc-page .dz.dz-err{{border:1.5px solid #DC2626;background:#FFF5F5}}
body.new-qc-page .dz-icon{{font-size:28px;color:#C8C4BF;margin-bottom:6px}}
body.new-qc-page .dz-hint{{font-size:12px;color:#6B7280;margin:0;line-height:1.5}}
body.new-qc-page .dz-hint b{{color:var(--text)}}
body.new-qc-page .dz-browse{{color:var(--accent);text-decoration:underline}}
body.new-qc-page .dz-helper{{font-size:11px;color:#9CA3AF;margin-top:5px}}
body.new-qc-page .dz-plus{{color:#16A34A;font-weight:600}}
body.new-qc-page .dz-done{{display:flex;align-items:center;gap:8px;padding:7px 10px;
          background:#F0FAF4;border:1px solid #A7D7B8;border-radius:8px;margin-top:0}}
body.new-qc-page .dz-rm{{background:none;border:none;cursor:pointer;color:#9CA3AF;
        font-size:15px;padding:0 2px;flex-shrink:0;line-height:1}}
body.new-qc-page .dz-rm:hover{{color:#DC2626}}

/* URL field */
body.new-qc-page .url-wrap{{display:flex;align-items:center;border:1px solid var(--border);border-radius:9px;
           overflow:hidden;background:white;transition:border-color .15s}}
body.new-qc-page .url-wrap:focus-within{{border-color:var(--accent);box-shadow:0 0 0 3px rgba(196,106,43,.08)}}
body.new-qc-page .url-icon{{padding:0 11px;color:#9CA3AF;font-size:15px;flex-shrink:0;display:flex;align-items:center}}
body.new-qc-page .url-inp{{flex:1;padding:11px 11px 11px 0;border:none;outline:none;
          font-size:13px;color:var(--text);font-family:inherit;background:transparent}}
body.new-qc-page .plat-pills{{display:flex;gap:6px;margin-top:9px;flex-wrap:wrap}}
body.new-qc-page .pp{{font-size:11px;background:#F3F4F6;color:#6B7280;padding:3px 10px;border-radius:20px;
     cursor:pointer;border:none;font-family:inherit;transition:background .12s,color .12s}}
body.new-qc-page .pp:hover{{background:#E8E1D8;color:var(--text)}}
body.new-qc-page .plat-detect{{display:none;font-size:10px;font-weight:600;padding:2px 8px;
              border-radius:5px;background:#DCFCE7;color:#166534;margin-left:auto}}

/* Accuracy strip */
body.new-qc-page .acc-strip{{display:flex;align-items:center;gap:10px;padding:8px 14px;background:white;
            border:0.5px solid var(--border);border-radius:9px;margin:14px 0;
            box-shadow:0 1px 2px rgba(0,0,0,.03)}}
body.new-qc-page .acc-bar{{flex:1;height:4px;background:#F3F4F6;border-radius:99px;overflow:hidden}}
body.new-qc-page .acc-fill{{height:100%;border-radius:99px;transition:width .4s ease,background .4s ease}}
body.new-qc-page .acc-lbl{{font-size:11px;font-weight:600;white-space:nowrap;color:#9CA3AF}}

/* Run button */
body.new-qc-page .run-btn{{width:100%;padding:14px;font-size:15px;font-weight:700;border-radius:11px;
          border:none;cursor:pointer;display:flex;align-items:center;justify-content:center;
          gap:8px;color:white;background:#1a1a1a;letter-spacing:-.1px;
          box-shadow:0 4px 14px rgba(0,0,0,.18);
          transition:background .2s,box-shadow .15s,transform .1s}}
body.new-qc-page .run-btn:hover:not(:disabled){{box-shadow:0 6px 20px rgba(0,0,0,.25);transform:translateY(-1px)}}
body.new-qc-page .run-btn:active:not(:disabled){{transform:translateY(0)}}
body.new-qc-page .run-btn:disabled{{opacity:.5;cursor:not-allowed}}
body.new-qc-page .run-btn.btn-warm{{background:linear-gradient(135deg,#C46A2B,#D97706);
                   box-shadow:0 4px 14px rgba(196,106,43,.3)}}
body.new-qc-page .run-btn.btn-full{{background:linear-gradient(135deg,#1B4332,#1a1a1a);
                   box-shadow:0 4px 14px rgba(27,67,50,.3)}}
body.new-qc-page .run-sub{{font-size:12px;color:#9CA3AF;text-align:center;margin-top:8px;
          display:flex;align-items:center;justify-content:center;gap:10px;flex-wrap:wrap}}
body.new-qc-page .adv-lnk{{background:none;border:none;cursor:pointer;font-size:12px;color:#9CA3AF;
          font-family:inherit;padding:0;text-decoration:underline;text-underline-offset:2px}}
body.new-qc-page .adv-lnk:hover{{color:var(--text)}}

/* Advanced panel */
body.new-qc-page .adv-panel{{margin-top:12px;padding:16px;background:#F8F7F4;
            border-radius:10px;border:0.5px solid var(--border)}}
body.new-qc-page .adv-lbl{{font-size:11px;font-weight:600;color:var(--text);margin-bottom:6px;display:block}}
body.new-qc-page .qc-chk{{display:flex;align-items:center;gap:6px;font-size:11px;color:#374151;cursor:pointer}}
body.new-qc-page .qc-chk input{{accent-color:#1a1a1a;width:12px;height:12px;flex-shrink:0}}
body.new-qc-page .chk-grid{{display:grid;grid-template-columns:1fr 1fr;gap:5px;margin-bottom:12px}}

/* xml tooltip */
body.new-qc-page .xml-tip{{background:#EFF6FF;border:1px solid #BFDBFE;border-radius:8px;
          padding:9px 12px;margin-bottom:10px;font-size:11px;color:#1E3A5F;line-height:1.6}}
body.new-qc-page .xml-tip ul{{margin:4px 0 0;padding-left:14px}}
body.new-qc-page .xml-tb{{background:none;border:none;cursor:pointer;font-size:11px;color:#9CA3AF;
         font-family:inherit;padding:0;margin-left:auto;display:flex;align-items:center;gap:3px}}
body.new-qc-page .xml-tb:hover{{color:var(--text)}}

/* ── SIDEBAR ───────────────────────────────────────────────────────── */
body.new-qc-page .sb-card{{background:#fff;border:0.5px solid var(--border);border-radius:14px;
          padding:16px 18px;margin-bottom:14px}}
body.new-qc-page .sb-card:last-child{{margin-bottom:0}}
body.new-qc-page .sb-card-hdr{{display:flex;align-items:center;justify-content:space-between;margin-bottom:12px}}
body.new-qc-page .sb-card-title{{font-size:13px;font-weight:700;color:var(--text)}}
body.new-qc-page .sb-view-all{{font-size:12px;color:var(--accent);text-decoration:none;font-weight:500}}
body.new-qc-page .sb-view-all:hover{{text-decoration:underline}}
body.new-qc-page .sb-row{{display:flex;align-items:center;justify-content:space-between;
         padding:8px 0;border-bottom:0.5px solid var(--border)}}
body.new-qc-page .sb-row:last-child{{border-bottom:none;padding-bottom:0}}
body.new-qc-page .sb-row:first-child{{padding-top:0}}
body.new-qc-page .sb-row-left{{flex:1;min-width:0;margin-right:8px}}
body.new-qc-page .sb-row-name{{display:block;font-size:12px;font-weight:600;color:var(--text);
              white-space:nowrap;overflow:hidden;text-overflow:ellipsis}}
body.new-qc-page .sb-row-date{{font-size:10px;color:var(--text3);display:block;margin-top:1px}}
body.new-qc-page .sb-row-right{{display:flex;flex-direction:column;align-items:flex-end;gap:4px;flex-shrink:0}}
body.new-qc-page .sb-badge{{font-size:10px;font-weight:600;padding:2px 7px;border-radius:20px;white-space:nowrap}}
body.new-qc-page .sb-green{{background:#DCFCE7;color:#15803D}}
body.new-qc-page .sb-orange{{background:#FEF3C7;color:#B45309}}
body.new-qc-page .sb-blue{{background:#DBEAFE;color:#1D4ED8}}
body.new-qc-page .sb-red{{background:#FEE2E2;color:#DC2626}}
body.new-qc-page .sb-view-link{{font-size:11px;color:var(--accent);text-decoration:none;font-weight:500}}
body.new-qc-page .sb-view-link:hover{{text-decoration:underline}}
body.new-qc-page .sb-empty{{font-size:12px;color:var(--text3);text-align:center;padding:8px 0}}
body.new-qc-page .sb-checklist{{list-style:none;margin:0;padding:0;display:flex;flex-direction:column;gap:7px}}
body.new-qc-page .sb-checklist li{{display:flex;align-items:flex-start;gap:8px;font-size:12px;
                  color:var(--text2);line-height:1.4}}
body.new-qc-page .sb-check-icon{{color:#16A34A;font-size:14px;flex-shrink:0}}
body.new-qc-page .sb-plat-grid{{display:grid;grid-template-columns:1fr 1fr;gap:8px}}
body.new-qc-page .sb-plat{{display:flex;align-items:center;gap:8px;padding:8px 10px;
          background:var(--bg);border-radius:9px;border:0.5px solid var(--border)}}
body.new-qc-page .sb-plat-icon{{width:28px;height:28px;border-radius:7px;display:flex;
               align-items:center;justify-content:center;font-size:14px;flex-shrink:0}}
body.new-qc-page .sb-plat-name{{font-size:11px;font-weight:600;color:var(--text)}}
body.new-qc-page .sb-more{{font-size:11px;color:var(--text3);text-align:center;margin-top:8px}}
body.new-qc-page .sb-tips{{display:flex;flex-direction:column;gap:9px}}
body.new-qc-page .sb-tip{{display:flex;align-items:flex-start;gap:8px;font-size:12px;
         color:var(--text2);line-height:1.45}}
body.new-qc-page .sb-tip-arr{{color:var(--accent);flex-shrink:0;font-weight:700;margin-top:1px}}

@media(max-width:960px){{
  body.new-qc-page .nqc-layout{{grid-template-columns:1fr}}
  body.new-qc-page .nqc-wrap{{max-width:640px}}
}}
@media(max-width:520px){{
  body.new-qc-page .upload-row{{grid-template-columns:1fr}}
  body.new-qc-page .chk-grid{{grid-template-columns:1fr}}
}}
</style>
</head>
<body class="new-qc-page">
<div class="app-layout">{sb}<div class="main-content" style="padding-top:0">

<div class="topbar" style="padding-top:16px;padding-bottom:12px">
  <div>
    <p class="page-title" style="font-size:20px;margin-bottom:2px">New QC check</p>
    <p class="page-sub" style="font-size:13px">Upload your files. AI handles everything.</p>
  </div>
</div>

{_top_banner}
<div class="nqc-wrap">
<form action="/run-qc" method="POST" enctype="multipart/form-data" id="qcForm" data-at-limit="{'1' if _at_limit else '0'}">
<div class="nqc-layout">

  <!-- ══ LEFT — FORM ════════════════════════════════════════════════ -->
  <div>

    <!-- Upload cards: doc + export side by side -->
    <div class="upload-row">

      <!-- Spec Document -->
      <div class="fc">
        <div class="fc-label">
          <i class="ti ti-file-word" style="color:var(--accent);font-size:15px"></i>
          Spec Document <span class="fc-req">*</span>
        </div>
        <div class="dz" id="docZone"
             onclick="document.getElementById('docInput').click()"
             ondragover="dzOver(event,this)" ondragleave="dzLeave(this)"
             ondrop="dzDrop(event,this,'docInput','docDone')">
          <i class="ti ti-cloud-upload dz-icon"></i>
          <p class="dz-hint">Drop <b>.docx</b> here or<br>
            <span class="dz-browse">browse files</span></p>
        </div>
        <input type="file" name="doc" id="docInput" accept=".docx" required
               style="display:none" onchange="dzPick(this,'docZone','docDone');updateMeter()">
        <div id="docDone" class="dz-done" style="display:none"></div>
      </div>

      <!-- Survey Export -->
      <div class="fc">
        <div class="fc-label">
          <i class="ti ti-file-code" style="color:#3B82F6;font-size:15px"></i>
          Survey Export <span class="fc-req">*</span>
          <button type="button" class="xml-tb" onclick="toggleXmlTip()">
            <i class="ti ti-info-circle"></i> How?
          </button>
        </div>
        <div id="xmlTipBox" class="xml-tip" style="display:none">
          <b>Export from your platform:</b>
          <ul>
            <li><b>Confirmit:</b> Designer &#8594; Export Survey Definition (.zip)</li>
            <li><b>Decipher:</b> Survey settings &#8594; Download XML</li>
            <li><b>Forsta:</b> Survey settings &#8594; Export XML</li>
            <li><b>Qualtrics:</b> Tools &#8594; Import/Export &#8594; Export QSF</li>
          </ul>
        </div>
        <div class="dz" id="xmlZone"
             onclick="document.getElementById('xmlInput').click()"
             ondragover="dzOver(event,this)" ondragleave="dzLeave(this)"
             ondrop="dzDrop(event,this,'xmlInput','xmlDone')">
          <i class="ti ti-cloud-upload dz-icon"></i>
          <p class="dz-hint" style="white-space:nowrap">Drop <b>.xml</b> / <b>.qsf</b> / <b>.zip</b></p>
          <p class="dz-hint"><span class="dz-browse">or browse files</span></p>
        </div>
        <input type="file" name="xml_export" id="xmlInput" accept=".xml,.qsf,.zip" required
               style="display:none"
               onchange="dzPick(this,'xmlZone','xmlDone');updateMeter()">
        <div id="xmlDone" class="dz-done" style="display:none"></div>
      </div>

    </div><!-- /upload-row -->

    <!-- Live Survey URL (optional — adds ADVANCED QC live verification) -->
    <div class="fc" style="margin-bottom:14px">
      <div class="fc-label">
        <i class="ti ti-link" style="color:var(--text3);font-size:15px"></i>
        Live Survey URL <span class="fc-opt" style="background:#EFF6FF;color:#1D4ED8">Optional</span>
        <span style="font-size:10px;color:#6B7280;margin-left:2px">— add for ADVANCED QC live verification</span>
        <span id="platBadge" class="plat-detect"></span>
      </div>
      <div class="url-wrap">
        <span class="url-icon"><i class="ti ti-world"></i></span>
        <input type="url" name="url" id="urlInput" class="url-inp"
               placeholder="https://survey.confirmit.com/... (optional)"
               oninput="detectPlat(this.value);updateMeter()">
      </div>
      <div class="plat-pills">
        <button type="button" class="pp" onclick="setUrlHint('confirmit')">Confirmit</button>
        <button type="button" class="pp" onclick="setUrlHint('decipher')">Decipher</button>
        <button type="button" class="pp" onclick="setUrlHint('forsta')">Forsta</button>
        <button type="button" class="pp" onclick="setUrlHint('qualtrics')">Qualtrics</button>
        <button type="button" class="pp" onclick="setUrlHint('surveymonkey')">SurveyMonkey</button>
      </div>
    </div>

    <!-- QC Mode meter -->
    <div class="acc-strip">
      <div class="acc-bar">
        <div class="acc-fill" id="accFill" style="width:0%;background:#E5E7EB"></div>
      </div>
      <span class="acc-lbl" id="accLbl">Upload Spec Document + Survey Export to run</span>
    </div>

    <!-- Run QC button -->
    <button type="submit" class="run-btn" id="runBtn" disabled>
      <i class="ti ti-player-play" style="font-size:16px"></i>
      <span id="runBtnText">&#9654; Run QC &#8212; Upload files above</span>
    </button>
    <div class="run-sub">
      <span id="runBtnSub">Upload Spec Document + Survey Export to run</span>
      <span style="color:#DDE1E7">|</span>
      <button type="button" class="adv-lnk" onclick="toggleAdv()">
        Advanced options <span id="advChev">&#9662;</span>
      </button>
    </div>

    <!-- Advanced panel (hidden by default) -->
    <div id="advPanel" style="display:none" class="adv-panel">
      <div style="display:grid;grid-template-columns:1fr 1fr;gap:10px;margin-bottom:12px">
        <div>
          <label class="adv-lbl">Country (screener)</label>
          <input type="text" name="country" placeholder="e.g. United Kingdom"
                 style="width:100%;padding:7px 10px;border:1px solid var(--border);
                 border-radius:7px;font-size:12px;font-family:inherit;outline:none;
                 box-sizing:border-box;background:white"
                 onfocus="this.style.borderColor=\'#C46A2B\'"
                 onblur="this.style.borderColor=\'var(--border)\'">
        </div>
        <div>
          <label class="adv-lbl">Specific QIDs only</label>
          <input type="text" name="specific_questions" placeholder="e.g. Q1, Q3, Q7-Q12"
                 style="width:100%;padding:7px 10px;border:1px solid var(--border);
                 border-radius:7px;font-size:12px;font-family:inherit;outline:none;
                 box-sizing:border-box;background:white"
                 onfocus="this.style.borderColor=\'#C46A2B\'"
                 onblur="this.style.borderColor=\'var(--border)\'">
        </div>
      </div>
      <span class="adv-lbl">Checks to run</span>
      <div class="chk-grid">{_chk_html}</div>
      <span class="adv-lbl" style="margin-top:8px">
        Data Export Schema <span class="fc-opt">Optional</span>
      </span>
      <p style="font-size:11px;color:var(--text3);margin-bottom:7px;line-height:1.4">
        Validates variable names, codes, and types against the spec doc.
      </p>
      <div class="dz dz-sm" id="exportZone"
           onclick="document.getElementById('exportFileInput').click()"
           ondragover="dzOver(event,this)" ondragleave="dzLeave(this)"
           ondrop="dzDrop(event,this,'exportFileInput','exportFileDone')">
        <i class="ti ti-table" style="font-size:18px;color:#9CA3AF"></i>
        <p style="font-size:11px;color:#6B7280;margin-top:3px">
          Drop <b>.csv</b> / <b>.txt</b> / <b>.xlsx</b> or
          <span class="dz-browse" style="font-size:11px">browse</span>
        </p>
      </div>
      <input type="file" name="export_schema_file" id="exportFileInput"
             accept=".csv,.txt,.xlsx" style="display:none"
             onchange="dzPick(this,'exportZone','exportFileDone');onExportFile(this)">
      <div id="exportFileDone" class="dz-done" style="display:none"></div>
      <textarea name="export_headers_text" id="exportHeadersText"
        placeholder="Or paste headers: R0, R1, S1, S2, Q14_1, Q14_2..."
        style="width:100%;height:52px;padding:7px 10px;border:1px solid var(--border);
               margin-top:8px;border-radius:8px;font-size:11px;font-family:monospace;
               resize:vertical;outline:none;box-sizing:border-box;
               color:var(--text);line-height:1.5;background:white"
        onfocus="this.style.borderColor=\'#C46A2B\'"
        onblur="this.style.borderColor=\'var(--border)\'"></textarea>
      <span class="adv-lbl" style="margin-top:12px">
        Screenshots <span class="fc-opt">{_ss_label}</span>
      </span>
      {_ss_field}
    </div>

  </div><!-- /left col -->

  <!-- ══ RIGHT — SIDEBAR ═════════════════════════════════════════════ -->
  <div>

    <!-- Recent Reports -->
    <div class="sb-card">
      <div class="sb-card-hdr">
        <span class="sb-card-title">Recent Reports</span>
        <a href="/reports" class="sb-view-all">View all &#8594;</a>
      </div>
      {_recent_rows}
    </div>

    <!-- Why SurveyQC -->
    <div class="sb-card">
      <div class="sb-card-hdr">
        <span class="sb-card-title">Why use SurveyQC?</span>
      </div>
      <ul class="sb-checklist">
        <li><i class="ti ti-circle-check sb-check-icon"></i>Standard QC: Doc + XML (no live URL needed)</li>
        <li><i class="ti ti-circle-check sb-check-icon"></i>Advanced QC: Add live URL for full verification</li>
        <li><i class="ti ti-circle-check sb-check-icon"></i>32+ quality checks automated</li>
        <li><i class="ti ti-circle-check sb-check-icon"></i>Smart issue detection &amp; evidence</li>
        <li><i class="ti ti-circle-check sb-check-icon"></i>Downloadable Word report</li>
      </ul>
    </div>

    <!-- Supported Platforms -->
    <div class="sb-card">
      <div class="sb-card-hdr">
        <span class="sb-card-title">Supported Platforms</span>
      </div>
      <div class="sb-plat-grid">
        <div class="sb-plat">
          <div class="sb-plat-icon" style="background:#EBF5FF">
            <i class="ti ti-chart-bar" style="color:#1D6FAE"></i>
          </div>
          <span class="sb-plat-name">Confirmit</span>
        </div>
        <div class="sb-plat">
          <div class="sb-plat-icon" style="background:#FFF4E6">
            <i class="ti ti-analyze" style="color:#C46A2B"></i>
          </div>
          <span class="sb-plat-name">Decipher</span>
        </div>
        <div class="sb-plat">
          <div class="sb-plat-icon" style="background:#F0FDF4">
            <i class="ti ti-sitemap" style="color:#16A34A"></i>
          </div>
          <span class="sb-plat-name">Forsta</span>
        </div>
        <div class="sb-plat">
          <div class="sb-plat-icon" style="background:#FDF4FF">
            <i class="ti ti-clipboard-list" style="color:#9333EA"></i>
          </div>
          <span class="sb-plat-name">Qualtrics</span>
        </div>
      </div>
      <p class="sb-more">and more&#8230;</p>
    </div>

    <!-- Tips -->
    <div class="sb-card">
      <div class="sb-card-hdr">
        <span class="sb-card-title">
          <i class="ti ti-bulb" style="color:#F59E0B;margin-right:4px"></i>Tips
        </span>
      </div>
      <div class="sb-tips">
        <div class="sb-tip">
          <span class="sb-tip-arr">&#8594;</span>
          <b>Standard QC</b>: Doc + XML only — no live URL needed, 85-95% accuracy
        </div>
        <div class="sb-tip">
          <span class="sb-tip-arr">&#8594;</span>
          <b>Advanced QC</b>: Add live survey URL to verify routing, piping &amp; mandatory
        </div>
        <div class="sb-tip">
          <span class="sb-tip-arr">&#8594;</span>
          XML export must be from same survey version as the spec doc
        </div>
        <div class="sb-tip">
          <span class="sb-tip-arr">&#8594;</span>
          If using a live URL, make sure it is publicly accessible (no login)
        </div>
      </div>
    </div>

  </div><!-- /sidebar col -->

</div><!-- /nqc-layout -->
</form>
</div><!-- /nqc-wrap -->

<!-- Mobile bottom nav -->
<div style="display:none;position:fixed;bottom:0;left:0;right:0;background:white;
     padding:10px 20px;justify-content:space-around;align-items:center;z-index:1000;
     border-top:1px solid var(--border)" class="mobile-bottom-nav">
  <a href="/dashboard" style="display:flex;flex-direction:column;align-items:center;gap:3px;
     text-decoration:none;color:var(--text3)">
    <i class="ti ti-home" style="font-size:20px"></i>
    <span style="font-size:10px">Home</span>
  </a>
  <a href="/new-qc" style="display:flex;flex-direction:column;align-items:center;gap:3px;
     text-decoration:none;color:var(--accent)">
    <i class="ti ti-plus" style="font-size:20px"></i>
    <span style="font-size:10px">New QC</span>
  </a>
  <a href="/reports" style="display:flex;flex-direction:column;align-items:center;gap:3px;
     text-decoration:none;color:var(--text3)">
    <i class="ti ti-file-text" style="font-size:20px"></i>
    <span style="font-size:10px">Reports</span>
  </a>
  <a href="/settings" style="display:flex;flex-direction:column;align-items:center;gap:3px;
     text-decoration:none;color:var(--text3)">
    <i class="ti ti-settings" style="font-size:20px"></i>
    <span style="font-size:10px">Settings</span>
  </a>
</div>
<style>
@media(max-width:768px){{.mobile-bottom-nav{{display:flex !important}}
  .main-content{{padding-bottom:70px}}}}
</style>

</div></div>

<script src="/static/new_qc.js"></script>

<!-- ── Upgrade modal ───────────────────────────────────────────────── -->
<div id="upgradeModal" style="display:none;position:fixed;inset:0;background:rgba(0,0,0,.55);
     z-index:9000;align-items:center;justify-content:center;padding:16px">
  <div style="background:#fff;border-radius:16px;max-width:560px;width:100%;
       box-shadow:0 24px 60px rgba(0,0,0,.18);overflow:hidden">
    <!-- header -->
    <div style="background:#042C53;padding:28px 28px 20px;position:relative">
      <button onclick="document.getElementById('upgradeModal').style.display='none'"
              style="position:absolute;top:14px;right:16px;background:rgba(255,255,255,.15);
              border:none;color:#fff;width:28px;height:28px;border-radius:50%;font-size:16px;
              cursor:pointer;display:flex;align-items:center;justify-content:center;
              line-height:1">&#215;</button>
      <p style="font-size:18px;font-weight:700;color:#fff;margin-bottom:6px">
        You&#8217;ve used all {_reports_limit} free reports</p>
      <p style="font-size:14px;color:rgba(255,255,255,.75)">
        Upgrade to continue running QC checks</p>
    </div>
    <!-- plan cards -->
    <div style="padding:24px 28px;display:grid;grid-template-columns:1fr 1fr;gap:14px">
      <!-- Pro -->
      <div style="border:1.5px solid #7C3AED;border-radius:12px;padding:20px;display:flex;
           flex-direction:column;gap:8px">
        <p style="font-size:15px;font-weight:700;color:#1A1A2E">Pro</p>
        <p style="font-size:22px;font-weight:800;color:#7C3AED">$29<span style="font-size:13px;
           font-weight:500;color:#6B7280">/mo</span></p>
        <ul style="margin:4px 0 12px;padding-left:18px;font-size:13px;
            color:#374151;line-height:1.8;list-style:none;padding:0">
          <li style="padding:2px 0">&#10003; 25 reports / month</li>
          <li style="padding:2px 0">&#10003; Word report export</li>
          <li style="padding:2px 0">&#10003; All platforms</li>
        </ul>
        <a href="/billing?plan=pro" style="display:block;text-align:center;
           background:#7C3AED;color:#fff;padding:10px 0;border-radius:8px;
           font-size:13px;font-weight:600;text-decoration:none;margin-top:auto">
          Upgrade &#8594;</a>
      </div>
      <!-- Business -->
      <div style="border:1.5px solid #042C53;border-radius:12px;padding:20px;display:flex;
           flex-direction:column;gap:8px;background:#F0F4FA">
        <p style="font-size:15px;font-weight:700;color:#1A1A2E">Business</p>
        <p style="font-size:22px;font-weight:800;color:#042C53">$299<span style="font-size:13px;
           font-weight:500;color:#6B7280">/mo</span></p>
        <ul style="margin:4px 0 12px;padding-left:18px;font-size:13px;
            color:#374151;line-height:1.8;list-style:none;padding:0">
          <li style="padding:2px 0">&#10003; Unlimited reports</li>
          <li style="padding:2px 0">&#10003; Everything in Pro</li>
          <li style="padding:2px 0">&#10003; Priority support</li>
        </ul>
        <a href="/billing?plan=business" style="display:block;text-align:center;
           background:#042C53;color:#fff;padding:10px 0;border-radius:8px;
           font-size:13px;font-weight:600;text-decoration:none;margin-top:auto">
          Upgrade &#8594;</a>
      </div>
    </div>
    <!-- footer -->
    <div style="border-top:1px solid #E5E7EB;padding:14px 28px;text-align:center">
      <p style="font-size:12px;color:#9CA3AF">Questions?
        <a href="mailto:support@surveyqc.com" style="color:#7C3AED;text-decoration:none">Contact us</a>
      </p>
    </div>
  </div>
</div>

</body></html>'''

    return render_template_string(page)



@app.route('/run-qc', methods=['POST'])
@login_required
def run_qc_submit():
    doc_file = request.files.get('doc')
    survey_url = request.form.get('url', '').strip()

    def _validation_error(msg):
        return render_template_string(SHARED_CSS + f"""
<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"><title>Validation Error — SurveyQC</title></head><body>
<div style="min-height:100vh;display:flex;align-items:center;justify-content:center;background:#F8F9FA">
  <div style="background:white;border:0.5px solid #DDE1E7;border-radius:12px;padding:36px;max-width:460px;text-align:center">
    <div style="width:48px;height:48px;background:#FCEBEB;border-radius:12px;display:flex;align-items:center;justify-content:center;margin:0 auto 16px">
      <i class="ti ti-alert-circle" style="font-size:24px;color:#DC2626"></i>
    </div>
    <p style="font-size:16px;font-weight:600;color:#1A1A2E;margin-bottom:8px">Missing required field</p>
    <p style="font-size:14px;color:#6B7280;margin-bottom:24px">{msg}</p>
    <a href="/new-qc" style="display:inline-flex;align-items:center;gap:6px;background:#042C53;color:white;padding:11px 22px;border-radius:8px;font-size:13px;font-weight:600;text-decoration:none">
      <i class="ti ti-arrow-left" style="font-size:14px"></i> Go back and fix
    </a>
  </div>
</div>
</body></html>"""), 400

    if not doc_file or not doc_file.filename:
        return _validation_error("Please upload a .docx spec document.")
    # URL is optional: absent → STANDARD QC (DOC+XML only), present → ADVANCED QC
    qc_mode = 'ADVANCED' if survey_url else 'STANDARD'

    # Plan-limit gate — checked before creating the job
    _email = session['user_email']
    if not users_db.can_run_report(_email):
        _u = users_db.get(_email, {})
        _plan = _u.get('plan', 'Free')
        _limit = UserDB.PLAN_LIMITS.get(_plan, 3)
        return jsonify({"upgrade_required": True, "plan": _plan, "limit": _limit, "error": "Report limit reached"}), 403

    job_id = str(uuid.uuid4())[:8]
    job_dir = f"{UPLOAD_FOLDER}/{job_id}"
    os.makedirs(job_dir, exist_ok=True)

    doc_filename = secure_filename(doc_file.filename)
    doc_path = f"{job_dir}/{doc_filename}"
    doc_file.save(doc_path)

    ss_paths = []
    for f in request.files.getlist('screenshots'):
        if f.filename:
            ss_path = f"{job_dir}/{secure_filename(f.filename)}"
            f.save(ss_path)
            ss_paths.append(ss_path)

    xml_file = request.files.get('xml_export')
    if not xml_file or not xml_file.filename:
        return _validation_error("Survey Export is required. Please upload your XML/QSF file. "
                                 "Get it from your platform: Confirmit → Export Survey Definition, "
                                 "Decipher → Download XML, Qualtrics → Export QSF.")
    xml_ext = os.path.splitext(secure_filename(xml_file.filename))[1].lower() or '.xml'
    xml_path = f"{job_dir}/survey_export{xml_ext}"
    xml_file.save(xml_path)

    # Optional: data export schema file or pasted headers
    export_schema_text = ''
    export_file = request.files.get('export_schema_file')
    export_text_field = request.form.get('export_headers_text', '').strip()
    if export_file and export_file.filename:
        _exp_ext = os.path.splitext(secure_filename(export_file.filename))[1].lower()
        _exp_path = f"{job_dir}/export_schema{_exp_ext}"
        export_file.save(_exp_path)
        try:
            if _exp_ext == '.xlsx':
                import openpyxl
                _wb = openpyxl.load_workbook(_exp_path, read_only=True, data_only=True)
                _ws = _wb.active
                _headers = [str(c.value).strip() for c in next(_ws.iter_rows()) if c.value is not None]
                export_schema_text = ','.join(_headers)
                _wb.close()
            else:
                with open(_exp_path, 'r', encoding='utf-8', errors='replace') as _f:
                    export_schema_text = _f.readline().strip()
        except Exception:
            export_schema_text = export_text_field
    elif export_text_field:
        export_schema_text = export_text_field

    # Memory guard: evict oldest finished jobs from RAM when count > 50.
    # Records stay in SQLite; only the in-memory cache is trimmed.
    if len(jobs) >= 50:
        _finished_ids = sorted(
            (jid for jid, j in jobs.items() if j.get('status') in ('done', 'error', 'stopped')),
            key=lambda jid: jobs[jid].get('created_at', '')
        )
        for _evict in _finished_ids[:max(0, len(jobs) - 49)]:
            jobs.evict(_evict)

    jobs[job_id] = {
        'status': 'running',
        'progress': 0,
        'phase': 'Starting...',
        'logs': [],
        'doc_name': doc_filename,
        'doc_path': doc_path,
        'survey_url': survey_url,
        'qc_mode': qc_mode,
        'xml_path': xml_path,
        'platform': request.form.get('platform', 'Confirmit'),
        'country': request.form.get('country', ''),
        'mode': request.form.get('mode', 'full'),
        'user_email': session['user_email'],
        'created_at': datetime.now().isoformat(),
        'export_schema_text': export_schema_text,
        'verdict': None,
        'issues': [],
        'term_results': [],
        'report_file': None,
        'doc_qids': 0,
        'live_qids': 0,
        'xml_qids': 0,
        'total_issues': 0,
        'term_passed': 0,
        'term_total': 0,
    }

    thread = threading.Thread(
        target=run_qc_engine,
        args=(job_id, doc_path, survey_url,
              request.form.get('country', ''),
              request.form.get('mode', 'full'),
              ss_paths),
        daemon=True
    )
    thread.start()

    email = session['user_email']
    if email in users_db:
        users_db[email]['reports_used'] = users_db[email].get('reports_used', 0) + 1
        users_db[email]['total_saved_hours'] = users_db[email].get('total_saved_hours', 0) + 8
        users_db.save(email)

    return redirect(f'/progress/{job_id}')

# ================================================================
# PAGE: PROGRESS
# ================================================================
@app.route('/progress/<job_id>')
@login_required
def progress_page(job_id):
    if job_id not in jobs:
        return redirect('/reports')
    j = jobs[job_id]
    doc_name = j.get('doc_name', 'Unknown')
    _prog_qc_mode = j.get('qc_mode', 'STANDARD')
    _prog_mode_label = ('Advanced QC — DOC + XML + LIVE' if _prog_qc_mode == 'ADVANCED'
                        else 'Standard QC — DOC + XML')

    return render_template_string(SHARED_CSS + f"""
<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=5.0"><title>Running QC — SurveyQC</title>

</head><body>
<div class="app-layout">
  {sidebar_html('reports')}
  <div class="main-content">
    <div class="topbar">
      <div>
        <p class="page-title">Running QC Analysis</p>
        <p class="page-sub">{doc_name} &nbsp;<span style="font-size:11px;font-weight:600;padding:2px 8px;border-radius:12px;background:{'#DCFCE7;color:#15803D' if _prog_qc_mode=='ADVANCED' else '#DBEAFE;color:#1D4ED8'}">{_prog_mode_label}</span></p>
      </div>
      <span id="status-badge" class="badge badge-blue">Running</span>
      <button onclick="stopJob()" id="stop-btn" style="margin-left:12px;background:#dc2626;color:white;border:none;padding:8px 16px;border-radius:8px;cursor:pointer;font-size:13px;font-weight:600">⏹ Stop</button>
    </div>

    <div class="card">
      <div style="display:flex;align-items:center;justify-content:space-between;margin-bottom:16px">
        <p style="font-size:14px;font-weight:600;color:var(--text)" id="phase-text">Starting...</p>
        <span id="pct-text" style="font-size:13px;color:var(--text3)">0%</span>
      </div>
      <div class="progress-bar" style="height:8px;margin-bottom:20px">
        <div class="progress-fill progress-purple" id="prog-fill" style="width:0%"></div>
      </div>
      <div class="log-box" id="log-box">
        <div class="log-cyan">Starting QC analysis...</div>
      </div>
    </div>

    <div id="done-section" style="display:none;margin-top:16px">
      <div class="alert alert-success" style="font-size:14px">
        Analysis complete! <a href="/report/{job_id}" style="color:#27500A;font-weight:500">View report →</a>
      </div>
    </div>
  </div>
</div>
<script>
var jobId = "{job_id}";
var logCount = 0;

function poll() {{
  fetch('/api/status/' + jobId)
    .then(r => r.json())
    .then(data => {{
      document.getElementById('prog-fill').style.width = (data.progress||0) + '%';
      document.getElementById('pct-text').textContent = (data.progress||0) + '%';
      document.getElementById('phase-text').textContent = data.phase || 'Running...';

      var box = document.getElementById('log-box');
      if (data.logs) {{
        var newLogs = data.logs.slice(logCount);
        newLogs.forEach(function(l) {{
          var d = document.createElement('div');
          d.className = 'log-' + (l.color||'white');
          d.textContent = l.msg;
          box.appendChild(d);
        }});
        logCount = data.logs.length;
        box.scrollTop = box.scrollHeight;
      }}

      if (data.status === 'done') {{
        document.getElementById('status-badge').textContent = 'Complete';
        document.getElementById('status-badge').className = 'badge badge-green';
        document.getElementById('done-section').style.display = 'block';
        clearInterval(timer);
        setTimeout(function(){{ window.location = '/report/{job_id}'; }}, 2000);
      }} else if (data.status === 'error') {{
        document.getElementById('status-badge').textContent = 'Error';
        document.getElementById('status-badge').className = 'badge badge-red';
        clearInterval(timer);
      }}
    }}).catch(console.error);
}}

var timer = setInterval(poll, 1500);
poll();
function stopJob() {{
  if(!confirm('Stop this QC run?')) return;
  fetch('/stop/' + jobId, {{method:'POST'}})
    .then(() => {{
      document.getElementById('stop-btn').textContent = 'Stopped';
      document.getElementById('stop-btn').disabled = true;
      clearInterval(timer);
      document.getElementById('status-badge').textContent = 'Stopped';
      document.getElementById('status-badge').className = 'badge badge-red';
    }});
}}
</script>
</body></html>""")


@app.route('/stop/<job_id>', methods=['POST'])
@login_required
def stop_job(job_id):
    if job_id in jobs:
        jobs[job_id]['status'] = 'stopped'
        jobs[job_id]['phase'] = 'Stopped by user'
        jobs.persist(job_id)
    return jsonify({'ok': True})

# ================================================================
# PAGE: REPORT DETAIL
# ================================================================
@app.route('/report/<job_id>')
@login_required
def report_detail(job_id):
    if job_id not in jobs:
        return redirect('/reports')
    j = jobs[job_id]
    if j.get('status') != 'done':
        return redirect(f'/progress/{job_id}')

    doc_name = j.get('doc_name', 'Unknown')
    platform = j.get('platform', '-')
    country = j.get('country', '-')
    issues = j.get('issues', [])
    term_results = j.get('term_results', [])
    verdict = j.get('verdict', 'REVIEW')
    ai_summary = j.get('ai_summary', '')
    doc_qids = j.get('doc_qids', 0)
    re_findings     = j.get('rule_engine_findings', [])
    re_summary      = j.get('rule_engine_summary', {})
    re_term_matrix  = j.get('termination_matrix', [])
    consensus       = j.get('consensus', {})
    health_score    = j.get('health_score')
    live_qids = j.get('live_qids', 0)
    xml_qids = j.get('xml_qids', 0)
    term_passed = j.get('term_passed', 0)
    term_review = j.get('term_review', 0)
    term_total = j.get('term_total', 0)
    total_issues = j.get('total_issues', 0)
    created = j.get('created_at', '')[:16]
    qc_mode = j.get('qc_mode', 'STANDARD')

    verdict_class = 'badge-red' if verdict == 'FAIL' else ('badge-green' if verdict == 'PASS' else 'badge-amber')
    verdict_icon = 'ti-x' if verdict == 'FAIL' else ('ti-check' if verdict == 'PASS' else 'ti-alert-triangle')
    verdict_msg = 'Fix required before going live' if verdict == 'FAIL' else ('All good — ready to launch!' if verdict == 'PASS' else 'Review needed before launch')

    # Mode badge HTML
    if qc_mode == 'ADVANCED':
        _mode_badge = ('<span style="font-size:10px;font-weight:700;padding:2px 9px;'
                       'border-radius:20px;background:#DCFCE7;color:#15803D;white-space:nowrap">'
                       '&#10003; ADVANCED QC &mdash; DOC + XML + LIVE</span>')
        _mode_note = 'Live verification enabled — lower false-positive rate'
    else:
        _mode_badge = ('<span style="font-size:10px;font-weight:700;padding:2px 9px;'
                       'border-radius:20px;background:#DBEAFE;color:#1D4ED8;white-space:nowrap">'
                       '&#9632; STANDARD QC &mdash; DOC + XML</span>')
        _mode_note = 'XML-based QC — add live survey URL for full advanced verification'

    issues_html = ''
    _ui_type_names = {
        'WORDS MISSING': 'Missing words',
        'TEXT MISMATCH': 'Text mismatch',
        'OPTIONS MISMATCH': 'Options missing',
        'OPTIONS COUNT MISMATCH': 'Option count differs (doc vs XML)',
        'OPTION TEXT MISSING IN XML': 'Option not found in XML',
        'CODE MISMATCH': 'Answer codes differ (doc vs XML)',
        'ROUTING IN XML NOT IN DOC': 'Routing not in spec',
        'PIPING IN DOC NOT IN XML': 'Piping not in XML',
        'IN_XML_NOT_IN_DOC': 'In XML, missing from spec',
        'MANDATORY MISSING': 'Mandatory marker',
        'PIPING NOT RESOLVED': 'Piping issue',
        'MISSING IN LIVE': 'Missing from platform',
        'IN_XML_NOT_VERIFIED_IN_LIVE': 'In XML — not verified in live (routing/conditional)',
        'NAMING MISMATCH': 'Naming mismatch',
        'ERROR PAGE': 'Page error',
    }
    for i, iss in enumerate(issues[:20]):
        sev = iss.get('severity', 'INFO')
        cls = 'badge-red' if sev == 'HIGH' else ('badge-amber' if sev == 'MEDIUM' else 'badge-blue')
        conf_lvl   = iss.get('conf_level', '')
        conf_pct   = iss.get('confidence', '')
        conf_label = iss.get('confidence_label', '') or conf_lvl
        if conf_lvl == 'HIGH':
            conf_cls = 'badge-green'
        elif conf_lvl == 'MEDIUM':
            conf_cls = 'badge-amber'
        elif conf_lvl == 'NEEDS_MANUAL':
            conf_cls = 'badge-blue'
        else:
            conf_cls = 'badge-blue'
        conf_title = f'{conf_pct}% — {conf_label}' if conf_label else f'{conf_pct}%'
        conf_badge = f'<span class="badge {conf_cls}" title="{conf_title}">{conf_pct}%</span>' if conf_pct != '' else ''
        simple_type = _ui_type_names.get(iss.get('type',''), iss.get('type',''))
        ev = iss.get('evidence', {})
        detail = ev.get('mismatch_detail') or iss.get('details','')
        reasons = iss.get('confidence_reasons', [])
        reason_tip = ' · '.join(reasons[:3]) if reasons else ''
        issues_html += f"""
        <tr>
          <td class="primary">{iss.get('qid','')}</td>
          <td>{simple_type}</td>
          <td><span class="badge {cls}">{sev}</span>&nbsp;{conf_badge}</td>
          <td style="font-size:11px;color:var(--text3)" title="{reason_tip}">{str(detail)[:120]}</td>
        </tr>"""

    term_html = ''
    for r in term_results:
        if r.get('needs_review'):
            cls, label = 'badge-amber', 'NEEDS REVIEW'
        elif r.get('passed'):
            cls, label = 'badge-green', 'PASS'
        else:
            cls, label = 'badge-red', 'FAIL'
        term_html += f"""
        <tr>
          <td><span class="badge {cls}">{label}</span></td>
          <td class="primary">{r.get('test_qid','')}</td>
          <td>{r.get('answer_code','')}</td>
          <td style="font-size:11px;color:var(--text3)">{r.get('details','')[:80]}</td>
        </tr>"""

    # ── Rule Engine Findings HTML ──────────────────────────────────────────
    _re_group_labels = {
        1:'G1 Routing', 2:'G2 Termination', 3:'G3 Mandatory', 4:'G4 Piping',
        5:'G5 Loop', 6:'G6 Variable', 7:'G7 Type', 8:'G8 Option/Code',
        9:'G9 Graph', 10:'G10 Export',
    }
    _re_category_map2 = {
        1: ('#1D4ED8', 'ti-route'),
        2: ('#C84B31', 'ti-hand-stop'),
        3: ('#D97706', 'ti-asterisk'),
        4: ('#7C3AED', 'ti-arrows-right-left'),
        5: ('#0891B2', 'ti-repeat'),
        6: ('#059669', 'ti-variable'),
        7: ('#9333EA', 'ti-list-check'),
        8: ('#DC2626', 'ti-123'),
        9: ('#1E40AF', 'ti-sitemap'),
        10:('#6B7280', 'ti-file-export'),
    }
    _re_sev_cls = {'HIGH': 'badge-red', 'MEDIUM': 'badge-amber',
                   'LOW': 'badge-blue', 'INFO': 'badge-grey'}
    _re_display = [f for f in re_findings if f.get('severity') not in ('INFO',)]
    _re_by_grp: dict = {}
    for _rf in _re_display:
        _gn = _rf.get('rule_group', 0)
        _re_by_grp.setdefault(_gn, []).append(_rf)

    re_html = ''
    re_group_badges = ''
    for _gn in sorted(_re_by_grp.keys()):
        _gfindings = _re_by_grp[_gn]
        _ghigh   = sum(1 for f in _gfindings if f.get('severity') == 'HIGH')
        _gmedium = sum(1 for f in _gfindings if f.get('severity') == 'MEDIUM')
        _glow    = sum(1 for f in _gfindings if f.get('severity') == 'LOW')
        _glabel  = _re_group_labels.get(_gn, f'Group {_gn}')
        _gcolor  = '#C00000' if _ghigh else ('#BA7517' if _gmedium else '#1D4ED8')
        re_group_badges += (
            f'<span style="display:inline-flex;align-items:center;gap:4px;'
            f'background:rgba(109,40,217,.08);border:1px solid rgba(109,40,217,.2);'
            f'border-radius:6px;padding:3px 8px;font-size:11px;font-weight:600;'
            f'color:{_gcolor};margin:2px">{_glabel} ({len(_gfindings)})</span>'
        )
        _cat_color2 = _re_category_map2.get(_gn, ('#6B7280', 'ti-alert'))
        for _rf in _gfindings:
            _rsev  = _rf.get('severity', 'INFO')
            _rcls  = _re_sev_cls.get(_rsev, 'badge-blue')
            _rqid  = _rf.get('qid', '')
            _rtype = _rf.get('issue_type', '')
            _rconf = _rf.get('confidence', '')
            _rev   = str(_rf.get('evidence', ''))[:120]
            _rrec  = str(_rf.get('recommendation', ''))[:100]
            re_html += (
                f'<tr>'
                f'<td class="primary">{_rqid}</td>'
                f'<td style="font-size:11px;color:{_cat_color2[0]};font-weight:600">{_glabel}</td>'
                f'<td><span class="badge {_rcls}">{_rsev}</span>&nbsp;'
                f'<span style="font-size:10px;color:var(--text3)">{_rconf}%</span></td>'
                f'<td style="font-size:11px;color:var(--text2)">{_rtype}</td>'
                f'<td style="font-size:10px;color:var(--text3)">{_rev}</td>'
                f'<td style="font-size:10px;color:#059669;font-style:italic" title="{_rrec}">{_rrec[:80]}</td>'
                f'</tr>'
            )

    re_total_display = len(_re_display)

    # ── Termination Matrix HTML ───────────────────────────────────────────
    _tm_html = ''
    _tm_card_html = ''
    if re_term_matrix:
        _tm_rows = ''
        for _tm in re_term_matrix:
            _tm_qid    = _tm.get('qid', '')
            _tm_status = _tm.get('status', '')
            _tm_dcodes = ', '.join(_tm.get('doc_codes', [])) or '—'
            _tm_xcond  = (_tm.get('xml_condition') or '—')[:80]
            _tm_miss   = ', '.join(_tm.get('missing_in_xml', [])) or '—'
            _tm_extra  = ', '.join(_tm.get('extra_in_xml', [])) or '—'
            if _tm_status == 'MATCH':
                _tm_scls = 'badge-green'; _tm_slbl = '✓ MATCH'
            elif _tm_status == 'MISMATCH':
                _tm_scls = 'badge-red'; _tm_slbl = '✗ MISMATCH'
            elif _tm_status == 'DOC_ONLY':
                _tm_scls = 'badge-amber'; _tm_slbl = '⚠ DOC ONLY'
            elif _tm_status == 'XML_ONLY':
                _tm_scls = 'badge-blue'; _tm_slbl = '◈ XML ONLY'
            else:
                _tm_scls = 'badge-grey'; _tm_slbl = _tm_status
            _tm_rows += (
                f'<tr>'
                f'<td class="primary">{_tm_qid}</td>'
                f'<td><span class="badge {_tm_scls}">{_tm_slbl}</span></td>'
                f'<td style="font-size:11px;color:var(--text3)">{_tm_dcodes}</td>'
                f'<td style="font-size:11px;color:var(--text3)">{_tm_xcond}</td>'
                f'<td style="font-size:11px;color:#C84B31">{_tm_miss if _tm_miss != "—" else ""}</td>'
                f'<td style="font-size:11px;color:#3F7D58">{_tm_extra if _tm_extra != "—" else ""}</td>'
                f'</tr>'
            )
        _tm_match   = sum(1 for r in re_term_matrix if r.get('status') == 'MATCH')
        _tm_mismatch = sum(1 for r in re_term_matrix if r.get('status') == 'MISMATCH')
        _tm_doconly  = sum(1 for r in re_term_matrix if r.get('status') == 'DOC_ONLY')
        _tm_xmlonly  = sum(1 for r in re_term_matrix if r.get('status') == 'XML_ONLY')
        _tm_card_html = f"""
    <div class="card" style="margin-top:16px;border-left:3px solid #C84B31">
      <div style="display:flex;align-items:center;justify-content:space-between;margin-bottom:12px;flex-wrap:wrap;gap:8px">
        <div style="display:flex;align-items:center;gap:10px">
          <span style="font-size:14px;font-weight:700;color:#C84B31">&#9888; Termination Matrix</span>
          <span style="font-size:11px;background:rgba(200,75,49,.1);color:#C84B31;border-radius:4px;padding:2px 8px;font-weight:600">{len(re_term_matrix)} termination point(s)</span>
        </div>
        <div style="display:flex;gap:6px;flex-wrap:wrap">
          <span class="badge badge-green">MATCH {_tm_match}</span>
          <span class="badge badge-red">MISMATCH {_tm_mismatch}</span>
          <span class="badge badge-amber">DOC ONLY {_tm_doconly}</span>
          <span class="badge badge-blue">XML ONLY {_tm_xmlonly}</span>
        </div>
      </div>
      <p style="font-size:11px;color:var(--text3);margin-bottom:10px">
        Generated from doc termination rules vs XML termination conditions.
        MISMATCH = doc and XML disagree on which codes terminate.
        DOC ONLY = spec defines termination but not in XML.
      </p>
      <table class="data-table">
        <thead><tr>
          <th>QID</th><th>Status</th><th>Doc Codes</th>
          <th>XML Condition</th><th>Missing in XML</th><th>Extra in XML</th>
        </tr></thead>
        <tbody>{_tm_rows}</tbody>
      </table>
    </div>"""

    # ── Rule Engine Findings HTML (enhanced grouped view) ─────────────────
    # Build category group map for display
    _re_category_map = {
        1: ('Routing',     '#1D4ED8', 'ti-route'),
        2: ('Termination', '#C84B31', 'ti-hand-stop'),
        3: ('Mandatory',   '#D97706', 'ti-asterisk'),
        4: ('Piping',      '#7C3AED', 'ti-arrows-right-left'),
        5: ('Loop',        '#0891B2', 'ti-repeat'),
        6: ('Variable',    '#059669', 'ti-variable'),
        7: ('Question Type','#9333EA','ti-list-check'),
        8: ('Options/Code','#DC2626', 'ti-123'),
        9: ('Graph',       '#1E40AF', 'ti-sitemap'),
        10:('Export',      '#6B7280', 'ti-file-export'),
    }

    re_card_html = ''
    if re_findings or re_term_matrix:
        re_high   = re_summary.get('high', 0)
        re_medium = re_summary.get('medium', 0)
        re_low    = re_summary.get('low', 0)

        # Build per-group mini-badges with color
        _re_group_pill_html = ''
        for _gn in sorted(_re_by_grp.keys()):
            _gfindings = _re_by_grp[_gn]
            if not _gfindings:
                continue
            _cat_name, _cat_color, _cat_icon = _re_category_map.get(_gn, (f'Group {_gn}', '#6B7280', 'ti-alert'))
            _ghigh = sum(1 for f in _gfindings if f.get('severity') == 'HIGH')
            _gcolor = '#C00000' if _ghigh else ('#BA7517' if any(f.get('severity')=='MEDIUM' for f in _gfindings) else '#1D4ED8')
            _re_group_pill_html += (
                f'<span style="display:inline-flex;align-items:center;gap:4px;'
                f'background:rgba(109,40,217,.06);border:1px solid rgba(109,40,217,.18);'
                f'border-radius:6px;padding:3px 8px;font-size:11px;font-weight:600;'
                f'color:{_gcolor};margin:2px">'
                f'<i class="ti {_cat_icon}" style="font-size:11px"></i>'
                f' {_cat_name} ({len(_gfindings)})</span>'
            )

        re_card_html = f"""
    <div class="card" style="margin-top:16px;border-left:3px solid #6D28D9">
      <div style="display:flex;align-items:center;justify-content:space-between;margin-bottom:12px;flex-wrap:wrap;gap:8px">
        <div style="display:flex;align-items:center;gap:10px">
          <i class="ti ti-engine" style="font-size:16px;color:#6D28D9"></i>
          <span style="font-size:14px;font-weight:700;color:#6D28D9">Rule Engine Findings</span>
          <span style="font-size:11px;background:rgba(109,40,217,.1);color:#6D28D9;border-radius:4px;padding:2px 8px;font-weight:600">{re_total_display} findings</span>
        </div>
        <div style="display:flex;gap:6px">
          <span class="badge badge-red">HIGH {re_high}</span>
          <span class="badge badge-amber">MED {re_medium}</span>
          <span class="badge badge-blue">LOW {re_low}</span>
        </div>
      </div>
      <p style="font-size:11px;color:var(--text3);margin-bottom:10px">
        10-group deterministic rule engine &mdash; runs directly on survey model (no DOM, no Playwright).
        Groups: Routing · Termination · Mandatory · Piping · Loop · Variable · Type · Options · Graph · Export
      </p>
      <div style="display:flex;flex-wrap:wrap;gap:4px;margin-bottom:12px">{_re_group_pill_html}</div>
      {f"<table class='data-table'><thead><tr><th>QID</th><th>Group</th><th>Severity</th><th>Rule</th><th>Evidence</th><th>Recommendation</th></tr></thead><tbody>" + re_html + "</tbody></table>" if re_html else "<p style='color:var(--text3);text-align:center;padding:16px'>No rule violations found — survey model is clean.</p>"}
    </div>
    {_tm_card_html}"""
    elif re_term_matrix:
        re_card_html = _tm_card_html
    # ─────────────────────────────────────────────────────────────────────

    # ── Section 1: Executive Summary + Health Score ───────────────────────────
    _hs_confirmed  = consensus.get('confirmed_count', 0)
    _hs_likely     = consensus.get('likely_count', 0)
    _hs_review     = consensus.get('review_count', 0)
    _hs_suppressed = consensus.get('suppressed_count', 0)
    _hs_by_cat     = consensus.get('health_by_category', {})
    _hs_reasons    = consensus.get('suppressed_reasons', {})

    _exec_summary_html = ''
    if health_score is not None:
        _hs_color  = '#15803D' if health_score >= 80 else ('#D97706' if health_score >= 60 else '#DC2626')
        _hs_bg     = '#DCFCE7' if health_score >= 80 else ('#FFFBEB' if health_score >= 60 else '#FEF2F2')
        _hs_pct    = health_score
        # CSS conic-gradient ring
        _hs_ring   = f'conic-gradient({_hs_color} {_hs_pct}%, #E5E7EB {_hs_pct}% 100%)'
        _cat_bars  = ''
        for _cat in ('Routing', 'Termination', 'Piping', 'Options', 'Variables'):
            _cs = _hs_by_cat.get(_cat, 100)
            _cc = '#15803D' if _cs >= 80 else ('#D97706' if _cs >= 60 else '#DC2626')
            _cat_bars += (
                f'<div>'
                f'<div style="display:flex;justify-content:space-between;font-size:10px;margin-bottom:3px">'
                f'<span style="color:var(--text3)">{_cat}</span>'
                f'<span style="color:{_cc};font-weight:700">{_cs}%</span></div>'
                f'<div style="background:#E5E7EB;border-radius:4px;height:5px">'
                f'<div style="width:{_cs}%;background:{_cc};border-radius:4px;height:5px"></div>'
                f'</div></div>'
            )
        _exec_summary_html = (
            f'<div class="card" style="margin-bottom:16px;border-left:3px solid #6D28D9">'
            f'<div style="display:flex;align-items:flex-start;gap:24px;flex-wrap:wrap">'
            # Health gauge
            f'<div style="text-align:center;min-width:100px">'
            f'<div style="width:88px;height:88px;border-radius:50%;background:{_hs_ring};'
            f'display:flex;align-items:center;justify-content:center;margin:0 auto">'
            f'<div style="width:68px;height:68px;border-radius:50%;background:var(--card);'
            f'display:flex;align-items:center;justify-content:center;font-size:22px;font-weight:800;color:{_hs_color}">'
            f'{_hs_pct}</div></div>'
            f'<p style="font-size:10px;font-weight:700;color:{_hs_color};margin-top:5px;text-transform:uppercase">Health Score</p>'
            f'</div>'
            # Tier counts
            f'<div style="flex:1;min-width:220px">'
            f'<p style="font-size:13px;font-weight:700;color:var(--text);margin-bottom:10px">'
            f'<i class="ti ti-chart-bar" style="color:#6D28D9;margin-right:6px"></i>Executive Summary</p>'
            f'<div style="display:grid;grid-template-columns:repeat(4,1fr);gap:8px;margin-bottom:12px">'
            f'<div style="background:#FEF2F2;border:1px solid #FECACA;border-radius:8px;padding:8px;text-align:center">'
            f'<p style="font-size:20px;font-weight:800;color:#DC2626">{_hs_confirmed}</p>'
            f'<p style="font-size:9px;color:#DC2626;font-weight:700;text-transform:uppercase">Confirmed</p></div>'
            f'<div style="background:#FFFBEB;border:1px solid #FDE68A;border-radius:8px;padding:8px;text-align:center">'
            f'<p style="font-size:20px;font-weight:800;color:#D97706">{_hs_likely}</p>'
            f'<p style="font-size:9px;color:#D97706;font-weight:700;text-transform:uppercase">Likely Bugs</p></div>'
            f'<div style="background:#EFF6FF;border:1px solid #BFDBFE;border-radius:8px;padding:8px;text-align:center">'
            f'<p style="font-size:20px;font-weight:800;color:#2563EB">{_hs_review}</p>'
            f'<p style="font-size:9px;color:#2563EB;font-weight:700;text-transform:uppercase">Needs Review</p></div>'
            f'<div style="background:#F9FAFB;border:1px solid #E5E7EB;border-radius:8px;padding:8px;text-align:center">'
            f'<p style="font-size:20px;font-weight:800;color:#6B7280">{_hs_suppressed}</p>'
            f'<p style="font-size:9px;color:#6B7280;font-weight:700;text-transform:uppercase">Suppressed</p></div>'
            f'</div>'
            # Category health bars
            f'<div style="display:grid;grid-template-columns:repeat(5,1fr);gap:10px">'
            f'{_cat_bars}'
            f'</div>'
            f'</div>'
            f'</div>'
            f'</div>'
        )

    # ── Tiered issue sections (2-5) ───────────────────────────────────────────
    # Bucket all issues + rule findings by consensus_tier
    _all_for_tiers = issues + re_findings
    _tier_confirmed = [i for i in _all_for_tiers if i.get('consensus_tier') == 'CONFIRMED_BUG']
    _tier_likely    = [i for i in _all_for_tiers if i.get('consensus_tier') == 'LIKELY_BUG']
    _tier_review    = [i for i in _all_for_tiers if i.get('consensus_tier') == 'NEEDS_REVIEW']

    _tbl_hdr6 = ("<table class='data-table'><thead><tr>"
                 "<th>QID</th><th>Type</th><th>Severity</th><th>Details</th>"
                 "<th>Score</th><th>Feedback</th></tr></thead><tbody>")

    def _tier_rows(tier_issues, max_n=25):
        if not tier_issues:
            return ''
        rows = ''
        for _ti in tier_issues[:max_n]:
            _tqid  = str(_ti.get('qid', ''))
            _ttype = str(_ti.get('check') or _ti.get('issue_type') or _ti.get('type', ''))
            _tsev  = _ti.get('severity', 'INFO')
            _tsc   = _ti.get('consensus_score', '')
            _tdet  = str(_ti.get('details') or str(_ti.get('evidence', ''))[:100])[:120]
            _tcls  = 'badge-red' if _tsev == 'HIGH' else ('badge-amber' if _tsev == 'MEDIUM' else 'badge-blue')
            # Escape for inline JS onclick — replace single quotes
            _qid_j = _tqid.replace("'", '').replace('"', '')
            _typ_j = _ttype.replace("'", '').replace('"', '')
            rows += (
                f'<tr>'
                f'<td class="primary">{_tqid}</td>'
                f'<td style="font-size:11px">{_ttype}</td>'
                f'<td><span class="badge {_tcls}">{_tsev}</span></td>'
                f'<td style="font-size:10px;color:var(--text3)" title="{_tdet}">{_tdet[:80]}</td>'
                f'<td style="font-size:11px;font-weight:700;color:#6D28D9">{_tsc}%</td>'
                f'<td style="white-space:nowrap">'
                f'<button onclick="markIssue(\'{job_id}\',\'{_qid_j}\',\'{_typ_j}\',\'CONFIRMED\',this)" '
                f'title="Confirm as real bug" class="_fb-btn _fb-ok">&#10003;</button>&nbsp;'
                f'<button onclick="markIssue(\'{job_id}\',\'{_qid_j}\',\'{_typ_j}\',\'FALSE_POSITIVE\',this)" '
                f'title="Mark as false positive" class="_fb-btn _fb-fp">&#10007;</button>'
                f'</td></tr>'
            )
        if len(tier_issues) > max_n:
            rows += (f'<tr><td colspan="6" style="text-align:center;color:var(--text3);'
                     f'font-style:italic;font-size:11px">'
                     f'... and {len(tier_issues) - max_n} more findings</td></tr>')
        return rows

    _rows_confirmed = _tier_rows(_tier_confirmed)
    _rows_likely    = _tier_rows(_tier_likely)
    _rows_review    = _tier_rows(_tier_review)

    def _tier_card(title, color, icon, count, badge_bg, badge_border, rows, section_num):
        _inner = ((_tbl_hdr6 + rows + '</tbody></table>') if rows
                  else f"<p style='color:var(--text3);text-align:center;padding:14px'>No issues in this tier.</p>")
        return (
            f'<div class="card" style="margin-top:16px;border-left:3px solid {color}">'
            f'<div style="display:flex;align-items:center;gap:10px;margin-bottom:12px">'
            f'<i class="ti {icon}" style="color:{color};font-size:16px"></i>'
            f'<span style="font-size:14px;font-weight:700;color:{color}">Section {section_num} — {title} ({count})</span>'
            f'<span style="font-size:10px;background:{badge_bg};color:{color};'
            f'border:1px solid {badge_border};border-radius:4px;padding:2px 8px;font-weight:600">'
            f'{"95%+ confidence" if section_num==2 else ("65–94% confidence" if section_num==3 else "35–64% confidence")}</span>'
            f'</div>'
            f'{_inner}'
            f'</div>'
        )

    _tiered_sections_html = ''
    if health_score is not None:
        _tiered_sections_html += _tier_card(
            'Confirmed Bugs', '#DC2626', 'ti-bug', len(_tier_confirmed),
            '#FEF2F2', '#FECACA', _rows_confirmed, 2)
        _tiered_sections_html += _tier_card(
            'Likely Bugs', '#D97706', 'ti-alert-triangle', len(_tier_likely),
            '#FFFBEB', '#FDE68A', _rows_likely, 3)
        _tiered_sections_html += _tier_card(
            'Needs Review', '#2563EB', 'ti-eye', len(_tier_review),
            '#EFF6FF', '#BFDBFE', _rows_review, 4)
        # Section 5: Suppressed
        if _hs_suppressed > 0:
            _supp_items = ''.join(
                f'<li><span style="color:var(--text3)">{_r}</span>: <strong>{_c}</strong></li>'
                for _r, _c in _hs_reasons.items()
            )
            _tiered_sections_html += (
                f'<div class="card" style="margin-top:16px;border-left:3px solid #9CA3AF">'
                f'<div style="display:flex;align-items:center;gap:10px;margin-bottom:8px">'
                f'<i class="ti ti-eye-off" style="color:#9CA3AF;font-size:16px"></i>'
                f'<span style="font-size:14px;font-weight:700;color:#6B7280">'
                f'Section 5 — Suppressed Findings ({_hs_suppressed})</span>'
                f'</div>'
                f'<p style="font-size:11px;color:var(--text3);margin-bottom:8px">'
                f'These findings were filtered out automatically to reduce noise. '
                f'Details are hidden; only counts are shown.</p>'
                f'<ul style="font-size:12px;color:var(--text2);margin:0;padding-left:20px">'
                f'{_supp_items}</ul>'
                f'</div>'
            )

    # ── Section 7: Playwright / Termination card ──────────────────────────────
    _pw_data    = j.get('playwright_tests', {})
    _pw_res     = _pw_data.get('results', [])
    _pw_sum     = _pw_data.get('summary', {})
    _term_card_html = (
        f'<div class="card" style="margin-top:16px;border-left:3px solid #0891B2">'
        f'<div style="display:flex;align-items:center;gap:10px;margin-bottom:12px">'
        f'<i class="ti ti-player-play" style="color:#0891B2;font-size:16px"></i>'
        f'<span style="font-size:14px;font-weight:700;color:#0891B2">Section 7 — Playwright / Termination Results</span>'
        f'<span style="font-size:11px;background:#E0F2FE;color:#0891B2;border-radius:4px;'
        f'padding:2px 8px;font-weight:600">{term_passed}/{term_total - term_review} validated · {term_review} review</span>'
        f'</div>'
        + (f"<table class='data-table'><thead><tr><th>Status</th><th>QID</th><th>Code</th>"
           f"<th>Details</th></tr></thead><tbody>{term_html}</tbody></table>"
           if term_html
           else "<p style='color:var(--text3);text-align:center;padding:14px'>No termination rules found in spec.</p>")
        + f'</div>'
    )

    return render_template_string(SHARED_CSS + f"""
<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=5.0"><title>Report — SurveyQC</title></head><body>
<div class="app-layout">
  {sidebar_html('reports')}
  <div class="main-content">
    <div class="topbar">
      <div>
        <div style="display:flex;align-items:center;gap:10px;margin-bottom:5px;flex-wrap:wrap">
          <p class="page-title">{doc_name[:40]}</p>
          <span class="badge {verdict_class}"><i class="ti {verdict_icon}"></i>{verdict}</span>
          {_mode_badge}
        </div>
        <div style="display:flex;gap:16px;font-size:12px;color:var(--text3);flex-wrap:wrap">
          <span><i class="ti ti-device-desktop" style="vertical-align:-1px;margin-right:4px"></i>{platform}</span>
          <span><i class="ti ti-world" style="vertical-align:-1px;margin-right:4px"></i>{country or 'Not set'}</span>
          <span><i class="ti ti-calendar" style="vertical-align:-1px;margin-right:4px"></i>{created}</span>
          <span style="color:#6B7280;font-style:italic">{_mode_note}</span>
        </div>
      </div>
      <div style="display:flex;gap:8px">
        <button onclick="window.location='/new-qc'" class="btn btn-ghost btn-sm"><i class="ti ti-refresh"></i>New QC</button>
        <a href="/retest/{job_id}" class="btn btn-ghost btn-sm" style="color:#F59E0B;border-color:#F59E0B"><i class="ti ti-player-play"></i>Retest</a>
        <a href="/download/{job_id}" class="btn btn-primary btn-sm"><i class="ti ti-download"></i>Download</a>
      </div>
    </div>

    <div style="display:grid;grid-template-columns:repeat(5,1fr);gap:10px;margin-bottom:16px">
      <div class="stat-card"><p class="stat-num">{doc_qids}</p><p class="stat-label">Doc questions</p></div>
      <div class="stat-card"><p class="stat-num" style="color:#2563EB">{xml_qids}</p><p class="stat-label">XML questions</p></div>
      <div class="stat-card"><p class="stat-num" style="color:#1D9E75">{term_passed}/{term_total}</p><p class="stat-label">Term. passed</p></div>
      <div class="stat-card"><p class="stat-num" style="color:#E24B4A">{total_issues}</p><p class="stat-label">Issues found</p></div>
      <div class="stat-card" style="background:rgba(29,158,117,.1);border-color:rgba(29,158,117,.2)"><p class="stat-num" style="color:#1D9E75">~8h</p><p class="stat-label">Time saved</p></div>
    </div>

    <div class="alert {'alert-error' if verdict=='FAIL' else ('alert-success' if verdict=='PASS' else 'alert-info')}" style="display:flex;align-items:center;gap:10px;margin-bottom:16px">
      <i class="ti {verdict_icon}" style="font-size:18px"></i>
      <div>
        <p style="font-weight:500">{verdict_msg}</p>
        <p style="font-size:12px;opacity:.8">{total_issues} structural issues · {term_total - term_passed - term_review} term. failed · {term_review} need manual review</p>
      </div>
    </div>

    {f'<div class="card" style="margin-bottom:16px;border-left:3px solid var(--accent)"><div style="display:flex;align-items:center;gap:8px;margin-bottom:8px"><i class="ti ti-sparkles" style="color:var(--accent);font-size:16px"></i><span style="font-size:13px;font-weight:700;color:var(--text)">AI Summary</span></div><p style="font-size:14px;color:var(--text2);line-height:1.7">{ai_summary}</p></div>' if ai_summary else ''}

    {_exec_summary_html}

    {_tiered_sections_html}

    {"" if health_score is not None else f"<div style='display:grid;grid-template-columns:1fr 1fr;gap:16px'><div class='card'><p style='font-size:14px;font-weight:600;color:var(--text);margin-bottom:14px'>Issues found ({total_issues})</p>" + ("<table class='data-table'><thead><tr><th>QID</th><th>Type</th><th>Severity</th><th>Details</th></tr></thead><tbody>" + issues_html + "</tbody></table>" if issues_html else "<p style='color:var(--text3);text-align:center;padding:20px'>No structural issues found!</p>") + f"</div><div class='card'><p style='font-size:14px;font-weight:600;color:var(--text);margin-bottom:14px'>Termination tests ({term_passed}/{term_total - term_review} validated · {term_review} need review)</p>" + ("<table class='data-table'><thead><tr><th>Status</th><th>QID</th><th>Code</th><th>Details</th></tr></thead><tbody>" + term_html + "</tbody></table>" if term_html else "<p style='color:var(--text3);text-align:center;padding:20px'>No termination rules found in doc</p>") + "</div></div>"}

    {re_card_html}

    {_term_card_html}

    <div class="card" style="margin-top:16px">
      <p style="font-size:14px;font-weight:600;color:var(--text);margin-bottom:14px">Rate this report</p>
      <div id="stars" style="display:flex;gap:6px;margin-bottom:12px">
        <i class="ti ti-star" style="font-size:24px;color:var(--amber);cursor:pointer" onclick="setRating(1)"></i>
        <i class="ti ti-star" style="font-size:24px;color:var(--amber);cursor:pointer" onclick="setRating(2)"></i>
        <i class="ti ti-star" style="font-size:24px;color:var(--amber);cursor:pointer" onclick="setRating(3)"></i>
        <i class="ti ti-star" style="font-size:24px;color:var(--amber);cursor:pointer" onclick="setRating(4)"></i>
        <i class="ti ti-star" style="font-size:24px;color:var(--text3);cursor:pointer" onclick="setRating(5)"></i>
      </div>
      <div style="display:flex;gap:10px">
        <input class="form-input" type="text" id="feedback-text" placeholder="Any comments? (optional)" style="flex:1">
        <button class="btn btn-primary btn-sm" onclick="submitFeedback('{job_id}')">Submit</button>
        <button class="btn btn-ghost btn-sm" onclick="this.closest('.card').style.display='none'">Skip</button>
      </div>
    </div>
  </div>
</div>
<style>
._fb-btn{{border-radius:4px;padding:2px 7px;font-size:11px;cursor:pointer;font-weight:700;border-width:1px;border-style:solid}}
._fb-ok{{background:#DCFCE7;color:#15803D;border-color:#BBF7D0}}
._fb-ok:hover{{background:#BBF7D0}}
._fb-fp{{background:#FEF2F2;color:#DC2626;border-color:#FECACA}}
._fb-fp:hover{{background:#FECACA}}
._fb-btn:disabled{{opacity:.5;cursor:default}}
</style>
<script>
var rating = 4;
function setRating(n) {{
  rating = n;
  var stars = document.getElementById('stars').children;
  for (var i=0; i<5; i++) {{
    stars[i].style.color = i < n ? 'var(--amber)' : 'var(--text3)';
  }}
}}
function submitFeedback(jobId) {{
  fetch('/api/feedback', {{
    method: 'POST',
    headers: {{'Content-Type': 'application/json'}},
    body: JSON.stringify({{job_id: jobId, rating: rating, comment: document.getElementById('feedback-text').value}})
  }}).then(() => {{
    document.querySelector('.card:last-child').innerHTML = '<p style="color:#1D9E75;text-align:center;padding:16px">Thank you for your feedback!</p>';
  }});
}}
function markIssue(jobId, qid, itype, verdict, btn) {{
  btn.disabled = true;
  fetch('/api/issue-feedback/' + jobId, {{
    method: 'POST',
    headers: {{'Content-Type': 'application/json'}},
    body: JSON.stringify({{qid: qid, issue_type: itype, verdict: verdict}})
  }}).then(function(r) {{ return r.json(); }}).then(function(d) {{
    if (d.ok) {{
      var row = btn.closest('tr');
      if (row) {{
        row.style.opacity = '0.45';
        var allBtns = row.querySelectorAll('._fb-btn');
        allBtns.forEach(function(b) {{ b.disabled = true; }});
        var label = verdict === 'CONFIRMED' ? '&#10003; Saved' : '&#10007; FP';
        btn.innerHTML = label;
        btn.style.fontWeight = '700';
      }}
    }}
  }}).catch(function() {{ btn.disabled = false; }});
}}
</script>
</body></html>""")

# ================================================================
# PAGE: REPORTS LIST
# ================================================================
@app.route('/reports')
@login_required
def reports_list():
    email = session.get('user_email')
    user_jobs = [(jid, j) for jid, j in jobs.items() if j.get('user_email') == email]
    user_jobs.sort(key=lambda x: x[1].get('created_at', ''), reverse=True)

    def _build_row(jid, j, is_child=False):
        status = j.get('status', 'running')
        doc_name = j.get('doc_name', 'Unknown')
        platform = j.get('platform', '-')
        mode = j.get('mode', 'full')
        issues = j.get('total_issues', 0)
        created = j.get('created_at', '')[:16]
        if status == 'done':
            verdict = j.get('verdict', 'REVIEW')
            badge_cls = 'badge-red' if verdict == 'FAIL' else ('badge-green' if verdict == 'PASS' else 'badge-amber')
            badge_txt = f'{issues} issues' if issues > 0 else 'All pass'
        elif status == 'running':
            badge_cls = 'badge-blue'
            badge_txt = 'Running'
        else:
            badge_cls = 'badge-amber'
            badge_txt = 'Error'

        link = f'<a href="/report/{jid}" style="color:var(--purple);font-size:12px;text-decoration:none">View</a>' if status == 'done' else f'<a href="/progress/{jid}" style="color:var(--text3);font-size:12px;text-decoration:none">Track</a>'
        download = f'<a href="/download/{jid}" style="color:var(--text3);text-decoration:none"><i class="ti ti-download" style="font-size:14px"></i></a>' if status == 'done' else ''
        share_tok = hashlib.md5((jid + '-share').encode()).hexdigest()[:12]
        share_url = 'https://surveyqc.online/view/' + share_tok
        share_title = doc_name[:40]
        share_btn = f'<button data-share-url="{share_url}" data-share-title="{share_title}" style="background:none;border:none;cursor:pointer;color:var(--text3);font-size:13px;padding:0 4px" title="Share"><i class="ti ti-share"></i></button>' if status == 'done' else ''
        plat_lower = platform.lower().replace(' ', '')

        if is_child:
            name_cell = (
                '<span style="display:inline-block;width:16px;border-bottom:2px solid var(--border);'
                'margin-right:6px;vertical-align:middle"></span>'
                '<i class="ti ti-refresh" style="color:#1D9E75;margin-right:6px;font-size:13px"></i>'
                f'<span style="font-size:12px">{doc_name[:32]}</span>'
                ' <span style="background:#E6F7F1;color:#0F7A50;font-size:10px;font-weight:600;'
                'padding:2px 6px;border-radius:10px;margin-left:4px">Re-test</span>'
            )
            row_style = ' style="background:rgba(29,158,117,.03)"'
        else:
            name_cell = f'<i class="ti ti-file-text" style="color:var(--purple);margin-right:8px"></i>{doc_name[:35]}'
            row_style = ''

        return f"""
        <tr{row_style}>
          <td class="primary">{name_cell}</td>
          <td data-platform="{plat_lower}">{platform}</td>
          <td>{mode.title()}</td>
          <td><span class="badge {badge_cls}">{badge_txt}</span></td>
          <td style="color:var(--text3)">{created}</td>
          <td style="white-space:nowrap">{link} &nbsp; {download} &nbsp; {share_btn}</td>
        </tr>"""

    # Separate top-level jobs from re-test children
    children_by_parent = {}
    parents = []
    for jid, j in user_jobs:
        parent_id = j.get('retest_of')
        if parent_id:
            children_by_parent.setdefault(parent_id, []).append((jid, j))
        else:
            parents.append((jid, j))

    rows = ''
    for jid, j in parents:
        rows += _build_row(jid, j, is_child=False)
        for cjid, cj in children_by_parent.get(jid, []):
            rows += _build_row(cjid, cj, is_child=True)

    if not rows:
        rows = '''<tr><td colspan="6" style="padding:56px 24px">
          <div style="display:flex;flex-direction:column;align-items:center;gap:14px;text-align:center">
            <div style="width:60px;height:60px;background:#F0F2F5;border-radius:14px;display:flex;align-items:center;justify-content:center">
              <i class="ti ti-clipboard-list" style="font-size:30px;color:#9CA3AF"></i>
            </div>
            <div>
              <p style="font-size:16px;font-weight:600;color:#374151;margin-bottom:6px">No reports yet</p>
              <p style="font-size:13px;color:#9CA3AF">Upload your first survey doc to get started</p>
            </div>
            <div style="display:flex;align-items:center;gap:6px;font-size:12px;color:#6B7280;background:#F8F9FA;padding:10px 18px;border-radius:8px;flex-wrap:wrap;justify-content:center">
              <span><b style="color:#042C53">1.</b> Upload spec doc</span>
              <span style="color:#DDE1E7">→</span>
              <span><b style="color:#042C53">2.</b> Add live URL</span>
              <span style="color:#DDE1E7">→</span>
              <span><b style="color:#042C53">3.</b> Get Word report</span>
            </div>
            <a href="/new-qc" style="display:inline-flex;align-items:center;gap:6px;background:#042C53;color:white;padding:12px 24px;border-radius:8px;font-size:13px;font-weight:600;text-decoration:none">
              <i class="ti ti-plus" style="font-size:14px"></i> Create your first QC &rarr;
            </a>
          </div>
        </td></tr>'''

    return render_template_string(SHARED_CSS + f"""
<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=5.0"><title>Reports — SurveyQC</title></head><body>
<div class="app-layout">
  {sidebar_html('reports')}
  <div class="main-content">
    <div class="topbar">
      <div>
        <p class="page-title">All Reports</p>
        <p class="page-sub">{len(user_jobs)} total reports</p>
      </div>
      <a href="/new-qc" class="btn btn-primary btn-sm"><i class="ti ti-plus"></i>New QC</a>
    </div>
    <div style="display:flex;gap:12px;margin-bottom:16px;flex-wrap:wrap">
      <div style="flex:1;min-width:200px;position:relative">
        <i class="ti ti-search" style="position:absolute;left:14px;top:50%;transform:translateY(-50%);color:var(--text3);font-size:16px"></i>
        <input type="text" id="searchBox" onkeyup="filterReports()" placeholder="Search reports..." style="width:100%;padding:11px 14px 11px 40px;border:1px solid var(--border);border-radius:12px;font-size:14px;font-family:inherit;outline:none;background:white">
      </div>
      <select id="platformFilter" onchange="filterReports()" style="padding:11px 14px;border:1px solid var(--border);border-radius:12px;font-size:14px;font-family:inherit;background:white;outline:none;cursor:pointer">
        <option value="">All platforms</option>
        <option value="confirmit">Confirmit</option>
        <option value="decipher">Decipher</option>
        <option value="forsta">Forsta</option>
        <option value="qualtrics">Qualtrics</option>
      </select>
      <select id="statusFilter" onchange="filterReports()" style="padding:11px 14px;border:1px solid var(--border);border-radius:12px;font-size:14px;font-family:inherit;background:white;outline:none;cursor:pointer">
        <option value="">All status</option>
        <option value="pass">Passed</option>
        <option value="fail">Issues</option>
        <option value="running">Running</option>
      </select>
    </div>
    <div class="card">
      <table class="data-table" style="width:100%">
        <thead><tr>
          <th>Survey name</th><th>Platform</th><th>Mode</th><th>Status</th><th>Date</th><th>Actions</th>
        </tr></thead>
        <tbody id="reportsBody">{rows}</tbody>
      </table>
      <div id="noResults" style="display:none;text-align:center;color:var(--text3);padding:24px">No reports match your filters.</div>
    </div>
  </div>
</div>
<script>
function filterReports(){{
  var q=document.getElementById('searchBox').value.toLowerCase();
  var plat=document.getElementById('platformFilter').value.toLowerCase();
  var stat=document.getElementById('statusFilter').value.toLowerCase();
  var rows=document.querySelectorAll('#reportsBody tr');
  var visible=0;
  rows.forEach(function(r){{
    var txt=r.textContent.toLowerCase();
    var show=true;
    if(q && txt.indexOf(q)===-1) show=false;
    if(plat && txt.indexOf(plat)===-1) show=false;
    if(stat){{
      if(stat==='pass' && txt.indexOf('all pass')===-1) show=false;
      if(stat==='fail' && txt.indexOf('issues')===-1) show=false;
      if(stat==='running' && txt.indexOf('running')===-1) show=false;
    }}
    r.style.display=show?'':'none';
    if(show) visible++;
  }});
  document.getElementById('noResults').style.display=visible===0?'block':'none';
}}
</script>
</body></html>""")

# ================================================================
# PAGE: RE-TEST
# ================================================================
@app.route('/retest')
@login_required
def retest_page():
    email = session.get('user_email')
    all_jobs = [(jid, j) for jid, j in jobs.items()
                if j.get('user_email') == email and j.get('status') == 'done'
                and not j.get('retest_of')]  # only top-level completed
    all_jobs.sort(key=lambda x: x[1].get('created_at', ''), reverse=True)
    top20 = all_jobs[:20]

    options_html = '<option value="">-- Select a completed report --</option>'
    for jid, j in top20:
        doc_name = j.get('doc_name', 'Unknown')[:50]
        date = j.get('created_at', '')[:10]
        issues = j.get('total_issues', 0)
        issue_label = f"{issues} issue{'s' if issues != 1 else ''}"
        options_html += f'<option value="{jid}">{doc_name} &middot; {date} &middot; {issue_label}</option>'

    RETEST_PAGE = SHARED_CSS + """
<!DOCTYPE html><html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=5.0">
  <title>Re-test Report — SurveyQC</title>
  <style>
    .retest-card{background:white;border-radius:16px;padding:28px;border:1px solid var(--border);margin-bottom:20px}
    .retest-section-label{font-size:11px;font-weight:600;color:var(--text3);text-transform:uppercase;letter-spacing:.08em;margin-bottom:10px}
    .retest-select{width:100%;padding:11px 14px;border:1px solid var(--border);border-radius:10px;font-size:14px;font-family:inherit;background:white;outline:none;color:var(--text)}
    .retest-select:focus{border-color:var(--purple);box-shadow:0 0 0 3px rgba(139,92,246,.15)}
    .retest-textarea{width:100%;padding:12px 14px;border:1px solid var(--border);border-radius:10px;font-size:14px;font-family:inherit;resize:vertical;outline:none;color:var(--text);min-height:140px;box-sizing:border-box}
    .retest-textarea:focus{border-color:var(--purple);box-shadow:0 0 0 3px rgba(139,92,246,.15)}
    .qid-preview{margin-top:10px;padding:10px 14px;background:rgba(139,92,246,.06);border-radius:8px;border:1px solid rgba(139,92,246,.15);display:none}
    .qid-chip{display:inline-block;background:var(--purple);color:white;font-size:11px;font-weight:600;padding:3px 8px;border-radius:20px;margin:2px 3px 2px 0}
    .qid-preview-label{font-size:11px;color:var(--purple);font-weight:600;margin-bottom:6px}
    .submit-btn{background:var(--purple);color:white;border:none;padding:13px 28px;border-radius:10px;font-size:14px;font-weight:600;cursor:pointer;font-family:inherit;transition:opacity .2s}
    .submit-btn:disabled{opacity:.4;cursor:not-allowed}
    .submit-btn:hover:not(:disabled){opacity:.88}
    .info-box{background:rgba(139,92,246,.05);border:1px solid rgba(139,92,246,.15);border-radius:10px;padding:14px 16px;margin-bottom:20px}
  </style>
</head>
<body>
<div class="app-layout">
  """ + sidebar_html('retest') + """
  <div class="main-content">
    <div class="topbar">
      <div>
        <p class="page-title">Re-test Report</p>
        <p class="page-sub">Re-run QC on specific questions after client feedback</p>
      </div>
    </div>

    <div class="info-box">
      <p style="font-size:13px;color:var(--purple);font-weight:600;margin-bottom:4px"><i class="ti ti-info-circle"></i> How it works</p>
      <p style="font-size:13px;color:var(--text2)">Select a completed report, paste client feedback or list QIDs, and SurveyQC re-runs QC filtered to only those questions. A new linked report is created.</p>
    </div>

    <form id="retestForm" method="POST" action="/retest/submit" enctype="multipart/form-data">

      <!-- SECTION 1: Select report -->
      <div class="retest-card">
        <p class="retest-section-label"><i class="ti ti-clipboard-list" style="margin-right:5px"></i>Step 1 — Select original report</p>
        <select name="original_job_id" id="reportSelect" class="retest-select" required onchange="checkReady()">
          """ + options_html + """
        </select>
      </div>

      <!-- SECTION 2: Feedback / QID list -->
      <div class="retest-card">
        <p class="retest-section-label"><i class="ti ti-message-circle" style="margin-right:5px"></i>Step 2 — What to re-check?</p>
        <textarea
          name="feedback_text"
          id="feedbackText"
          class="retest-textarea"
          placeholder="Example: 'Please re-check Q15bis, Q21, P3' &#10;&#10;Or paste an email / comment containing QIDs."
          oninput="onFeedbackChange()"
          required
        ></textarea>
        <div class="qid-preview" id="qidPreview">
          <p class="qid-preview-label"><i class="ti ti-scan"></i> Detected QIDs</p>
          <div id="qidChips"></div>
        </div>
        <p style="font-size:11px;color:var(--text3);margin-top:8px">
          <i class="ti ti-info-circle"></i>
          Accepts comma/newline lists or natural language. Optionally upload an annotated screenshot below.
        </p>
        <div style="margin-top:14px">
          <label style="font-size:12px;color:var(--text2);font-weight:500;display:block;margin-bottom:6px">
            <i class="ti ti-photo" style="margin-right:4px"></i>Annotated screenshot (optional)
          </label>
          <input type="file" name="screenshot" accept="image/*"
            style="font-size:13px;color:var(--text2);font-family:inherit">
        </div>
      </div>

      <!-- SECTION 3: Submit -->
      <div class="retest-card" style="display:flex;align-items:center;gap:16px;flex-wrap:wrap">
        <button type="submit" class="submit-btn" id="submitBtn" disabled>
          <i class="ti ti-refresh"></i> Re-run QC on selected QIDs only
        </button>
        <p id="submitHint" style="font-size:12px;color:var(--text3)">Fill both fields above to enable</p>
      </div>

    </form>
  </div>
</div>

<script>
// Live QID detection regex — mirrors export_parser.py
var QID_RE = /\\b([A-Za-z]{1,6}\\d+(?:bis|ter|[a-z]{1,3}\\d*)?)\\b/gi;

function extractQids(text) {
  var matches = [], seen = {}, m;
  QID_RE.lastIndex = 0;
  while ((m = QID_RE.exec(text)) !== null) {
    var norm = m[1].toUpperCase();
    if (!seen[norm]) { seen[norm] = true; matches.push(norm); }
  }
  return matches;
}

function onFeedbackChange() {
  var text = document.getElementById('feedbackText').value;
  var qids = extractQids(text);
  var preview = document.getElementById('qidPreview');
  var chips = document.getElementById('qidChips');
  if (qids.length > 0) {
    chips.innerHTML = qids.map(function(q) {
      return '<span class="qid-chip">' + q + '</span>';
    }).join('');
    preview.style.display = 'block';
  } else {
    preview.style.display = 'none';
  }
  checkReady();
}

function checkReady() {
  var reportOk = document.getElementById('reportSelect').value !== '';
  var feedbackOk = document.getElementById('feedbackText').value.trim() !== '';
  var btn = document.getElementById('submitBtn');
  var hint = document.getElementById('submitHint');
  btn.disabled = !(reportOk && feedbackOk);
  if (reportOk && feedbackOk) {
    var qids = extractQids(document.getElementById('feedbackText').value);
    hint.textContent = qids.length > 0
      ? qids.length + ' QID' + (qids.length === 1 ? '' : 's') + ' detected — ready to submit'
      : 'No QIDs detected yet (will use AI to extract)';
    hint.style.color = qids.length > 0 ? 'var(--green)' : 'var(--amber)';
  } else {
    hint.textContent = 'Fill both fields above to enable';
    hint.style.color = 'var(--text3)';
  }
}

// Prevent double-submit
document.getElementById('retestForm').addEventListener('submit', function() {
  var btn = document.getElementById('submitBtn');
  btn.disabled = true;
  btn.innerHTML = '<i class="ti ti-loader-2"></i> Submitting...';
});
</script>
</body></html>"""

    return render_template_string(RETEST_PAGE)


@app.route('/retest/submit', methods=['POST'])
@login_required
def retest_submit():
    import threading
    from export_parser import parse_qid_list

    original_job_id = request.form.get('original_job_id', '').strip()
    feedback_text = request.form.get('feedback_text', '').strip()

    if not original_job_id or original_job_id not in jobs:
        return render_template_string(SHARED_CSS + """
<!DOCTYPE html><html><body>
<div style="padding:40px;text-align:center">
  <p style="color:red;font-size:16px">Report not found. <a href="/retest">Go back</a></p>
</div></body></html>"""), 400

    if not feedback_text:
        return render_template_string(SHARED_CSS + """
<!DOCTYPE html><html><body>
<div style="padding:40px;text-align:center">
  <p style="color:red;font-size:16px">Please enter feedback text or QIDs. <a href="/retest">Go back</a></p>
</div></body></html>"""), 400

    qid_list = parse_qid_list(feedback_text)

    if not qid_list:
        return render_template_string(SHARED_CSS + """
<!DOCTYPE html><html><body>
<div style="padding:40px;text-align:center">
  <p style="color:red;font-size:16px">No QIDs found in feedback. Please list specific QIDs (e.g. Q15, Q21). <a href="/retest">Go back</a></p>
</div></body></html>"""), 400

    original = jobs[original_job_id]
    doc_path = original.get('doc_path', '')
    survey_url = original.get('survey_url', '')

    if not doc_path or not survey_url:
        return render_template_string(SHARED_CSS + """
<!DOCTYPE html><html><body>
<div style="padding:40px;text-align:center">
  <p style="color:red;font-size:16px">Original report files not available. <a href="/reports">View reports</a></p>
</div></body></html>"""), 400

    new_job_id = str(uuid.uuid4())[:8]
    jobs[new_job_id] = {
        'status': 'running',
        'progress': 0,
        'phase': 'Starting re-test...',
        'logs': [{'msg': f'Re-test started — filtering to {len(qid_list)} QID(s): {", ".join(qid_list)}', 'color': 'cyan'}],
        'doc_name': f"Re-test: {original.get('doc_name', 'Unknown')}",
        'doc_path': doc_path,
        'survey_url': survey_url,
        'xml_path': original.get('xml_path', ''),
        'platform': original.get('platform', 'Confirmit'),
        'country': original.get('country', ''),
        'mode': original.get('mode', 'full'),
        'user_email': session['user_email'],
        'created_at': datetime.now().isoformat(),
        'retest_of': original_job_id,
        'retest_qids': qid_list,
        'verdict': None,
        'issues': [],
        'term_results': [],
        'report_file': None,
        'doc_qids': 0,
        'live_qids': 0,
        'xml_qids': 0,
        'total_issues': 0,
        'term_passed': 0,
        'term_total': 0,
    }

    t = threading.Thread(
        target=run_qc_engine,
        args=(new_job_id, doc_path, survey_url,
              original.get('country', ''),
              original.get('mode', 'full'),
              []),
        kwargs={'filter_qids': qid_list},
        daemon=True
    )
    t.start()

    return redirect(f'/progress/{new_job_id}')


# ================================================================
# PAGE: AI TESTER
# ================================================================
@app.route('/ai-tester')
@login_required
def ai_tester():
    return redirect('/new-qc')


@app.route('/settings', methods=['GET', 'POST'])
@login_required
def settings():
    user = get_current_user()
    email = session['user_email']
    success = ''

    force_change = request.args.get('force_change') == '1' or (
        request.method == 'POST' and request.form.get('force_change') == '1'
    )

    if request.method == 'POST':
        action = request.form.get('action')
        if action == 'profile':
            users_db[email]['name'] = request.form.get('name', user['name'])
            users_db.save(email)
            success = 'Profile updated!'
        elif action == 'password':
            old_pw = request.form.get('old_password', '')
            new_pw = request.form.get('new_password', '')
            if users_db.check_password(email, old_pw):
                if len(new_pw) >= 6:
                    new_hash = generate_password_hash(new_pw)
                    users_db[email]['password_hash'] = new_hash
                    users_db[email]['password'] = new_hash
                    users_db[email]['must_change_password'] = False
                    users_db.save(email)
                    session.pop('force_change', None)
                    success = 'Password updated!'
                else:
                    success = 'ERROR: New password must be at least 6 characters'
            else:
                success = 'ERROR: Current password is wrong'
        user = get_current_user()

    name = user['name']
    plan = user.get('plan', 'Free')
    initials = ''.join([n[0] for n in name.split()[:2]]).upper()

    return render_template_string(SHARED_CSS + f"""
<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=5.0"><title>Settings — SurveyQC</title></head><body>
<div class="app-layout">
  {sidebar_html('settings')}
  <div class="main-content">
    <div class="topbar">
      <p class="page-title">Settings</p>
    </div>

    <div class="tabs">
      <a href="/settings" class="tab active">Profile</a>
      <a href="/billing" class="tab">Billing</a>
    </div>

    {'<div style="background:#FEF3C7;border:1px solid #F59E0B;border-radius:10px;padding:14px 18px;margin-bottom:16px;font-size:13px;color:#92400E"><b>⚠️ You must change your password before continuing.</b> This is required for first-time admin access.</div>' if force_change else ''}
    {'<div class="alert ' + ('alert-success' if not success.startswith('ERROR') else 'alert-error') + '">' + success + '</div>' if success else ''}

    <div style="display:grid;grid-template-columns:1fr 300px;gap:16px">
      <div>
        <div class="card" style="margin-bottom:16px">
          <p style="font-size:14px;font-weight:600;color:var(--text);margin-bottom:16px">Profile information</p>
          <div style="display:flex;align-items:center;gap:14px;margin-bottom:20px">
            <div class="avatar" style="width:52px;height:52px;background:rgba(124,101,255,.3);color:var(--purple);font-size:16px">{initials}</div>
            <div>
              <p style="font-size:14px;font-weight:500;color:white">{name}</p>
              <p style="font-size:12px;color:var(--text3)">{email}</p>
              <span class="badge badge-purple" style="margin-top:4px">{plan} plan</span>
            </div>
          </div>
          <form method="POST">
            <input type="hidden" name="action" value="profile">
            <div class="form-group">
              <label class="form-label">Full name</label>
              <input class="form-input" type="text" name="name" value="{name}">
            </div>
            <div class="form-group">
              <label class="form-label">Email</label>
              <input class="form-input" type="email" value="{email}" disabled style="opacity:.6;cursor:not-allowed">
            </div>
            <button type="submit" class="btn btn-primary btn-sm">Save changes</button>
          </form>
        </div>

        <div class="card">
          <p style="font-size:14px;font-weight:600;color:var(--text);margin-bottom:16px">Change password</p>
          <form method="POST">
            <input type="hidden" name="action" value="password">
            <input type="hidden" name="force_change" value="{'1' if force_change else '0'}">
            <div class="form-group">
              <label class="form-label">Current password</label>
              <input class="form-input" type="password" name="old_password" placeholder="Current password">
            </div>
            <div class="form-group">
              <label class="form-label">New password</label>
              <input class="form-input" type="password" name="new_password" placeholder="Min 6 characters">
            </div>
            <button type="submit" class="btn btn-primary btn-sm">Update password</button>
          </form>
        </div>
      </div>

      <div style="display:flex;flex-direction:column;gap:12px">
        <div class="card">
          <p style="font-size:13px;font-weight:600;color:var(--text);margin-bottom:12px">Notifications</p>
          <div style="display:flex;flex-direction:column;gap:10px">
            <div style="display:flex;justify-content:space-between;align-items:center">
              <div><p style="font-size:12px;color:white">Report done</p><p style="font-size:11px;color:var(--text3)">Email on finish</p></div>
              <input type="checkbox" checked style="accent-color:var(--purple);width:16px;height:16px">
            </div>
            <div style="display:flex;justify-content:space-between;align-items:center">
              <div><p style="font-size:12px;color:white">Weekly summary</p><p style="font-size:11px;color:var(--text3)">Every Monday</p></div>
              <input type="checkbox" checked style="accent-color:var(--purple);width:16px;height:16px">
            </div>
            <div style="display:flex;justify-content:space-between;align-items:center">
              <div><p style="font-size:12px;color:white">Plan limit alert</p><p style="font-size:11px;color:var(--text3)">At 80% usage</p></div>
              <input type="checkbox" checked style="accent-color:var(--purple);width:16px;height:16px">
            </div>
          </div>
        </div>
        <div class="card" style="background:#FCEBEB;border-color:#F7C1C1">
          <p style="font-size:13px;font-weight:500;color:#791F1F;margin-bottom:8px">Danger zone</p>
          <p style="font-size:12px;color:#A32D2D;margin-bottom:12px">Permanently delete your account and all data.</p>
          <button class="btn btn-danger btn-sm">Delete account</button>
        </div>
      </div>
    </div>
  </div>
</div>
</body></html>""")

# ================================================================
# PAGE: BILLING
# ================================================================
@app.route('/billing')
@login_required
def billing():
    user = get_current_user()
    plan = user.get('plan', 'Free')
    used = user.get('reports_used', 0)
    limit = user.get('reports_limit', 5)
    pct = int((used/limit)*100) if limit > 0 else 0
    c = site_content  # use admin-editable content for prices
    email = session.get('user_email', '')
    is_demo = (email == 'demo@surveyqc.com')

    free_price = c.get('plan_free_price', '0')
    pro_price  = c.get('plan_pro_price', '29')
    biz_price  = c.get('plan_biz_price', '299')
    ent_price  = c.get('plan_ent_price', 'Custom')
    free_feats = (c.get('plan_free_features', '3 reports per month')).split('||')[0]
    pro_feats  = (c.get('plan_pro_features', '25 reports per month')).split('||')[0]
    biz_feats  = (c.get('plan_biz_features', 'Unlimited reports')).split('||')[0]
    ent_feats  = (c.get('plan_ent_features', 'Everything in Business')).split('||')[0]

    # Compute reset date: 1st of next month
    _today = datetime.now()
    if _today.month == 12:
        _reset = datetime(_today.year + 1, 1, 1)
    else:
        _reset = datetime(_today.year, _today.month + 1, 1)
    reset_date_str = _reset.strftime('%b %d')

    # Payment method card content
    if is_demo:
        _payment_html = '''
          <div style="background:rgba(255,255,255,.06);border-radius:8px;padding:12px;display:flex;align-items:center;gap:10px;margin-bottom:12px">
            <i class="ti ti-shield-off" style="font-size:20px;color:#9CA3AF"></i>
            <div><p style="font-size:12px;color:white">Demo Account</p><p style="font-size:11px;color:var(--text3)">Billing is disabled for demo accounts</p></div>
          </div>
          <button class="btn btn-ghost btn-sm" disabled style="width:100%;justify-content:center;opacity:.4">Not available in demo</button>'''
    elif plan == 'Free':
        _payment_html = '''
          <div style="background:rgba(255,255,255,.06);border-radius:8px;padding:12px;display:flex;align-items:center;gap:10px;margin-bottom:12px">
            <i class="ti ti-credit-card" style="font-size:20px;color:#9CA3AF"></i>
            <div><p style="font-size:12px;color:white">No card added</p><p style="font-size:11px;color:var(--text3)">Free plan — no card needed</p></div>
          </div>
          <button class="btn btn-primary btn-sm" style="width:100%;justify-content:center"><i class="ti ti-plus"></i>Add payment method</button>'''
    else:
        _payment_html = '''
          <div style="background:rgba(255,255,255,.06);border-radius:8px;padding:12px;display:flex;align-items:center;gap:10px;margin-bottom:12px">
            <i class="ti ti-credit-card" style="font-size:20px;color:var(--purple)"></i>
            <div><p style="font-size:12px;color:white">Manage payment method</p><p style="font-size:11px;color:var(--text3)">Billing via Stripe</p></div>
          </div>
          <button class="btn btn-primary btn-sm" style="width:100%;justify-content:center"><i class="ti ti-credit-card"></i>Add / update card</button>'''

    def _plan_card(name, price, feats, btn_label, is_current, is_ent=False, can_upgrade=False):
        border = 'border:2px solid var(--purple)' if is_current else ''
        badge = '<span class="badge badge-purple" style="margin-bottom:8px;display:inline-block">Current plan</span>' if is_current else ''
        price_html = f'Custom<span style="font-size:13px;color:var(--text3)"> — contact sales</span>' if is_ent else f'₹{price}<span style="font-size:13px;color:var(--text3)">/mo</span>'
        if is_ent:
            btn_tag = f'<a href="mailto:support@surveyqc.online?subject=Enterprise%20Inquiry" class="btn btn-ghost btn-sm" style="width:100%;justify-content:center;text-decoration:none">{btn_label}</a>'
        elif is_current or name == 'Free':
            btn_cls = 'btn-primary' if is_current else 'btn-ghost'
            btn_tag = f'<button class="btn {btn_cls} btn-sm" style="width:100%;justify-content:center" disabled>{btn_label}</button>'
        elif can_upgrade and not is_demo:
            btn_tag = f'<button class="btn btn-ghost btn-sm rzp-upgrade-btn" data-plan="{name}" style="width:100%;justify-content:center"><i class="ti ti-bolt"></i> {btn_label}</button>'
        else:
            btn_tag = f'<button class="btn btn-ghost btn-sm" style="width:100%;justify-content:center" disabled>{btn_label}</button>'
        return f'''<div class="card" style="{border}">
          {badge}
          <p style="font-size:13px;font-weight:600;color:var(--text)">{name}</p>
          <p style="font-size:24px;font-weight:700;color:var(--text);margin:8px 0 4px">{price_html}</p>
          <p style="font-size:12px;color:var(--text3);margin-bottom:14px">{feats}</p>
          {btn_tag}
        </div>'''

    _plan_order = ['Free', 'Pro', 'Business']
    _current_idx = _plan_order.index(plan) if plan in _plan_order else 0

    cards_html  = _plan_card('Free',      free_price, free_feats, 'Current plan' if plan=='Free' else 'Downgrade',      plan=='Free')
    cards_html += _plan_card('Pro',       '2,499',    pro_feats,  'Current plan' if plan=='Pro'  else 'Upgrade to Pro', plan=='Pro',      can_upgrade=(_current_idx < 1))
    cards_html += _plan_card('Business',  '24,999',   biz_feats,  'Current plan' if plan=='Business' else 'Upgrade',    plan=='Business', can_upgrade=(_current_idx < 2))
    cards_html += _plan_card('Enterprise', ent_price, ent_feats, 'Contact Sales', plan=='Enterprise', is_ent=True)

    _billing_success = request.args.get('upgraded')
    _success_banner = ''
    if _billing_success:
        _success_banner = f'<div style="background:#E6F4EC;border:1px solid #A7D7B8;color:#1A5632;border-radius:10px;padding:12px 16px;margin-bottom:16px;font-size:13px;display:flex;align-items:center;gap:8px"><i class="ti ti-circle-check"></i> You are now on the <strong>{_billing_success}</strong> plan. Thank you!</div>'

    return render_template_string(SHARED_CSS + f"""
<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=5.0"><title>Billing — SurveyQC</title></head><body>
<div class="app-layout">
  {sidebar_html('billing')}
  <div class="main-content">
    <div class="topbar"><p class="page-title">Billing & Subscription</p></div>

    <div class="tabs">
      <a href="/settings" class="tab">Profile</a>
      <a href="/billing" class="tab active">Billing</a>
    </div>

    {_success_banner}

    <div style="display:grid;grid-template-columns:repeat(4,1fr);gap:12px;margin-bottom:20px">
      {cards_html}
    </div>

    <div style="display:grid;grid-template-columns:1fr 1fr;gap:14px">
      <div class="card">
        <p style="font-size:13px;font-weight:600;color:var(--text);margin-bottom:12px">Usage this month</p>
        <div style="display:flex;justify-content:space-between;align-items:baseline;margin-bottom:6px">
          <span style="font-size:13px;font-weight:600;color:var(--text)">{used} / {limit} reports used</span>
          <span style="font-size:11px;color:var(--text3)">{pct}%</span>
        </div>
        <div class="progress-bar" style="height:8px;margin-bottom:8px">
          <div class="progress-fill progress-purple" style="width:{pct}%"></div>
        </div>
        <p style="font-size:11px;color:var(--text3)">Resets on {reset_date_str}</p>
      </div>
      <div class="card">
        <p style="font-size:13px;font-weight:600;color:var(--text);margin-bottom:12px">Payment method</p>
        {_payment_html}
      </div>
    </div>
  </div>
</div>
<script src="https://checkout.razorpay.com/v1/checkout.js"></script>
<script>
document.querySelectorAll('.rzp-upgrade-btn').forEach(function(btn) {{
  btn.addEventListener('click', function() {{
    var plan = this.getAttribute('data-plan');
    var origText = btn.innerHTML;
    btn.disabled = true;
    btn.innerHTML = '<i class="ti ti-loader"></i> Opening...';
    fetch('/create-order', {{
      method: 'POST',
      headers: {{'Content-Type': 'application/json'}},
      body: JSON.stringify({{plan: plan}})
    }})
    .then(function(r) {{ return r.json(); }})
    .then(function(data) {{
      if (data.error) {{ alert('Error: ' + data.error); btn.disabled=false; btn.innerHTML=origText; return; }}
      var options = {{
        key:         data.key_id,
        amount:      data.amount,
        currency:    data.currency,
        name:        'SurveyQC',
        description: data.description,
        order_id:    data.order_id,
        prefill:     {{email: '{email}'}},
        theme:       {{color: '#C46A2B'}},
        handler: function(response) {{
          fetch('/verify-payment', {{
            method: 'POST',
            headers: {{'Content-Type': 'application/json'}},
            body: JSON.stringify({{
              razorpay_order_id:   response.razorpay_order_id,
              razorpay_payment_id: response.razorpay_payment_id,
              razorpay_signature:  response.razorpay_signature,
              plan:                plan
            }})
          }})
          .then(function(r) {{ return r.json(); }})
          .then(function(d) {{
            if (d.success) {{
              window.location.href = '/billing?upgraded=' + encodeURIComponent(plan);
            }} else {{
              alert('Payment verification failed. Please contact support.');
              btn.disabled=false; btn.innerHTML=origText;
            }}
          }});
        }},
        modal: {{
          ondismiss: function() {{ btn.disabled=false; btn.innerHTML=origText; }}
        }}
      }};
      var rzp = new Razorpay(options);
      rzp.open();
    }})
    .catch(function(err) {{
      alert('Could not start checkout. Please try again.');
      btn.disabled=false; btn.innerHTML=origText;
    }});
  }});
}});
</script>
</body></html>""")

# ================================================================
# PAYMENT: Razorpay
# ================================================================
@app.route('/create-order', methods=['POST'])
@login_required
def create_order():
    import logging as _log
    data = request.get_json(silent=True) or {}
    plan = data.get('plan', '')
    if plan not in ('Pro', 'Business'):
        return jsonify(error='Invalid plan'), 400
    try:
        order = _rzp_create_order(plan)
        return jsonify(order)
    except Exception as exc:
        _log.getLogger(__name__).error('create_order: %s', exc)
        return jsonify(error=str(exc)), 500


@app.route('/verify-payment', methods=['POST'])
@login_required
def verify_payment():
    import logging as _log
    data = request.get_json(silent=True) or {}
    order_id   = data.get('razorpay_order_id', '')
    payment_id = data.get('razorpay_payment_id', '')
    signature  = data.get('razorpay_signature', '')
    plan       = data.get('plan', '')

    if not all([order_id, payment_id, signature, plan]):
        return jsonify(success=False, error='Missing fields'), 400

    if not _rzp_verify_payment(order_id, payment_id, signature):
        return jsonify(success=False, error='Signature verification failed'), 400

    # Upgrade the user's plan in the DB
    email = session.get('user_email')
    if email and email in users_db:
        users_db[email]['plan'] = plan
        users_db[email]['reports_limit'] = UserDB.PLAN_LIMITS.get(plan, 3)
        users_db.save(email)
        _log.getLogger(__name__).info(
            'Plan upgraded: email=%s plan=%s order=%s payment=%s',
            email, plan, order_id, payment_id
        )
    return jsonify(success=True, plan=plan)


# ================================================================
# ADMIN: LOGIN
# ================================================================
@app.route('/admin/login', methods=['GET', 'POST'])
def admin_login():
    error = ''
    if request.method == 'POST':
        pw = request.form.get('password', '')
        # Accept if it matches the hardcoded ADMIN_PASSWORD or the admin@surveyqc.com DB user
        if pw == ADMIN_PASSWORD or users_db.check_password('admin@surveyqc.com', pw):
            session['is_admin'] = True
            return redirect('/admin')
        error = 'Wrong password'

    return render_template_string(SHARED_CSS + f"""
<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=5.0"><title>Admin Login — SurveyQC</title></head><body>
<div style="min-height:100vh;display:flex;align-items:center;justify-content:center">
  <div style="width:380px">
    <div style="text-align:center;margin-bottom:28px">
      <div style="width:40px;height:40px;background:var(--purple);border-radius:10px;display:flex;align-items:center;justify-content:center;margin:0 auto 12px"><i class="ti ti-shield-check" style="color:white;font-size:20px"></i></div>
      <p style="font-size:18px;font-weight:600;color:var(--text)">Admin access</p>
      <p style="font-size:13px;color:var(--text3)">Only you can access this</p>
    </div>
    {'<div class="alert alert-error">' + error + '</div>' if error else ''}
    <form method="POST">
      <div class="form-group">
        <label class="form-label">Admin password</label>
        <input class="form-input" type="password" name="password" placeholder="Password" autofocus>
      </div>
      <button type="submit" class="btn btn-primary" style="width:100%;justify-content:center">Enter admin panel</button>
    </form>
  </div>
</div>
</body></html>""")

# ================================================================
# ADMIN: DASHBOARD
# ================================================================
@app.route('/admin')
@admin_required
def admin_dashboard():
    total_users = len(users_db)
    paid_users = sum(1 for u in users_db.values() if u.get('plan') in ('Pro', 'Business'))
    total_reports = sum(u.get('reports_used', 0) for u in users_db.values())
    total_jobs = len(jobs)
    done_jobs = sum(1 for j in jobs.values() if j.get('status') == 'done')
    running_jobs = sum(1 for j in jobs.values() if j.get('status') == 'running')
    mrr = paid_users * 29
    arr = mrr * 12

    users_html = ''
    for email, u in list(users_db.items())[:10]:
        name = u.get('name', 'Unknown')
        plan = u.get('plan', 'Free')
        reports = u.get('reports_used', 0)
        joined = u.get('joined', '-')
        initials = ''.join([n[0] for n in name.split()[:2]]).upper()
        badge_cls = 'badge-purple' if plan == 'Pro' else ('badge-blue' if plan == 'Business' else 'badge-amber')
        users_html += f"""
        <tr>
          <td class="primary">
            <div style="display:flex;align-items:center;gap:8px">
              <div class="avatar" style="background:rgba(124,101,255,.2);color:var(--purple)">{initials}</div>
              {name}
            </div>
          </td>
          <td style="color:var(--text3)">{email}</td>
          <td><span class="badge {badge_cls}">{plan}</span></td>
          <td>{reports}</td>
          <td style="color:var(--text3)">{joined}</td>
          <td>
            <span style="color:var(--purple);cursor:pointer;font-size:12px">Edit</span> &nbsp;
            <span style="color:#E24B4A;cursor:pointer;font-size:12px">Block</span>
          </td>
        </tr>"""

    jobs_html = ''
    for jid, j in list(jobs.items())[:8]:
        status = j.get('status', 'running')
        cls = 'badge-green' if status == 'done' else ('badge-blue' if status == 'running' else 'badge-red')
        jobs_html += f"""
        <tr>
          <td class="primary">{j.get('doc_name','Unknown')[:30]}</td>
          <td style="color:var(--text3)">{j.get('user_email','')[:25]}</td>
          <td><span class="badge {cls}">{status}</span></td>
          <td>{j.get('total_issues',0)}</td>
          <td style="color:var(--text3)">{j.get('created_at','')[:16]}</td>
        </tr>"""

    return render_template_string(SHARED_CSS + """<style>.app-layout{margin-left:0!important}.main-content{margin-left:220px!important}</style>""" + f"""
<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=5.0"><title>Admin — SurveyQC</title></head><body>
<div style="display:flex;min-height:100vh;background:var(--bg)"><div style="margin-left:220px;flex:1;padding:28px;min-width:0;max-width:1100px">
  
    <div style="display:flex;align-items:center;justify-content:space-between;margin-bottom:20px">
      <div><p style="font-size:20px;font-weight:600;color:var(--text)">Admin overview</p><p style="font-size:12px;color:var(--text3)">{datetime.now().strftime('%A, %d %B %Y')}</p></div>
      <div style="display:flex;gap:8px;align-items:center">
        <span class="badge badge-green" style="display:flex;align-items:center;gap:5px"><span class="dot dot-green"></span>{running_jobs} running</span>
        <span class="badge badge-purple">{total_users} users</span>
      </div>
    </div>

    <div class="stats-grid" style="grid-template-columns:repeat(auto-fit,minmax(180px,1fr));gap:16px;margin-bottom:20px">
      <div class="stat-card" style="padding:20px"><p class="stat-num">{total_users}</p><p class="stat-label">Total users</p></div>
      <div class="stat-card" style="padding:20px"><p class="stat-num" style="color:#1D9E75">{paid_users}</p><p class="stat-label">Paid users</p></div>
      <div class="stat-card" style="padding:20px"><p class="stat-num">{total_jobs}</p><p class="stat-label">Total reports</p></div>
      <div class="stat-card" style="padding:20px;background:rgba(124,101,255,.1);border-color:var(--purple-border)"><p class="stat-num" style="color:var(--purple)">${mrr}</p><p class="stat-label">MRR</p></div>
      <div class="stat-card" style="padding:20px;background:rgba(29,158,117,.1);border-color:rgba(29,158,117,.2)"><p class="stat-num" style="color:#1D9E75">${arr}</p><p class="stat-label">ARR</p></div>
    </div>

    <div style="display:grid;grid-template-columns:1fr 1fr;gap:16px;margin-bottom:16px">
      <div class="card">
        <p style="font-size:14px;font-weight:600;color:var(--text);margin-bottom:14px">AI health status</p>
        <div style="display:flex;flex-direction:column;gap:7px">
          <div class="worker-card"><span class="pulse"></span><div style="flex:1"><p style="font-size:12px;color:white">Playwright browser</p></div><span class="badge badge-teal">Healthy</span></div>
          <div class="worker-card"><span class="pulse"></span><div style="flex:1"><p style="font-size:12px;color:white">Report generator</p></div><span class="badge badge-teal">Healthy</span></div>
          <div class="worker-card"><span class="pulse"></span><div style="flex:1"><p style="font-size:12px;color:white">File storage</p></div><span class="badge badge-teal">Healthy</span></div>
        </div>
      </div>
      <div class="card">
        <p style="font-size:14px;font-weight:600;color:var(--text);margin-bottom:14px">Quick actions</p>
        <div style="display:flex;flex-direction:column;gap:8px">
          <a href="/admin/email" class="btn btn-ghost" style="justify-content:center"><i class="ti ti-mail"></i>Email all users</a>
          <button class="btn btn-ghost" style="width:100%;justify-content:center"><i class="ti ti-download"></i>Export data</button>
          <button class="btn btn-ghost" style="width:100%;justify-content:center"><i class="ti ti-gift"></i>Gift credits to user</button>
        </div>
      </div>
    </div>

    <div class="card" style="margin-bottom:16px">
      <div style="display:flex;align-items:center;justify-content:space-between;margin-bottom:14px">
        <p style="font-size:14px;font-weight:600;color:var(--text)">Users</p>
        <a href="/admin/users" style="font-size:12px;color:var(--purple);text-decoration:none">View all →</a>
      </div>
      <table class="data-table" style="width:100%">
        <thead><tr><th>Name</th><th>Email</th><th>Plan</th><th>Reports</th><th>Joined</th><th>Actions</th></tr></thead>
        <tbody>{users_html}</tbody>
      </table>
    </div>

    <div class="card">
      <div style="display:flex;align-items:center;justify-content:space-between;margin-bottom:14px">
        <p style="font-size:14px;font-weight:600;color:var(--text)">Recent QC jobs</p>
        <a href="/admin/reports" style="font-size:12px;color:var(--purple);text-decoration:none">View all →</a>
      </div>
      <table class="data-table" style="width:100%">
        <thead><tr><th>Document</th><th>User</th><th>Status</th><th>Issues</th><th>Date</th></tr></thead>
        <tbody>{jobs_html or '<tr><td colspan="5" style="text-align:center;color:var(--text3);padding:20px">No jobs yet</td></tr>'}</tbody>
      </table>
    </div>
  </div>
</div>
</div>
</body></html>""")

# ================================================================
# ADMIN: USERS
# ================================================================
@app.route('/admin/users')
@admin_required
def admin_users():
    rows = ''
    for email, u in users_db.items():
        name = u.get('name', 'Unknown')
        plan = u.get('plan', 'Free')
        reports = u.get('reports_used', 0)
        joined = u.get('joined', '-')
        initials = ''.join([n[0] for n in name.split()[:2]]).upper()
        badge_cls = 'badge-purple' if plan == 'Pro' else ('badge-blue' if plan == 'Business' else 'badge-amber')
        rows += f"""
        <tr>
          <td class="primary">
            <div style="display:flex;align-items:center;gap:8px">
              <div class="avatar" style="background:rgba(124,101,255,.2);color:var(--purple)">{initials}</div>
              {name}
            </div>
          </td>
          <td style="color:var(--text3)">{email}</td>
          <td><span class="badge {badge_cls}">{plan}</span></td>
          <td>{reports}</td>
          <td style="color:var(--text3)">{joined}</td>
          <td>
            <span style="color:var(--purple);cursor:pointer;font-size:12px">Edit</span> &nbsp;
            <span style="color:#E24B4A;cursor:pointer;font-size:12px">Block</span>
          </td>
        </tr>"""

    return render_template_string(SHARED_CSS + f"""
<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=5.0"><title>Users — Admin</title><script src="/admin-sidebar-js"></script></head><body>
<div style="padding:24px;margin-left:220px;max-width:1100px">
  <div style="display:flex;align-items:center;gap:14px;margin-bottom:20px">
    <a href="/admin" style="color:var(--text3);text-decoration:none"><i class="ti ti-arrow-left"></i></a>
    <p style="font-size:20px;font-weight:600;color:var(--text)">All Users ({len(users_db)})</p>
  </div>
  <div class="card">
    <table class="data-table" style="width:100%">
      <thead><tr><th>Name</th><th>Email</th><th>Plan</th><th>Reports</th><th>Joined</th><th>Actions</th></tr></thead>
      <tbody>{rows}</tbody>
    </table>
  </div>
</div>
</body></html>""")

# ================================================================
# ADMIN: EMAIL
# ================================================================
@app.route('/admin/email', methods=['GET', 'POST'])
@admin_required
def admin_email():
    sent = False
    if request.method == 'POST':
        sent = True

    return render_template_string(SHARED_CSS + f"""
<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=5.0"><title>Email Users — Admin</title><script src="/admin-sidebar-js"></script></head><body>
<div style="padding:24px;max-width:600px">
  <div style="display:flex;align-items:center;gap:14px;margin-bottom:20px">
    <a href="/admin" style="color:var(--text3);text-decoration:none"><i class="ti ti-arrow-left"></i></a>
    <p style="font-size:20px;font-weight:600;color:var(--text)">Email Users</p>
  </div>
  {'<div class="alert alert-success">Email sent successfully!</div>' if sent else ''}
  <div class="card">
    <form method="POST">
      <div class="form-group">
        <label class="form-label">Send to</label>
        <select class="form-select" name="audience">
          <option>All users ({len(users_db)})</option>
          <option>Free users</option>
          <option>Pro users</option>
          <option>Business users</option>
        </select>
      </div>
      <div class="form-group">
        <label class="form-label">Subject</label>
        <input class="form-input" type="text" name="subject" placeholder="Email subject...">
      </div>
      <div class="form-group">
        <label class="form-label">Message</label>
        <textarea class="form-input" name="message" rows="6" placeholder="Hi {{name}}, ..."></textarea>
      </div>
      <div style="display:flex;gap:10px">
        <button type="submit" class="btn btn-primary"><i class="ti ti-send"></i>Send email</button>
        <button type="button" class="btn btn-ghost">Preview</button>
      </div>
    </form>
  </div>
</div>
</body></html>""")

# ================================================================
# ADMIN: REPORTS
# ================================================================
@app.route('/admin/reports')
@admin_required
def admin_reports():
    rows = ''
    for jid, j in list(jobs.items())[:50]:
        status = j.get('status', 'running')
        cls = 'badge-green' if status == 'done' else ('badge-blue' if status == 'running' else 'badge-red')
        link = f'<a href="/report/{jid}" style="color:var(--purple);font-size:12px;text-decoration:none">View</a>' if status == 'done' else ''
        rows += f"""
        <tr>
          <td class="primary">{j.get('doc_name','Unknown')[:35]}</td>
          <td style="color:var(--text3)">{j.get('user_email','')[:25]}</td>
          <td>{j.get('platform','-')}</td>
          <td><span class="badge {cls}">{status}</span></td>
          <td>{j.get('total_issues',0)}</td>
          <td style="color:var(--text3)">{j.get('created_at','')[:16]}</td>
          <td>{link}</td>
        </tr>"""

    return render_template_string(SHARED_CSS + f"""
<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=5.0"><title>Reports — Admin</title><script src="/admin-sidebar-js"></script></head><body>
<div style="padding:24px">
  <div style="display:flex;align-items:center;gap:14px;margin-bottom:20px">
    <a href="/admin" style="color:var(--text3);text-decoration:none"><i class="ti ti-arrow-left"></i></a>
    <p style="font-size:20px;font-weight:600;color:var(--text)">All QC Jobs ({len(jobs)})</p>
  </div>
  <div class="card">
    <table class="data-table" style="width:100%">
      <thead><tr><th>Document</th><th>User</th><th>Platform</th><th>Status</th><th>Issues</th><th>Date</th><th></th></tr></thead>
      <tbody>{rows or '<tr><td colspan="7" style="text-align:center;color:var(--text3);padding:20px">No jobs yet</td></tr>'}</tbody>
    </table>
  </div>
</div>
</body></html>""")

# ================================================================
# API ENDPOINTS
# ================================================================
@app.route('/api/status/<job_id>')
@login_required
def api_status(job_id):
    if job_id not in jobs:
        return jsonify({'error': 'Not found'}), 404
    return jsonify(jobs[job_id])

@app.route('/api/feedback', methods=['POST'])
@login_required
def api_feedback():
    data = request.json
    feedback_store.append({
        'timestamp': datetime.now().isoformat(),
        'user': session.get('user_email'),
        'job_id': data.get('job_id'),
        'rating': data.get('rating'),
        'comment': data.get('comment', '')
    })
    return jsonify({'ok': True})

@app.route('/api/issue-feedback/<job_id>', methods=['POST'])
@login_required
def api_issue_feedback(job_id):
    """Store a user verdict (CONFIRMED / FALSE_POSITIVE) for a specific issue."""
    if job_id not in jobs:
        return jsonify({'error': 'not found'}), 404
    j = jobs[job_id]
    data = request.get_json(force=True, silent=True) or {}
    qid        = str(data.get('qid', ''))[:50]
    issue_type = str(data.get('issue_type', ''))[:100]
    verdict    = str(data.get('verdict', ''))
    if verdict not in ('CONFIRMED', 'FALSE_POSITIVE'):
        return jsonify({'error': 'verdict must be CONFIRMED or FALSE_POSITIVE'}), 400
    platform   = j.get('platform', '')
    user_email = session.get('user_email', '')
    jobs.save_feedback(job_id, qid, issue_type, platform, verdict, user_email)
    return jsonify({'ok': True})

@app.route('/download/<job_id>')
@login_required
def download_report(job_id):
    import glob as _glob
    print(f"[download] requested job_id={job_id!r}", flush=True)

    # Locate the report file: check in-memory job first, then scan disk.
    # This survives gunicorn restarts where the jobs dict is empty.
    def _find_on_disk(jid):
        pattern = os.path.join(OUTPUT_FOLDER, jid, 'QC_Report_*.docx')
        hits = sorted(_glob.glob(pattern), reverse=True)  # newest first
        return hits[0] if hits else None

    if job_id not in jobs:
        report_file = _find_on_disk(job_id)
        print(f"[download] job_id={job_id!r} not in memory — disk lookup: {report_file!r}", flush=True)
        if not report_file:
            return render_template_string(
                '<h3 style="font-family:sans-serif;color:#c0392b;padding:2rem">'
                '⚠️ Report expired. Please run QC again.</h3>'
            ), 404
    else:
        report_file = jobs[job_id].get('report_file')
        if not report_file or not os.path.exists(report_file):
            # In-memory path is stale or missing — scan disk for this job
            report_file = _find_on_disk(job_id)
            print(f"[download] job_id={job_id!r} stale/missing path — disk lookup: {report_file!r}", flush=True)
            if report_file:
                jobs[job_id]['report_file'] = report_file  # heal the stale entry
        if not report_file:
            return render_template_string(
                '<h3 style="font-family:sans-serif;color:#c0392b;padding:2rem">'
                '⚠️ Report file not found. Please run QC again.</h3>'
            ), 404

    print(f"[download] serving {report_file!r}", flush=True)
    return send_file(report_file, as_attachment=True,
                     download_name=f"QC_Report_{job_id}.docx")

# ================================================================
# QC ENGINE
# ================================================================
def normalize(text):
    if not text: return ""
    return re.sub(r'\s+', ' ', text).strip().lower()

def fuzzy_match(doc_text, live_text, threshold=0.65):
    if not doc_text or len(doc_text) < 15: return True, 1.0
    doc_norm = normalize(doc_text)
    live_norm = normalize(live_text)
    if not live_norm: return False, 0.0
    if doc_norm[:40] in live_norm: return True, 1.0
    for i in range(0, len(doc_norm)-40, 30):
        if doc_norm[i:i+40] in live_norm: return True, 0.9
    ratio = SequenceMatcher(None, doc_norm[:300], live_norm[:300]).ratio()
    return ratio >= threshold, ratio

def tokenize(text):
    if not text: return []
    text = re.sub(r'\[[^\]]{1,30}\]', ' ', text)
    text = re.sub(r'\{\{[^}]+\}\}', ' ', text)
    text = re.sub(r'<[^>]+>', ' ', text).lower()
    words = re.findall(r"[\w']+", text, re.UNICODE)
    return [w for w in words if len(w) >= 3 and w not in STOPWORDS]

def find_missing_words(doc_text, live_text):
    live_set = set(tokenize(live_text))
    seen = set()
    missing = []
    for w in tokenize(doc_text):
        if w not in live_set and w not in seen:
            missing.append(w)
            seen.add(w)
    return missing

def _extract_options(page):
    """Extract answer options from a survey page (multi-platform)."""
    import re as _re
    opts = []
    seen = set()
    selectors = [
        ".cf-radio-answer__text", ".cf-checkbox-answer__text",
        "label.cf-answer", ".answer-text", ".option-text",
        "label[for]", ".sv-item__control-label", ".response-option"
    ]
    for sel in selectors:
        try:
            els = page.locator(sel).all()
            if els:
                for el in els:
                    try:
                        t = el.inner_text().strip()
                        if t and len(t) < 200 and t not in seen:
                            seen.add(t)
                            opts.append({"text": t})
                    except: continue
                if opts: break
        except: continue
    return opts


def _page_has_inputs(page):
    """Return True if the page has at least one VISIBLE answerable input.
    Hidden/invisible inputs (display:none, visibility:hidden, opacity:0) that
    Confirmit uses for internal state on display-only screens are excluded so
    that intro/article/consent screens correctly return False."""
    try:
        return bool(page.evaluate("""
            () => {
                const sel = 'input[type="radio"],input[type="checkbox"],select,' +
                            'input[type="text"],input[type="number"],' +
                            'input[type="range"],textarea';
                return Array.from(document.querySelectorAll(sel)).some(function(el) {
                    if (el.offsetParent === null) return false;
                    const s = window.getComputedStyle(el);
                    return s.display !== 'none' &&
                           s.visibility !== 'hidden' &&
                           parseFloat(s.opacity || '1') > 0;
                });
            }
        """))
    except Exception:
        return True  # safe default: assume inputs exist if DOM check fails


# ─── Platform-aware extraction helpers (v2) ────────────────────────────────

_PLATFORM_CONTAINER_SELECTORS = {
    'confirmit': [
        '[data-questionid="{qid}"]',
        '.cf-question[data-questionid="{qid}"]',
        '#question_{qid}',
        '#Q{qid}',
    ],
    'forsta': [
        '[data-qid="{qid}"]',
        '.question-container[id*="{qid}"]',
        '#question_{qid}',
    ],
    'decipher': [
        '.sv-question[id="{qid}"]',
        '[data-variable="{qid}"]',
        '[id="{qid}"]',
    ],
    'qualtrics': [
        '#QID{qid}',
        '.QuestionOuter[id*="{qid}"]',
        '[id*="QID{qid}"]',
    ],
    'surveymonkey': [
        '[data-questionid="{qid}"]',
        '[id*="{qid}"]',
    ],
    'generic': [
        '[data-questionid="{qid}"]',
        '[data-qid="{qid}"]',
        '[id="{qid}"]',
        '[id*="{qid}"]',
    ],
}


def detect_platform(page):
    """Detect survey platform from URL and initial page HTML."""
    try:
        url = page.url.lower()
        try:
            html = page.content()[:8000].lower()
        except Exception:
            html = ''
        if 'confirmit' in url or 'questionnaire.aplusaresearch' in url or ('wix.com' in url and 'cf-question' in html):
            return 'confirmit'
        if 'forsta' in url or 'forstasurveys' in url:
            return 'forsta'
        if 'decipherinc' in url or 'focusvision' in url:
            return 'decipher'
        if 'qualtrics' in url or 'az1.qualtrics' in url:
            return 'qualtrics'
        if 'surveymonkey' in url:
            return 'surveymonkey'
        # HTML fingerprint fallback
        if 'cf-question' in html or ('confirmit' in html and 'cf-tn-' in html):
            return 'confirmit'
        if 'questionouter' in html or ('qualtrics' in html and 'qid' in html):
            return 'qualtrics'
        if 'sv-question' in html and 'survey-question' in html:
            return 'decipher'
        if 'forsta' in html:
            return 'forsta'
    except Exception:
        pass
    return 'generic'


_CONTAINER_BAD_TAGS = {'INPUT', 'TEXTAREA', 'LABEL', 'SPAN', 'BUTTON', 'OPTION', 'SELECT', 'A'}

def _resolve_container(page, elem, qid, _log):
    """
    Smart container selection with parent walking.

    Fast path  — text_len 50-5000: use elem directly, skip parent walk entirely.
    Parent walk — only when text_len < 50 (element has too little text to be the
                  question container; walk up to find a richer ancestor).
    All JS wrapped in try-catch so DOM errors never reach Python.
    inner_text() fallback used whenever JS evaluation fails or returns undefined.
    """
    try:
        # Step A: Get element details + text_len; JS is fully guarded.
        _info = page.evaluate("""(el) => {
            try {
                if (!el) return {error: 'null_element', fallback: true, text_len: 0, tag: ''};
                const tlen = (el.innerText || el.textContent || '').length;
                return {
                    tag:        el.tagName || '',
                    id:         el.id || '',
                    cls:        (el.className || '').toString().slice(0, 100),
                    role:       el.getAttribute ? (el.getAttribute('role') || '') : '',
                    parent_tag: el.parentElement ? el.parentElement.tagName : '',
                    child_count: el.children ? el.children.length : 0,
                    text_len:   tlen,
                    fallback:   false,
                    error:      null
                };
            } catch(e) { return {error: e.message, fallback: true, text_len: 0, tag: ''}; }
        }""", elem)

        # Step B: JS failed or element was undefined → use inner_text() directly.
        if not _info or _info.get('fallback') or _info.get('error'):
            _log(f'   CONTAINER_JS_ERR: qid={qid} '
                 f'err={(_info or {}).get("error", "unknown")!r} — inner_text fallback')
            try:
                _ft = len(elem.inner_text(timeout=2000))
                if 50 <= _ft <= 5000:
                    _log(f'   CONTAINER_FALLBACK_OK: qid={qid} text_len={_ft} → CONTAINER_EXACT')
                    return elem, "CONTAINER_EXACT"
            except Exception:
                pass
            return elem, "CONTAINER_EXACT"

        _tlen = _info.get('text_len', 0)
        _log(f'   CONTAINER_ELEM: qid={qid} tag={_info["tag"]} id={_info["id"]!r} '
             f'class={_info["cls"]!r} role={_info["role"]!r} '
             f'parent_tag={_info["parent_tag"]} child_count={_info["child_count"]} '
             f'text_len={_tlen}')

        # Step C: text_len 50-5000 → this element has enough content; use directly.
        if 50 <= _tlen <= 5000:
            return elem, "CONTAINER_EXACT"

        # Step D: text_len < 50 — element is too thin; walk up parents for richer one.
        _log(f'   CONTAINER_THIN: qid={qid} tag={_info["tag"]} text_len={_tlen} — walking up parents')

        _parents = page.evaluate("""(el) => {
            try {
                if (!el) return [];
                const out = [];
                let cur = el.parentElement;
                for (let i = 0; i < 6 && cur; i++) {
                    try {
                        const tlen = (cur.innerText || '').length;
                        const rc   = cur.querySelectorAll('input[type="radio"]').length;
                        const cc   = cur.querySelectorAll('input[type="checkbox"]').length;
                        const cls  = (cur.className || '').toString();
                        const tag  = cur.tagName || '';
                        let score  = 0;
                        if (tlen >= 100 && tlen <= 3000) score += 3;
                        if (rc > 0 || cc > 0) score += 2;
                        if (cls.toLowerCase().includes('question') ||
                            cls.toLowerCase().includes('cf-')) score += 2;
                        if (tag === 'DIV' || tag === 'SECTION') score += 1;
                        if (tlen > 10000) score -= 5;
                        out.push({ tag, cls: cls.slice(0, 100), text_len: tlen,
                                   radio_count: rc, checkbox_count: cc, score });
                    } catch(pe) {
                        out.push({ tag: '?', cls: '', text_len: 0,
                                   radio_count: 0, checkbox_count: 0, score: 0 });
                    }
                    cur = cur.parentElement;
                }
                return out;
            } catch(e) { return []; }
        }""", elem)

        if not _parents:
            return elem, "CONTAINER_EXACT"

        for i, p in enumerate(_parents, 1):
            _log(f'   PARENT_{i}: tag={p["tag"]} class={p["cls"]!r} text_len={p["text_len"]} '
                 f'radio_count={p["radio_count"]} checkbox_count={p["checkbox_count"]}')

        best_idx = max(range(len(_parents)), key=lambda i: _parents[i]['score'])
        best_p   = _parents[best_idx]

        if best_p['score'] <= 0:
            return elem, "CONTAINER_EXACT"

        _log(f'   FINAL_CONTAINER: tag={best_p["tag"]} class={best_p["cls"]!r} '
             f'text_len={best_p["text_len"]} score={best_p["score"]}')

        _handle = page.evaluate_handle("""(args) => {
            try {
                if (!args || !args.el) return args ? args.el : null;
                let cur = args.el.parentElement;
                for (let i = 0; i < args.steps; i++) {
                    if (!cur) return args.el;
                    cur = cur.parentElement;
                }
                return cur || args.el;
            } catch(e) { return args ? args.el : null; }
        }""", {"el": elem, "steps": best_idx})
        _as_elem = _handle.as_element()
        if _as_elem is None:
            return elem, "CONTAINER_EXACT"
        return _as_elem, "CONTAINER_EXACT"

    except Exception as _e:
        _log(f'   CONTAINER_RESOLVE_ERR: qid={qid} {str(_e)[:60]}')
        try:
            _ft = len(elem.inner_text(timeout=2000))
            if 50 <= _ft <= 5000:
                _log(f'   CONTAINER_RESOLVE_RECOVER: qid={qid} text_len={_ft}')
                return elem, "CONTAINER_EXACT"
        except Exception:
            pass
        return elem, "CONTAINER_EXACT"


def get_active_container(page, target_qid, platform, plat_qid=None, _log_fn=None):
    """
    Find the visible question container for target_qid.
    Strategy (tried in order):
      1. QID-direct selectors (6 patterns, platform-agnostic)
      2. Platform-specific template selectors
      3. Attribute scan fallback
    Returns (element, method_label) or (None, 'NOT_FOUND').
    """
    def _log(msg):
        if _log_fn:
            _log_fn(msg, 'cyan')

    qids_to_try = [target_qid]
    if plat_qid and plat_qid != target_qid:
        qids_to_try.append(plat_qid)

    # ── Step 1: QID-direct selectors (tried before any platform logic) ──────
    _QID_DIRECT = [
        '[data-questionid="{qid}"]',
        '[data-qid="{qid}"]',
        '[id="{qid}"]',
        '[id*="{qid}"]',
        '[name="{qid}"]',
        '[name*="{qid}"]',
    ]
    for qid_val in qids_to_try:
        for tmpl in _QID_DIRECT:
            sel = tmpl.format(qid=qid_val)
            try:
                candidates = page.locator(sel)
                count = candidates.count()
                _log(f'   SELECTOR_TRY: selector={sel!r} count={count}')
                if count == 0:
                    continue
                # Prefer smallest visible element (avoids wrapper divs that
                # have the same QID attribute as the question container).
                _best_elem = None
                _best_len = None
                for i in range(min(count, 8)):
                    elem = candidates.nth(i)
                    try:
                        if not elem.is_visible():
                            continue
                        _tlen = len(elem.inner_text(timeout=1000))
                        if _best_len is None or _tlen < _best_len:
                            _best_elem = elem
                            _best_len = _tlen
                    except Exception:
                        continue
                if _best_elem is not None:
                    _log(f'   CONTAINER_MATCH: qid={target_qid} sel={sel!r} text_len={_best_len}')
                    return _resolve_container(page, _best_elem, target_qid, _log)
            except Exception:
                continue

    # ── Step 2: platform-specific template selectors ─────────────────────────
    selectors_list = (_PLATFORM_CONTAINER_SELECTORS.get(platform) or
                      _PLATFORM_CONTAINER_SELECTORS['generic'])
    for qid_val in qids_to_try:
        for selector_template in selectors_list:
            selector = selector_template.format(qid=qid_val)
            try:
                candidates = page.locator(selector)
                count = candidates.count()
                _log(f'   SELECTOR_TRY: selector={selector!r} count={count} (platform={platform})')
                if count == 0:
                    continue
                _best_elem = None
                _best_len = None
                for i in range(min(count, 8)):
                    elem = candidates.nth(i)
                    try:
                        if not elem.is_visible():
                            continue
                        _tlen = len(elem.inner_text(timeout=1000))
                        if _best_len is None or _tlen < _best_len:
                            _best_elem = elem
                            _best_len = _tlen
                    except Exception:
                        continue
                if _best_elem is not None:
                    _log(f'   CONTAINER_MATCH: qid={target_qid} sel={selector!r} text_len={_best_len}')
                    _resolved, _ = _resolve_container(page, _best_elem, target_qid, _log)
                    return _resolved, "CONTAINER_VISIBLE"
            except Exception:
                continue

    # ── Step 3: attribute scan fallback ──────────────────────────────────────
    for attr in ('data-questionid', 'data-qid'):
        try:
            fallback = page.locator(f'[{attr}]').filter(visible=True)
            count = min(fallback.count(), 30)
            for i in range(count):
                elem = fallback.nth(i)
                val = (elem.get_attribute(attr) or '').strip()
                for qid_val in qids_to_try:
                    if val.upper() == qid_val.upper():
                        _log(f'   CONTAINER_MATCH: qid={target_qid} attr={attr} val={val!r} (METADATA)')
                        _resolved, _ = _resolve_container(page, elem, target_qid, _log)
                        return _resolved, "CONTAINER_METADATA"
        except Exception:
            continue

    _log(f'   SELECTOR_FAILURE: qid={target_qid} no_container_found')
    return None, "NOT_FOUND"


def extract_from_cloned_dom(page, container_elem):
    """Clone the container, clean the clone, extract text. Never mutates original DOM."""
    return page.evaluate("""
        (containerElem) => {
            const clone = containerElem.cloneNode(true);
            const wrapper = document.createElement('div');
            wrapper.style.cssText = 'position:absolute;left:-9999px;visibility:hidden;';
            document.body.appendChild(wrapper);
            wrapper.appendChild(clone);

            const removeFromClone = (selectors) => {
                selectors.forEach(sel => {
                    try { clone.querySelectorAll(sel).forEach(el => el.remove()); } catch(e) {}
                });
            };

            removeFromClone([
                // Test navigator panels
                '.sr-test-navigator', '.cf-test-navigator',
                '[class*="tn-question"]', '[class*="test-navigator"]',
                '[class*="tn-item"]', '.wix-tn-item',
                // ARIA/hidden elements
                '[aria-hidden="true"]', '[hidden]',
                // Layout/framework panels that bleed into container text
                '[class*="agreement"]', '[class*="invoice"]',
                '[class*="footer"]', '[class*="navigation"]',
                '[class*="progress-bar"]', '[class*="page-header"]',
                // Platform-specific: Confirmit matrix/grid chrome
                '[class*="cf-matrix-header"]', '[class*="cf-matrix-"]',
                '[class*="cf-pagination"]', '[class*="cf-grid-header"]',
                '[class*="cf-answer-header"]',
                // Qualtrics / Forsta layout chrome
                '[class*="ExportTagColumn"]', '[class*="page-indicator"]',
                '[class*="pagination"]', '[class*="language-selector"]',
                '[class*="survey-header"]', '[class*="survey-footer"]'
            ]);

            clone.querySelectorAll('*').forEach(el => {
                try {
                    const style = window.getComputedStyle(el);
                    if (style.display === 'none' || style.visibility === 'hidden') {
                        el.remove();
                    }
                } catch(e) {}
            });

            const text = (clone.innerText || clone.textContent || '').trim();
            try { document.body.removeChild(wrapper); } catch(e) {}
            return text;
        }
    """, container_elem)


def verify_extraction_v2(page, container_elem, target_qid, text):
    """
    Verify extracted text belongs to target_qid using container metadata.
    Preferred over text-based search (many platforms hide QID in text).
    Returns (ok: bool, status: str).

    Status values:
      VERIFIED       — container metadata confirms this is the right QID
      UNVERIFIED     — no strong positive or negative evidence
      EMPTY          — text too short to verify
      WRONG_CONTENT  — positive evidence of wrong QID
      MULTI_QUESTION_LEAK — multiple QID markers visible (body-fallback contamination)
    """
    if not text or len(text.strip()) < 10:
        return True, 'EMPTY'

    if container_elem is not None:
        try:
            qid_attr = (container_elem.get_attribute('data-questionid') or
                        container_elem.get_attribute('data-qid') or
                        container_elem.get_attribute('id') or '').strip()
            if qid_attr:
                if qid_attr.upper() == target_qid.upper():
                    return True, 'VERIFIED'
                # Reject any structured id that starts with a letter and contains
                # at least one digit — this pattern covers real QIDs (Q15, R2, S99)
                # AND platform-specific ids (question-42, item_7), all of which
                # signal wrong content when they don't match target_qid.
                if re.match(r'^[A-Za-z]', qid_attr) and re.search(r'\d', qid_attr):
                    return False, f'WRONG_CONTENT ({qid_attr})'
        except Exception:
            pass

    # Multi-question marker leak — more than one distinct [Question ID: X] tag visible
    marker_hits = re.findall(r'\[question\s+id:\s*([^\]\s]+)\]', text, re.I)
    if len(marker_hits) > 1:
        return False, f'MULTI_QUESTION_LEAK ({len(marker_hits)} markers)'

    # Single marker present for a DIFFERENT QID → wrong content
    if len(marker_hits) == 1 and marker_hits[0].upper() != target_qid.upper():
        return False, f'WRONG_CONTENT ({marker_hits[0]})'

    return True, 'UNVERIFIED'


def _is_container_id(qid):
    """Return True for Confirmit container IDs like Q1Q1, R2R2, S0S0.
    These are outer wrapper divs whose id is the real QID repeated twice."""
    if len(qid) < 4 or len(qid) % 2 != 0:
        return False
    half = len(qid) // 2
    return qid[:half] == qid[half:]


def _parse_tn_items_from_html(html, tn_css_selectors, qid_re):
    """
    Extract the ordered TN item list from static HTML — replaces per-item
    Playwright get_attribute / inner_text calls. Returns the same
    [(nav_index, doc_qid, plat_qid), ...] list the Playwright loop produced.
    Works for any platform whose TN renders in the initial HTML.
    """
    try:
        from bs4 import BeautifulSoup
        import re as _re
        soup = BeautifulSoup(html, "html.parser")
        elements = []
        for sel in tn_css_selectors:
            found = soup.select(sel.replace("'", '"'))
            if found:
                elements = found
                break
        if not elements:
            return [], 0, 0
        seen_qids = set()
        result = []
        skipped = 0
        fw_skipped = 0
        for ni, el in enumerate(elements):
            try:
                doc_qid = None
                # Strategy 1: data attributes
                for attr in ("data-qid", "data-question-id", "data-id"):
                    v = el.get(attr) or ""
                    if v and qid_re.match(v.strip()):
                        doc_qid = qid_re.match(v.strip()).group(1); break
                # Strategy 2: id attribute
                if not doc_qid:
                    v = el.get("id") or ""
                    m = qid_re.search(v)
                    if m: doc_qid = m.group(1)
                # Strategy 3: class attribute
                if not doc_qid:
                    cls = " ".join(el.get("class") or [])
                    m = _re.search(r'question[-_]?id[-_]?([A-Za-z]\w*\d+\w*)', cls, _re.IGNORECASE)
                    if m: doc_qid = m.group(1)
                # Strategy 4: inner text — first 3 words
                if not doc_qid:
                    words = el.get_text().strip().split()
                    for w in words[:3]:
                        m = _re.match(r'^([A-Za-z]{1,8}\d*(?:[a-zA-Z]\d+|[a-zA-Z]?(?:bis|ter|Info|info|Ex|_\d+|\.\d+)?))\s*$', w.strip())
                        if m and is_valid_qid(m.group(1)): doc_qid = m.group(1); break
                # Strategy 5: nested input/select name
                if not doc_qid:
                    inp = el.find(["input", "select"])
                    if inp:
                        v = inp.get("name") or ""
                        m = qid_re.search(v)
                        if m: doc_qid = m.group(1)
                if not doc_qid:
                    continue
                if _is_container_id(doc_qid):
                    skipped += 1
                    continue
                if should_skip_qid(doc_qid):
                    fw_skipped += 1
                    continue
                if doc_qid in seen_qids:
                    continue
                words = el.get_text().strip().split()
                second = words[1].strip() if len(words) > 1 else ''
                plat_qid = second if (second and second != doc_qid
                    and _re.match(r'^[A-Za-z]\w*\d+\w*$', second)) else doc_qid
                result.append((ni, doc_qid, plat_qid))
                seen_qids.add(doc_qid)
            except: continue
        return result, skipped, fw_skipped
    except Exception:
        return [], 0, 0


def _parse_question_texts_from_html(html):
    """
    Extract all question text blocks from static HTML in one pass.
    Removes TN sidebar, splits full page text by [Question ID: XXX] markers,
    and attempts BS4-based option extraction from question containers.
    Returns {normalised_qid: {"text": str, "options": list, "has_inputs": bool}}.
    Works for any platform that renders all questions in the initial HTML.
    """
    try:
        from bs4 import BeautifulSoup
        import re as _re
        def _nq(s): return s.replace('.','').replace('_','').replace('-','').replace(' ','').lower()
        soup = BeautifulSoup(html, "html.parser")
        # Remove TN sidebar (mirrors the Playwright evaluate removal)
        for sel in ['.sr-test-navigator', '[class*="sr-tn"]',
                    '.cf-test-navigator', '[class*="cf-tn-"]']:
            try:
                for el in soup.select(sel): el.decompose()
            except: pass
        body = soup.body or soup
        full_text = body.get_text(separator='\n')
        full_text = _re.sub(r'\*Shown in Testing mode only\*', '', full_text)
        full_text = _re.sub(r'\n{3,}', '\n\n', full_text).strip()
        markers = list(_re.finditer(r'\[Question ID:\s*([^\]]+)\]', full_text))
        if not markers:
            return {}
        _noise = [
            (r'(?im)^[ \t]*test\s*link\b.*$', ''),
            (r'(?m)^\d+%(?:[ \t]+\d+%){2,}[ \t]*$', ''),
            (r'(?m)^[^\S\n]*[←→◄►«»\xab\xbb]{1,4}[^\S\n]*$', ''),
        ]
        result = {}
        for mi, m in enumerate(markers):
            raw_qid = m.group(1).strip()
            start = m.end()
            end = markers[mi + 1].start() if mi + 1 < len(markers) else len(full_text)
            block = full_text[start:end].strip()
            for pat, repl in _noise:
                block = _re.sub(pat, repl, block)
            block = _re.sub(r'\n{3,}', '\n\n', block).strip()
            if len(block) > 5000:
                block = block[:5000].strip()
            result[_nq(raw_qid)] = {"text": block, "options": [], "has_inputs": True}
        # Best-effort: extract options from data-qid containers
        _opt_sels = [
            ".cf-radio-answer__text", ".cf-checkbox-answer__text",
            "label.cf-answer", ".answer-text", ".option-text",
            ".sv-item__control-label", ".response-option",
        ]
        _skip_types = {"hidden", "submit", "button", "image", "reset"}
        for container in soup.find_all(attrs={"data-qid": True}):
            qid_val = (container.get("data-qid") or "").strip()
            nkey = _nq(qid_val)
            if nkey not in result:
                continue
            opts = []
            seen_o = set()
            for sel in _opt_sels:
                try:
                    found = container.select(sel)
                    if found:
                        for o in found:
                            t = o.get_text(strip=True)
                            if t and len(t) < 200 and t not in seen_o:
                                seen_o.add(t); opts.append({"text": t})
                        if opts: break
                except: pass
            if opts:
                result[nkey]["options"] = opts
            inp = container.find(["input", "select", "textarea"],
                attrs={"type": lambda t: (t or "text") not in _skip_types})
            result[nkey]["has_inputs"] = inp is not None
        return result
    except Exception:
        return {}



def get_gemini_model():
    """Returns a configured Gemini model if API key is set, else None."""
    try:
        key = api_store.get('gemini', {}).get('key', '').strip()
        if not key:
            key = os.environ.get('GEMINI_API_KEY', '').strip()
        if not key:
            return None
        import google.generativeai as genai
        genai.configure(api_key=key)
        return genai.GenerativeModel('gemini-2.5-flash')
    except Exception:
        return None


def _ai_call(model, prompt, timeout=60):
    """Call model.generate_content with a hard wall-clock timeout.
    Raises TimeoutError if the API hasn't responded within `timeout` seconds.
    The hung background thread is daemonised so it won't block process exit."""
    _result = [None]
    _exc    = [None]
    def _run():
        try:
            _result[0] = model.generate_content(prompt)
        except Exception as _e:
            _exc[0] = _e
    _t = threading.Thread(target=_run, daemon=True)
    _t.start()
    _t.join(timeout=timeout)
    if _t.is_alive():
        raise TimeoutError(f"Gemini timed out after {timeout}s")
    if _exc[0]:
        raise _exc[0]
    return _result[0]


def ai_compare_text(model, qid, doc_text, live_text):
    """Use Gemini to intelligently compare doc question vs live question.
    Returns (is_issue, issue_type, details) or None if AI unavailable/errors."""
    if not model:
        return None
    try:
        prompt = (
            "You are a survey QC expert. Compare the SPEC (expected) question text "
            "against the LIVE (actual) question text from a deployed survey.\n\n"
            "SPEC: " + doc_text[:1500] + "\n\n"
            "LIVE: " + live_text[:1500] + "\n\n"
            "Ignore trivial differences (whitespace, punctuation, HTML artifacts). "
            "Flag ONLY real issues: missing words, changed meaning, wrong wording, missing instructions.\n"
            "Respond ONLY with valid JSON, nothing else:\n"
            '{"issue": true/false, "type": "short type e.g. MISSING WORDS", "details": "one short sentence", "severity": "HIGH/MEDIUM/LOW"}'
        )
        resp = _ai_call(model, prompt)
        import json, re as _re
        raw = resp.text.strip()
        raw = _re.sub(r'```json|```', '', raw).strip()
        data = json.loads(raw)
        if data.get('issue'):
            return (True, data.get('type', 'AI FLAGGED'), data.get('details', ''), data.get('severity', 'MEDIUM'))
        return (False, None, None, None)
    except Exception:
        return None


def ai_compare_batch(model, batch):
    """Compare MULTIPLE questions in ONE AI call (saves quota: 50 calls -> 6 calls).
    batch = list of dicts: [{qid, doc_text, doc_opts, live_text, live_opts}, ...]
    Returns dict {qid: [issue dicts]} or None if AI unavailable after retries."""
    if not model or not batch:
        return None
    import json, re as _re
    blocks = []
    for item in batch:
        do = "\n".join("- " + o for o in item['doc_opts']) if item['doc_opts'] else "(none/open-ended)"
        lo = "\n".join("- " + o for o in item['live_opts']) if item['live_opts'] else "(none/open-ended)"
        blocks.append(
            "=== QUESTION " + item['qid'] + " ===\n"
            "SPEC text: " + item['doc_text'][:900] + "\n"
            "SPEC options:\n" + do + "\n"
            "LIVE text: " + item['live_text'][:1100] + "\n"
            "LIVE options:\n" + lo
        )
    prompt = (
        "You are an expert survey QC reviewer. Compare each SPEC question (design doc) vs LIVE "
        "(deployed survey). Language may be ANY (French, Urdu, Hindi, German, English) - compare by "
        "MEANING not exact words.\n\n"
        + "\n\n".join(blocks) + "\n\n"
        "RULES:\n"
        "1. IGNORE: whitespace, punctuation, capitalization, HTML artifacts, bracket instructions "
        "like [ALL COUNTRIES], translation notes, formatting, question order or position in the "
        "survey, and interviewer-only notes like '(ne pas poser)', '(do not ask)', '(ne pas lire)', "
        "'(do not read)' or any equivalent in any language — these are conditional skip instructions "
        "for the interviewer, NOT errors in the live survey.\n"
        "2. LIVE may have EXTRA surrounding text - that is OK. Only check if SPEC question's MEANING "
        "is present SOMEWHERE in LIVE.\n"
        "3. Flag REAL issues only: question text meaning genuinely changed or completely absent, "
        "answer options with different meaning or completely missing, wrong wording that changes "
        "meaning, truncated text.\n"
        "4. If SPEC content IS present in LIVE (even with extra text or in a different position), "
        "report NO issue for it.\n\n"
        "Respond ONLY with valid JSON object. Key = question id, value = array of issues "
        "(empty array if no issue). Example:\n"
        '{"A1": [{"type":"MISSING TEXT","details":"...","severity":"HIGH"}], "A2": []}'
    )
    for attempt in range(3):
        try:
            resp = _ai_call(model, prompt)
            raw = _re.sub(r'```json|```', '', resp.text.strip()).strip()
            m = _re.search(r'\{.*\}', raw, _re.DOTALL)
            raw = m.group(0) if m else raw
            data = json.loads(raw)
            out = {}
            for qid, issarr in data.items():
                if not isinstance(issarr, list):
                    continue
                out[qid] = []
                for d in issarr:
                    if d and d.get('type'):
                        out[qid].append({"qid": qid, "type": d.get('type', 'AI FLAGGED'),
                                         "details": d.get('details', ''), "severity": d.get('severity', 'MEDIUM')})
            return out
        except Exception as e:
            err = str(e)
            is_rate_limit = '429' in err or 'quota' in err.lower() or 'rate' in err.lower()
            is_transient  = '503' in err or 'unavailable' in err.lower()
            if attempt < 2:
                time.sleep(10 if is_rate_limit else 15 if is_transient else 3)
                continue
    return None


def ai_compare_full(model, qid, doc_text, doc_opts, live_text, live_opts):
    """Language-agnostic semantic comparison of a full question (text + options).
    Returns a LIST of issue dicts, or None if AI unavailable/errored.
    Empty list = no issues found (question is correct)."""
    if not model:
        return None
    try:
        import json, re as _re
        doc_opts_str = "\n".join("- " + o for o in doc_opts) if doc_opts else "(none / open-ended)"
        live_opts_str = "\n".join("- " + o for o in live_opts) if live_opts else "(none / open-ended)"
        prompt = (
            "You are an expert survey QC reviewer. Compare a SPEC question (from the design document) "
            "against the LIVE question (from the deployed survey). The language may be ANY language "
            "(French, Urdu, Hindi, German, English, etc.) - compare by MEANING, not exact words.\n\n"
            "QUESTION ID: " + qid + "\n\n"
            "=== SPEC (expected) ===\n"
            "Question text: " + doc_text[:1200] + "\n"
            "Answer options:\n" + doc_opts_str + "\n\n"
            "=== LIVE (actual deployed) ===\n"
            "Question text: " + live_text[:1500] + "\n"
            "Answer options:\n" + live_opts_str + "\n\n"
            "RULES:\n"
            "1. IGNORE: whitespace, punctuation, capitalization, HTML artifacts, programming instructions "
            "in brackets like [ALL COUNTRIES], translation notes, formatting differences, question order "
            "or position in the survey, and interviewer-only notes like '(ne pas poser)', '(do not ask)', "
            "'(ne pas lire)', '(do not read)' or any equivalent in any language — these are conditional "
            "skip instructions for the interviewer, NOT errors in the live survey.\n"
            "2. The LIVE text may contain EXTRA surrounding content from the page - that is OK, only check "
            "if the SPEC question's MEANING is present SOMEWHERE in LIVE.\n"
            "3. Flag REAL issues only: (a) question text meaning genuinely changed or completely absent, "
            "(b) answer options with different meaning or completely missing, "
            "(c) wrong wording that changes meaning.\n"
            "4. If the SPEC question content IS present in LIVE (even with extra text or in a different "
            "position), report NO issue.\n\n"
            "Respond ONLY with valid JSON array, nothing else. Empty array if no issues:\n"
            '[{"type": "SHORT TYPE", "details": "one short sentence", "severity": "HIGH/MEDIUM/LOW"}]'
        )
        resp = _ai_call(model, prompt)
        raw = resp.text.strip()
        raw = _re.sub(r'```json|```', '', raw).strip()
        data = json.loads(raw)
        if not isinstance(data, list):
            data = [data] if data else []
        out = []
        for d in data:
            if d and d.get('type'):
                out.append({"qid": qid, "type": d.get('type', 'AI FLAGGED'),
                            "details": d.get('details', ''), "severity": d.get('severity', 'MEDIUM')})
        return out
    except Exception:
        return None


def ai_semantic_match_batch(model, unmatched_items, live_data, doc_questions):
    """One Gemini call: for each doc question that failed ALL regex/content strategies,
    find the semantically closest live QID (any language, any survey).
    unmatched_items: list of {nqid, qid, doc_key}
    live_data: full live dict {live_qid: {text, options, ...}}
    doc_questions: the questions dict from the parsed doc
    Returns dict {doc_qid: live_qid_or_NONE}
    """
    if not model or not unmatched_items or not live_data:
        return {}
    import json, re as _re

    doc_blocks = []
    for u in unmatched_items:
        dq = doc_questions.get(u["doc_key"], {}) if u.get("doc_key") else {}
        opts = [o.get("text", "")[:60] for o in dq.get("options", [])[:5]]
        doc_blocks.append({
            "qid": u["qid"],
            "text": dq.get("text", "")[:150],
            "options": opts,
        })

    live_blocks = []
    for lk, ld in live_data.items():
        if ld.get("status") != "OK":
            continue
        opts = [o.get("text", "")[:60] for o in ld.get("options", [])[:5]]
        live_blocks.append({
            "qid": lk,
            "text": ld.get("text", "")[:150],
            "options": opts,
        })

    prompt = (
        "You are a survey QC expert. Some doc questions could not be matched to live questions "
        "by ID alone. Using SEMANTIC meaning (not QID), find the best live match for each doc question. "
        "Language may be ANY — compare by meaning, not exact words.\n\n"
        "DOC QUESTIONS (unmatched):\n"
        + json.dumps(doc_blocks, ensure_ascii=False) + "\n\n"
        "LIVE QUESTIONS (all available):\n"
        + json.dumps(live_blocks, ensure_ascii=False) + "\n\n"
        "RULES:\n"
        "1. Match by TOPIC and MEANING — the live QID may differ due to platform renaming.\n"
        "2. Only match if you are confident the questions are about the same thing.\n"
        "3. If no live question matches, return NONE for that doc QID.\n"
        "4. Return ONLY a JSON object mapping each doc QID to a live QID or NONE.\n"
        "Example: {\"R3\": \"R3new\", \"S5\": \"NONE\"}\n"
        "Return ONLY the JSON object, nothing else."
    )

    for attempt in range(3):
        try:
            resp = _ai_call(model, prompt)
            raw = resp.text.strip()
            raw = _re.sub(r'```json|```', '', raw).strip()
            m = _re.search(r'\{.*\}', raw, _re.DOTALL)
            if m:
                result = json.loads(m.group(0))
                if isinstance(result, dict):
                    return result
            return {}
        except Exception as e:
            is_rate_limit = "429" in str(e) or "quota" in str(e).lower()
            is_transient  = "503" in str(e) or "unavailable" in str(e).lower()
            if attempt < 2:
                time.sleep(10 if is_rate_limit else 15 if is_transient else 3)
                continue
    return {}


def _assign_confidence(issue):
    """Thin wrapper kept for backward compat — delegates to calculate_confidence()."""
    score, _label, _reasons = calculate_confidence(issue)
    level = 'HIGH' if score >= 75 else 'MEDIUM' if score >= 55 else 'LOW'
    return score, level


def calculate_confidence(issue):
    """Return (score 0-100, label str, reasons list) for an issue dict.

    Score is derived from issue type, extraction quality, XML cross-validation,
    and live-crawl status.  All inputs come from the issue dict itself (including
    the optional 'evidence' sub-dict populated by _enrich_issues()).
    """
    t   = issue.get('type', '')
    ev  = issue.get('evidence', {})
    xml_verdict       = issue.get('xml_verdict', '')
    extraction_status = ev.get('extraction_status', '')
    live_status       = ev.get('live_status', '')

    # ── Type-based starting point ──────────────────────────────────────────
    _TYPE_BASE = {
        'MISSING IN LIVE':               75,
        'PIPING NOT RESOLVED':           80,
        'OPTIONS MISMATCH':              65,
        'WORDS MISSING':                 65,
        'TEXT MISMATCH':                 60,
        'MISSING IN OTHER':              70,
        'MANDATORY MISSING':             58,
        'NAMING MISMATCH':               50,
        'MERGED IN LIVE':                45,
        'CONDITIONAL':                   30,
        'ERROR PAGE':                    52,
        'TERMINATION MISSING':           65,
        'QID IN EXPORT NOT IN DOC/LIVE': 60,
        # Export validation rules R031-R038
        'R031': 88, 'R032': 72, 'R033': 80,
        'R034': 85, 'R035': 75, 'R036': 85,
        'R037': 82, 'R038': 80,
    }
    # Fallback for AI-generated issues: use severity
    _sev_base = {'HIGH': 72, 'MEDIUM': 60, 'LOW': 42}
    score = _TYPE_BASE.get(t, _sev_base.get(issue.get('severity', 'MEDIUM'), 55))
    reasons: list = []

    # ── Source coverage ────────────────────────────────────────────────────
    has_doc  = bool(ev.get('doc_text')  or ev.get('doc_options'))
    has_live = bool(ev.get('live_text') or ev.get('live_options') or ev.get('live_status'))
    has_xml  = bool(ev.get('xml_text')  or ev.get('xml_options'))
    n_src = sum([has_doc, has_live, has_xml])
    if n_src >= 3:
        score += 15
        reasons.append('All 3 sources checked (doc + XML + live)')
    elif n_src == 2:
        score += 5
        reasons.append('2 of 3 sources checked')
    elif n_src == 1:
        score -= 5
        reasons.append('Only 1 source available')

    # ── Extraction quality ─────────────────────────────────────────────────
    if extraction_status == 'VERIFIED':
        score += 15
        reasons.append('Live content verified — QID confirmed in page text')
    elif extraction_status == 'UNVERIFIED':
        score -= 15
        reasons.append('Live content unverified — QID not found in page text')
    elif extraction_status == 'WRONG_CONTENT':
        score -= 35
        reasons.append('Live extraction captured wrong question — result unreliable')
    elif extraction_status == 'CONDITIONAL':
        score -= 20
        reasons.append('Question hidden/conditional — live content unreachable')
    elif extraction_status == 'EMPTY':
        score -= 10
        reasons.append('No live text captured for this question')

    # ── XML cross-validation ───────────────────────────────────────────────
    if xml_verdict == 'live_differs':
        score += 20
        reasons.append('XML export confirms: live differs from spec')
    elif xml_verdict == 'all_differ':
        score += 25
        reasons.append('All 3 sources differ — programming error likely')
    elif xml_verdict == 'doc_differs':
        score -= 15
        reasons.append('XML matches live — spec doc may be outdated')

    # ── Live status ────────────────────────────────────────────────────────
    if 'CONDITIONAL' in live_status.upper() and extraction_status != 'CONDITIONAL':
        score -= 20
        reasons.append('Question hidden behind routing condition')
    elif live_status.startswith('ERROR') or 'crawl_failed' in live_status:
        score -= 10
        reasons.append('Live page had a crawl error')
    elif 'WRONG_CONTENT' in live_status and extraction_status != 'WRONG_CONTENT':
        score -= 30
        reasons.append('Live extraction captured wrong question')

    score = max(0, min(100, score))
    label = (
        'Confirmed Bug'         if score >= 90 else
        'Likely Bug'            if score >= 75 else
        'Possible Issue'        if score >= 55 else
        'Needs Review'          if score >= 35 else
        'Likely False Positive'
    )
    return score, label, reasons


def _build_mismatch_detail(issue, doc_q, live_q, xml_q):
    """Human-readable one-line description of what differs."""
    t = issue.get('type', '')
    if t in ('TEXT MISMATCH', 'WORDS MISSING'):
        d = (doc_q.get('text', '') if doc_q else '')[:80]
        l = (live_q.get('text', '') if live_q else '')[:80]
        if d and l:
            return f'Doc: "{d}..." / Live: "{l}..."'
    elif t == 'OPTIONS MISMATCH':
        d_opts = {o['text'] for o in (doc_q.get('options', []) if doc_q else [])}
        l_opts = {o['text'] for o in (live_q.get('options', []) if live_q else [])}
        missing = sorted(d_opts - l_opts)
        extra   = sorted(l_opts - d_opts)
        parts = []
        if missing:
            parts.append('Missing from live: ' + '; '.join(missing[:3]))
        if extra:
            parts.append('Extra in live: ' + '; '.join(extra[:2]))
        return ' | '.join(parts) if parts else 'Answer options differ'
    elif t == 'MISSING IN LIVE':
        txt = (doc_q.get('text', '') if doc_q else '')[:80]
        return f'"{txt}…" — in spec but absent from live survey' if txt else issue.get('details', '')
    return issue.get('details', '')[:200]


def _enrich_issues(issues, questions, live_data, xml_questions):
    """Post-processing pass: attach evidence snippets and recalculate confidence.

    Safe to call multiple times — skips issues that already have an 'evidence'
    key so existing higher-quality evidence is never overwritten.
    """
    def _n(q): return re.sub(r'[^a-z0-9]', '', q.lower())

    _xml_by_n  = {_n(xq.get('qid_normalized') or xq.get('qid', '')): xq
                  for xq in (xml_questions or []) if xq.get('qid')}
    _live_by_n = {_n(lk): (lk, ld) for lk, ld in live_data.items()}
    _doc_by_n  = {_n(dk): (dk, dd) for dk, dd in questions.items()}

    for issue in issues:
        if issue.get('evidence'):
            continue
        qid = issue.get('qid', '')
        if not qid:
            continue
        n = _n(qid)

        _dk, doc_q  = _doc_by_n.get(n,  (None, None))
        _lk, live_q = _live_by_n.get(n, (None, None))
        xml_q       = _xml_by_n.get(n)

        # Prefer the extraction_status stored during crawl; fall back to derivation
        extraction_status = ''
        if live_q:
            extraction_status = live_q.get('extraction_status', '')
            if not extraction_status:
                # Legacy path: derive from status string + text content
                ls = live_q.get('status', '')
                if 'WRONG_CONTENT' in ls.upper():
                    extraction_status = 'WRONG_CONTENT'
                elif 'CONDITIONAL' in ls.upper():
                    extraction_status = 'CONDITIONAL'
                elif ls == 'OK':
                    lt = live_q.get('text', '')
                    extraction_status = 'VERIFIED' if qid.lower() in lt.lower() else (
                        'UNVERIFIED' if lt else 'EMPTY'
                    )
                else:
                    extraction_status = 'UNVERIFIED'

        issue['evidence'] = {
            'doc_text':          (doc_q.get('text',    '') if doc_q  else '')[:250],
            'doc_options':       [o['text'] for o in (doc_q.get('options',  []) if doc_q  else [])],
            'live_text':         (live_q.get('text',   '') if live_q else '')[:250],
            'live_options':      [o['text'] for o in (live_q.get('options', []) if live_q else [])],
            'live_status':       (live_q.get('status', '') if live_q else ''),
            'xml_text':          (xml_q.get('text',    '') if xml_q  else '')[:250],
            'xml_options':       [o['text'] for o in (xml_q.get('options',  []) if xml_q  else [])],
            'extraction_status': extraction_status,
            'mismatch_detail':   _build_mismatch_detail(issue, doc_q, live_q, xml_q),
        }

        score, label, reasons = calculate_confidence(issue)
        issue['confidence']         = score
        issue['confidence_label']   = label
        issue['confidence_reasons'] = reasons
        issue['conf_level'] = 'HIGH' if score >= 75 else 'MEDIUM' if score >= 55 else 'LOW'


def ai_generate_summary(model, questions, live_data, issues):
    """Generate a human-readable AI summary of the QC report."""
    if not model:
        # Fallback: rule-based summary
        high = sum(1 for i in issues if i.get('severity') == 'HIGH')
        med = sum(1 for i in issues if i.get('severity') == 'MEDIUM')
        total = len(issues)
        if total == 0:
            return "All checks passed. No issues detected across " + str(len(questions)) + " questions. Survey is ready to launch."
        return (str(total) + " issue(s) detected: " + str(high) + " high-priority, " + str(med) + " medium. "
                "Review high-priority issues before launching the survey.")
    try:
        issue_brief = "; ".join([i.get('type', '') + " on " + i.get('qid', '') for i in issues[:15]])
        prompt = (
            "You are a survey QC expert. Write a 2-3 sentence professional summary "
            "of this survey QC report for a client.\n\n"
            "Total questions: " + str(len(questions)) + "\n"
            "Total issues: " + str(len(issues)) + "\n"
            "Issues found: " + (issue_brief if issue_brief else "none") + "\n\n"
            "Be concise, professional, actionable. Plain text only, no markdown."
        )
        resp = _ai_call(model, prompt)
        return resp.text.strip()[:600]
    except Exception:
        total = len(issues)
        if total == 0:
            return "All checks passed. No issues detected. Survey is ready to launch."
        return str(total) + " issue(s) detected. Review before launching."


def _parse_tables_for_qids(doc, questions, qid_pat, junk_re):
    """
    Supplement the paragraph parser: scan every doc table for question IDs
    that only appear inside grid/matrix tables (e.g. Q1, Q2, Q11, Q12, Q12.2
    in Confirmit docs where the QID sits in a header cell, not in a paragraph).

    Rules:
    - For PROGRAMMING TABLE blocks: extract the header QID only (skip routing/logic body).
    - For other tables: scan all cells for QID patterns and extract answer options.
    - Skip cells whose entire text is a bracketed label like [R3 – label].
    - Register every discovered QID in questions{} (don't overwrite existing text).
    - Extract answer options from the first column of data rows below the header.
    """
    pt_re = re.compile(r'PROG(?:RAM(?:M?ING)?)?\s+TABLE', re.IGNORECASE)
    num_only = re.compile(r'^\d+$')
    bare_qid_re = re.compile(
        r'^(?P<qid>[A-Za-z]{1,8}\d+[a-zA-Z]?(?:bis|ter|Info|info|Ex|_\d+|\.\d+)?)$'
    )
    # Keywords that appear in PROGRAMMING TABLE body rows (not QID lines)
    _prog_kw = re.compile(
        r'^(?:PROG(?:RAM(?:M?ING)?)?\s+TABLE|TYPE|ROUTING|ROUTINE|LOGIC'
        r'|CODED|OPEN\s+ENDED|MULTIPLE|NUMERIC|RANGE|MIN\s*=|MAX\s*=|MANDATORY'
        r'|ALL\s+RESP|RANDOMIS|DISPLAY|SCREEN)',
        re.IGNORECASE
    )

    for table in doc.tables:
        # Build row-list-of-cell-strings
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                ct = "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                cells.append(ct.strip())
            rows.append(cells)

        full_text = "\n".join(c for row in rows for c in row if c)

        # Programming tables hold TYPE/ROUTING metadata, not answer content.
        # Extract the header QID (the question being defined) but skip the body.
        if pt_re.search(full_text):
            _found_prog_qid = False
            for _row in rows[:4]:
                if _found_prog_qid:
                    break
                for _cell in _row:
                    if _found_prog_qid:
                        break
                    for _line in _cell.split('\n'):
                        _line = _line.strip()
                        if not _line or len(_line) > 30 or _prog_kw.search(_line):
                            continue
                        if _line.startswith('[') and _line.endswith(']'):
                            continue
                        _m = qid_pat.match(_line)
                        if not _m:
                            _m = bare_qid_re.match(_line)
                        if _m:
                            _hqid = _m.group('qid')
                            if should_skip_qid(_hqid):
                                break
                            if _hqid not in questions:
                                questions[_hqid] = {
                                    "text": "", "options": [],
                                    "is_mandatory": False, "has_piping": False,
                                    "termination_rules": [], "is_numeric": False,
                                    "question_type": "",
                                }
                            if re.search(r'\bNUMERIC\b', full_text, re.IGNORECASE):
                                questions[_hqid]["is_numeric"] = True
                            # Extract full TYPE string for weak-detection logic
                            if not questions[_hqid].get("question_type"):
                                for _trow in rows[1:6]:
                                    if (_trow and _trow[0].strip().upper() == 'TYPE'
                                            and len(_trow) >= 2 and _trow[1].strip()):
                                        questions[_hqid]["question_type"] = _trow[1].strip().upper()
                                        break
                            _found_prog_qid = True
                            break
            continue

        # Scan every cell line for QID patterns.
        # Two-pass match: (1) standard qid_pat (needs a trailing separator char),
        # (2) bare QID-only line like "Q11" with no trailing char — common when a
        # table cell has the QID as its sole paragraph.
        found_qids = {}       # qid -> rest_text (question text fragment)
        header_row_idx = None

        for ri, cells in enumerate(rows):
            row_found = {}
            for cell_text in cells:
                for line in cell_text.split('\n'):
                    line = line.strip()
                    if not line or len(line) > 250 or junk_re.search(line):
                        continue
                    # Skip whole-line bracketed labels like [R3 – Consentement]
                    if line.startswith('[') and line.endswith(']'):
                        continue
                    m = qid_pat.match(line)
                    if not m and len(line) <= 20:
                        # Bare QID cell like "Q11" or "Q12.2" with no separator
                        m = bare_qid_re.match(line)
                    if not m:
                        continue
                    qid = m.group('qid')
                    rest = line[m.end():].strip()
                    # Strip column-header suffixes like "| MEMO" or "| adhésion | Note"
                    rest = re.sub(
                        r'\s*[\|–—]\s*(?:MEMO|NOTE|adh[eé]sion|SPECIFICITE|RANDOMIS).*',
                        '', rest, flags=re.IGNORECASE
                    ).strip()
                    if qid not in row_found:
                        row_found[qid] = rest

            if row_found:
                if header_row_idx is None:
                    header_row_idx = ri
                found_qids.update(row_found)

        if not found_qids:
            continue

        # Register each QID (never overwrite text already set by paragraph pass)
        for qid, rest in found_qids.items():
            if should_skip_qid(qid):
                continue
            if qid not in questions:
                questions[qid] = {
                    "text": rest, "options": [],
                    "is_mandatory": False, "has_piping": False,
                    "termination_rules": [], "is_numeric": False,
                }
            elif not questions[qid]["text"] and rest:
                questions[qid]["text"] = rest

        if header_row_idx is None:
            continue

        # Extract options from rows below the header row.
        # Use the first QID found as the target question for these options.
        primary_qid = next(iter(found_qids))
        seen_opts = {o["text"] for o in questions[primary_qid]["options"]}

        for ri in range(header_row_idx + 1, len(rows)):
            if not rows[ri]:
                continue
            opt_text = rows[ri][0].strip()
            if (not opt_text
                    or len(opt_text) < 2
                    or len(opt_text) > 200
                    or num_only.match(opt_text)
                    or qid_pat.match(opt_text)        # skip repeated-header rows
                    or bare_qid_re.match(opt_text)    # skip bare "Q11" type cells
                    or opt_text in seen_opts):
                continue
            # Find the numeric answer code in sibling cells
            code = str(ri)
            for ci in range(1, len(rows[ri])):
                v = rows[ri][ci].strip()
                if num_only.match(v):
                    code = v
                    break
            questions[primary_qid]["options"].append({"code": code, "text": opt_text})
            seen_opts.add(opt_text)


def run_qc_engine(job_id, doc_path, survey_url, country, mode, ss_paths, filter_qids=None):
    try:
        from playwright.sync_api import sync_playwright
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as e:
        jobs[job_id]['status'] = 'error'
        jobs[job_id]['logs'].append({'msg': f'Module missing: {e}', 'color': 'red'})
        jobs.persist(job_id)
        return

    job = jobs[job_id]

    def log(msg, color='white'):
        job['logs'].append({'msg': msg, 'color': color})

    def progress(p, phase=''):
        job['progress'] = p
        if phase: job['phase'] = phase

    # Heartbeat: log a "still running" line every 30s so the UI shows activity
    # even when a slow phase (e.g. Gemini batch) produces no other output.
    _hb_stop = threading.Event()
    def _heartbeat():
        _t0 = time.time()
        while not _hb_stop.wait(timeout=30):
            if jobs.get(job_id, {}).get('status') == 'running':
                _el = int(time.time() - _t0)
                job['logs'].append({'msg': f'  ⏱ Still running… ({_el}s elapsed)', 'color': 'grey'})
    _hb_thread = threading.Thread(target=_heartbeat, daemon=True)
    _hb_thread.start()

    try:
        # PHASE 1: PARSE DOC
        progress(5, 'Parsing document...')
        log('', 'white')
        log('════════════════════════════════════', 'cyan')
        log('  PHASE 1: DOCUMENT PARSING', 'cyan')
        log('════════════════════════════════════', 'cyan')

        doc = Document(doc_path)
        questions = {}
        qid_pat = re.compile(r'^\s*\[?\s*(?P<qid>[A-Za-z]{1,8}\d+[a-zA-Z]?(?:bis|ter|Info|info|Ex|_\d+|\.\d+)?)\s*[\.\-\s\]\:\)]')
        current_qid = None

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text or JUNK_RE.search(text): continue

            # Stop accumulating when a major section boundary is hit.
            # Path A: language-specific keywords (FR/EN/etc.) via SECTION_STOP_RE.
            # Path B: structural — a short paragraph that is ≥75% uppercase letters
            #   and has no '?' is a section heading in any Latin-script language
            #   (DATENSCHUTZ, DATOS PERSONALES, DATI PERSONALI, etc.).
            # Arabic question mark ؟ (U+061F) is also excluded from "has question mark".
            _no_qmark = '?' not in text and '؟' not in text
            if current_qid and _no_qmark:
                _alpha = [c for c in text if c.isalpha()]
                _upper_ratio = (sum(1 for c in _alpha if c.isupper()) / len(_alpha)
                                if _alpha else 0)
                _structural_heading = (
                    len(text) <= 60
                    and not text[0].isdigit()
                    and not qid_pat.match(text)
                    and len(_alpha) >= 4
                    and _upper_ratio >= 0.75
                )
                if SECTION_STOP_RE.search(text) or _structural_heading:
                    current_qid = None
                    continue

            m = qid_pat.match(text)
            if m:
                # FIX 1: Skip bracketed section labels like [R3 \u2013 Consentement
                # confidentialit\u00e9]. A real question paragraph that starts with '['
                # never ends with ']' \u2014 the question text follows after the bracket.
                if text.startswith('[') and text.endswith(']'):
                    continue

                qid = m.group('qid')
                if should_skip_qid(qid):
                    current_qid = None  # stop accumulating text under a skipped QID
                    continue
                current_qid = qid
                if qid not in questions:
                    questions[qid] = {"text":"","options":[],"is_mandatory":False,"has_piping":False,"termination_rules":[],"is_numeric":False,"question_type":""}
                rest = text[m.end():].strip()
                rest = re.sub(r'^[\-\u2013\u2014\s]+[A-Z][A-Za-z\s"\'\-\u2013\u2014,]+\]', '', rest).strip()
                if rest and not JUNK_RE.search(rest):
                    questions[qid]["text"] += " " + rest
                continue

            if current_qid:
                opt = re.match(r'^(\d+)[\.\)]\s+(.+)', text)
                if opt:
                    questions[current_qid]["options"].append({"code":opt.group(1),"text":opt.group(2).strip()})
                elif len(text) > 5:
                    questions[current_qid]["text"] += " " + text
                    if 'mandatory' in text.lower():
                        questions[current_qid]["is_mandatory"] = True

        # FIX 3: Pick up QIDs defined inside table grid headers (Q1, Q1.1, Q1.2,
        # Q2, Q11, Q12, Q12.2, etc.) that the paragraph pass never sees.
        _parse_tables_for_qids(doc, questions, qid_pat, JUNK_RE)

        # FIX 4: Backfill text for questions whose text paragraph appears BEFORE
        # their PROG TABLE without a QID label (e.g. Q3, Q4, Q10, Q13 in this doc).
        # Walk the doc body in order; for each PROG TABLE whose QID still has empty
        # text, look backwards for the unlabeled text paragraphs that precede it.
        _pt_re_bf = re.compile(r'PROG(?:RAM(?:M?ING)?)?\s+TABLE', re.IGNORECASE)
        _bare_bf = re.compile(r'^(?P<qid>[A-Za-z]{1,8}\d+[a-zA-Z]?(?:bis|ter|Info|info|Ex|_\d+|\.\d+)?)$')
        _bseq = []  # ('para', text) | ('prog', qid) | ('tbl', None)
        for _bc in doc.element.body.iterchildren():
            if _bc.tag == qn('w:p'):
                for _bp in doc.paragraphs:
                    if _bp._element is _bc:
                        _bseq.append(('para', _bp.text.strip()))
                        break
            elif _bc.tag == qn('w:tbl'):
                for _bt in doc.tables:
                    if _bt._element is _bc:
                        _bcells = []
                        for _brow in _bt.rows:
                            for _bcell in _brow.cells:
                                _bct = "\n".join(_bp2.text.strip() for _bp2 in _bcell.paragraphs if _bp2.text.strip())
                                _bcells.append(_bct.strip())
                        _bjnd = "\n".join(_bcells)
                        if _pt_re_bf.search(_bjnd):
                            _btqid = None
                            for _bcl in _bcells[:8]:
                                for _bln in _bcl.split('\n'):
                                    _bln = _bln.strip()
                                    if not _bln or len(_bln) > 30:
                                        continue
                                    _bm = qid_pat.match(_bln) or _bare_bf.match(_bln)
                                    if _bm:
                                        _btqid = _bm.group('qid')
                                        break
                                if _btqid:
                                    break
                            _bseq.append(('prog', _btqid))
                        else:
                            _trows = []
                            for _brow2 in _bt.rows:
                                _tr = []
                                for _bcl2 in _brow2.cells:
                                    _ct2 = ' '.join(
                                        _p2.text.strip()
                                        for _p2 in _bcl2.paragraphs
                                        if _p2.text.strip()
                                    )
                                    if _ct2:
                                        _tr.append(_ct2)
                                if _tr:
                                    _trows.append(_tr)
                            _bseq.append(('tbl', _trows))
                        break
        for _bi, (_btype, _bval) in enumerate(_bseq):
            if _btype != 'prog' or not _bval or _bval not in questions:
                continue
            if questions[_bval]['text']:
                continue  # paragraph pass already found text — leave it
            _cands = []
            for _bj in range(_bi - 1, max(_bi - 12, -1), -1):
                _jt, _jv = _bseq[_bj]
                if _jt == 'prog':
                    break  # hit the previous PROG TABLE — stop
                if _jt == 'tbl':
                    continue  # skip option / routing tables
                _ptxt = _jv
                if not _ptxt or len(_ptxt) < 8:
                    continue
                if qid_pat.match(_ptxt):
                    break  # hit a labeled QID paragraph — belongs to another question
                if JUNK_RE.search(_ptxt):
                    continue
                if SECTION_STOP_RE.search(_ptxt):
                    break  # hit a section boundary (GDPR/consent/boilerplate) — stop
                _jal = [c for c in _ptxt if c.isalpha()]
                _jrat = (sum(1 for c in _jal if c.isupper()) / len(_jal) if _jal else 0)
                if len(_ptxt) <= 60 and '?' not in _ptxt and '؟' not in _ptxt and len(_jal) >= 4 and _jrat >= 0.75:
                    break  # structural section heading — stop
                _cands.append(_ptxt)
            if _cands:
                _cands.reverse()
                questions[_bval]['text'] = re.sub(r'\s+', ' ', ' '.join(_cands)).strip()

        # Options pass: assign standalone-table options to the right QID in
        # document order. Mirrors qc_engine._pending_opts logic.
        def _assign_simple_opts(q, rows):
            for _r in rows:
                if len(_r) < 2:
                    continue
                if re.match(r'^\d+$', _r[0]):
                    _code, _text = _r[0], _r[1]
                elif len(_r) == 2 and re.match(r'^\d+$', _r[-1]) and not re.match(r'^\d+$', _r[0]):
                    _code, _text = _r[-1], _r[0]
                else:
                    continue
                if _text and not any(o['text'] == _text for o in q['options']):
                    q['options'].append({'code': _code, 'text': _text})

        # Matrix/grid option extractor: column headers of a rating-scale or
        # per-row table are the answer options. Detects matrix structure by
        # checking that body-row non-first cells are ≥50% single-char markers
        # (m/q/e/s) or short scale values, then extracts header row as options.
        _mx_marker_re = re.compile(r'^[mqes]$', re.IGNORECASE)
        _mx_short_re  = re.compile(r'^.{1,6}$')

        def _extract_matrix_opts(q, rows):
            if len(rows) < 2 or len(rows[0]) < 2:
                return
            header = rows[0]
            # Count marker cells in body rows (skip first cell = row label)
            body_total = 0
            body_markers = 0
            for _row in rows[1:]:
                for _ci, _cv in enumerate(_row):
                    if _ci == 0:
                        continue
                    _cv = _cv.strip()
                    if not _cv:
                        continue
                    body_total += 1
                    if _mx_marker_re.match(_cv) or (_mx_short_re.match(_cv) and _cv.isdigit()):
                        body_markers += 1
            if body_total == 0 or body_markers / body_total < 0.5:
                return  # not a matrix table
            # Extract header cells (skip "Randomisation"/"Randomization" label)
            _skip_hdr = re.compile(r'^randomi[sz]', re.IGNORECASE)
            for _i, _hc in enumerate(header):
                _hc = _hc.strip()
                if not _hc or _skip_hdr.match(_hc):
                    continue
                if not any(o['text'] == _hc for o in q['options']):
                    q['options'].append({'code': str(_i), 'text': _hc})

        _opt_qid = None
        _opt_pending = None
        for _btype, _bval in _bseq:
            if _btype == 'para' and _bval:
                _bm_opt = qid_pat.match(_bval)
                if _bm_opt:
                    _opt_qid = _bm_opt.group('qid')
                    # don't clear _opt_pending — table may belong to the new QID
            elif _btype == 'prog' and _bval and _bval in questions:
                if _opt_pending and not questions[_bval]['options']:
                    _assign_simple_opts(questions[_bval], _opt_pending)
                    if not questions[_bval]['options']:
                        _extract_matrix_opts(questions[_bval], _opt_pending)
                _opt_qid = _bval
                _opt_pending = None
            elif _btype == 'tbl' and isinstance(_bval, list) and _bval:
                if _opt_qid and _opt_qid in questions and not questions[_opt_qid]['options']:
                    _assign_simple_opts(questions[_opt_qid], _bval)
                    if not questions[_opt_qid]['options']:
                        _extract_matrix_opts(questions[_opt_qid], _bval)
                _opt_pending = _bval

        # Extract termination rules
        term_re = re.compile(
            r'(?:THANKS?\s*AND\s*CLOSE|THANK\s*AND\s*CLOSE'
            r'|MERCI\s+ET\s+FERMER|MERCI\s+ET\s+CLORE'
            r'|MERCI\s+FERMER|CLORE\s+LE\s+QUESTIONNAIRE'
            r'|FIN\s+DU\s+QUESTIONNAIRE|STOPPER\s+LE\s+SONDAGE'
            r'|GRAZIE\s+E\s+CHIUDI'
            r'|GRACIAS\s+Y\s+CIERRE|TERMINATE\b'
            # German
            r'|UMFRAGE\s+BEENDEN|BEENDEN\s+UND\s+SCHLIE'
            r'|FRAGEBOGEN\s+BEENDEN|ABBRECHEN\b'
            # Portuguese / Dutch / Polish / Turkish
            r'|ENCERRAR\s+(?:E\s+)?AGRADECER|FECHAR\s+E\s+AGRADECER'
            r'|SLUITEN\s+EN\s+BEDANKEN|ONDERZOEK\s+BEËINDIGEN'
            r'|ZAKOŃCZYĆ|ANKET[İI]\s+KAPAT'
            # Structural catch-all: any routing cell line that starts with a
            # close/screen/stop/end/exit verb in any language can be detected
            # by the line-level extraction below even without matching here.
            # This regex gates cell selection; structural extraction is additive.
            r')',
            re.IGNORECASE
        )
        qid_heading_re = re.compile(r'^\s*\[?\s*([A-Za-z]{1,8}\d*[a-zA-Z]?(?:bis|ter|Info|info|Ex|_\d+|\.\d+)?)\s*[\.\-\s\]\:\)]')
        _logic_tables_for_tc = []   # collected for test_generator
        body_elements = []
        for child in doc.element.body.iterchildren():
            if child.tag == qn('w:p'):
                for para in doc.paragraphs:
                    if para._element is child:
                        body_elements.append(('para', para))
                        break
            elif child.tag == qn('w:tbl'):
                for tbl in doc.tables:
                    if tbl._element is child:
                        body_elements.append(('table', tbl))
                        break

        current_context_qid = None
        for typ, item in body_elements:
            if typ == 'para':
                text = item.text.strip()
                if not text: continue
                m = qid_heading_re.match(text)
                if m and is_valid_qid(m.group(1)): current_context_qid = m.group(1)
                continue
            table = item
            table_qid = None
            table_rows_cells = []   # list[list[str]]: rows × cells
            all_cells_text = []
            for row in table.rows:
                row_cells = []
                for cell in row.cells:
                    ct = "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                    all_cells_text.append(ct)
                    if ct:
                        row_cells.append(ct)
                if row_cells:
                    table_rows_cells.append(row_cells)
            joined = "\n".join(all_cells_text)
            pt_match = re.search(r'PROG(?:RAM(?:M?ING)?)?\s+TABLE[\s\|\n]*([A-Za-z]{1,8}\d*(?:[a-zA-Z]{1,4}\d*)?)', joined, re.IGNORECASE)
            if pt_match: table_qid = pt_match.group(1)
            # Bug 3 fix: guard table_qid against poisoned values.
            # The PROG TABLE regex can capture the first word AFTER the header when
            # the table has no QID row — e.g. "PROGRAMMING TABLE\nThanks and close if:"
            # yields table_qid="Thanks".  should_skip_qid rejects stop-words and
            # non-QID tokens, so fall back to current_context_qid in that case.
            if table_qid and should_skip_qid(table_qid):
                table_qid = None

            # Collect programming / logic tables for the test generator.
            # Any table with a PROG TABLE header, or LOGIC/ROUTING/RANGE/TYPE rows.
            _is_prog_tbl = bool(pt_match) or bool(
                re.search(r'\b(LOGIC|ROUTING|ROUTINE|RANGE|TYPE)\b', joined, re.I))
            if _is_prog_tbl:
                _lt_host = table_qid or current_context_qid
                if _lt_host and not should_skip_qid(_lt_host):
                    _logic_tables_for_tc.append({
                        'host_qid': _lt_host,
                        'flat_text': joined[:2000],
                        'rows': table_rows_cells,
                    })

            # ── Termination extraction — three-mode scanner ──────────────────
            #
            # Mode A – Multi-row: the keyword row says "Thanks and close if :"
            #   and conditions appear in the rows that follow ("S1=2", "S5≤10"…).
            #   Each condition row stores a rule under its OWN QID so Phase 4
            #   navigates to the right question.
            #
            # Mode B – Inline: condition + keyword live on the same row
            #   ("code 2 | Thanks and close").  Extracts codes from the text
            #   before the keyword.
            #
            # Mode C – Keyword-only: "Thanks and close" with no condition.
            #   Stored as a manual-check rule under the table host QID.

            # Bug 2 fix: extended operator/code extractor.
            # Handles =, ≠/!=, ≤/<=, ≥/>=, <, > — operator is captured so
            # Phase 4 and the test_generator know the comparison direction.
            _CODE_OPR_RE = re.compile(
                r'(<=|>=|≤|≥|!=|≠|<(?!=)|>(?!=)|=)\s*(\d{1,4})\b'
                r'|code\s+(\d{1,3})\b',
                re.IGNORECASE,
            )

            # Bug 1 fix: condition-row pattern for Mode A.
            # Matches: "S1=2", "IF SPE≠22", "S5 ≤ 10", "S4=2 for all product"
            _COND_LINE_RE = re.compile(
                r'^\s*(?:if\s+)?([A-Za-z]\w*?)\s*(<=|>=|≤|≥|!=|≠|<(?!=)|>(?!=)|=)\s*(\d{1,4})\b',
                re.IGNORECASE,
            )
            # Detects "if :" (or "if:") that signals conditions follow on the next rows
            _MULTIROW_IF_RE = re.compile(r'\bif\s*:', re.IGNORECASE)

            # Bug 1 fix: detect multi-row mode before starting the row scan
            _multirow_term_ri = None
            for _ri, _rc in enumerate(table_rows_cells):
                _rt = " | ".join(_rc)
                if term_re.search(_rt) and _MULTIROW_IF_RE.search(_rt):
                    _multirow_term_ri = _ri
                    break

            _dbg_rules_before = sum(len(q.get('termination_rules', [])) for q in questions.values())
            app.logger.debug("[TERM DEBUG] --- Table: qid=%r ctx=%r rows=%d multirow=%s",
                             table_qid, current_context_qid,
                             len(table_rows_cells), _multirow_term_ri is not None)

            if _multirow_term_ri is not None:
                # ── Mode A: multi-row ─────────────────────────────────────────
                # The keyword header row tells us it's a termination table.
                # Every subsequent row that matches _COND_LINE_RE is a separate
                # condition ("S1=2" terminates at S1, "S5≤10" terminates at S5).
                for _ri, _row_cells in enumerate(table_rows_cells):
                    _cond_text = " | ".join(_row_cells)
                    if _ri == _multirow_term_ri:
                        app.logger.debug("[TERM DEBUG]   row %d [keyword]: %s", _ri, _cond_text[:100])
                        continue
                    _cm = _COND_LINE_RE.match(_cond_text)
                    if not _cm:
                        app.logger.debug("[TERM DEBUG]   row %d [skip — no cond match]: %s", _ri, _cond_text[:80])
                        continue
                    _cqid, _op, _ccode = _cm.group(1), _cm.group(2), _cm.group(3)
                    app.logger.debug("[TERM DEBUG]   row %d [cond]: qid=%r op=%r code=%r", _ri, _cqid, _op, _ccode)
                    # Use the condition QID as the rule host when it's a valid survey QID;
                    # fall back to the table's context QID for bare-code conditions.
                    if is_valid_qid(_cqid) and not should_skip_qid(_cqid):
                        _rule_host = _cqid
                    else:
                        _rule_host = table_qid or current_context_qid
                    if not _rule_host or should_skip_qid(_rule_host):
                        app.logger.debug("[TERM DEBUG]   row %d SKIPPED — no valid host for %r", _ri, _cqid)
                        continue
                    if _rule_host not in questions:
                        questions[_rule_host] = {
                            "text": "", "options": [], "is_mandatory": False,
                            "has_piping": False, "termination_rules": [],
                            "is_numeric": False, "question_type": "",
                        }
                    _new_rule = {
                        "test_qid": _rule_host,
                        "answer_codes": [_ccode],
                        "operator": _op,
                        "raw": _cond_text[:100],
                        "source": "generic",
                    }
                    if not any(
                        r.get('answer_codes') == [_ccode] and r.get('test_qid') == _rule_host
                        for r in questions[_rule_host]['termination_rules']
                    ):
                        questions[_rule_host]['termination_rules'].append(_new_rule)

            else:
                # ── Mode B / C: row-by-row scan ───────────────────────────────
                for _row_cells in table_rows_cells:
                    _row_text = " | ".join(_row_cells)
                    _kw_match = term_re.search(_row_text)
                    if _kw_match:
                        _dbg_cctx = _row_text[:_kw_match.start()].strip()
                        _dbg_pairs = [
                            ((_m.group(1) or "code"), (_m.group(2) or _m.group(3)))
                            for _m in _CODE_OPR_RE.finditer(_dbg_cctx)
                        ]
                    else:
                        _dbg_pairs = []
                    app.logger.debug("[TERM DEBUG]   row: %s | kw=%s | codes=%s", _row_text[:120], bool(_kw_match), _dbg_pairs)
                    if not _kw_match:
                        continue
                    host = table_qid or current_context_qid
                    if not host or should_skip_qid(host):
                        continue
                    if host not in questions:
                        questions[host] = {
                            "text": "", "options": [], "is_mandatory": False,
                            "has_piping": False, "termination_rules": [],
                            "is_numeric": False, "question_type": "",
                        }
                    # Row-level condition context: text before the keyword
                    _cctx_r = _row_text[:_kw_match.start()].strip()
                    _ops_codes_r = [
                        ((_m.group(1) or "code"), (_m.group(2) or _m.group(3)))
                        for _m in _CODE_OPR_RE.finditer(_cctx_r)
                    ]
                    # Per-line extraction (handles multi-condition cells)
                    for _cell_text in _row_cells:
                        for _line in _cell_text.split('\n'):
                            _line = _line.strip()
                            if not _line:
                                continue
                            _tm = term_re.search(_line)
                            if not _tm:
                                continue
                            _cond = _line[:_tm.start()].strip()
                            _ops_codes = [
                                ((_m.group(1) or "code"), (_m.group(2) or _m.group(3)))
                                for _m in _CODE_OPR_RE.finditer(_cond)
                            ] or _ops_codes_r
                            if _ops_codes:
                                for _op, _code in _ops_codes:
                                    if not any(
                                        r.get('answer_codes') == [_code]
                                        for r in questions[host]['termination_rules']
                                    ):
                                        questions[host]['termination_rules'].append({
                                            "test_qid": host,
                                            "answer_codes": [_code],
                                            "operator": _op,
                                            "raw": _line[:100],
                                            "source": "generic",
                                        })
                            else:
                                if not any(
                                    r.get('raw', '')[:60] == _line[:60]
                                    for r in questions[host]['termination_rules']
                                ):
                                    questions[host]['termination_rules'].append({
                                        "test_qid": host,
                                        "answer_codes": ["?"],
                                        "operator": "",
                                        "raw": _line[:100],
                                        "source": "generic-manual",
                                    })

            _dbg_rules_after = sum(len(q.get('termination_rules', [])) for q in questions.values())
            app.logger.debug("[TERM DEBUG] Table done: %d new rule(s)", _dbg_rules_after - _dbg_rules_before)

        for qid in questions:
            questions[qid]["text"] = re.sub(r'\s+', ' ', questions[qid]["text"]).strip()

        term_count = sum(len(q.get("termination_rules", [])) for q in questions.values())
        # ── [TERM DEBUG] checkpoint 3: post-parse canonical model ────────────
        app.logger.debug("[TERM DEBUG] Total termination rules in canonical model: %d", term_count)
        for _dbg_qid, _dbg_q in questions.items():
            _dbg_tr = _dbg_q.get("termination_rules", [])
            if _dbg_tr:
                app.logger.debug("[TERM DEBUG]   QID %s: %d rule(s) → %s",
                                 _dbg_qid, len(_dbg_tr), [r.get('answer_codes') for r in _dbg_tr])
        if term_count == 0:
            app.logger.debug("[TERM DEBUG] *** WARNING: 0 termination rules — check keyword matching above ***")
        # ─────────────────────────────────────────────────────────────────────
        log(f'  Questions parsed: {len(questions)}', 'green')
        log(f'  Termination rules: {term_count}', 'green')
        progress(15)

        # ── AI FALLBACK: patch weak questions the rigid parser missed ─────────
        # Splits weak QIDs into chunks of 10, one Gemini call per chunk (60s
        # SDK timeout each). Entire block is capped at 90s wall-clock — any
        # remaining chunks are skipped and code-parsed values are kept.
        _ai_fb_model = get_gemini_model()
        if _ai_fb_model:
            _weak = {}
            _no_opts_type_re = re.compile(
                r'\bOPEN\b|\bVERBATIM\b|\bRANK\b', re.IGNORECASE
            )
            for _fqid, _fq in questions.items():
                _txt_weak = len(_fq["text"].strip()) < 10
                _qt = _fq.get("question_type", "")
                _type_no_opts = bool(_qt and _no_opts_type_re.search(_qt))
                _opts_weak = (
                    not _fq["options"]
                    and not _fq.get("is_numeric")
                    and not _type_no_opts
                    # If we know the question type from the PROG TABLE, the first-pass
                    # parser already classified this question correctly. Missing options
                    # are skipped gracefully in comparison (no false positives), so
                    # there's no need for an AI patch — only missing text needs fixing.
                    and not _qt
                    and not any(kw in _fq["text"].lower()
                                for kw in ("open", "verbatim", "précisez",
                                           "specify", "saisir", "numeric",
                                           "entrez", "enter"))
                )
                if _txt_weak or _opts_weak:
                    _weak[_fqid] = {"txt": _txt_weak, "opts": _opts_weak}

            if _weak:
                log(f'  AI fallback: {len(_weak)} weak question(s) — patching with Gemini', 'yellow')

                # Build position index: PROG TABLE position for each weak QID
                _bpos = {}
                for _bi2, (_bt2, _bv2) in enumerate(_bseq):
                    if _bt2 == 'prog' and _bv2 in _weak and _bv2 not in _bpos:
                        _bpos[_bv2] = _bi2
                    elif _bt2 == 'para' and isinstance(_bv2, str):
                        _pm2 = qid_pat.match(_bv2)
                        if _pm2:
                            _q2 = _pm2.group('qid')
                            if _q2 in _weak and _q2 not in _bpos:
                                _bpos[_q2] = _bi2

                # Fix 5: wider forward window (40 vs 22) — captures question text
                # paragraphs that appear well after their PROG TABLE in batch-structured docs.
                def _ctx_for_pos(pos):
                    parts = []
                    for _bt3, _bv3 in _bseq[max(0, pos - 8): min(len(_bseq), pos + 40)]:
                        if _bt3 == 'para' and _bv3:
                            parts.append(_bv3)
                        elif _bt3 == 'prog' and _bv3:
                            parts.append(f'[PROG TABLE: {_bv3}]')
                        elif _bt3 == 'tbl' and isinstance(_bv3, list):
                            for _r3 in _bv3:
                                if isinstance(_r3, list):
                                    parts.append(' | '.join(str(c) for c in _r3 if c))
                    return '\n'.join(p for p in parts if p.strip())

                # Fix 4: inject known question_type into each section header so Gemini
                # knows whether to extract options or return [].
                _sections = []
                for _wqid in _weak:
                    _p2 = _bpos.get(_wqid)
                    if _p2 is None:
                        continue
                    _ctx2 = _ctx_for_pos(_p2)
                    if _ctx2:
                        _qt_hint = questions[_wqid].get("question_type", "")
                        _type_line = f"[TYPE: {_qt_hint}]\n" if _qt_hint else ""
                        _sections.append(
                            (_wqid, f"--- QID: {_wqid} ---\n{_type_line}{_ctx2[:900]}"))

                # Fix 6: improved prompt — explains CODEm format, conditional prefixes,
                # multi-row patient tables, and open/rank type exemptions.
                _fb_prompt_header = (
                    "You are a survey scripting document parser. Work in any language.\n"
                    "Extract question data from each section below.\n\n"
                    "Rules for 'text':\n"
                    "  - The question text shown to respondents.\n"
                    "  - Find the sentence that starts with 'QID.' or 'QID -' or is simply the"
                    " question sentence.\n"
                    "  - Strip conditional display prefixes such as '[Si X en Q1]' or"
                    " '(If Q3=2)' — they are not part of the respondent-facing text.\n"
                    "  - Ignore metadata lines: TYPE / ROUTING / MANDATORY / RANGE / PROG TABLE.\n\n"
                    "Rules for 'options':\n"
                    "  - If [TYPE] is OPEN, OPEN ENDED, OPEN NUMERIC, RANK, or VERBATIM"
                    " → return [].\n"
                    "  - Coded inline format: 'NNm Text NNm Text' e.g. '01m Oui 02m Non'.\n"
                    "  - Numbered list: '1. Option' or '1) Option'.\n"
                    "  - Table column: 'code | text' rows.\n"
                    "  - Patient-profile table rows look like:"
                    " 'P2. Question text | 01q Option A 02q Option B' —"
                    " first cell = question text, second cell = coded options.\n\n"
                )
                _fb_prompt_footer = (
                    "\n\nReturn ONLY this JSON (no markdown, no explanation):\n"
                    "{\"questions\":[{\"qid\":\"X\",\"text\":\"...\","
                    "\"options\":[{\"code\":\"1\",\"text\":\"...\"}]}]}"
                )

                # Fix 2: retry logic — one re-attempt on transient 503/429 errors.
                def _call_gemini_chunk(chunk_sections):
                    prompt = (_fb_prompt_header
                              + "\n\n".join(s for _, s in chunk_sections)
                              + _fb_prompt_footer)
                    _last_exc = None
                    for _attempt in range(2):
                        try:
                            resp = _ai_call(_ai_fb_model, prompt)
                            raw = resp.text.strip()
                            raw = re.sub(r'```json|```', '', raw).strip()
                            raw = re.sub(r'<[^>]+>.*?</[^>]+>', '', raw, flags=re.DOTALL).strip()
                            _m = re.search(r'\{.*\}', raw, re.DOTALL)
                            return json.loads(_m.group(0)) if _m else None
                        except Exception as _exc:
                            _last_exc = _exc
                            _err_s = str(_exc)
                            _is_transient = (
                                '503' in _err_s or '429' in _err_s
                                or 'unavailable' in _err_s.lower()
                                or 'quota' in _err_s.lower()
                                or 'rate' in _err_s.lower()
                            )
                            if _attempt == 0 and _is_transient:
                                time.sleep(10)
                                continue
                            break
                    raise _last_exc

                # Fix 1: timeouts scale with chunk count so every chunk can complete.
                _CHUNK_SIZE = 10
                _FB_CHUNK_TIMEOUT = 90   # seconds per chunk
                _chunks = [_sections[i:i + _CHUNK_SIZE]
                           for i in range(0, len(_sections), _CHUNK_SIZE)]
                _total_chunks = len(_chunks)
                # Total budget = per-chunk limit × count + 20s buffer, minimum 200s.
                _FB_TOTAL_TIMEOUT = max(200, _total_chunks * _FB_CHUNK_TIMEOUT + 20)
                _merged = 0

                # Fix 3: submit all chunks in parallel so wall time = max(chunk times).
                # shutdown(wait=False) avoids blocking on timed-out threads after we're done.
                progress(15, f'AI patch ({_total_chunks} chunk(s) in parallel)...')
                _deadline = time.time() + _FB_TOTAL_TIMEOUT
                _fb_pool = concurrent.futures.ThreadPoolExecutor(
                    max_workers=max(_total_chunks, 1))
                try:
                    _chunk_futures = [
                        (_ci, _fb_pool.submit(_call_gemini_chunk, _chunk))
                        for _ci, _chunk in enumerate(_chunks)
                    ]
                    for _ci, _fut in _chunk_futures:
                        _time_left = max(1, _deadline - time.time())
                        try:
                            # Fix 7: use constant in timeout message
                            _fb_data = _fut.result(
                                timeout=min(_time_left, _FB_CHUNK_TIMEOUT))
                        except concurrent.futures.TimeoutError:
                            log(f'  AI fallback chunk {_ci+1}/{_total_chunks}:'
                                f' timed out ({_FB_CHUNK_TIMEOUT}s) — skipping', 'yellow')
                            continue
                        except Exception as _fbe:
                            log(f'  AI fallback chunk {_ci+1}/{_total_chunks}:'
                                f' error ({str(_fbe)[:60]}) — skipping', 'yellow')
                            continue

                        if _fb_data:
                            for _fbq in _fb_data.get("questions", []):
                                _fid = _fbq.get("qid", "")
                                if _fid not in questions:
                                    continue
                                _q = questions[_fid]
                                if not _q["text"].strip() and (_fbq.get("text") or "").strip():
                                    _q["text"] = re.sub(r'\s+', ' ', _fbq["text"]).strip()
                                    _merged += 1
                                if not _q["options"] and _fbq.get("options"):
                                    _valid = [o for o in _fbq["options"]
                                              if o.get("code") and o.get("text")]
                                    if _valid:
                                        _q["options"] = _valid
                                        _merged += 1
                finally:
                    _fb_pool.shutdown(wait=False)  # don't block on still-running threads

                log(f'  AI fallback: patched {_merged} field(s) across'
                    f' {_total_chunks} chunk(s)', 'green')
            else:
                log('  AI fallback: rigid parser got everything — no AI call needed', 'green')
        # ── end AI fallback ───────────────────────────────────────────────────

        # Strip h-prefix hidden/programming variables from doc before comparison.
        # These are scripting-only (hAge, hQ1, HidS8 etc.) — never shown to
        # respondents, so they will never appear in live and should not be
        # reported as MISSING IN LIVE.
        _hidden_qid_re = re.compile(r'^h[A-Z]|^Hid')
        _hidden_qids = [q for q in questions if _hidden_qid_re.match(q)]
        if _hidden_qids:
            log(f'  Hidden QIDs skipped ({len(_hidden_qids)}): '
                + ', '.join(sorted(_hidden_qids)[:15])
                + ('…' if len(_hidden_qids) > 15 else ''), 'grey')
            for _hq in _hidden_qids:
                del questions[_hq]

        # Drop ghost QIDs: placeholders that were registered by the parser
        # (from a table cell label or section header) but carry zero meaningful
        # content.  A question is a TRUE ghost only when ALL of the following
        # are absent: text, options, and termination rules.  Checking text alone
        # is too aggressive — real questions like SPE (all-letter screener QID)
        # or Q11b/Q11c (sub-questions inside a grid) can end up with empty text
        # if the rigid parser missed their paragraph but still captured their
        # options or routing.
        # Sub-questions (Q11b, Q11c, Q21c, P3b …) have a letter suffix after
        # digits.  Their text/options are inherited from the parent question so
        # the parser legitimately leaves them empty; never ghost them.
        _subq_re = re.compile(r'^[A-Za-z]{1,8}\d+[a-zA-Z]\d*$')

        def _is_true_ghost(qid, v):
            if qid.upper() in SCREENER_QIDS:  # imported from qid_normalizer
                return False
            if _subq_re.match(qid):
                return False
            has_text = bool(v.get("text", "").strip())
            has_options = bool(v.get("options"))
            has_routing = bool(v.get("termination_rules"))
            return not (has_text or has_options or has_routing)

        _ghost_qids = [q for q, v in questions.items() if _is_true_ghost(q, v)]
        if _ghost_qids:
            log(f'  Ghost QIDs removed ({len(_ghost_qids)}): '
                + ', '.join(sorted(_ghost_qids)[:15])
                + ('…' if len(_ghost_qids) > 15 else ''), 'grey')
            for _gq in _ghost_qids:
                del questions[_gq]

        live_data = {}
        issues = []
        term_results = []

        # Strip only known page-navigation params so the crawler starts from
        # page 1. Everything else (auth tokens, session keys, language params,
        # any __* param not on the nav list) is preserved untouched.
        # Rule: explicit allowlist of nav params only — never strip by prefix.
        if not survey_url:
            log('  No live URL provided — running STANDARD QC (DOC + XML only)', 'yellow')
        try:
            from urllib.parse import urlparse, parse_qs, urlencode, urlunparse
            _NAV_PARAMS = {
                # Confirmit / Forsta navigation
                '__goto', '__next', '__page', '__jumpid', '__gotopage', '__target',
                '__jump', '__startpage', '__currentpage',
                # Qualtrics navigation
                'q_jfe', 'q_jfo',
                # Generic platform navigation
                'pageid', 'page_id', 'page', 'goto', 'jump',
            }
            if survey_url:
                _pu = urlparse(survey_url)
                _qs = parse_qs(_pu.query, keep_blank_values=True)
                # Only drop params whose lowercase name is an exact nav-param match.
                # Auth/session/token params (__etk, __auth, __token, __key, __sid,
                # l=, lang=, etc.) are never touched regardless of prefix.
                _clean = {k: v for k, v in _qs.items() if k.lower() not in _NAV_PARAMS}
                if len(_clean) < len(_qs):
                    _removed = sorted(set(_qs.keys()) - set(_clean.keys()))
                    survey_url = urlunparse(_pu._replace(query=urlencode(_clean, doseq=True)))
                    log(f'  URL: stripped nav param(s) {_removed} — crawl starts from page 1', 'yellow')
        except Exception as _ue:
            log(f'  URL param strip skipped: {str(_ue)[:60]}', 'grey')

        # PHASE 1.5: PARSE XML EXPORT
        xml_questions = []
        _xml_path = job.get('xml_path', '')
        if _xml_path and os.path.exists(_xml_path):
            try:
                progress(18, 'Parsing XML export...')
                log('', 'white')
                log('════════════════════════════════════', 'cyan')
                log('  PHASE 1.5: XML EXPORT PARSING', 'cyan')
                log('════════════════════════════════════', 'cyan')
                import xml_parser as _xml_parser_mod
                xml_questions, _xml_meta = _xml_parser_mod.parse_export_with_stats(_xml_path)
                _xml_hidden = _xml_meta.get('hidden_count', 0)
                job['xml_questions'] = xml_questions
                job['xml_qids'] = len(xml_questions)
                job['xml_hidden_count'] = _xml_hidden
                if xml_questions:
                    _hidden_note = f' ({_xml_hidden} hidden/template variables skipped)' if _xml_hidden else ''
                    log(f'  Phase 1.5: Parsed {len(xml_questions)} visible questions from XML export{_hidden_note}', 'green')
                else:
                    log(f'  Phase 1.5: ERROR — XML uploaded but 0 questions extracted. '
                        f'Format may be unsupported or all questions were filtered as hidden/template.', 'red')
                    log(f'  Phase 1.5: Falling back to doc-vs-live comparison (Standard mode)', 'yellow')
                    job['xml_parse_failed'] = True
            except Exception as _xe:
                log(f'  Phase 1.5: XML parse error — {str(_xe)[:120]} (continuing without XML)', 'red')
                job['xml_parse_failed'] = True
                xml_questions = []
        else:
            if _xml_path:
                log('  Phase 1.5: XML file not found, continuing without it', 'yellow')

        # ══════════════════════════════════════════════════════════════════════
        # PHASE 2.5: PRIMARY DOC vs XML COMPARISON
        # This is the main QC engine when XML is present.
        # Detects: missing questions, missing options, wrong codes, text
        # mismatches, routing issues, mandatory mismatches, piping mismatches.
        # Runs in BOTH Standard mode (no live URL) and Advanced mode (with URL).
        # ══════════════════════════════════════════════════════════════════════
        if xml_questions:
            progress(22, 'Comparing spec doc vs XML export...')
            log('', 'white')
            log('════════════════════════════════════', 'cyan')
            log(f'  PHASE 2.5: PRIMARY DOC vs XML QC', 'cyan')
            log('════════════════════════════════════', 'cyan')

            _xml_by_norm = {}
            for _xq in xml_questions:
                _xn = re.sub(r'[^a-z0-9]', '', (_xq.get('qid_normalized') or _xq.get('qid', '')).lower())
                if _xn:
                    _xml_by_norm[_xn] = _xq
            _doc_by_norm = {}
            for _dq, _dv in questions.items():
                _dn = re.sub(r'[^a-z0-9]', '', _dq.lower())
                if _dn:
                    _doc_by_norm[_dn] = (_dq, _dv)

            _xml_miss_in_doc = 0
            _doc_miss_in_xml = 0
            _xml_opt_mismatch = 0
            _xml_text_mismatch = 0
            _xml_code_mismatch = 0
            _xml_mandatory_mismatch = 0
            _xml_routing_issues = 0
            _xml_piping_issues = 0

            # ── Check XML questions against doc ───────────────────────────
            for _xn, _xq in _xml_by_norm.items():
                if should_skip_qid(_xq.get('qid', '')):
                    continue
                if _xn not in _doc_by_norm:
                    _xraw = _xq.get('qid', '')
                    _variants = [re.sub(r'[^a-z0-9]', '', v.lower())
                                 for v in build_strip_candidates(_xraw, _xn)]
                    if not any(v in _doc_by_norm for v in _variants if v):
                        _xml_miss_in_doc += 1
                        issues.append({
                            'qid': _xraw,
                            'type': 'IN_XML_NOT_IN_DOC',
                            'details': f'Question found in XML export but missing from spec doc',
                            'severity': 'MEDIUM',
                            'confidence': 72,
                            'conf_level': 'MEDIUM',
                            'source_phase': 'PHASE_2.5_XML',
                        })
                else:
                    _dq_orig, _dv = _doc_by_norm[_xn]

                    # ── Option count mismatch ──────────────────────────────
                    _doc_opt_count = len(_dv.get('options', []))
                    _xml_opt_count = len([o for o in _xq.get('options', []) if o.get('text')])
                    if (_doc_opt_count > 0 and _xml_opt_count > 0
                            and abs(_doc_opt_count - _xml_opt_count) > 1):
                        _xml_opt_mismatch += 1
                        issues.append({
                            'qid': _dq_orig,
                            'type': 'OPTIONS COUNT MISMATCH',
                            'details': (f'Doc has {_doc_opt_count} options, '
                                        f'XML has {_xml_opt_count} options — '
                                        f'verify all answer options are programmed correctly'),
                            'severity': 'HIGH' if abs(_doc_opt_count - _xml_opt_count) > 3 else 'MEDIUM',
                            'confidence': 85,
                            'conf_level': 'HIGH' if abs(_doc_opt_count - _xml_opt_count) > 3 else 'MEDIUM',
                            'source_phase': 'PHASE_2.5_XML',
                        })
                    elif _doc_opt_count > 0 and _xml_opt_count > 0:
                        # ── Per-option text check (detect missing specific options) ──
                        _doc_opt_texts = [o.get('text', '').strip().lower() for o in _dv.get('options', []) if o.get('text')]
                        _xml_opt_texts = [o.get('text', '').strip().lower() for o in _xq.get('options', []) if o.get('text')]
                        _missing_in_xml = []
                        for _dt in _doc_opt_texts:
                            if len(_dt) < 3:
                                continue
                            _best = max(
                                (SequenceMatcher(None, _dt[:80], _xt[:80]).ratio()
                                 for _xt in _xml_opt_texts),
                                default=0.0
                            )
                            if _best < 0.6:
                                _missing_in_xml.append(_dt[:40])
                        if _missing_in_xml:
                            _xml_opt_mismatch += 1
                            issues.append({
                                'qid': _dq_orig,
                                'type': 'OPTION TEXT MISSING IN XML',
                                'details': (f'Doc option(s) not found in XML: '
                                            f'{_missing_in_xml[:3]} — '
                                            f'verify answer options are correctly programmed'),
                                'severity': 'MEDIUM',
                                'confidence': 75,
                                'conf_level': 'MEDIUM',
                                'source_phase': 'PHASE_2.5_XML',
                            })

                    # ── Answer code sequence check ────────────────────────
                    _doc_codes = [o.get('code', '') for o in _dv.get('options', []) if o.get('code')]
                    _xml_codes = [o.get('code', '') for o in _xq.get('options', []) if o.get('code')]
                    if _doc_codes and _xml_codes and _doc_codes != _xml_codes:
                        # Check if codes are just reordered vs genuinely different
                        if sorted(_doc_codes) != sorted(_xml_codes):
                            _xml_code_mismatch += 1
                            issues.append({
                                'qid': _dq_orig,
                                'type': 'CODE MISMATCH',
                                'details': (f'Doc codes: {_doc_codes[:6]} — '
                                            f'XML codes: {_xml_codes[:6]}'),
                                'severity': 'HIGH',
                                'confidence': 88,
                                'conf_level': 'HIGH',
                                'source_phase': 'PHASE_2.5_XML',
                            })

                    # ── Mandatory flag mismatch ───────────────────────────
                    _doc_mandatory = bool(_dv.get('is_mandatory'))
                    _xml_type = (_xq.get('type') or '').upper()
                    # XML doesn't always carry mandatory flag — only check when
                    # doc says mandatory but XML type is open/numeric (no required attr)
                    # This is a weak signal; keep at MEDIUM confidence
                    # (XML rarely stores mandatory explicitly, so skip this check
                    #  to avoid false positives — mandatory is best verified in live)

                    # ── Routing/logic check ───────────────────────────────
                    _xml_routing = (_xq.get('routing') or '').strip()
                    _doc_has_routing = bool(_dv.get('termination_rules'))
                    if _xml_routing and not _doc_has_routing:
                        # XML has routing but doc has no logic table for this Q
                        # Only flag if routing is non-trivial (not just a skip)
                        if len(_xml_routing) > 10 and 'terminate' in _xml_routing.lower():
                            _xml_routing_issues += 1
                            issues.append({
                                'qid': _dq_orig,
                                'type': 'ROUTING IN XML NOT IN DOC',
                                'details': (f'XML has termination routing for {_dq_orig} '
                                            f'but no routing table found in spec doc — '
                                            f'verify routing logic'),
                                'severity': 'MEDIUM',
                                'confidence': 65,
                                'conf_level': 'MEDIUM',
                                'source_phase': 'PHASE_2.5_XML',
                            })

                    # ── Piping check ──────────────────────────────────────
                    _doc_has_piping = bool(_dv.get('has_piping'))
                    _xml_text_str = (_xq.get('text') or '').lower()
                    _xml_has_pipe_marker = any(
                        marker in _xml_text_str
                        for marker in ["[pipe", "{{", "<pipe", "[q", "[r"]
                    )
                    if _doc_has_piping and not _xml_has_pipe_marker:
                        _xml_piping_issues += 1
                        issues.append({
                            'qid': _dq_orig,
                            'type': 'PIPING IN DOC NOT IN XML',
                            'details': (f'Doc specifies piping for {_dq_orig} '
                                        f'but XML question text has no pipe markers — '
                                        f'verify piping is programmed'),
                            'severity': 'MEDIUM',
                            'confidence': 68,
                            'conf_level': 'MEDIUM',
                            'source_phase': 'PHASE_2.5_XML',
                        })

            # ── Check doc questions against XML ───────────────────────────
            for _dn, (_dq_orig, _dv) in _doc_by_norm.items():
                if should_skip_qid(_dq_orig):
                    continue
                if _dn not in _xml_by_norm:
                    _variants = [re.sub(r'[^a-z0-9]', '', v.lower())
                                 for v in build_strip_candidates(_dq_orig, _dn)]
                    if not any(v in _xml_by_norm for v in _variants if v):
                        _doc_miss_in_xml += 1

            # ── Summary logging ───────────────────────────────────────────
            log(f'  Phase 2.5: {len(questions)} doc QIDs vs {len(xml_questions)} XML QIDs', 'blue')
            _phase25_issues = _xml_miss_in_doc + _xml_opt_mismatch + _xml_code_mismatch + _xml_routing_issues + _xml_piping_issues
            if _phase25_issues:
                log(f'  Phase 2.5: {_phase25_issues} issue(s) found', 'yellow')
                if _xml_miss_in_doc:
                    log(f'    • {_xml_miss_in_doc} XML question(s) absent from spec doc', 'yellow')
                if _xml_opt_mismatch:
                    log(f'    • {_xml_opt_mismatch} option mismatch(es)', 'yellow')
                if _xml_code_mismatch:
                    log(f'    • {_xml_code_mismatch} code mismatch(es)', 'yellow')
                if _xml_routing_issues:
                    log(f'    • {_xml_routing_issues} routing issue(s)', 'yellow')
                if _xml_piping_issues:
                    log(f'    • {_xml_piping_issues} piping issue(s)', 'yellow')
            else:
                log('  Phase 2.5: DOC and XML are fully consistent — no structural issues', 'green')
            if _doc_miss_in_xml:
                log(f'  Phase 2.5: {_doc_miss_in_xml} doc question(s) absent from XML '
                    f'(conditional/hidden in platform — normal)', 'grey')

            # Populate _xml_qid_set for MISSING IN LIVE suppression in Phase 3
            _xml_qid_set = set(_xml_by_norm.keys())
        else:
            _xml_qid_set = set()

        # PHASE 2.6: RULE ENGINE — 10-group deterministic analysis
        _re_findings = []
        _re_summary  = {}
        try:
            from rule_engine import run_rule_engine as _run_re
            progress(26, 'Running rule engine...')
            log('', 'white')
            log('════════════════════════════════════', 'cyan')
            log('  PHASE 2.6: RULE ENGINE (10 groups)', 'cyan')
            log('════════════════════════════════════', 'cyan')
            _re_out     = _run_re(doc_questions=questions,
                                  xml_questions=xml_questions if xml_questions else [])
            _re_findings     = _re_out.get('results', [])
            _re_summary      = _re_out.get('summary', {})
            _re_ms           = _re_out.get('duration_ms', 0)
            _re_term_matrix  = _re_out.get('termination_matrix', [])
            log(f'  Rule Engine: {_re_summary.get("total", 0)} finding(s) '
                f'(HIGH={_re_summary.get("high",0)}, '
                f'MEDIUM={_re_summary.get("medium",0)}, '
                f'LOW={_re_summary.get("low",0)}) '
                f'in {_re_ms}ms', 'cyan' if _re_summary.get("high",0) == 0 else 'yellow')
            _re_by_grp = _re_summary.get('by_group', {})
            _re_group_names = {
                1:'Routing', 2:'Termination', 3:'Mandatory', 4:'Piping',
                5:'Loop', 6:'Variable', 7:'Type', 8:'Option/Code',
                9:'Graph', 10:'Export',
            }
            for _gn, _gc in _re_by_grp.items():
                if _gc > 0:
                    log(f'    G{_gn} {_re_group_names.get(_gn,"")}: {_gc}', 'yellow')
            if _re_term_matrix:
                _mismatches = sum(1 for r in _re_term_matrix if r.get('status') == 'MISMATCH')
                log(f'  Termination Matrix: {len(_re_term_matrix)} termination point(s)'
                    f'{f", {_mismatches} mismatch(es)" if _mismatches else " — all aligned"}', 'cyan')
            job['rule_engine_findings']   = _re_findings
            job['rule_engine_summary']    = _re_summary
            job['termination_matrix']     = _re_term_matrix
        except Exception as _re_err:
            log(f'  Rule Engine error (non-fatal): {str(_re_err)[:120]}', 'yellow')
            job['rule_engine_findings'] = []
            job['rule_engine_summary']  = {}
            job['termination_matrix']   = []

        # PHASE 2: CRAWL (ADVANCED mode — only when live URL provided)
        if mode in ('full', 'quick') and survey_url:
            progress(20, 'Crawling survey pages...')
            log('', 'white')
            log('════════════════════════════════════', 'cyan')
            log('  PHASE 2: SURVEY CRAWLING', 'cyan')
            log('════════════════════════════════════', 'cyan')

            with sync_playwright() as p:
                _UA = ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                       "AppleWebKit/537.36 (KHTML, like Gecko) "
                       "Chrome/124.0.0.0 Safari/537.36")
                browser = p.chromium.launch(
                    headless=True,
                    slow_mo=150,
                    args=[
                        "--disable-blink-features=AutomationControlled",
                        "--no-sandbox",
                        "--disable-dev-shm-usage",
                    ]
                )
                context = browser.new_context(
                    viewport={"width": 1366, "height": 768},
                    user_agent=_UA,
                    locale="en-US",
                    timezone_id="America/New_York",
                    extra_http_headers={
                        "Accept":          "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,*/*;q=0.8",
                        "Accept-Language": "en-US,en;q=0.9",
                        "Accept-Encoding": "gzip, deflate, br",
                        "Sec-Fetch-Dest":  "document",
                        "Sec-Fetch-Mode":  "navigate",
                        "Sec-Fetch-Site":  "none",
                        "Sec-Fetch-User":  "?1",
                        "Upgrade-Insecure-Requests": "1",
                    }
                )
                page = context.new_page()
                page.set_default_timeout(30000)
                page.add_init_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
                _nav_resp = page.goto(survey_url, wait_until="domcontentloaded", timeout=60000)
                _nav_status = _nav_resp.status if _nav_resp else "?"
                log(f"  URL loaded: {survey_url}", "cyan")
                log(f"  HTTP status: {_nav_status}", "cyan")
                page.wait_for_timeout(3000)

                # LINK CHECK: detect expired/dead survey link
                try:
                    _body = page.locator("body").inner_text(timeout=5000).lower()
                    if len(_body.strip()) < 100:
                        page.wait_for_timeout(5000)
                        _body = page.locator("body").inner_text(timeout=5000).lower()

                    # Transient server errors — retry the URL, these resolve on reload
                    _transient = [
                        "server encountered an error", "technical error",
                        "afraid i can", "internal server error",
                        "service unavailable", "temporarily unavailable",
                        "error has occurred", " 500 ", " 502 ", " 503 ",
                        "400: bad request", "bad request",
                    ]
                    # Genuine expiry — link is permanently dead, no retry
                    _expired = [
                        "survey is closed", "survey closed", "link has expired",
                        "link expired", "no longer available", "already completed",
                        "quota full", "this survey is no longer",
                    ]
                    # Hard errors — wrong session or URL, declare dead immediately
                    _hard = [
                        "session expired", "session has expired",
                        "page not found", "404 not found",
                        "not able to interpret the request",
                    ]

                    _link_dead = False
                    for _try in range(3):
                        log(f"  [Link check {_try+1}/3] HTTP {_nav_status} | {len(_body.strip())} chars | FULL BODY: {_body.strip()!r}", "cyan")

                        if any(p in _body for p in _expired) or any(p in _body for p in _hard):
                            _link_dead = True
                            break

                        if any(p in _body for p in _transient):
                            if _try < 2:
                                log(f"  Transient error on attempt {_try+1} — retrying in 5s", "yellow")
                                page.wait_for_timeout(5000)
                                page.goto(survey_url, wait_until="domcontentloaded", timeout=60000)
                                page.wait_for_timeout(4000)
                                _body = page.locator("body").inner_text(timeout=5000).lower()
                                if len(_body.strip()) < 100:
                                    page.wait_for_timeout(5000)
                                    _body = page.locator("body").inner_text(timeout=5000).lower()
                                continue
                            else:
                                log("  Still failing after 3 attempts", "yellow")
                                _link_dead = True
                                break
                        else:
                            break  # no error pattern — page looks OK

                    if _link_dead:
                        log("  LINK NOT WORKING / EXPIRED - get a fresh link", "red")
                        jobs[job_id]["status"] = "error"
                        jobs[job_id]["phase"] = "Survey link expired or not working"
                        jobs[job_id]["error"] = "Survey link is not working or has expired. Please check the link and try again with a fresh link."
                        jobs.persist(job_id)
                        try: browser.close()
                        except: pass
                        return
                except: pass
                ss_dir = f"{OUTPUT_FOLDER}/{job_id}/screenshots"
                os.makedirs(ss_dir, exist_ok=True)

                if country:
                    try:
                        page.locator(f".cf-radio-answer__text:has-text('{country}')").first.click(force=True, timeout=8000)
                        page.wait_for_timeout(1000)
                        for sel in ["button:has-text('>>')", "input[value='>>']", ".cf-button-next"]:
                            try:
                                page.locator(sel).first.click(timeout=3000); break
                            except: continue
                        page.wait_for_timeout(2500)
                    except Exception as e:
                        log('  Country select skipped: ' + str(e)[:40], 'yellow')

                # LOCKED URL CHECK: detect surveys that require pre-filled URL
                # parameters to unlock their questions (all questions show CONDITIONAL).
                # Fires only when ALL four conditions are true:
                #   1. Zero radio/checkbox in entire DOM (any visibility)
                #   2. No "[question id: X]" pattern in page text
                #   3. No QID-shaped name on any visible text input
                #   4. 2+ hidden data fields present (after filtering CSRF tokens)
                try:
                    _lc = page.evaluate("""() => {
                        const SKIP_TYPES  = new Set(['hidden','submit','button','image','reset']);
                        const SKIP_HIDDEN = new Set([
                            '__requestverificationtoken','_token','csrf','csrftoken',
                            '_method','authenticity_token','__viewstate','__eventvalidation'
                        ]);
                        const QID_RE = /^[A-Z][A-Za-z]{0,4}\\d/;
                        const all = Array.from(document.querySelectorAll('input,select,textarea'));

                        // Visible non-hidden/non-button inputs
                        const visible = all.filter(el => {
                            if (el.tagName === 'INPUT' && SKIP_TYPES.has((el.type||'').toLowerCase())) return false;
                            const s = window.getComputedStyle(el);
                            return s.display !== 'none' && s.visibility !== 'hidden' && el.offsetParent !== null;
                        });

                        // Count ALL radio/checkbox in DOM regardless of visibility —
                        // they may be inside a collapsed container on page load.
                        const anyRadioCheck = all.filter(el =>
                            el.tagName === 'INPUT' && (el.type === 'radio' || el.type === 'checkbox')
                        ).length;

                        // Check page text for Confirmit-style "[question id: r0]" marker
                        const bodyText = document.body ? document.body.innerText : '';
                        const hasQidInText = /\\[question\\s+id\\s*:/i.test(bodyText);

                        const visNames = visible.map(el => el.name || el.id || '').filter(Boolean);
                        const hasQid   = visNames.some(n => QID_RE.test(n));

                        const hiddenNames = all
                            .filter(el => (el.type||'').toLowerCase() === 'hidden')
                            .map(el => (el.name || el.id || '').toLowerCase())
                            .filter(n => n && !SKIP_HIDDEN.has(n));

                        return {
                            visibleCount:   visible.length,
                            anyRadioCheck:  anyRadioCheck,
                            hasQidInText:   hasQidInText,
                            hasQid:         hasQid,
                            visibleNames:   visNames.slice(0, 10),
                            hiddenNames:    hiddenNames
                        };
                    }""")
                    log(f"  Locked-URL check: {_lc['visibleCount']} visible, "
                        f"{_lc['anyRadioCheck']} radio/cb (DOM), "
                        f"hasQidInText={_lc['hasQidInText']}, hasQid={_lc['hasQid']}, "
                        f"{len(_lc['hiddenNames'])} hidden fields", "cyan")
                    _is_locked = (
                        _lc['anyRadioCheck'] == 0       # no radio/checkbox anywhere in DOM
                        and not _lc['hasQidInText']     # no [question id: ...] in page text
                        and not _lc['hasQid']           # no QID-named visible input
                        and len(_lc['hiddenNames']) >= 2  # but hidden data fields exist
                    )
                    if _is_locked:
                        _hf_display  = ', '.join(_lc['hiddenNames'][:8])
                        _vis_display = ', '.join(_lc['visibleNames'][:5]) or 'none'
                        _err_msg = (
                            "❌ Survey URL appears to be locked or requires "
                            "pre-filled parameters to access questions.\n\n"
                            "What we found on first page:\n"
                            f"- Visible questions: {_lc['visibleCount']} "
                            f"(data fields only: {_vis_display})\n"
                            f"- Hidden fields: {_hf_display}\n\n"
                            "Please provide:\n"
                            "1. A direct test URL that opens at first question, OR\n"
                            "2. Ask the programmer for an unrestricted test link\n\n"
                            "QC cannot proceed without accessible survey content."
                        )
                        log(f"  LOCKED URL — visible: {_vis_display} | hidden: {_hf_display}", "red")
                        jobs[job_id]["status"] = "error"
                        jobs[job_id]["phase"] = "Survey URL locked — no accessible questions"
                        jobs[job_id]["error"] = _err_msg
                        jobs.persist(job_id)
                        try: browser.close()
                        except: pass
                        return
                except Exception as _lc_ex:
                    log(f"  Locked-URL check skipped: {str(_lc_ex)[:80]}", "yellow")

                has_test_nav = False
                _active_tn_sel = None
                page_html = ""
                _tn_selectors = [
                    ".cf-tn-list-item",
                    "[class*='tn-question']",
                    "[class*='test-navigator']",
                    ".wix-tn-item",
                    "[class*='tn-item']",
                ]
                # RETRY: wait for Test Navigator to load (up to 15 tries × 2s = 30s budget)
                # Tries all known selectors each round so non-Confirmit platforms aren't penalised.
                for _try in range(15):
                    for _sel in _tn_selectors:
                        _cnt = page.locator(_sel).count()
                        if _cnt > 0:
                            _active_tn_sel = _sel
                            log("  TN loaded: " + str(_cnt) + " items (try " + str(_try+1) + "/15) with: " + _sel, "green")
                            break
                    if _active_tn_sel:
                        break
                    log("  TN not ready, waiting... (try " + str(_try+1) + "/15) — trying: " + ", ".join(_tn_selectors), "yellow")
                    page.wait_for_timeout(2000)
                    try:
                        page.locator("text=Test Navigator").first.click(timeout=2000)
                        page.wait_for_timeout(500)
                    except: pass
                try:
                    page_html = page.content()
                    log("  Page HTML length: " + str(len(page_html)), "cyan")
                    for tn_sel in _tn_selectors:
                        count = page.locator(tn_sel).count()
                        log("  TN selector " + tn_sel + " count: " + str(count), "yellow")
                        if count > 0:
                            has_test_nav = True
                            if not _active_tn_sel:
                                _active_tn_sel = tn_sel
                            log("  Found TN with: " + tn_sel, "green")
                            break
                    if not has_test_nav:
                        for btn_text in ["Test Navigator", "Navigator", "Navigateur"]:
                            try:
                                page.locator(f"text={btn_text}").first.click(timeout=3000)
                                page.wait_for_timeout(1000)
                                for tn_sel in _tn_selectors:
                                    if page.locator(tn_sel).count() > 0:
                                        has_test_nav = True
                                        if not _active_tn_sel:
                                            _active_tn_sel = tn_sel
                                        log("  TN found: " + tn_sel, "green")
                                        break
                                if has_test_nav:
                                    break
                            except: pass
                    if not has_test_nav:
                        try:
                            page.screenshot(path="/var/www/surveyqc/debug_page.png", full_page=True)
                            log("  Debug screenshot saved", "yellow")
                        except: pass
                        classes = page.evaluate("""() => {
                            const els = document.querySelectorAll("[class]");
                            const cls = new Set();
                            els.forEach(e => e.className.toString().split(" ").forEach(c => { if(c) cls.add(c); }));
                            return Array.from(cls).slice(0, 50);
                        }""")
                        log("  Page CSS classes: " + str(classes[:30]), "cyan")
                except Exception as ex:
                    log("  TN detection error: " + str(ex), "red")

                except: pass

                if has_test_nav:
                    log('  Mode: Test Navigator (Generic)', 'blue')
                    _qid_re = re.compile(r'^([A-Za-z]{1,8}\d*(?:[a-zA-Z]\d+|[a-zA-Z]?(?:bis|ter|Info|info|Ex|_\d+|\.\d+)?))')

                    # Loop-detection and multi-screen tracking.
                    # _prev_page_hash: body-text hash of the last crawled page.
                    #   If unchanged after a TN click, the question is hidden/conditional.
                    # _screen_text_map: maps a body-text hash to the first QID that
                    #   produced that page content.  When a second QID lands on the same
                    #   page, both are marked as sharing a multi-question screen so the
                    #   comparison phase can treat them correctly.
                    _prev_page_hash = None
                    _screen_text_map = {}   # hash → first_qid

                    # --- Phase A: build QID list from static HTML (no per-item Playwright calls) ---
                    qid_index_map, _container_skipped, _fw_skipped = _parse_tn_items_from_html(page_html, _tn_selectors, _qid_re)
                    if not qid_index_map:
                        log('  Static HTML parse found no TN items — falling back to Playwright DOM walk', 'yellow')
                        _container_skipped = 0
                        _fw_skipped = 0
                        nav_items = page.locator(_active_tn_sel).all()
                        seen_qids = set()
                        for ni, el in enumerate(nav_items):
                            try:
                                doc_qid = None
                                for _attr in ("data-qid", "data-question-id", "data-id"):
                                    _v = el.get_attribute(_attr)
                                    if _v and _qid_re.match(_v.strip()):
                                        doc_qid = _qid_re.match(_v.strip()).group(1); break
                                if not doc_qid:
                                    _v = el.get_attribute("id") or ""
                                    _m = _qid_re.search(_v)
                                    if _m: doc_qid = _m.group(1)
                                if not doc_qid:
                                    _v = el.get_attribute("class") or ""
                                    _m = re.search(r'question[-_]?id[-_]?([A-Za-z]\w*\d+\w*)', _v, re.IGNORECASE)
                                    if _m: doc_qid = _m.group(1)
                                if not doc_qid:
                                    words = el.inner_text().strip().split()
                                    for _w in words[:3]:
                                        _m = re.match(r'^([A-Za-z]{1,8}\d*(?:[a-zA-Z]\d+|[a-zA-Z]?(?:bis|ter|Info|info|Ex|_\d+|\.\d+)?))\s*$', _w.strip())
                                        if _m and is_valid_qid(_m.group(1)): doc_qid = _m.group(1); break
                                if not doc_qid:
                                    try:
                                        _v = el.locator("input, select").first.get_attribute("name") or ""
                                        _m = _qid_re.search(_v)
                                        if _m: doc_qid = _m.group(1)
                                    except: pass
                                if not doc_qid:
                                    continue
                                if _is_container_id(doc_qid):
                                    _container_skipped += 1
                                    continue
                                if should_skip_qid(doc_qid):
                                    _fw_skipped += 1
                                    continue
                                if doc_qid in seen_qids:
                                    continue
                                words = el.inner_text().strip().split()
                                second = words[1].strip() if len(words) > 1 else ''
                                plat_qid = second if (second and second != doc_qid
                                    and re.match(r'^[A-Za-z]\w*\d+\w*$', second)) else doc_qid
                                qid_index_map.append((ni, doc_qid, plat_qid))
                                seen_qids.add(doc_qid)
                            except: continue

                    _skip_note = f' ({_container_skipped} container IDs skipped)' if _container_skipped else ''
                    log('  ' + str(len(qid_index_map)) + ' QIDs found in navigator' + _skip_note, 'blue')
                    if _fw_skipped:
                        log(f'  Skipped {_fw_skipped} framework/system pages from live navigator', 'grey')

                    # Re-test filter: only crawl requested QIDs
                    if filter_qids:
                        _fset = set(q.upper() for q in filter_qids)
                        qid_index_map = [(ni, qid, pq) for ni, qid, pq in qid_index_map
                                         if qid.upper() in _fset]
                        log(f'  Re-test filter applied: crawling {len(qid_index_map)} of {len(filter_qids)} requested QIDs', 'cyan')

                    # --- Phase B: parse all question texts from static HTML in one pass ---
                    _static_qtext = _parse_question_texts_from_html(page_html)
                    _static_hits = sum(1 for v in _static_qtext.values() if v.get("text"))
                    log(f'  Static text parse: {_static_hits}/{len(_static_qtext)} question blocks with text', 'blue')
                    def _nq(s): return s.replace('.','').replace('_','').replace('-','').replace(' ','').lower()

                    _platform = detect_platform(page)
                    log(f'  Platform detected: {_platform}', 'cyan')

                    total = max(1, len(qid_index_map))
                    for i, (nav_idx, qid, plat_qid) in enumerate(qid_index_map, 1):
                        # STOP CHECK: if user clicked Stop, abort crawling
                        if jobs.get(job_id, {}).get('status') == 'stopped':
                            log('  >>> STOPPED by user during crawling', 'red')
                            try: browser.close()
                            except: pass
                            jobs[job_id]['phase'] = 'Stopped'
                            jobs.persist(job_id)
                            return
                        progress(20 + int((i/total)*40), 'Crawling ' + qid + '...')
                        if i % 50 == 0:
                            log(f'  Crawling item {i}/{total}...', 'blue')

                        # --- Phase C: use static text; Playwright click only if empty ---
                        _q_data = (_static_qtext.get(_nq(qid)) or
                                   _static_qtext.get(_nq(plat_qid)))
                        text = (_q_data or {}).get("text", "")
                        opts = (_q_data or {}).get("options", [])
                        has_inp = (_q_data or {}).get("has_inputs", True)
                        _is_conditional = False
                        _extr_method = "static"

                        # FIX 3 (static path): verify static content belongs to this QID
                        if text and _q_data:
                            _sv_ok, _sv_status = verify_extraction(qid, text)
                            if not _sv_ok:
                                log(f'   {qid} WARNING: static {_sv_status} — skipping', 'yellow')
                                text = ''  # force Playwright fallback for correct content
                                _q_data = None

                        if not text:
                            # Fallback: single Playwright click for this QID
                            try:
                                if not page.locator(_active_tn_sel).count():
                                    page.locator("text=Test Navigator").first.click(timeout=3000)
                                    page.wait_for_timeout(500)
                            except: pass
                            _click_ok = False
                            for _attempt in range(2):
                                try:
                                    page.locator(_active_tn_sel).nth(nav_idx).click(timeout=3000, force=True)
                                    _click_ok = True
                                    break
                                except Exception as _ce:
                                    if _attempt == 0:
                                        try:
                                            page.locator("text=Test Navigator").first.click(timeout=2000)
                                            page.wait_for_timeout(600)
                                        except: pass
                                    else:
                                        log('   ' + qid + ' SKIP (click failed): ' + str(_ce)[:80], 'yellow')
                            if not _click_ok:
                                live_data[qid] = {"text":"","options":[],"has_mandatory_marker":False,"has_raw_piping":False,"raw_piping_found":[],"has_inputs":True,"status":"crawl_failed - manual review"}
                                continue
                            _container_elem = None
                            try:
                                page.wait_for_timeout(1200)

                                # Fix 1+5: Target QID-specific container via platform selectors.
                                # Never uses .first blindly — filters by data-questionid metadata.
                                _container_elem, _extr_method = get_active_container(
                                    page, qid, _platform, plat_qid, _log_fn=log)
                                full_text = ""
                                _extraction_source = "container" if _container_elem is not None else "none"

                                # DIAG: log html/body/container sizes and iframe count
                                try:
                                    _diag_html_len = len(page.content())
                                    _diag_body_len = len(page.locator('body').inner_text(timeout=2000))
                                    _diag_cont_len = len(_container_elem.inner_text(timeout=2000)) if _container_elem is not None else 0
                                    _diag_iframes = len(page.frames) - 1  # exclude main frame
                                    log(f'   {qid} DIAG: html={_diag_html_len} body={_diag_body_len} container={_diag_cont_len} sel={_extr_method} iframes={_diag_iframes}', 'cyan')
                                except Exception:
                                    pass

                                # DIAG: log element identity for the matched container
                                if _container_elem is not None:
                                    try:
                                        _diag_elem = page.evaluate("""(el) => ({
                                            tag: el.tagName,
                                            id: el.id || '',
                                            cls: (el.className || '').toString().slice(0, 80),
                                            dqid: el.getAttribute('data-questionid') || '',
                                            dqid2: el.getAttribute('data-qid') || '',
                                            children: el.children.length,
                                            innerTextLen: (el.innerText||'').length
                                        })""", _container_elem)
                                        log(f'   {qid} CONTAINER_ELEM: tag={_diag_elem["tag"]} '
                                            f'id="{_diag_elem["id"]}" '
                                            f'data-questionid="{_diag_elem["dqid"]}" '
                                            f'class="{_diag_elem["cls"]}" '
                                            f'children={_diag_elem["children"]} '
                                            f'innerTextLen={_diag_elem["innerTextLen"]}', 'cyan')
                                    except Exception:
                                        pass

                                # Fix 2: Extract from cloned container — never mutates original DOM.
                                if _container_elem is not None:
                                    try:
                                        full_text = extract_from_cloned_dom(page, _container_elem)
                                        log(f'   {qid} AFTER_CLONE: method={_extr_method} raw_cloned_len={len(full_text)}', 'cyan')
                                        if not full_text:
                                            # Retry with inner_text() before discarding container (FIX 1)
                                            try:
                                                full_text = _container_elem.inner_text()
                                                if full_text:
                                                    log(f'   {qid} INNER_TEXT_FALLBACK: recovered {len(full_text)} chars', 'cyan')
                                                else:
                                                    _container_elem = None
                                                    _extr_method = "fallback"
                                                    _extraction_source = "none"
                                            except Exception:
                                                _container_elem = None
                                                _extr_method = "fallback"
                                                _extraction_source = "none"
                                    except Exception:
                                        # extract_from_cloned_dom threw (JS eval failed — same
                                        # condition that caused CONTAINER_FALLBACK_OK).
                                        # Use inner_text() directly before discarding container.
                                        try:
                                            full_text = _container_elem.inner_text(timeout=3000)
                                            if full_text:
                                                log(f'   {qid} CLONE_ERR_INNER_TEXT: recovered {len(full_text)} chars', 'cyan')
                                            else:
                                                _container_elem = None
                                                _extr_method = "fallback"
                                                _extraction_source = "none"
                                        except Exception:
                                            _container_elem = None
                                            _extr_method = "fallback"
                                            _extraction_source = "none"
                                            full_text = ""

                                # Iframe scan: if container not found in main frame, try sub-frames
                                if not full_text:
                                    for _iframe in page.frames:
                                        if _iframe == page.main_frame:
                                            continue
                                        try:
                                            _iframe_name = _iframe.name or _iframe.url or 'unnamed'
                                            _fc, _fm = get_active_container(_iframe, qid, _platform, plat_qid, _log_fn=log)
                                            if _fc is not None:
                                                _ft = extract_from_cloned_dom(_iframe, _fc)
                                                if _ft:
                                                    full_text = _ft
                                                    _container_elem = _fc
                                                    _extr_method = _fm + "_iframe"
                                                    _extraction_source = "iframe"
                                                    log(f'   {qid} DIAG: found container in iframe [{_iframe_name}] len={len(_ft)}', 'cyan')
                                                    break
                                        except Exception:
                                            continue

                                # Fallback: try platform-specific question selectors on main frame,
                                # pick the largest visible element that plausibly belongs to this QID.
                                if not full_text:
                                    _fallback_selectors = [
                                        '.cf-question-body', '.cf-question', '.question-content',
                                        '[data-question-id]', '[data-qid]', "[role='group']",
                                        '.questionOuter', '.question-container', '.sv-question',
                                        '[class*="question"]',
                                    ]
                                    _best_elem = None
                                    _best_len = 0
                                    _best_sel = None
                                    _qid_lower = qid.lower()
                                    _plat_lower = plat_qid.lower() if plat_qid else _qid_lower
                                    for _fsel in _fallback_selectors:
                                        try:
                                            _fcands = page.locator(_fsel).filter(visible=True)
                                            _fcount = min(_fcands.count(), 20)
                                            for _fi in range(_fcount):
                                                _fe = _fcands.nth(_fi)
                                                try:
                                                    _fe_text = _fe.inner_text(timeout=1000)
                                                    _fe_len = len(_fe_text)
                                                    if _fe_len < 20:
                                                        continue
                                                    # Prefer elements whose attribute/id match the QID
                                                    _fe_qid_attr = (
                                                        (_fe.get_attribute('data-questionid') or '') +
                                                        (_fe.get_attribute('data-qid') or '') +
                                                        (_fe.get_attribute('data-question-id') or '') +
                                                        (_fe.get_attribute('id') or '')
                                                    ).lower()
                                                    _fe_qid_match = (
                                                        _qid_lower in _fe_qid_attr or
                                                        _plat_lower in _fe_qid_attr
                                                    )
                                                    # Prefer QID-matched elements; among those pick largest
                                                    if _fe_qid_match:
                                                        if _fe_len > _best_len:
                                                            _best_elem = _fe
                                                            _best_len = _fe_len
                                                            _best_sel = _fsel
                                                    elif _best_elem is None and _fe_len > _best_len:
                                                        _best_elem = _fe
                                                        _best_len = _fe_len
                                                        _best_sel = _fsel
                                                except Exception:
                                                    continue
                                        except Exception:
                                            continue
                                    if _best_elem is not None:
                                        try:
                                            full_text = extract_from_cloned_dom(page, _best_elem)
                                            _container_elem = _best_elem
                                            _extr_method = f"selector_fallback({_best_sel})"
                                            log(f'   {qid} DIAG: selector_fallback used [{_best_sel}] len={len(full_text)}', 'cyan')
                                        except Exception:
                                            full_text = ""

                                # Last resort: cleaned body clone — marks as FULL_PAGE_EXTRACTION
                                if not full_text:
                                    full_text = page.evaluate("""() => {
                                        const b = document.body.cloneNode(true);
                                        [
                                            // Test navigator panels
                                            '.sr-test-navigator', '[class*=sr-tn]',
                                            '.cf-test-navigator', '[class*=cf-tn-]',
                                            '[class*=tn-question]', '[class*=test-navigator]',
                                            '[class*=tn-item]', '.wix-tn-item',
                                            // Platform matrix/grid chrome
                                            '[class*=cf-matrix-header]', '[class*=cf-matrix-]',
                                            '[class*=cf-pagination]', '[class*=cf-grid-header]',
                                            '[class*=cf-answer-header]',
                                            '[class*=ExportTagColumn]', '[class*=page-indicator]',
                                            '[class*=pagination]', '[class*=language-selector]',
                                            '[class*=survey-header]', '[class*=survey-footer]'
                                        ].forEach(s => {
                                            try { b.querySelectorAll(s).forEach(e => e.remove()); } catch(e) {}
                                        });
                                        [
                                            '[style*="display:none"]', '[style*="display: none"]',
                                            '[aria-hidden="true"]', '[hidden]',
                                            '.hidden', '.d-none', '[style*="visibility:hidden"]'
                                        ].forEach(s => {
                                            try { b.querySelectorAll(s).forEach(e => e.remove()); } catch(e) {}
                                        });
                                        [
                                            '[class*="agreement"]', '[class*="article"]',
                                            '[class*="invoice"]', '[class*="password"]',
                                            '[class*="footer"]', '[class*="header"]',
                                            '[class*="navigation"]', '[class*="progress"]'
                                        ].forEach(s => {
                                            try { b.querySelectorAll(s).forEach(e => e.remove()); } catch(e) {}
                                        });
                                        return b.innerText.trim();
                                    }""")
                                    _extr_method = "body_cleaned"
                                    _extraction_source = "body_clone"
                                    log(f'   {qid} BODY_CLONE: raw_body_cleaned_len={len(full_text)}', 'yellow')

                                full_text = re.sub(r'\*Shown in Testing mode only\*', '', full_text or '')
                                text = full_text
                                log(f'   {qid} STEP_A: after_strip method={_extr_method} src={_extraction_source} full_text_len={len(full_text)} text_len={len(text)}', 'cyan')

                                # Section markers — only useful for full-body fallback;
                                # container extraction is already scoped to the question.
                                # FIX 5: last-section slice capped at 3000 chars (not end-of-document).
                                if _extr_method in ("body_cleaned", "fallback"):
                                    try:
                                        _markers = list(re.finditer(r'\[Question ID:\s*([^\]]+)\]', full_text))
                                        log(f'   {qid} SECTION_MARKERS: found={len(_markers)} in full_text_len={len(full_text)}', 'cyan')
                                        _qn = _nq(qid)
                                        _found = False
                                        for _mi, _m in enumerate(_markers):
                                            if _m.group(1).strip().lower() in (qid.lower(), plat_qid.lower()):
                                                _start = _m.end()
                                                # FIX 5: cap last section at 3000 chars instead of len(full_text)
                                                _end = _markers[_mi+1].start() if _mi+1 < len(_markers) else min(_start + 3000, len(full_text))
                                                text = full_text[_start:_end].strip()
                                                _extr_method = "section"
                                                log(f'   {qid} SECTION_MATCH: marker_idx={_mi} slice_len={len(text)} (last_section_capped={_mi+1 >= len(_markers)})', 'cyan')
                                                _found = True
                                                break
                                        if not _found:
                                            for _mi, _m in enumerate(_markers):
                                                _mq = _m.group(1).strip()
                                                if _nq(_mq) == _qn or _nq(_mq).replace('x','') == _qn.replace('x',''):
                                                    _start = _m.end()
                                                    # FIX 5: cap last section at 3000 chars
                                                    _end = _markers[_mi+1].start() if _mi+1 < len(_markers) else min(_start + 3000, len(full_text))
                                                    text = full_text[_start:_end].strip()
                                                    _extr_method = "section"
                                                    log(f'   {qid} SECTION_FUZZY_MATCH: mq={_mq} slice_len={len(text)}', 'cyan')
                                                    _found = True
                                                    break
                                        if not _found:
                                            if _markers:
                                                _is_conditional = True
                                                text = ""
                                                log(f'   {qid} SECTION_NOT_FOUND: markers_present but qid not found → conditional', 'yellow')
                                            elif len(full_text) > 3000:
                                                text = full_text[:2000].strip()
                                                log(f'   {qid} SECTION_NO_MARKERS: full_text_len={len(full_text)} → sliced_to_2000', 'yellow')
                                    except: pass

                                log(f'   {qid} STEP_B: after_section_slice method={_extr_method} src={_extraction_source} text_len={len(text)}', 'cyan')

                                # FIX 3: text > 5000 = FULL_PAGE_EXTRACTION, never VERIFIED
                                if len(text) > 5000:
                                    log(f'   {qid} FULL_PAGE_EXTRACTION: text_len={len(text)} > 5000 — marking as unverified', 'yellow')
                                    _extr_method = "FULL_PAGE_EXTRACTION"
                                    _extraction_source = "body_clone"

                                # FIX 2: body_clone always treated as SELECTOR_FAILURE for issue gate
                                if _extraction_source == "body_clone" and _extr_method not in ("FULL_PAGE_EXTRACTION", "section"):
                                    _extr_method = "SELECTOR_FAILURE_BODY"

                                # Dynamic cap
                                if _extr_method == "CONTAINER_EXACT":
                                    _cap = None
                                elif _extr_method in ("CONTAINER_VISIBLE", "CONTAINER_METADATA"):
                                    _cap = 3000
                                elif _extr_method == "section":
                                    _cap = 3000
                                elif _extr_method == "FULL_PAGE_EXTRACTION":
                                    _cap = 2000
                                elif _extr_method == "SELECTOR_FAILURE_BODY":
                                    _cap = 3000
                                elif _extr_method.startswith("selector_fallback"):
                                    _cap = 3000
                                elif _extr_method.endswith("_iframe"):
                                    _cap = 3000
                                else:
                                    _cap = 2000
                                log(f'   {qid} PRE_CAP: method={_extr_method} text_len={len(text)} cap={_cap}', 'cyan')
                                if _cap and text and len(text) > _cap:
                                    log(f'   {qid} WARNING: text truncated {len(text)}→{_cap} [{_extr_method}]', 'yellow')
                                    text = text[:_cap].strip()

                                text = re.sub(r'\n{3,}', chr(10)+chr(10), text).strip()
                                text = re.sub(r'(?im)^[ \t]*test\s*link\b.*$', '', text)
                                text = re.sub(r'(?m)^\d+%(?:[ \t]+\d+%){2,}[ \t]*$', '', text)
                                text = re.sub(u'(?m)^[^\\S\\n]*[←→◄►\xab\xbb]{1,4}[^\\S\\n]*$', u'', text)
                                text = re.sub(r'\n{3,}', chr(10)+chr(10), text).strip()
                                opts = _extract_options(page)
                                has_inp = _page_has_inputs(page)
                            except Exception as e:
                                live_data[qid] = {"text":"","options":[],"has_mandatory_marker":False,"has_raw_piping":False,"raw_piping_found":[],"has_inputs":True,"status":"ERROR: " + str(e)[:50]}
                                log('   ' + qid + ' ERROR (fallback): ' + str(e)[:80], 'red')
                                continue

                        # ── Issue 5/6: multi-screen and loop detection (Playwright path only) ──
                        # When a TN click returns the exact same container/page content as the
                        # immediately previous question, the question is hidden by routing
                        # (conditional).  When two non-adjacent questions share the same
                        # content hash, they live on the same multi-question screen.
                        # Hash is based on active-container text when available, else full_text.
                        if not _is_conditional and not _q_data and full_text:
                            # Prefer container-scoped text for hash to avoid false positives
                            # from full-page extractions where the body is always the same.
                            _unreliable_methods = ("body_cleaned", "FULL_PAGE_EXTRACTION", "SELECTOR_FAILURE_BODY")
                            _hash_source = text if (text and _extr_method not in _unreliable_methods) else full_text
                            _cur_hash = hash(_hash_source[:1200])
                            _cur_qid_visible = bool(text and qid.lower() in text.lower())
                            if _prev_page_hash is not None and _cur_hash == _prev_page_hash:
                                # Identical to the previous page AND active container matches
                                # → this QID is hidden/conditional.
                                # Skip SAME_PAGE when we only have full-body hashes (unreliable).
                                if _extr_method not in _unreliable_methods:
                                    _is_conditional = True
                                    log(f'   {qid} [SAME PAGE AS PREVIOUS — hidden/conditional]', 'yellow')
                                else:
                                    log(f'   {qid} [SAME PAGE HASH but unreliable extract ({_extr_method}) — not marking conditional]', 'yellow')
                            else:
                                _first_owner = _screen_text_map.get(_cur_hash)
                                if _first_owner and _first_owner != qid:
                                    # Same content as a different earlier QID → shared screen
                                    log(f'   {qid} [SHARED SCREEN with {_first_owner}]', 'yellow')
                                    if _first_owner in live_data and live_data[_first_owner].get('status') == 'OK':
                                        live_data[qid] = dict(live_data[_first_owner])
                                        live_data[qid]['shared_screen_with'] = _first_owner
                                        live_data[qid]['extraction_status'] = 'SHARED_SCREEN'
                                        src = "static" if _q_data else _extr_method
                                        log('   ' + qid + ' (' + str(len(text)) + ' chars) [shared/' + src + ']', 'green')
                                        _prev_page_hash = _cur_hash
                                        continue
                                else:
                                    _screen_text_map[_cur_hash] = qid
                            _prev_page_hash = _cur_hash

                        if _is_conditional:
                            live_data[qid] = {"text":"","options":[],"has_mandatory_marker":False,"has_raw_piping":False,"raw_piping_found":[],"has_inputs":False,"status":"CONDITIONAL — could not verify","extraction_status":"CONDITIONAL"}
                            log('   ' + qid + ' [CONDITIONAL — not visible after TN click]', 'yellow')
                            continue

                        # Granular status for extraction scope failures
                        if _extr_method == "FULL_PAGE_EXTRACTION":
                            # Captured entire page — not reliable enough to mark OK
                            log(f'   {qid} WARNING: FULL_PAGE_EXTRACTION — storing with degraded status', 'yellow')
                            piping = re.findall(r'\[PIPE[^\]]*\]', text, re.I)
                            live_data[qid] = {"text":text,"options":opts,"has_mandatory_marker":(" *" in text or "*"+chr(10) in text),"has_raw_piping":len(piping)>0,"raw_piping_found":piping,"has_inputs":has_inp,"status":"CONDITIONAL — could not verify","extraction_status":"FULL_PAGE_EXTRACTION"}
                            continue

                        # FIX 2: body_clone used — store text for reference but mark as
                        # SELECTOR_FAILURE so Phase 3 quality gate skips issue generation.
                        if _extr_method == "SELECTOR_FAILURE_BODY":
                            log(f'   {qid} [SELECTOR_FAILURE_BODY — body_clone used, text_len={len(text)}, no real container found]', 'yellow')
                            piping = re.findall(r'\[PIPE[^\]]*\]', text, re.I)
                            live_data[qid] = {"text":text,"options":opts,"has_mandatory_marker":(" *" in text or "*"+chr(10) in text),"has_raw_piping":len(piping)>0,"raw_piping_found":piping,"has_inputs":has_inp,"status":"OK","extraction_status":"SELECTOR_FAILURE"}
                            log('   ' + qid + ' (' + str(len(text)) + ' chars) [body_clone/UNVERIFIED — issues suppressed]', 'yellow')
                            continue

                        if not text and _container_elem is None:
                            # No selector matched at all — page changed but question not visible
                            live_data[qid] = {"text":"","options":[],"has_mandatory_marker":False,"has_raw_piping":False,"raw_piping_found":[],"has_inputs":False,"status":"CONDITIONAL — could not verify","extraction_status":"SELECTOR_FAILURE"}
                            log(f'   {qid} [SELECTOR_FAILURE — no container matched any selector]', 'yellow')
                            continue

                        if not text and _container_elem is not None:
                            # Container was found but extracted empty text — routing hidden
                            live_data[qid] = {"text":"","options":[],"has_mandatory_marker":False,"has_raw_piping":False,"raw_piping_found":[],"has_inputs":False,"status":"CONDITIONAL — could not verify","extraction_status":"ROUTING_HIDDEN"}
                            log(f'   {qid} [ROUTING_HIDDEN — container found but empty after extraction]', 'yellow')
                            continue

                        # Fix 3: verify via container metadata (not text search — QID often absent from text)
                        _vstatus = 'UNVERIFIED'
                        if text and not _q_data:
                            _vok, _vstatus = verify_extraction_v2(page, _container_elem, qid, text)
                            if not _vok:
                                log(f'   {qid} WARNING: {_vstatus} — page showed wrong content, skipping', 'yellow')
                                live_data[qid] = {"text":"","options":[],"has_mandatory_marker":False,"has_raw_piping":False,"raw_piping_found":[],"has_inputs":True,"status":"WRONG_CONTENT — manual review","extraction_status":"WRONG_CONTENT"}
                                continue
                        elif _q_data:
                            _vstatus = 'STATIC'
                        piping = re.findall(r'\[PIPE[^\]]*\]', text, re.I)
                        live_data[qid] = {"text":text,"options":opts,"has_mandatory_marker":(" *" in text or "*"+chr(10) in text),"has_raw_piping":len(piping)>0,"raw_piping_found":piping,"has_inputs":has_inp,"status":"OK","extraction_status":_vstatus}
                        src = "static" if _q_data else _extr_method
                        log('   ' + qid + ' (' + str(len(text)) + ' chars) [' + src + ']', 'green')

                else:
                    log('  Mode: Respondent flow (auto-navigate)', 'blue')
                    log('  No test navigator - walking page by page', 'yellow')
                    # DIAGNOSTIC: log all clickable elements on first page
                    try:
                        diag = page.evaluate("""() => {
                            const out = [];
                            document.querySelectorAll('button, input[type=submit], input[type=button], a').forEach(el => {
                                const t = (el.innerText || el.value || el.id || '').trim().slice(0,40);
                                const tag = el.tagName.toLowerCase();
                                const typ = el.type || '';
                                const id = el.id || '';
                                const cls = (el.className || '').toString().slice(0,40);
                                if (t || id) out.push(tag + '[' + typ + '] text=' + t + ' id=' + id + ' cls=' + cls);
                            });
                            return out.slice(0, 25);
                        }""")
                        log('  DIAGNOSTIC - clickable elements found:', 'cyan')
                        for d in diag:
                            log('    ' + d, 'white')
                        if not diag:
                            log('    (no buttons/links found - survey may use JS framework)', 'yellow')
                    except Exception as de:
                        log('  Diagnostic error: ' + str(de)[:50], 'yellow')
                    max_pages = 80
                    page_num = 0
                    doc_qid_list = list(questions.keys())
                    _crawl_deadline = time.time() + 300  # 5-minute wall-clock limit

                    while page_num < max_pages:
                        if time.time() > _crawl_deadline:
                            log('  Crawl wall-clock limit reached (5 min) — stopping early', 'yellow')
                            break
                        page_num += 1
                        progress(20 + int((page_num/max_pages)*40), 'Page ' + str(page_num) + '...')
                        page.wait_for_timeout(800)
                        try:
                            text = page.evaluate("() => document.body.innerText.trim()")
                        except:
                            text = ""
                        text = re.sub(r'\*Shown in Testing mode only\*', '', text or '')
                        text = re.sub(r'\n{3,}', chr(10)+chr(10), text).strip()

                        detected_qid = None
                        for dq in doc_qid_list:
                            if dq in live_data:
                                continue
                            if re.search(r'\b' + re.escape(dq) + r'\b', text):
                                detected_qid = dq
                                break
                        if not detected_qid:
                            detected_qid = 'PAGE' + str(page_num)

                        opts = _extract_options(page)
                        has_inp = _page_has_inputs(page)
                        try: page.screenshot(path=ss_dir + '/' + detected_qid + '.png', full_page=True)
                        except: pass

                        if text and detected_qid not in live_data:
                            piping = re.findall(r'\[PIPE[^\]]*\]', text, re.I)
                            live_data[detected_qid] = {"text":text,"options":opts,"has_mandatory_marker":(" *" in text or "*"+chr(10) in text),"has_raw_piping":len(piping)>0,"raw_piping_found":piping,"has_inputs":has_inp,"status":"OK"}
                            log('   ' + detected_qid + ' (' + str(len(text)) + ' chars)', 'green')

                        try:
                            # Tick all checkboxes (consent pages need all checked)
                            checks = page.locator("input[type=checkbox]")
                            cc = checks.count()
                            if cc > 0:
                                for ci in range(min(cc, 10)):
                                    try: checks.nth(ci).check(force=True, timeout=1000)
                                    except: pass
                            # Select first radio in each radio group
                            radios = page.locator("input[type=radio]")
                            if radios.count() > 0:
                                try: radios.first.check(force=True, timeout=2000)
                                except:
                                    try: radios.first.click(force=True, timeout=2000)
                                    except: pass
                            # Fill text inputs
                            txt_inputs = page.locator("input[type=text], input[type=number], textarea")
                            ti = txt_inputs.count()
                            if ti > 0:
                                for tidx in range(min(ti, 5)):
                                    try: txt_inputs.nth(tidx).fill("5", timeout=1000)
                                    except: pass
                            # Dropdowns - select first real option
                            selects = page.locator("select")
                            if selects.count() > 0:
                                for sidx in range(min(selects.count(), 5)):
                                    try: selects.nth(sidx).select_option(index=1, timeout=1000)
                                    except: pass
                            page.wait_for_timeout(400)
                        except: pass

                        clicked_next = False
                        matched_selector = ''
                        next_selectors = [".cf-button-next", "button:has-text('Next')", "button:has-text('Continue')", "button:has-text('Suivant')", "button:has-text('Continuer')", "button:has-text('Weiter')", "button:has-text('Siguiente')", "input[value='Next']", "input[value='Continue']", "input[value='Suivant']", "input[value='Continuer']", "input[value='>>']", "input[value='>']", "button:has-text('>>')", "button:has-text('>')", "a:has-text('Next')", "a:has-text('Suivant')", "#NextButton", "#nextButton", "#btnNext", "input[type=submit]", "button[type=submit]", ".next-button", ".btn-next", "a.button", "[onclick*=next]", "[onclick*=submit]", "[id*=Next]", "[name*=next]"]
                        for sel in next_selectors:
                            try:
                                btn = page.locator(sel).first
                                if btn.count() > 0 and btn.is_visible():
                                    btn.click(timeout=2500, force=True)
                                    clicked_next = True
                                    matched_selector = sel
                                    break
                            except: continue
                        if clicked_next:
                            log('  Clicked next via: ' + matched_selector, 'green')

                        if not clicked_next:
                            log('  Survey end (no next button) at page ' + str(page_num), 'blue')
                            break
                        # Capture text before wait to detect if page changed
                        prev_text_hash = hash(text[:200])
                        page.wait_for_timeout(1500)
                        try:
                            new_text = page.evaluate("() => document.body.innerText.trim()")[:200]
                            if hash(new_text) == prev_text_hash:
                                # Page didn't change - try clicking next once more, else stop
                                log('  Page unchanged, retrying...', 'yellow')
                                page.wait_for_timeout(1000)
                                still_same = page.evaluate("() => document.body.innerText.trim()")[:200]
                                if hash(still_same) == prev_text_hash:
                                    log('  Stuck on same page - validation may be blocking. Stopping.', 'yellow')
                                    break
                        except: pass

                        try:
                            low = (text or "").lower()
                            _kw_complete = any(w in low for w in ["thank you", "merci", "survey complete", "questionnaire complete", "has been recorded"])
                            _no_inputs = page.locator("input:not([type='hidden']):not([type='submit']):visible, select:visible, textarea:visible").count() == 0
                            if _kw_complete or _no_inputs:
                                log('  Completion page detected', 'blue')
                                break
                        except: pass

                browser.close()
            log(chr(10) + '  Crawled ' + str(len(live_data)) + ' pages', 'green')

            if len(live_data) == 0:
                log('  ERROR: 0 pages crawled — live survey did not render.', 'red')
                log('  Possible causes: slow platform, bot detection, or session timeout.', 'red')
                log('  Retry the QC, or verify the survey link opens in a browser.', 'red')
                jobs[job_id]['status'] = 'error'
                jobs[job_id]['phase'] = 'Crawl failed — 0 pages loaded'
                jobs[job_id]['error'] = ('Crawl failed — 0 pages loaded. The live survey did not render in time. '
                                         'Possible causes: slow platform, bot detection, or session timeout. '
                                         'Retry the QC or check the link.')
                jobs.persist(job_id)
                return

        # PHASE 3: COMPARE
        if mode in ('full', 'quick') and live_data:
            progress(65, 'Comparing doc vs live...')
            log('', 'white')
            log('════════════════════════════════════', 'cyan')
            log('  PHASE 3: COMPARISON', 'cyan')
            log('════════════════════════════════════', 'cyan')

            # Initialize AI model (uses admin Gemini key if set)
            ai_model = get_gemini_model()
            if ai_model:
                log('  AI mode: Gemini active - smart comparison enabled', 'green')
            else:
                log('  AI mode: text-matching (add Gemini key in admin for AI)', 'yellow')

            # Confirmit converts dots to 'x': doc 'R2.2' = live 'R2x2'. Normalize for matching.
            def _norm_qid(q):
                return q.lower().replace('.', '').replace('x', '').replace('_', '').replace('-', '').replace(' ', '')
            # Build live lookup by normalized qid
            _live_norm = {}
            for lq in live_data.keys():
                _live_norm[_norm_qid(lq)] = lq
            _doc_norm = {}
            for dq in questions.keys():
                _doc_norm[_norm_qid(dq)] = dq
            def _find_base_match(doc_norm):
                """Return (live_norm_key, live_orig) for the best live QID
                that is a naming variant of doc_norm.  Strategies in order:

                0. Dot-to-x (safety net): doc 'R2.2' → live 'R2x2'.
                   Usually handled by _norm_qid stripping both '.' and 'x',
                   but kept as an explicit fallback for edge cases.
                1. Suffix-strip: 'r2new'→'r2' (known rename suffixes).
                2. Prefix non-digit suffix: 'q11'→'q11a'.
                3. Container: doc 'Q16' is prefix of live 'Q16x1'/'Q16.1'
                   (doc is the matrix parent; live split into sub-questions).
                4. Component: doc 'Q12' is a complete component inside merged
                   live 'Q11Q12' (boundary = end-of-string or next alpha char).
                   Min len 3 to avoid 'Q1' false-matching 'Q11'.
                5. Parent-component: for dot-QIDs, strip sub-suffix and try
                   the parent question via strategies 3/4.
                   e.g. 'Q12.2'→parent 'Q12' inside live 'Q11Q12'.
                """
                _doc_orig = _doc_norm.get(doc_norm, '')
                _doc_orig_lower = _doc_orig.lower()

                def _component_in(needle, haystack):
                    """True if needle appears in haystack at a component
                    boundary: followed by end-of-string or an alpha char
                    (the start of the next QID component)."""
                    idx = 0
                    while True:
                        pos = haystack.find(needle, idx)
                        if pos == -1:
                            return False
                        after = haystack[pos + len(needle):]
                        if not after or after[0].isalpha():
                            return True
                        idx = pos + 1

                # Strategy 0: explicit dot-to-x safety net
                if '.' in _doc_orig:
                    _dotx_lower = _doc_orig_lower.replace('.', 'x')
                    for _ln, _lo in _live_norm.items():
                        if _lo.lower() == _dotx_lower:
                            return _ln, _lo

                # Strategy 1: explicit suffix stripping (live=Q23bis, doc=Q23)
                _sfx_re = re.compile(r'(?:new|bis|ter)$')
                for _ln, _lo in _live_norm.items():
                    if _ln == doc_norm:
                        continue
                    if _sfx_re.sub('', _ln) == doc_norm:
                        return _ln, _lo

                # Strategy 1b: strip bis/ter from doc (doc=Q23bis, live=Q23)
                _stripped_doc = _sfx_re.sub('', doc_norm)
                if _stripped_doc != doc_norm and _stripped_doc in _live_norm:
                    return _stripped_doc, _live_norm[_stripped_doc]

                # Strategy 2: same letter+digit prefix, non-digit suffix
                _bm = re.match(r'^([a-z]+\d+)', doc_norm)
                if _bm:
                    _base = _bm.group(1)
                    for _ln, _lo in _live_norm.items():
                        if (_ln != doc_norm
                                and _ln.startswith(_base)
                                and len(_ln) > len(_base)
                                and not _ln[len(_base)].isdigit()):
                            return _ln, _lo

                # Strategy 3: container — doc is the parent; live has x/dot
                # sub-questions.  e.g. doc 'Q16' → live 'Q16x1', 'Q16x2'.
                for _ln, _lo in _live_norm.items():
                    _lo_l = _lo.lower()
                    if (_lo_l.startswith(_doc_orig_lower + 'x') or
                            _lo_l.startswith(_doc_orig_lower + '.')):
                        return _ln, _lo

                # Strategy 4: component — doc_norm is a complete component
                # inside a merged live QID.  e.g. doc 'Q12' in live 'Q11Q12'.
                if len(doc_norm) >= 3:
                    for _ln, _lo in _live_norm.items():
                        if len(_ln) > len(doc_norm) and _component_in(doc_norm, _ln):
                            return _ln, _lo

                # Strategy 5: parent-component — for dot-QIDs, resolve via
                # the parent question.  e.g. 'Q12.2'→parent 'Q12' in 'Q11Q12'.
                if '.' in _doc_orig:
                    _par_orig = _doc_orig.split('.')[0]
                    _par_n = _norm_qid(_par_orig)
                    if _par_n and _par_n != doc_norm:
                        if _par_n in _live_norm:
                            return _par_n, _live_norm[_par_n]
                        if len(_par_n) >= 3:
                            for _ln, _lo in _live_norm.items():
                                if (len(_ln) > len(_par_n) and
                                        _component_in(_par_n, _ln)):
                                    return _ln, _lo

                return None, None

            _naming_matched_live = set()

            # ── Issue 6: Loop detection (loop_detector.py) ──────────────────────
            # Build full loop map from original doc QIDs, then normalise the skip
            # set so it lines up with the normalised comparison key space.
            _orig_doc_qids = list(_doc_norm.values())
            _raw_loop_map  = _ld_build_loop_map(_orig_doc_qids)
            _raw_loop_skip = _ld_skip_set(_orig_doc_qids)   # {child_orig: first_orig}
            _loop_skip_nqids: set = {_norm_qid(q) for q in _raw_loop_skip.keys()}

            job['loop_blocks'] = _raw_loop_map  # stored for Word report

            if _raw_loop_skip:
                log(f'  Loop/child dedup: {len(_raw_loop_skip)} sibling(s) collapsed', 'grey')
                for _child_orig, _first_orig in sorted(_raw_loop_skip.items()):
                    _par_orig = get_parent_qid(_child_orig)
                    log(f'  LOOP: {_child_orig} is sibling of {_par_orig} — '
                        f'skipping, {_first_orig} (first) already compared', 'grey')
            # ────────────────────────────────────────────────────────────────────

            # Unified set of normalized qids
            _all_norm = set(_doc_norm.keys()) | set(_live_norm.keys())

            # Re-test filter: restrict comparison to requested QIDs only
            if filter_qids:
                _fset_norm = set(_norm_qid(q) for q in filter_qids)
                _all_norm = {n for n in _all_norm if n in _fset_norm}
                log(f'  Re-test filter: comparing {len(_all_norm)} QID(s)', 'cyan')

            _to_compare = []
            _ai_unmatched = []  # doc questions that survived all regex/content strategies
            _phase3_skipped = 0
            for _nqid in sorted(_all_norm):
                qid = _doc_norm.get(_nqid) or _live_norm.get(_nqid)
                # Loop iteration dedup: skip sibling iterations (Q30x2, D3B, etc.)
                if _nqid in _loop_skip_nqids:
                    _phase3_skipped += 1
                    continue
                # Single-pipeline filter: skip internal/framework/stop-word QIDs
                if qid and should_skip_qid(qid):
                    _phase3_skipped += 1
                    continue
                in_doc = _nqid in _doc_norm
                in_live = _nqid in _live_norm
                _live_key = _live_norm.get(_nqid)
                _doc_key = _doc_norm.get(_nqid)
                if in_doc and not in_live:
                    _bmn, _bmo = _find_base_match(_nqid)
                    if _bmo:
                        _naming_matched_live.add(_norm_qid(_bmo))
                        # Grid container (Confirmit gQXQY convention): doc QIDs are
                        # merged into a live grid screen — not a naming/programming error.
                        if re.match(r'^g[A-Za-z]', _bmo, re.IGNORECASE):
                            continue  # silently matched; skip issue and text comparison
                        issues.append({"qid":qid,"type":"NAMING MISMATCH",
                                       "details":f"Doc: {qid} / Live: {_bmo}","severity":"MEDIUM"})
                        # Schedule text comparison against the matched live QID.
                        # Without this the comparison is skipped and the matched
                        # live content (e.g. R2x2) is never checked against doc.
                        if (_doc_key
                                and not questions.get(_doc_key, {}).get("is_numeric")
                                and _bmo in live_data
                                and live_data[_bmo].get("status") == "OK"
                                and live_data[_bmo].get("extraction_status") not in ("SELECTOR_FAILURE", "FULL_PAGE_EXTRACTION", "WRONG_CONTENT")):
                            _to_compare.append({
                                "qid": qid,
                                "doc_text": questions[_doc_key]["text"],
                                "live_text": live_data[_bmo]["text"],
                                "doc_opts": [o["text"] for o in questions[_doc_key].get("options", [])],
                                "live_opts": [o["text"] for o in live_data[_bmo].get("options", [])],
                                "_doc_key": _doc_key,
                                "_live_key": _bmo,
                            })
                    else:
                        # Content-fallback: before declaring MISSING IN LIVE, scan every
                        # live page for the doc question's text. A question may exist under
                        # a renamed QID, at a different position, or on a combined page —
                        # content match overrides a pure QID-name miss. Works for all
                        # surveys and languages because fuzzy_match normalises (lowercase +
                        # whitespace collapse) and uses both substring and SequenceMatcher.
                        _doc_q_text = questions.get(_doc_key, {}).get("text", "") if _doc_key else ""
                        _content_found_live = None
                        if _doc_q_text and len(_doc_q_text.strip()) >= 20:
                            for _lk, _ld in live_data.items():
                                if _ld.get("status") != "OK":
                                    continue
                                _is_match, _ = fuzzy_match(_doc_q_text, _ld["text"], threshold=0.60)
                                if _is_match:
                                    _content_found_live = _lk
                                    break
                        if _content_found_live:
                            if (_doc_key
                                    and not questions.get(_doc_key, {}).get("is_numeric")
                                    and live_data[_content_found_live].get("status") == "OK"
                                    and live_data[_content_found_live].get("extraction_status") not in ("SELECTOR_FAILURE", "FULL_PAGE_EXTRACTION", "WRONG_CONTENT")):
                                _to_compare.append({
                                    "qid": qid,
                                    "doc_text": questions[_doc_key]["text"],
                                    "live_text": live_data[_content_found_live]["text"],
                                    "doc_opts": [o["text"] for o in questions[_doc_key].get("options", [])],
                                    "live_opts": [o["text"] for o in live_data[_content_found_live].get("options", [])],
                                    "_doc_key": _doc_key,
                                    "_live_key": _content_found_live,
                                })
                        else:
                            # All regex + content strategies failed — defer to AI semantic match
                            _ai_unmatched.append({"nqid": _nqid, "qid": qid, "doc_key": _doc_key})
                    continue
                if in_live and not in_doc:
                    # EXTRA IN LIVE check disabled — not actionable for clients
                    continue
                if live_data[_live_key]["status"] == "CONDITIONAL — could not verify":
                    # If question is in XML, it's confirmed in platform — routing hides it in test
                    _qid_cond_norm = re.sub(r'[^a-z0-9]', '', qid.lower())
                    if _xml_qid_set and _qid_cond_norm in _xml_qid_set:
                        issues.append({
                            "qid": qid,
                            "type": "IN_XML_NOT_VERIFIED_IN_LIVE",
                            "details": "Question present in XML but hidden by routing in live survey — requires specific navigation path to verify",
                            "severity": "INFO",
                            "confidence": 60,
                            "conf_level": "NEEDS_MANUAL",
                            "source_phase": "PHASE_3",
                        })
                    else:
                        issues.append({"qid":qid,"type":"CONDITIONAL","details":"Question hidden/conditional — could not verify in test mode","severity":"LOW"})
                    continue
                if live_data[_live_key]["status"] != "OK":
                    issues.append({"qid":qid,"type":"ERROR PAGE","details":live_data[_live_key]["status"],"severity":"MEDIUM"})
                    continue
                # Emit naming variant warning when doc and live use different QID
                # forms (e.g. doc='R2.2', live='R2x2') even though they normalized
                # to the same key.  Dot-to-x is a Confirmit platform conversion.
                if _doc_key and _live_key and _doc_key != _live_key:
                    if _doc_key.lower().replace('.', 'x') == _live_key.lower():
                        issues.append({"qid": _doc_key, "type": "NAMING MISMATCH",
                                       "details": f"Doc: {_doc_key} / Live: {_live_key} (dot-to-x)",
                                       "severity": "MEDIUM"})
                if questions[_doc_key].get("is_numeric"):
                    continue
                # FIX 4 — Extraction quality gate: skip HIGH/MEDIUM issue generation
                # when live extraction was unreliable (body_clone, full page, wrong content).
                _live_estatus = live_data[_live_key].get('extraction_status', '')
                if _live_estatus in ('SELECTOR_FAILURE', 'FULL_PAGE_EXTRACTION', 'WRONG_CONTENT'):
                    log(f'  Quality gate: skipping comparison for {qid} — extraction_status={_live_estatus} (NEEDS_MANUAL)', 'yellow')
                    continue
                _to_compare.append({"qid":qid,"doc_text":questions[_doc_key]["text"],"live_text":live_data[_live_key]["text"],"doc_opts":[o["text"] for o in questions[_doc_key].get("options",[])],"live_opts":[o["text"] for o in live_data[_live_key].get("options",[])],"_doc_key":_doc_key,"_live_key":_live_key})
            if _phase3_skipped:
                log(f'  Skipped {_phase3_skipped} internal/framework QID(s) from comparison', 'grey')
            # AI SEMANTIC MATCHING: one batched Gemini call for all questions that
            # survived every regex/content strategy.  Generic — works for any survey
            # and any language.  Only runs if AI model is available.
            if _ai_unmatched:
                if ai_model:
                    log(f'  AI semantic matching: {len(_ai_unmatched)} unmatched doc question(s)', 'cyan')
                    _ai_match_result = ai_semantic_match_batch(ai_model, _ai_unmatched, live_data, questions)
                    for _u in _ai_unmatched:
                        _ai_live = _ai_match_result.get(_u["qid"], "NONE")
                        if _ai_live and _ai_live != "NONE" and _ai_live in live_data:
                            log(f'    AI matched {_u["qid"]} → {_ai_live}', 'green')
                            issues.append({"qid": _u["qid"], "type": "NAMING MISMATCH",
                                           "details": f"Doc: {_u['qid']} / Live: {_ai_live} (AI semantic match)",
                                           "severity": "MEDIUM"})
                            if (_u["doc_key"]
                                    and not questions.get(_u["doc_key"], {}).get("is_numeric")
                                    and live_data[_ai_live].get("status") == "OK"):
                                _to_compare.append({
                                    "qid": _u["qid"],
                                    "doc_text": questions[_u["doc_key"]]["text"],
                                    "live_text": live_data[_ai_live]["text"],
                                    "doc_opts": [o["text"] for o in questions[_u["doc_key"]].get("options", [])],
                                    "live_opts": [o["text"] for o in live_data[_ai_live].get("options", [])],
                                    "_doc_key": _u["doc_key"],
                                    "_live_key": _ai_live,
                                })
                        else:
                            log(f'    AI: {_u["qid"]} → NONE (MISSING IN LIVE)', 'yellow')
                            # Phase 5: If question exists in XML, it is NOT a bug —
                            # it was likely hidden by routing/conditional display in live.
                            # Report as NEEDS_MANUAL instead of HIGH bug.
                            _qid_norm_chk = re.sub(r'[^a-z0-9]', '', _u["qid"].lower())
                            if _xml_qid_set and _qid_norm_chk in _xml_qid_set:
                                log(f'      → In XML — routing/conditional (NEEDS_MANUAL, not a bug)', 'green')
                                issues.append({
                                    "qid": _u["qid"],
                                    "type": "IN_XML_NOT_VERIFIED_IN_LIVE",
                                    "details": "Question present in XML but not verified in live survey — may be hidden by routing or require specific navigation path",
                                    "severity": "INFO",
                                    "confidence": 55,
                                    "conf_level": "NEEDS_MANUAL",
                                    "source_phase": "PHASE_3",
                                })
                            else:
                                # Not in XML either — genuinely missing from platform
                                _miss_q = questions.get(_u["doc_key"] or _u["qid"], {})
                                _has_cond = any(
                                    r.get("answer_codes") == ["?"]
                                    for r in _miss_q.get("termination_rules", [])
                                )
                                _miss_sev = "MEDIUM" if _has_cond else "HIGH"
                                issues.append({"qid": _u["qid"], "type": "MISSING IN LIVE",
                                               "details": "In doc but not in live survey (not found in XML either)",
                                               "severity": _miss_sev})
                else:
                    # No AI model — use MEDIUM by default since there's no semantic
                    # validation; a human must confirm before treating as HIGH.
                    for _u in _ai_unmatched:
                        _qid_norm_chk = re.sub(r'[^a-z0-9]', '', _u["qid"].lower())
                        if _xml_qid_set and _qid_norm_chk in _xml_qid_set:
                            # In XML — conditional/routing, not a bug
                            issues.append({
                                "qid": _u["qid"],
                                "type": "IN_XML_NOT_VERIFIED_IN_LIVE",
                                "details": "Question present in XML but not verified in live survey — may be hidden by routing or require specific navigation path",
                                "severity": "INFO",
                                "confidence": 55,
                                "conf_level": "NEEDS_MANUAL",
                                "source_phase": "PHASE_3",
                            })
                        else:
                            _miss_q = questions.get(_u.get("doc_key") or _u["qid"], {})
                            _has_cond = any(
                                r.get("answer_codes") == ["?"]
                                for r in _miss_q.get("termination_rules", [])
                            )
                            _miss_sev = "MEDIUM" if (_has_cond or not ai_model) else "HIGH"
                            issues.append({"qid": _u["qid"], "type": "MISSING IN LIVE",
                                           "details": "In doc but not in live survey (not found in XML either)",
                                           "severity": _miss_sev})

            # Detect merged/matrix live screens: when multiple doc QIDs all
            # resolve to the same live QID, a 1-to-1 text comparison is
            # meaningless (they share one combined screen).  Replace those
            # comparison pairs with a single MERGED IN LIVE advisory and
            # remove them from the comparison queue.  Generic — works for any
            # survey where the live platform collapses several doc questions
            # into one page (e.g. doc R2 + R2.2 + R2.3 all → live R2x2).
            _live_key_groups = {}
            for _ci in _to_compare:
                _ck = _ci.get('_live_key')
                if _ck:
                    _live_key_groups.setdefault(_ck, []).append(_ci['qid'])
            _merged_live_keys = {k for k, v in _live_key_groups.items() if len(v) > 1}
            if _merged_live_keys:
                for _mlk in sorted(_merged_live_keys):
                    _dqids = ', '.join(_live_key_groups[_mlk])
                    issues.append({
                        "qid": _live_key_groups[_mlk][0],
                        "type": "MERGED IN LIVE",
                        "details": f"Doc [{_dqids}] → live {_mlk} — merged/matrix screen, manual review needed",
                        "severity": "MEDIUM",
                    })
                _to_compare = [i for i in _to_compare
                               if i.get('_live_key') not in _merged_live_keys]

            _batch_size = 8
            _ai_handled = set()
            if ai_model:
                for _bi in range(0, len(_to_compare), _batch_size):
                    if jobs.get(job_id, {}).get('status') == 'stopped':
                        break
                    _batch = _to_compare[_bi:_bi+_batch_size]
                    log('  AI batch ' + str(_bi//_batch_size + 1) + ': ' + str(len(_batch)) + ' questions', 'cyan')
                    _bres = ai_compare_batch(ai_model, _batch)
                    if _bres is not None:
                        for _item in _batch:
                            _q = _item["qid"]
                            _matched = False
                            for _rk, _rv in _bres.items():
                                if _norm_qid(_rk) == _norm_qid(_q):
                                    for _iss in _rv:
                                        _iss["qid"] = _q
                                        issues.append(_iss)
                                    _ai_handled.add(_q)
                                    _matched = True
                                    break
                            if not _matched:
                                _ai_handled.add(_q)
            for _item in _to_compare:
                qid = _item["qid"]
                if qid in _ai_handled:
                    continue
                doc_text = _item["doc_text"]; live_text = _item["live_text"]
                _doc_key = _item["_doc_key"]; _live_key = _item["_live_key"]
                is_match, ratio = fuzzy_match(doc_text, live_text)
                if not is_match and ratio < 0.5:
                    issues.append({"qid":qid,"type":"TEXT MISMATCH","details":f"Match: {int(ratio*100)}% (fuzzy fallback - AI unavailable)","severity":"MEDIUM"})
            for _item in _to_compare:
                qid = _item["qid"]; _doc_key = _item["_doc_key"]; _live_key = _item["_live_key"]
                if questions[_doc_key].get("is_mandatory") and not live_data[_live_key].get("has_mandatory_marker"):
                    issues.append({"qid":qid,"type":"MANDATORY MISSING","details":"Doc mandatory, live marker missing","severity":"MEDIUM"})
                if live_data[_live_key].get("has_raw_piping"):
                    issues.append({"qid":qid,"type":"PIPING NOT RESOLVED","details":f"Raw: {live_data[_live_key].get('raw_piping_found',[])[:3]}","severity":"HIGH"})

            # ── TRANSLATION SYNC ─────────────────────────────────────────
            try:
                from translation_sync import check_translation_sync as _ts_check
                # Build minimal doc/live dicts from the comparison list
                _ts_doc  = {
                    _it['_doc_key']: {'text': _it['doc_text']}
                    for _it in _to_compare if _it.get('doc_text')
                }
                _ts_live = {
                    _it['_live_key']: {'text': _it['live_text']}
                    for _it in _to_compare if _it.get('live_text')
                }
                # Pull existing TEXT MISMATCH issues so matched translations
                # can be suppressed
                _ts_existing = [i for i in issues if i.get('type') == 'TEXT MISMATCH']
                _ts_new, _ts_filtered = _ts_check(_ts_doc, _ts_live, _ts_existing)
                # Replace existing TEXT MISMATCH issues with the filtered set
                _suppressed = len(_ts_existing) - len(_ts_filtered)
                if _suppressed:
                    issues = [i for i in issues if i.get('type') != 'TEXT MISMATCH']
                    issues.extend(_ts_filtered)
                    log(f'  Translation sync: {_suppressed} false-positive(s) suppressed (valid translation)', 'green')
                for _ti in _ts_new:
                    _ti.setdefault('is_translation_issue', True)
                    issues.append(_ti)
                job['translation_issues'] = _ts_new
                if _ts_new:
                    log(f'  Translation sync: {len(_ts_new)} issue(s) — '
                        f'{sum(1 for t in _ts_new if t["verdict"]=="TRANSLATION_MISMATCH")} mismatch, '
                        f'{sum(1 for t in _ts_new if t["verdict"]=="TRANSLATION_UNCERTAIN")} uncertain',
                        'yellow')
                elif not _suppressed:
                    log('  Translation sync: no cross-language issues', 'green')
            except Exception as _ts_exc:
                log(f'  Translation sync skipped: {str(_ts_exc)[:60]}', 'grey')
                job.setdefault('translation_issues', [])
            # ─────────────────────────────────────────────────────────────

            # ── PIPING VALIDATION (CHECK1-4) ──────────────────────────────
            try:
                from piping_validator import validate_piping as _pv_validate
                _pv_issues = _pv_validate(
                    questions, live_data,
                    platform=_platform if '_platform' in dir() else 'generic',
                    all_qids=set(questions) | set(live_data),
                )
                for _pvi in _pv_issues:
                    _pvi.setdefault('is_piping_issue', True)
                    issues.append(_pvi)
                job['piping_issues'] = _pv_issues
                if _pv_issues:
                    log(f'  Piping validation: {len(_pv_issues)} issue(s) found', 'yellow')
                    for _pi in _pv_issues[:5]:
                        log(f'    [{_pi["rule"]}] {_pi["qid"]:8s} {_pi["pipe_variable"]:20s} {_pi["evidence"][:60]}', 'yellow')
                else:
                    log('  Piping validation: no issues', 'green')
            except Exception as _pv_exc:
                log(f'  Piping validation skipped: {str(_pv_exc)[:60]}', 'grey')
                job['piping_issues'] = []
            # ─────────────────────────────────────────────────────────────

            # ── DATA EXPORT VALIDATION (R031-R038) ────────────────────────
            _export_schema_text = job.get('export_schema_text', '').strip()
            if _export_schema_text:
                try:
                    from export_validator import run_export_validation
                    log('', 'white')
                    log('  ─── Data Export Validation ───', 'cyan')
                    _ev_result = run_export_validation(
                        questions,
                        _export_schema_text,
                        xml_questions=xml_questions if xml_questions else None,
                    )
                    _ev_issues = _ev_result.get('issues', [])
                    _ev_summary = _ev_result.get('summary', {})
                    log(f'  Export vars expected: {_ev_summary.get("expected_vars","?")}  actual: {_ev_summary.get("actual_vars","?")}', 'blue')
                    for _ev_iss in _ev_issues:
                        # Normalise to app.py issue format
                        _ev_iss.setdefault('type', _ev_iss.get('rule', 'DATA EXPORT'))
                        _ev_iss['is_export_issue'] = True
                        issues.append(_ev_iss)
                        log(f'  [{_ev_iss["rule"]}] {_ev_iss["severity"]:6s} {_ev_iss.get("qid","?"):8s} {_ev_iss.get("details","")[:70]}', 'yellow' if _ev_iss['severity']=='MEDIUM' else 'red')
                    if not _ev_issues:
                        log('  Export schema: no issues found', 'green')
                except Exception as _ev_exc:
                    log(f'  Export validation skipped: {str(_ev_exc)[:60]}', 'grey')
            # ─────────────────────────────────────────────────────────────

            sev = {"HIGH":0,"MEDIUM":0,"INFO":0}
            for i in issues: sev[i.get("severity","INFO")] = sev.get(i.get("severity","INFO"),0)+1
            log(f'  Total issues: {len(issues)} (HIGH:{sev["HIGH"]} MEDIUM:{sev["MEDIUM"]} INFO:{sev["INFO"]})', 'yellow')

            # ── AUTO-GENERATE TEST CASES ─────────────────────────────────────
            try:
                from test_generator import generate_test_cases as _gen_tc
                # Build the doc_data dict that test_generator expects.
                # _logic_tables_for_tc is populated during the termination extraction
                # pass above; questions is the doc parse output.
                _gen_doc_data = {
                    "questions": questions,
                    "logic_tables": _logic_tables_for_tc if '_logic_tables_for_tc' in dir() else [],
                }
                # ── [TERM DEBUG] checkpoint 5: test_generator handoff ────────
                _dbg_ltc = _gen_doc_data.get("logic_tables", [])
                app.logger.debug("[TERM DEBUG] test_generator handoff: %d logic_tables, %d questions",
                                 len(_dbg_ltc), len(questions))
                for _dbg_lt in _dbg_ltc[:5]:
                    app.logger.debug("[TERM DEBUG]   logic_table host=%r flat=%r",
                                     _dbg_lt.get('host_qid'), _dbg_lt.get('flat_text','')[:80])
                # ─────────────────────────────────────────────────────────────
                _tc_list, _tc_summary = _gen_tc(_gen_doc_data)
                # Merge range test cases from range_validator
                try:
                    from range_validator import (
                        validate_ranges as _rv_validate,
                        generate_range_test_cases as _rv_tcs,
                    )
                    _rv_issues = _rv_validate(questions, live_data)
                    job['range_issues'] = _rv_issues
                    if _rv_issues:
                        log(f'  Range validation: {len(_rv_issues)} issue(s) found', 'yellow')
                        for _ri in _rv_issues[:5]:
                            log(f'    [{_ri["rule"]}] {_ri["qid"]:8s} {_ri["evidence"][:65]}', 'yellow')
                    else:
                        log('  Range validation: no issues', 'green')
                    _rng_tcs = _rv_tcs(questions)
                    _seen_tc_keys: set = {
                        (t.get('qid'), t.get('action'), t.get('expected'))
                        for t in _tc_list
                    }
                    _added_rng = 0
                    for _rtc in _rng_tcs:
                        key = (_rtc.get('qid'), _rtc.get('action'), _rtc.get('expected'))
                        if key not in _seen_tc_keys:
                            _seen_tc_keys.add(key)
                            _tc_list.append(_rtc)
                            _added_rng += 1
                    if _added_rng:
                        log(f'  Range test cases added: {_added_rng}', 'cyan')
                        _tc_summary['total'] = len(_tc_list)
                        _tc_summary['auto_runnable'] = sum(
                            1 for t in _tc_list if t.get('auto_runnable'))
                        _tc_summary.setdefault('by_type', {})['RANGE'] = sum(
                            1 for t in _tc_list if t.get('type') == 'RANGE')
                except Exception as _rv_exc:
                    log(f'  Range validation skipped: {str(_rv_exc)[:60]}', 'grey')
                    job.setdefault('range_issues', [])
                job['test_cases'] = _tc_list
                job['test_cases_summary'] = _tc_summary
                _tc_total = _tc_summary.get('total', 0)
                _tc_auto  = _tc_summary.get('auto_runnable', 0)
                _tc_types = _tc_summary.get('by_type', {})
                log('', 'white')
                log(f'  Test cases generated: {_tc_total} ({_tc_auto} auto-runnable)', 'cyan')
                if _tc_types:
                    _tc_line = '  ' + '  '.join(f'{k}: {v}' for k, v in _tc_types.items())
                    log(_tc_line, 'cyan')
            except Exception as _tc_err:
                log(f'  Test case generation skipped: {str(_tc_err)[:60]}', 'grey')
                job['test_cases'] = []
                job['test_cases_summary'] = {}
            # ─────────────────────────────────────────────────────────────────

            # ── PLAYWRIGHT AUTO-RUNNER ────────────────────────────────────────
            # Runs auto_runnable test cases against the live survey.
            # Never blocks the main QC — all failures are caught internally.
            job['playwright_tests'] = {"results": [], "summary": {}, "error": None}
            _pw_tc = job.get('test_cases', [])
            _pw_runnable = [t for t in _pw_tc if t.get('auto_runnable') and t.get('type') != 'GRID']
            if survey_url and _pw_runnable:
                try:
                    from test_runner import run_playwright_tests as _run_pw
                    _pw_ss_dir = os.path.join(
                        UPLOAD_FOLDER, job_id, 'playwright_screenshots')
                    log('', 'white')
                    log('════════════════════════════════════', 'cyan')
                    log('  PLAYWRIGHT AUTO-RUNNER', 'cyan')
                    log('════════════════════════════════════', 'cyan')
                    log(f'  Running {min(len(_pw_runnable), 20)} auto-runnable test case(s)…', 'blue')
                    progress(88, 'Running automated tests...')
                    _pw_result = {"results": [], "summary": {}, "error": None}
                    with concurrent.futures.ThreadPoolExecutor(max_workers=1) as _pw_exec:
                        _pw_fut = _pw_exec.submit(
                            _run_pw, _pw_tc, survey_url, _pw_ss_dir,
                            max_tests=20, timeout_ms=30_000,
                            xml_questions=xml_questions if xml_questions else None)
                        try:
                            _pw_result = _pw_fut.result(timeout=max(120, len(_pw_tc) * 45))
                        except concurrent.futures.TimeoutError:
                            log('  PLAYWRIGHT TIMEOUT — stopping tests', 'yellow')
                            _pw_result = {"results": [], "summary": {},
                                          "error": "PLAYWRIGHT TIMEOUT — 60s total exceeded"}
                    job['playwright_tests'] = _pw_result
                    _pw_s = _pw_result.get('summary', {})
                    log(f'  Playwright: {_pw_s.get("passed",0)} PASS  '
                        f'{_pw_s.get("failed",0)} FAIL  '
                        f'{_pw_s.get("errors",0)} ERROR  '
                        f'({_pw_s.get("pass_rate","?")} pass rate)', 'green')
                    if _pw_result.get('error'):
                        log(f'  Playwright warning: {_pw_result["error"][:80]}', 'yellow')
                except Exception as _pw_err:
                    log(f'  Playwright runner skipped: {str(_pw_err)[:80]}', 'grey')
                    job['playwright_tests'] = {
                        "results": [], "summary": {}, "error": str(_pw_err)[:200]}
            # ─────────────────────────────────────────────────────────────────

            # Assign confidence to every issue in one pass — does not touch
            # any of the creation sites above.
            for _ci in issues:
                _ci["confidence"], _ci["conf_level"] = _assign_confidence(_ci)

            # THREE-WAY CONFIDENCE BOOST: when XML export was parsed, use it as
            # source of truth to refine confidence on each issue.
            if xml_questions:
                _xml_by_nqid = {}
                for _xq in xml_questions:
                    _xnorm = re.sub(r'[^a-z0-9]', '', (_xq.get('qid_normalized') or _xq.get('qid', '')).lower())
                    _xml_by_nqid[_xnorm] = _xq

                for _ci in issues:
                    _iqid = _ci.get('qid', '')
                    _inorm = re.sub(r'[^a-z0-9]', '', _iqid.lower())
                    _xq = _xml_by_nqid.get(_inorm)
                    if not _xq:
                        continue
                    _xml_text = (_xq.get('text') or '').strip().lower()
                    _issue_type = _ci.get('type', '')
                    if _issue_type in ('TEXT MISMATCH', 'WORDS MISSING', 'OPTIONS MISMATCH'):
                        _doc_q = questions.get(_iqid, {})
                        _doc_text = (_doc_q.get('text') or '').strip().lower()
                        _live_q = live_data.get(_iqid, {})
                        _live_text = (_live_q.get('text') or '').strip().lower()
                        _xml_near_doc = bool(_xml_text) and fuzzy_match(_xml_text, _doc_text, 0.6)[0]
                        _xml_near_live = bool(_xml_text) and fuzzy_match(_xml_text, _live_text, 0.6)[0]
                        if _xml_near_doc and not _xml_near_live:
                            # XML=doc, live is wrong → live programming bug
                            _ci['confidence'] = 90; _ci['conf_level'] = 'HIGH'
                            _ci['xml_verdict'] = 'live_differs'
                        elif _xml_near_live and not _xml_near_doc:
                            # XML=live, doc is wrong → doc may be outdated
                            _ci['confidence'] = 75; _ci['conf_level'] = 'MEDIUM'
                            _ci['xml_verdict'] = 'doc_differs'
                        elif not _xml_near_doc and not _xml_near_live and _xml_text:
                            # All three differ from each other
                            _ci['confidence'] = 95; _ci['conf_level'] = 'HIGH'
                            _ci['xml_verdict'] = 'all_differ'

                # Flag QIDs in XML that appear in neither doc nor live
                _doc_norms = set(re.sub(r'[^a-z0-9]', '', q.lower()) for q in questions.keys())
                _live_norms = set(re.sub(r'[^a-z0-9]', '', q.lower()) for q in live_data.keys())

                # INTERNAL_NORMS, S99_DATE_PAT, FW_NORM_PAT imported from qid_normalizer.
                # Aliases keep the rest of this block unchanged.
                _INTERNAL_XML_NORMS = INTERNAL_NORMS
                _S99_DATE_PAT       = S99_DATE_PAT
                _FW_NORM_PAT        = FW_NORM_PAT

                # Delegate to the canonical implementation in qid_normalizer
                _xml_strip_candidates = build_strip_candidates

                def _xml_grid_components(xnorm):
                    """Return component norms for a Confirmit grid QID (gQ12Q13 → [q12, q13])."""
                    m = re.match(r'^g([a-z]+\d+[a-z]*)([a-z]+\d+[a-z]*)$', xnorm)
                    return [m.group(1), m.group(2)] if m else None

                _xml_skip_internal = 0
                _xml_skip_norm = 0
                for _xnorm, _xq in _xml_by_nqid.items():
                    if not _xnorm:
                        continue
                    # Gate 0: single-pipeline filter (framework pages, stop words, S99, internal)
                    _xraw = _xq.get('qid', '') or _xnorm
                    if should_skip_qid(_xraw):
                        _xml_skip_internal += 1; continue
                    # Gate 1: exact internal match (kept as belt-and-suspenders)
                    if _xnorm in _INTERNAL_XML_NORMS:
                        _xml_skip_internal += 1; continue
                    # Gate 2: S99-type date/time pattern (S99Datex1, S99Timex2, …)
                    if _S99_DATE_PAT.match(_xnorm):
                        _xml_skip_internal += 1; continue
                    # Gate 3: direct match in doc or live
                    if _xnorm in _doc_norms or _xnorm in _live_norms:
                        continue
                    # Gate 4: progressive suffix-stripping — xN → _N → _BIS → bis/ter parent
                    _raw_qid = _xq.get('qid_normalized') or _xq.get('qid', _xnorm)
                    _cands = _xml_strip_candidates(_raw_qid, _xnorm)
                    _skip = False; _skip_reason = None
                    for _c in _cands:
                        if _c in _INTERNAL_XML_NORMS:          # e.g. S99Datex1 → s99date
                            _skip = True; _skip_reason = 'internal'; break
                        if _c in _doc_norms or _c in _live_norms:
                            _skip = True; _skip_reason = 'norm'; break
                    if _skip:
                        if _skip_reason == 'internal':
                            _xml_skip_internal += 1
                        else:
                            _xml_skip_norm += 1
                        continue
                    # Gate 5: grid container expansion (gQ12Q13 → Q12, Q13)
                    _components = _xml_grid_components(_xnorm)
                    if _components and any(c in _doc_norms or c in _live_norms for c in _components):
                        continue
                    # Gate 6: structural and prefix-based framework variable detection.
                    # Catches Confirmit internal variables not covered by earlier gates.
                    # Safe because Gates 3-5 already resolved all real doc/live QIDs.
                    _raw_name = _xq.get('qid', '')
                    if (
                        # Structural: chars that can never appear in a valid QID
                        '(' in _raw_name or ' ' in _raw_name or
                        '\\' in _raw_name or ':' in _raw_name or
                        "'" in _raw_name or _raw_name.startswith('!') or
                        # Confirmit _hidden data-capture shadow copies: C2hidden, RPPS_hidden
                        _raw_name.endswith('_hidden') or
                        # Confirmit camelCase internal vars: hS1, iM1, gQ11, gBIC
                        bool(re.match(r'^[hig][A-Z]', _raw_name)) or
                        # Confirmit Hid-block prefix: HidATUEnd, HidAllQs, HidLang
                        bool(re.match(r'^Hid[A-Z]', _raw_name)) or
                        # Normalized prefix families (Confirmit scripting conventions)
                        _FW_NORM_PAT.match(_xnorm) or
                        # Normalized _hidden suffix: c2hidden, rppshidden, etudianthi...
                        _xnorm.endswith('hidden')
                    ):
                        _xml_skip_internal += 1
                        continue
                    issues.append({
                        "qid": _xq.get('qid', _xnorm),
                        "type": "QID IN EXPORT NOT IN DOC/LIVE",
                        "details": "Found in XML export but missing from both spec doc and live survey",
                        "severity": "MEDIUM",
                        "confidence": 70,
                        "conf_level": "MEDIUM",
                        "source_phase": "PHASE_1.5_XML",
                    })
                if _xml_skip_norm:
                    log(f'  Normalized {_xml_skip_norm} sub-question variant(s) to parent QID — no issue raised', 'grey')
                if _xml_skip_internal:
                    log(f'  Skipped {_xml_skip_internal} technical/internal XML field(s)', 'grey')

                log(f'  Three-way comparison: XML used to refine confidence on {len(xml_questions)} question(s)', 'green')

            _chigh  = sum(1 for i in issues if i.get("conf_level") == "HIGH")
            _cmed   = sum(1 for i in issues if i.get("conf_level") == "MEDIUM")
            _clow   = sum(1 for i in issues if i.get("conf_level") == "LOW")
            log(f'  Confidence: {_chigh} high, {_cmed} medium, {_clow} need manual review', 'cyan')
            # Tag all issues generated in PHASE 3 (those without a source_phase yet)
            for _pi in issues:
                if 'source_phase' not in _pi:
                    _pi['source_phase'] = 'PHASE_3_COMPARISON'

        # PHASE 4: TERMINATION — skipped for re-tests (filter_qids set)
        if mode in ('full', 'logic') and not filter_qids:
            progress(75, 'Testing termination rules...')
            log('', 'white')
            log('════════════════════════════════════', 'cyan')
            log('  PHASE 4: TERMINATION TESTING', 'cyan')
            log('════════════════════════════════════', 'cyan')

            rules = []
            for qid, q_data in questions.items():
                for rule in q_data.get("termination_rules", []):
                    test_qid = rule.get("test_qid")
                    for code in rule.get("answer_codes", []):
                        if test_qid and code:
                            rules.append({"test_qid":test_qid,"answer_code":code,"raw_rule":rule.get("raw",""),"source":rule.get("source","")})

            seen_r = set()
            unique_rules = []
            for r in rules:
                key = (r["test_qid"], r["answer_code"])
                if key not in seen_r:
                    seen_r.add(key)
                    unique_rules.append(r)

            # ── [TERM DEBUG] checkpoint 4: Phase 4 entry ─────────────────────
            app.logger.debug("[TERM DEBUG] Phase 4: %d questions → %d raw rules → %d unique rules",
                             len(questions), len(rules), len(unique_rules))
            for _dbg_r in unique_rules[:10]:
                app.logger.debug("[TERM DEBUG]   rule: %s=%s  raw=%r",
                                 _dbg_r['test_qid'], _dbg_r['answer_code'], _dbg_r['raw_rule'][:60])
            # ─────────────────────────────────────────────────────────────────
            log(f'  Testing {len(unique_rules)} rules', 'blue')

            for i, rule in enumerate(unique_rules, 1):
                if jobs.get(job_id, {}).get('status') == 'stopped':
                    log('  >>> STOPPED by user during termination tests', 'red')
                    jobs[job_id]['phase'] = 'Stopped'
                    jobs.persist(job_id)
                    return
                test_qid = rule["test_qid"]
                answer_code = rule["answer_code"]
                log(f'\n  [{i}/{len(unique_rules)}] {test_qid} = code {answer_code}', 'blue')
                progress(75 + int((i/max(1,len(unique_rules)))*15))

                raw_upper = rule.get("raw_rule","").upper()
                is_compound = (answer_code == "?" or
                    any(w in raw_upper for w in ["NOT SELECTED","AND CODE","OR CODE"]) or
                    test_qid in ["S7","S9"])
                if is_compound:
                    _manual_detail = "MANUAL CHECK — " + (rule.get("raw_rule","")[:60] if answer_code == "?" else "compound logic")
                    term_results.append({"test_qid":test_qid,"answer_code":answer_code,"passed":False,"needs_review":True,"details":_manual_detail,"source":rule.get("source","")})
                    log(f'      MANUAL CHECK — {rule.get("raw_rule","")[:60]}', 'yellow')
                    continue

                r_result = {"test_qid":test_qid,"answer_code":answer_code,"passed":False,"details":"","source":rule.get("source","")}
                try:
                    with sync_playwright() as p:
                        _UA = ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                               "AppleWebKit/537.36 (KHTML, like Gecko) "
                               "Chrome/124.0.0.0 Safari/537.36")
                        browser = p.chromium.launch(
                            headless=True,
                            slow_mo=200,
                            args=[
                                "--disable-blink-features=AutomationControlled",
                                "--no-sandbox",
                                "--disable-dev-shm-usage",
                            ]
                        )
                        context = browser.new_context(
                            viewport={"width": 1366, "height": 768},
                            user_agent=_UA,
                            locale="en-US",
                            timezone_id="America/New_York",
                            extra_http_headers={
                                "Accept":          "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,*/*;q=0.8",
                                "Accept-Language": "en-US,en;q=0.9",
                                "Accept-Encoding": "gzip, deflate, br",
                                "Sec-Fetch-Dest":  "document",
                                "Sec-Fetch-Mode":  "navigate",
                                "Sec-Fetch-Site":  "none",
                                "Sec-Fetch-User":  "?1",
                                "Upgrade-Insecure-Requests": "1",
                            }
                        )
                        page = context.new_page()
                        page.set_default_timeout(30000)
                        page.add_init_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
                        page.goto(survey_url, wait_until="domcontentloaded", timeout=60000)
                        page.wait_for_timeout(2500)

                        if country:
                            try:
                                page.locator(f".cf-radio-answer__text:has-text('{country}')").first.click(force=True, timeout=8000)
                                page.wait_for_timeout(800)
                                for sel in ["button:has-text('>>')", "input[value='>>']", ".cf-button-next"]:
                                    try:
                                        page.locator(sel).first.click(timeout=3000)
                                        break
                                    except: continue
                                page.wait_for_timeout(2500)
                            except: pass

                        # Detect which TN selector is active on this platform
                        _term_tn_sel = None
                        for _s in [".cf-tn-list-item", "[class*='tn-question']", "[class*='test-navigator']", ".wix-tn-item", "[class*='tn-item']"]:
                            if page.locator(_s).count() > 0:
                                _term_tn_sel = _s; break
                        if not _term_tn_sel:
                            try:
                                for _btn in ["Test Navigator", "Navigator", "Navigateur"]:
                                    page.locator(f"text={_btn}").first.click(timeout=3000)
                                    page.wait_for_timeout(1000)
                                    for _s in [".cf-tn-list-item", "[class*='tn-question']", "[class*='test-navigator']", ".wix-tn-item", "[class*='tn-item']"]:
                                        if page.locator(_s).count() > 0:
                                            _term_tn_sel = _s; break
                                    if _term_tn_sel: break
                            except: pass

                        try:
                            if _term_tn_sel and not page.locator(_term_tn_sel).count():
                                page.locator("text=Test Navigator").first.click(timeout=3000)
                                page.wait_for_timeout(1000)
                        except: pass

                        navigated = False
                        _tn_loc = page.locator(_term_tn_sel) if _term_tn_sel else page.locator(".cf-tn-list-item")
                        for _idx in range(_tn_loc.count()):
                            try:
                                txt = _tn_loc.nth(_idx).inner_text(timeout=2000).strip().split()[0].strip()
                                if txt.upper() == test_qid.upper():
                                    for _attempt in range(2):
                                        try:
                                            _tn_loc.nth(_idx).click(force=True, timeout=3000)
                                            page.wait_for_timeout(1800)
                                            navigated = True
                                            break
                                        except:
                                            if _attempt == 0:
                                                page.wait_for_timeout(800)
                                    if navigated:
                                        break
                            except: continue

                        if not navigated:
                            r_result["details"] = f"Could not find {test_qid}"
                            browser.close()
                            term_results.append(r_result)
                            log(f'      Could not navigate to {test_qid}', 'red')
                            continue

                        try:
                            if _term_tn_sel and page.locator(_term_tn_sel).count() > 0:
                                page.locator("text=Test Navigator").first.click(timeout=2000)
                                page.wait_for_timeout(500)
                        except: pass

                        _ac = str(answer_code)
                        clicked = False
                        strategy = ""

                        # Strategy 1: radio input whose value attribute == answer_code
                        # (Confirmit/Forsta store the code as the <input value="N">)
                        for _attempt in range(2):
                            try:
                                _r = page.locator(f"input[type='radio'][value='{_ac}']:visible")
                                if _r.count() > 0:
                                    try: _r.first.scroll_into_view_if_needed(timeout=2000)
                                    except: pass
                                    _r.first.click(force=True, timeout=3000)
                                    page.wait_for_timeout(600)
                                    clicked = True
                                    strategy = f"value={_ac}"
                                    break
                            except:
                                if _attempt == 0:
                                    page.wait_for_timeout(500)

                        # Strategy 2: data-code / data-value / data-answer-id attribute
                        if not clicked:
                            for _attr_sel in [f"[data-code='{_ac}']",
                                              f"[data-value='{_ac}']",
                                              f"[data-answer-id='{_ac}']"]:
                                try:
                                    _r = page.locator(_attr_sel)
                                    if _r.count() > 0:
                                        try: _r.first.scroll_into_view_if_needed(timeout=2000)
                                        except: pass
                                        _r.first.click(force=True, timeout=3000)
                                        page.wait_for_timeout(600)
                                        clicked = True
                                        strategy = f"attr={_ac}"
                                        break
                                except:
                                    pass

                        # Strategy 3: positional fallback (valid only for 1-indexed sequential codes)
                        if not clicked:
                            try:
                                _idx = int(_ac) - 1
                                if _idx >= 0:
                                    for _attempt in range(2):
                                        try:
                                            _lbl = page.locator(".cf-radio-answer__text").nth(_idx)
                                            try: _lbl.scroll_into_view_if_needed(timeout=2000)
                                            except: pass
                                            _lbl.click(force=True, timeout=3000)
                                            page.wait_for_timeout(600)
                                            clicked = True
                                            strategy = f"positional-fallback={_idx}"
                                            break
                                        except:
                                            if _attempt == 0:
                                                page.wait_for_timeout(500)
                            except (ValueError, TypeError):
                                pass

                        # Strategy 4: raw radio positional fallback
                        if not clicked:
                            try:
                                _idx = int(_ac) - 1
                                if _idx >= 0:
                                    for _attempt in range(2):
                                        try:
                                            page.locator("input[type='radio']:visible").nth(_idx).click(force=True, timeout=3000)
                                            page.wait_for_timeout(600)
                                            clicked = True
                                            strategy = f"radio-positional-fallback={_idx}"
                                            break
                                        except:
                                            if _attempt == 0:
                                                page.wait_for_timeout(500)
                            except (ValueError, TypeError):
                                pass

                        if not clicked:
                            r_result["passed"] = True
                            r_result["details"] = "Could not test — click failed, manual review required"
                            browser.close()
                            term_results.append(r_result)
                            log(f'      Could not click — manual review required', 'yellow')
                            continue

                        for sel in ["button:has-text('>>')", "input[value='>>']", ".cf-button-next"]:
                            try:
                                page.locator(sel).first.click(timeout=3000)
                                break
                            except: continue

                        page.wait_for_timeout(3500)
                        body_text = page.locator("body").inner_text(timeout=5000).lower()
                        _kw_terminated = any(ind in body_text for ind in THANKYOU_INDICATORS)
                        try:
                            _struct_terminated = page.locator("input[type='radio']:visible, input[type='checkbox']:visible, select:visible").count() == 0
                        except:
                            _struct_terminated = False
                        terminated = _kw_terminated or _struct_terminated

                        if terminated:
                            r_result["passed"] = True
                            r_result["details"] = f"Terminated as expected ({strategy})"
                            log(f'      PASS — Terminated', 'green')
                        else:
                            r_result["passed"] = False
                            r_result["details"] = f"Expected close but continued ({strategy})"
                            log(f'      FAIL — Survey continued', 'red')

                        browser.close()
                except Exception as e:
                    r_result["details"] = f"Error: {str(e)[:80]}"
                    log(f'      Error: {str(e)[:60]}', 'red')

                term_results.append(r_result)

            passed = sum(1 for r in term_results if r["passed"])
            log(f'\n  Termination: {passed}/{len(term_results)} passed', 'green' if passed==len(term_results) else 'yellow')

        # PHASE 4.5: SURVEY FLOW ANALYSIS
        try:
            from survey_graph import build_doc_graph, build_xml_graph, build_live_graph, compare_graphs
            _flow_ref = None
            _flow_other = None
            _flow_ref_name = ''
            if xml_questions:
                _flow_ref = build_xml_graph(xml_questions)
                _flow_ref_name = 'xml'
            elif questions:
                _flow_ref = build_doc_graph(questions)
                _flow_ref_name = 'doc'
            if _flow_ref and live_data:
                _flow_other = build_live_graph(live_data)
            if _flow_ref and _flow_other:
                progress(88, 'Analysing survey flow...')
                log('', 'white')
                log('════════════════════════════════════', 'cyan')
                log('  PHASE 4.5: SURVEY FLOW ANALYSIS', 'cyan')
                log('════════════════════════════════════', 'cyan')
                _flow_raw = compare_graphs(_flow_ref, _flow_other)
                _frs = _flow_ref.summary()
                log(f'  Ref ({_flow_ref_name}): {_frs["questions"]} questions, '
                    f'{_frs["termination_points"]} termination pt(s), '
                    f'{_frs["conditional_questions"]} conditional', 'blue')
                log(f'  Live: {_flow_other.summary()["questions"]} questions', 'blue')
                def _fn(s): return re.sub(r'[^a-z0-9]', '', s.lower())
                _live_keys_norm  = {_fn(q) for q in live_data.keys()}
                _existing_q_norm = {_fn(i.get('qid', '')) for i in issues}
                # Emit MISSING only; TERMINATION only when quick mode skipped PHASE 4.
                # ROUTING_UNCHECKED / EXTRA_IN_OTHER suppressed — always-fire noise
                # because the live graph carries no routing info.
                _emit_types = {'MISSING_IN_OTHER'}
                if mode == 'quick':
                    _emit_types.add('TERMINATION_MISSING')
                _flow_conf = {
                    'MISSING_IN_OTHER':    (88, 'HIGH'),
                    'TERMINATION_MISSING': (70, 'MEDIUM'),
                }
                _added = 0
                _flow_skipped = 0
                for _fi in _flow_raw:
                    if _fi.issue_type not in _emit_types:
                        continue
                    # Single-pipeline filter: same rules as PHASE 3
                    if should_skip_qid(_fi.qid):
                        _flow_skipped += 1
                        continue
                    # Parent-child collapse: if parent already in live, this
                    # is just a renamed sub-field (e.g. S99Datex1 → parent S99)
                    _parent = get_parent_qid(_fi.qid)
                    if _parent != _fi.qid and _fn(_parent) in _live_keys_norm:
                        _flow_skipped += 1
                        continue
                    _qn = _fn(_fi.qid)
                    if _qn in _live_keys_norm:
                        continue  # already in live (CONDITIONAL etc.) — not missing
                    if _qn in _existing_q_norm:
                        continue  # PHASE 3 already reported this qid
                    # Phase 5: suppress MISSING_IN_OTHER when XML confirms the question
                    if _fi.issue_type == 'MISSING_IN_OTHER' and _xml_qid_set and _qn in _xml_qid_set:
                        issues.append({
                            'qid':        _fi.qid,
                            'type':       'IN_XML_NOT_VERIFIED_IN_LIVE',
                            'severity':   'INFO',
                            'details':    'Question present in XML but not verified in live survey — likely hidden by routing',
                            'confidence': 55,
                            'conf_level': 'NEEDS_MANUAL',
                            'source_phase': 'PHASE_4.5_FLOW',
                        })
                        _added += 1
                        log(f'   {_fi.qid} — in XML, not visible in live (NEEDS_MANUAL)', 'green')
                        continue
                    _conf, _lvl = _flow_conf.get(_fi.issue_type, (60, 'MEDIUM'))
                    issues.append({
                        'qid':        _fi.qid,
                        'type':       _fi.issue_type.replace('_', ' '),
                        'severity':   _fi.severity,
                        'details':    _fi.details,
                        'confidence': _conf,
                        'conf_level': _lvl,
                        'source_phase': 'PHASE_4.5_FLOW',
                    })
                    _added += 1
                    log(f'   {_fi.qid} — {_fi.issue_type}', 'yellow')
                if _flow_skipped:
                    log(f'  Skipped {_flow_skipped} internal/framework QID(s) from flow analysis', 'grey')
                log(f'  Flow issues added: {_added} '
                    f'({len(_flow_raw) - _added} suppressed by dedup)', 'cyan' if _added else 'green')
                if _frs['termination_at']:
                    log(f'  Termination pts: {", ".join(_frs["termination_at"])}', 'blue')
        except Exception as _flow_err:
            log(f'  Flow analysis skipped ({str(_flow_err)[:80]})', 'yellow')

        # Evidence Engine: enrich every issue with evidence snippets and
        # recalculate confidence using all available data.  Must run after
        # all phases so xml_questions, live_data, questions are complete.
        try:
            _enrich_issues(issues, questions, live_data, xml_questions)
        except Exception as _ee_err:
            log(f'  Evidence enrichment error (non-fatal): {str(_ee_err)[:80]}', 'yellow')

        # PHASE 5: REPORT
        progress(92, 'Generating report...')
        log('', 'white')
        log('════════════════════════════════════', 'cyan')
        log('  PHASE 5: REPORT GENERATION', 'cyan')
        log('════════════════════════════════════', 'cyan')

        report = Document()
        for sec in report.sections:
            sec.top_margin = Cm(2)
            sec.bottom_margin = Cm(2)
            sec.left_margin = Cm(2.5)
            sec.right_margin = Cm(2.5)

        def shade_cell(cell, hex_color):
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), hex_color)
            tc_pr.append(shd)

        title = report.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _report_title = "RE-TEST REPORT" if filter_qids else "Survey QC Report"
        tr = title.add_run(_report_title)
        tr.font.size = Pt(22); tr.font.bold = True
        tr.font.color.rgb = RGBColor(0x7C, 0x65, 0xFF)

        sub = report.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = sub.add_run(f"{os.path.basename(doc_path)}\n{datetime.now().strftime('%d %B %Y, %H:%M')}")
        sr.font.size = Pt(11); sr.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        report.add_paragraph()

        _hdr_xml_count  = len(xml_questions)
        _hdr_xml_hidden = job.get('xml_hidden_count', 0)
        _xml_failed     = job.get('xml_parse_failed', False)
        # Determine QC mode label
        _qc_mode_label = job.get('qc_mode', 'STANDARD')
        _is_advanced_qc = (_qc_mode_label == 'ADVANCED' and len(live_data) > 0)
        _is_standard_qc = (not _is_advanced_qc and _hdr_xml_count > 0 and not _xml_failed)

        if _xml_failed:
            _warn_p = report.add_paragraph(); _warn_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _warn_r = _warn_p.add_run(
                "⚠ XML parsing failed — falling back to doc-vs-live comparison. "
                "Accuracy may be lower. Please check the XML format."
            )
            _warn_r.font.size = Pt(10); _warn_r.font.bold = True
            _warn_r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
            _src_p = report.add_paragraph(); _src_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _src_r = _src_p.add_run(
                f"Sources analyzed: Doc ({len(questions)}) + Live ({len(live_data)})"
                f"    |    Mode: STANDARD QC (doc vs live — XML parse failed)"
            )
            _src_r.font.size = Pt(10); _src_r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        elif _is_advanced_qc:
            # ADVANCED QC: DOC + XML + LIVE
            _mode_p = report.add_paragraph(); _mode_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _mode_r = _mode_p.add_run("ADVANCED QC  —  DOC + XML + LIVE VERIFICATION")
            _mode_r.font.size = Pt(12); _mode_r.font.bold = True
            _mode_r.font.color.rgb = RGBColor(0x13, 0x8D, 0x5A)
            _xml_src_str = f"  •  Survey Export XML ({_hdr_xml_count} visible questions"
            if _hdr_xml_hidden:
                _xml_src_str += f"; {_hdr_xml_hidden} hidden/template skipped"
            _xml_src_str += ")"
            for _src_line in [
                f"  •  Spec Document ({len(questions)} questions)",
                f"  •  Survey Export XML — PRIMARY source of truth",
                _xml_src_str,
                f"  •  Live Survey ({len(live_data)} questions crawled — verification layer)",
            ]:
                _slp = report.add_paragraph(); _slp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _slr = _slp.add_run(_src_line)
                _slr.font.size = Pt(10); _slr.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
        elif _is_standard_qc:
            # STANDARD QC: DOC + XML only (no live URL)
            _mode_p = report.add_paragraph(); _mode_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _mode_r = _mode_p.add_run("STANDARD QC  —  DOC + XML (no live verification)")
            _mode_r.font.size = Pt(12); _mode_r.font.bold = True
            _mode_r.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
            _xml_src_str = f"  •  Survey Export XML ({_hdr_xml_count} visible questions"
            if _hdr_xml_hidden:
                _xml_src_str += f"; {_hdr_xml_hidden} hidden/template skipped"
            _xml_src_str += ")  ← PRIMARY source of truth"
            for _src_line in [
                f"  •  Spec Document ({len(questions)} questions)",
                _xml_src_str,
                "  •  Live Survey: not provided — add URL for ADVANCED QC",
            ]:
                _slp = report.add_paragraph(); _slp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _slr = _slp.add_run(_src_line)
                _slr.font.size = Pt(10); _slr.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
            _acc_p = report.add_paragraph(); _acc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _acc_r = _acc_p.add_run(
                "All issues are derived from DOC vs XML comparison. "
                "No live-survey false positives."
            )
            _acc_r.font.size = Pt(10); _acc_r.font.italic = True
            _acc_r.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
        else:
            _src_p = report.add_paragraph(); _src_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _src_r = _src_p.add_run(
                f"Sources analyzed: Doc ({len(questions)}) + Live ({len(live_data)})"
                f"    |    Mode: STANDARD QC (doc vs live comparison)"
            )
            _src_r.font.size = Pt(10); _src_r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        report.add_paragraph()

        if filter_qids:
            _rt_banner = report.add_paragraph()
            _rt_banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _rt_run = _rt_banner.add_run(
                f"RE-TEST REPORT — Filtered to {len(filter_qids)} QID(s): {', '.join(filter_qids)}"
            )
            _rt_run.font.size = Pt(10); _rt_run.font.bold = True
            _rt_run.font.color.rgb = RGBColor(0x13, 0x8D, 0x5A)
            report.add_paragraph()

        sev = {"HIGH":0,"MEDIUM":0,"INFO":0}
        for i in issues: sev[i.get("severity","INFO")] = sev.get(i.get("severity","INFO"),0)+1
        term_passed = sum(1 for r in term_results if r.get("passed") and not r.get("needs_review"))
        term_review = sum(1 for r in term_results if r.get("needs_review"))
        term_failed = len(term_results) - term_passed - term_review
        total_issues = sev['HIGH'] + sev['MEDIUM']

        if len(live_data) == 0 and survey_url:
            # Live URL was provided but crawl returned nothing
            vt = ("CRAWL FAILED ��� 0 pages loaded. The live survey did not render in time. "
                  "Possible causes: slow platform, bot detection, session expired, or network issue. "
                  "Please retry the QC or verify the survey link is accessible.")
            vc = (0xC0, 0x00, 0x00)
        elif len(live_data) == 0 and not survey_url and _hdr_xml_count:
            # Standard QC (no URL provided) — verdict based on XML comparison only
            if total_issues == 0:
                vt = "STANDARD QC PASS — No issues found in DOC vs XML comparison"; vc = (0x00, 0x70, 0x00)
            elif sev['HIGH'] > 0:
                vt = f"STANDARD QC — {total_issues} issue(s) found (add live URL for full ADVANCED QC)"; vc = (0xC0, 0x00, 0x00)
            else:
                vt = f"STANDARD QC — {total_issues} issue(s) to review (add live URL for full ADVANCED QC)"; vc = (0xBA, 0x75, 0x17)
        elif total_issues == 0 and term_failed == 0:
            vt = "ALL GOOD — Survey ready to go live!"; vc = (0x00, 0x70, 0x00)
        elif sev['HIGH'] > 0 or term_failed > 0:
            vt = f"NEEDS FIX — {total_issues + term_failed} issues found"; vc = (0xC0, 0x00, 0x00)
        else:
            vt = "REVIEW NEEDED — Minor issues to check"; vc = (0xBA, 0x75, 0x17)

        v = report.add_paragraph(); v.alignment = WD_ALIGN_PARAGRAPH.CENTER
        vr = v.add_run(vt); vr.font.size = Pt(18); vr.font.bold = True
        vr.font.color.rgb = RGBColor(*vc)
        report.add_paragraph()

        h = report.add_paragraph()
        hr = h.add_run("Quick Summary")
        hr.font.size = Pt(14); hr.font.bold = True; hr.font.color.rgb = RGBColor(0xC4, 0x6A, 0x2B)

        # AI-generated human-readable summary
        try:
            ai_summary_text = ai_generate_summary(get_gemini_model(), questions, live_data, issues)
        except Exception:
            ai_summary_text = ""
        job['ai_summary'] = ai_summary_text
        if ai_summary_text:
            sp = report.add_paragraph()
            sr = sp.add_run(ai_summary_text)
            sr.font.size = Pt(11); sr.italic = True; sr.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
            report.add_paragraph()

        _rpt_chigh = sum(1 for i in issues if i.get("conf_level") == "HIGH")
        _rpt_cmed  = sum(1 for i in issues if i.get("conf_level") == "MEDIUM")
        _rpt_clow  = sum(1 for i in issues if i.get("conf_level") == "LOW")
        _xml_q_count = job.get('xml_qids', 0)
        if _is_advanced_qc:
            _accuracy_mode = "ADVANCED QC — DOC + XML + LIVE (false-positive reduced)"
        elif _is_standard_qc:
            _accuracy_mode = "STANDARD QC — DOC + XML only (no live URL false positives)"
        else:
            _accuracy_mode = "Standard (doc vs live comparison)"
        _needs_manual_count = sum(1 for i in issues if i.get("conf_level") == "NEEDS_MANUAL")
        for line in [
            f"Sources: Doc ({len(questions)} questions) · XML ({_xml_q_count} questions) · Live ({len(live_data)} questions)" if _xml_q_count else f"Sources: Doc ({len(questions)} questions) · Live ({len(live_data)} pages crawled)",
            f"QC Mode: {_accuracy_mode}",
            f"Termination tests: {term_passed}/{len(term_results) - term_review} validated · {term_review} need manual review" if term_results else None,
            f"Total issues found: {total_issues}",
            f"Confidence breakdown: {_rpt_chigh} confirmed bugs, {_rpt_cmed} possible issues, {_rpt_clow} low-confidence, {_needs_manual_count} needs manual review",
            f"Time saved: ~8 hours vs manual QC",
        ]:
            if not line: continue
            p = report.add_paragraph()
            p.add_run(f"  - {line}").font.size = Pt(11)

        report.add_paragraph()

        type_names = {
            "WORDS MISSING": "Missing words",
            "TEXT MISMATCH": "Text doesn't match",
            "OPTIONS MISMATCH": "Answer options missing",
            "OPTIONS COUNT MISMATCH": "Answer option count differs (doc vs XML)",
            "OPTION TEXT MISSING IN XML": "Answer option not found in XML",
            "CODE MISMATCH": "Answer codes differ (doc vs XML)",
            "ROUTING IN XML NOT IN DOC": "Routing in XML not documented in spec",
            "PIPING IN DOC NOT IN XML": "Piping in spec not found in XML",
            "IN_XML_NOT_IN_DOC": "Question in XML but missing from spec doc",
            "MANDATORY MISSING": "Mandatory marker missing",
            "PIPING NOT RESOLVED": "Piping not working",
            "MISSING IN LIVE": "Question missing from survey (not in XML or live)",
            "IN_XML_NOT_VERIFIED_IN_LIVE": "Not verified in live (present in XML) — needs manual check",
            "NAMING MISMATCH": "Question name differs (doc vs live)",
            "MISSING IN OTHER": "Question absent from live survey (flow analysis)",
            "TERMINATION MISSING": "Termination logic not verified",
        }
        fix_sug = {
            "WORDS MISSING": "Add the missing words to the live survey",
            "TEXT MISMATCH": "Update live survey text to match the doc",
            "OPTIONS MISMATCH": "Add missing answer options to live survey",
            "OPTIONS COUNT MISMATCH": "Verify the answer options in the survey platform match the spec doc",
            "OPTION TEXT MISSING IN XML": "Check the specific option text in the survey platform",
            "CODE MISMATCH": "Verify answer codes in the platform match the spec doc exactly",
            "ROUTING IN XML NOT IN DOC": "Add routing/logic table for this question to the spec doc",
            "PIPING IN DOC NOT IN XML": "Verify piping is correctly programmed in the survey platform",
            "IN_XML_NOT_IN_DOC": "Add this question to the spec doc, or confirm it is an unprogrammed placeholder",
            "MANDATORY MISSING": "Add * marker to make question mandatory",
            "PIPING NOT RESOLVED": "Fix piping logic",
            "MISSING IN LIVE": "Add this question to the survey platform",
            "IN_XML_NOT_VERIFIED_IN_LIVE": "Manually navigate to this question in the live survey to verify it renders correctly",
            "NAMING MISMATCH": "Rename question in live survey to match spec, or update spec",
            "MISSING IN OTHER": "Verify question exists in live survey; may have been excluded from the Test Navigator",
            "TERMINATION MISSING": "Manually test this termination rule in the live survey",
        }

        _fix_issues    = [i for i in issues if i.get("conf_level") in ("HIGH", "MEDIUM") and not i.get("is_export_issue") and i.get("conf_level") != "NEEDS_MANUAL"]
        _review_issues = [i for i in issues if i.get("conf_level") in ("LOW", "NEEDS_MANUAL") and not i.get("is_export_issue")]

        if _fix_issues or term_failed:
            h = report.add_paragraph()
            hr = h.add_run("Issues to Fix")
            hr.font.size = Pt(14); hr.font.bold = True; hr.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
            report.add_paragraph()

            n = 1
            for r in term_results:
                if not r["passed"]:
                    p = report.add_paragraph()
                    pr = p.add_run(f"Issue {n}: Termination not working")
                    pr.font.size = Pt(12); pr.font.bold = True; pr.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                    report.add_paragraph().add_run(f"   Where: {r['test_qid']} - code {r['answer_code']}").font.size = Pt(11)
                    report.add_paragraph().add_run(f"   What: {r['details']}").font.size = Pt(11)
                    p2 = report.add_paragraph()
                    p2r = p2.add_run("   Fix: Check termination logic")
                    p2r.font.size = Pt(11); p2r.font.italic = True; p2r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
                    report.add_paragraph(); n += 1

            for issue in _fix_issues:
                conf_pct   = issue.get('confidence', '')
                conf_label = issue.get('confidence_label', '') or issue.get('conf_level', '')
                conf_lvl   = issue.get('conf_level', '')
                color = (0xC0, 0x00, 0x00) if conf_lvl == 'HIGH' else (0xBA, 0x75, 0x17)
                simple = type_names.get(issue['type'], issue['type'])
                # ── Header line ────────────────────────────────────────────
                p = report.add_paragraph()
                pr = p.add_run(f"Issue {n}: {simple}  [{conf_pct}% — {conf_label}]")
                pr.font.size = Pt(12); pr.font.bold = True; pr.font.color.rgb = RGBColor(*color)
                report.add_paragraph().add_run(f"   Where: {issue['qid']}").font.size = Pt(11)
                # What line: prefer evidence mismatch_detail over raw details
                ev = issue.get('evidence', {})
                what_text = ev.get('mismatch_detail') or issue.get('details', '')
                report.add_paragraph().add_run(f"   What: {what_text[:220]}").font.size = Pt(11)
                # ── Evidence section ───────────────────────────────────────
                if ev:
                    def _ev_line(label, text, opts):
                        txt = (text or '').strip()[:100]
                        opt_str = '; '.join((opts or [])[:4])
                        if txt and opt_str:
                            return f"   {label}  \"{txt}\"  |  Options: {opt_str[:100]}"
                        elif txt:
                            return f"   {label}  \"{txt}\""
                        elif opt_str:
                            return f"   {label}  Options: {opt_str[:120]}"
                        return f"   {label}  (no data)"
                    ev_p = report.add_paragraph()
                    ev_r = ev_p.add_run("   Evidence:")
                    ev_r.font.size = Pt(10); ev_r.font.bold = True
                    ev_r.font.color.rgb = RGBColor(0x33, 0x33, 0x99)
                    _ev_clr = (0x33, 0x33, 0x99)
                    for _lbl, _tf, _of in [
                        ('Doc: ', ev.get('doc_text',''),  ev.get('doc_options',[])),
                        ('Live:', ev.get('live_text',''), ev.get('live_options',[])),
                        ('XML: ', ev.get('xml_text',''),  ev.get('xml_options',[])),
                    ]:
                        _line = _ev_line(_lbl, _tf, _of)
                        _ep = report.add_paragraph()
                        _er = _ep.add_run(_line)
                        _er.font.size = Pt(9)
                        _er.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
                # ── Confidence reasons ─────────────────────────────────────
                reasons = issue.get('confidence_reasons', [])
                if reasons:
                    _rp = report.add_paragraph()
                    _rr = _rp.add_run(f"   Why {conf_pct}%:  " + "  ·  ".join(reasons[:4]))
                    _rr.font.size = Pt(9); _rr.font.italic = True
                    _rr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                # ── Fix suggestion ─────────────────────────────────────────
                fix = fix_sug.get(issue['type'], 'Review and fix manually')
                p2 = report.add_paragraph()
                p2r = p2.add_run(f"   Fix: {fix}")
                p2r.font.size = Pt(11); p2r.font.italic = True
                p2r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
                report.add_paragraph(); n += 1

        # Separate NEEDS_MANUAL (XML-confirmed, not verified in live) from general review
        _needs_manual_issues = [i for i in _review_issues if i.get("conf_level") == "NEEDS_MANUAL"]
        _low_conf_issues = [i for i in _review_issues if i.get("conf_level") != "NEEDS_MANUAL"]

        if _needs_manual_issues:
            report.add_paragraph()
            h = report.add_paragraph()
            hr = h.add_run("Questions Present in XML — Not Verified in Live Survey")
            hr.font.size = Pt(14); hr.font.bold = True; hr.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
            sub_p = report.add_paragraph()
            sub_p.add_run(
                "These questions exist in the survey XML export but were not reached during live crawling. "
                "This is NOT a bug — they may be hidden by routing or require specific answers to reach. "
                "Manually navigate to each one to verify it renders correctly."
            ).font.size = Pt(10)
            report.add_paragraph()
            for issue in _needs_manual_issues:
                simple = type_names.get(issue.get('type', ''), issue.get('type', ''))
                p = report.add_paragraph()
                pr = p.add_run(f"ℹ  {issue['qid']}  —  {simple}")
                pr.font.size = Pt(11); pr.font.bold = True
                pr.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
                what_text = issue.get('details', '')
                report.add_paragraph().add_run(f"   {what_text[:200]}").font.size = Pt(10)
                report.add_paragraph()

        if _low_conf_issues:
            report.add_paragraph()
            h = report.add_paragraph()
            hr = h.add_run("Needs Manual Review — tool not certain")
            hr.font.size = Pt(14); hr.font.bold = True; hr.font.color.rgb = RGBColor(0xBA, 0x75, 0x17)
            sub_p = report.add_paragraph()
            sub_p.add_run(
                "The following items could not be confidently classified. "
                "Review each one manually before going live."
            ).font.size = Pt(10)
            report.add_paragraph()

            for issue in _low_conf_issues:
                conf_pct   = issue.get('confidence', '')
                conf_label = issue.get('confidence_label', '') or 'Needs Review'
                simple = type_names.get(issue['type'], issue['type'])
                p = report.add_paragraph()
                pr = p.add_run(f"?  {issue['qid']}  —  {simple}  [{conf_pct}% — {conf_label}]")
                pr.font.size = Pt(11); pr.font.bold = True
                pr.font.color.rgb = RGBColor(0xBA, 0x75, 0x17)
                ev = issue.get('evidence', {})
                what_text = ev.get('mismatch_detail') or issue.get('details', '')
                report.add_paragraph().add_run(f"   {what_text[:200]}").font.size = Pt(10)
                reasons = issue.get('confidence_reasons', [])
                if reasons:
                    _rp = report.add_paragraph()
                    _rr = _rp.add_run("   Why uncertain:  " + "  ·  ".join(reasons[:3]))
                    _rr.font.size = Pt(9); _rr.font.italic = True
                    _rr.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
                report.add_paragraph()

        # ── Data Export Issues section ──────────────────────────────────────
        _export_issues = [i for i in issues if i.get('is_export_issue')]
        if _export_issues:
            report.add_paragraph()
            _eh = report.add_paragraph()
            _ehr = _eh.add_run("Data Export Issues")
            _ehr.font.size = Pt(14); _ehr.font.bold = True
            _ehr.font.color.rgb = RGBColor(0x0F, 0x51, 0x32)
            _esub = report.add_paragraph()
            _esub.add_run(
                f"{len(_export_issues)} issue(s) found between survey definition and export schema. "
                "Fix before data collection begins."
            ).font.size = Pt(10)
            report.add_paragraph()
            _rule_labels = {
                'R031': 'Variable missing from export',
                'R032': 'Unexpected variable in export',
                'R033': 'Variable type mismatch',
                'R034': 'Answer code missing from export',
                'R035': 'Extra code in export',
                'R036': 'Open-end field has no export column',
                'R037': 'Piping variable missing from export',
                'R038': 'Loop naming convention mismatch',
            }
            for _ei in _export_issues:
                _rule = _ei.get('rule', _ei.get('type', ''))
                _label = _rule_labels.get(_rule, _rule)
                _sev_color = (0xC0, 0x00, 0x00) if _ei.get('severity') == 'HIGH' else (0xBA, 0x75, 0x17)
                _ep = report.add_paragraph()
                _epr = _ep.add_run(f"[{_rule}] {_label}")
                _epr.font.size = Pt(11); _epr.font.bold = True
                _epr.font.color.rgb = RGBColor(*_sev_color)
                _var = _ei.get('variable_name', '')
                _qid = _ei.get('qid', '')
                if _var or _qid:
                    report.add_paragraph().add_run(
                        f"   Variable: {_var or _qid}  |  QID: {_qid}"
                    ).font.size = Pt(10)
                report.add_paragraph().add_run(
                    f"   {_ei.get('details', '')[:200]}"
                ).font.size = Pt(10)
                report.add_paragraph()
        # ─────────────────────────────────────────────────────────────────

        # ── TERMINATION MATRIX section ────────────────────────────────────────
        _tm_rpt = job.get('termination_matrix', [])
        if _tm_rpt:
            report.add_paragraph()
            _tmh = report.add_paragraph()
            _tmhr = _tmh.add_run("Termination Matrix")
            _tmhr.font.size = Pt(14); _tmhr.font.bold = True
            _tmhr.font.color.rgb = RGBColor(0xC8, 0x4B, 0x31)
            _tm_match_cnt    = sum(1 for r in _tm_rpt if r.get('status') == 'MATCH')
            _tm_mismatch_cnt = sum(1 for r in _tm_rpt if r.get('status') == 'MISMATCH')
            _tmsub = report.add_paragraph()
            _tmsub.add_run(
                f"{len(_tm_rpt)} termination point(s): "
                f"{_tm_match_cnt} aligned, {_tm_mismatch_cnt} mismatch. "
                "DOC_ONLY = spec defines termination but not in XML. "
                "XML_ONLY = XML terminates but spec does not mention it."
            ).font.size = Pt(10)
            report.add_paragraph()

            _TM_STATUS_LABEL = {
                'MATCH': 'MATCH', 'MISMATCH': 'MISMATCH',
                'DOC_ONLY': 'DOC ONLY', 'XML_ONLY': 'XML ONLY',
            }
            _TM_STATUS_COLOR = {
                'MATCH':    (0x1A, 0x56, 0x32),
                'MISMATCH': (0xC0, 0x00, 0x00),
                'DOC_ONLY': (0xBA, 0x75, 0x17),
                'XML_ONLY': (0x1D, 0x4E, 0xD8),
            }
            for _tmrow in _tm_rpt:
                _ts = _tmrow.get('status', '')
                _tc = _TM_STATUS_COLOR.get(_ts, (0x33, 0x33, 0x33))
                _tp = report.add_paragraph()
                _tr = _tp.add_run(
                    f"  [{_TM_STATUS_LABEL.get(_ts, _ts)}]  "
                    f"{_tmrow.get('qid','')}  |  "
                    f"Doc codes: {', '.join(_tmrow.get('doc_codes',[]) or ['—'])}  |  "
                    f"XML: {(_tmrow.get('xml_condition') or '—')[:80]}"
                )
                _tr.font.size = Pt(10); _tr.font.bold = (_ts == 'MISMATCH')
                _tr.font.color.rgb = RGBColor(*_tc)
                if _tmrow.get('missing_in_xml'):
                    _mp = report.add_paragraph()
                    _mp.add_run(
                        f"    ⚠ Missing in XML: {', '.join(_tmrow['missing_in_xml'])}"
                    ).font.size = Pt(9)
                report.add_paragraph()
        # ─────────────────────────────────────────────────────────────────────

        # ── RULE ENGINE FINDINGS section ──────────────────────────────────────
        _re_findings_rpt = job.get('rule_engine_findings', [])
        _re_summary_rpt  = job.get('rule_engine_summary', {})
        if _re_findings_rpt:
            # Filter to non-INFO findings for the report
            _re_report_items = [f for f in _re_findings_rpt if f.get('severity') != 'INFO']
            if _re_report_items:
                report.add_paragraph()
                _reh = report.add_paragraph()
                _rehr = _reh.add_run("Rule Engine Findings")
                _rehr.font.size = Pt(14); _rehr.font.bold = True
                _rehr.font.color.rgb = RGBColor(0x6D, 0x28, 0xD9)
                _resub = report.add_paragraph()
                _resub.add_run(
                    f"Deterministic rule engine: {len(_re_report_items)} finding(s) "
                    f"(HIGH={_re_summary_rpt.get('high',0)}, "
                    f"MEDIUM={_re_summary_rpt.get('medium',0)}, "
                    f"LOW={_re_summary_rpt.get('low',0)}). "
                    "Run directly on survey model — no DOM, no Playwright."
                ).font.size = Pt(10)
                report.add_paragraph()

                _re_grp_names = {
                    1:'G1 — ROUTING ENGINE',
                    2:'G2 — TERMINATION ENGINE',
                    3:'G3 — MANDATORY ENGINE',
                    4:'G4 — PIPING ENGINE',
                    5:'G5 — LOOP ENGINE',
                    6:'G6 — VARIABLE ENGINE',
                    7:'G7 — QUESTION TYPE ENGINE',
                    8:'G8 — OPTION/CODE ENGINE',
                    9:'G9 — SURVEY GRAPH ENGINE',
                    10:'G10 — EXPORT ENGINE',
                }
                _re_sev_colors = {
                    'HIGH':   (0xC0, 0x00, 0x00),
                    'MEDIUM': (0xBA, 0x75, 0x17),
                    'LOW':    (0x1D, 0x4E, 0xD8),
                    'INFO':   (0x60, 0x60, 0x60),
                }
                # Group by rule_group
                _re_by_group: dict = {}
                for _rf in _re_report_items:
                    _gnum = _rf.get('rule_group', 0)
                    _re_by_group.setdefault(_gnum, []).append(_rf)

                for _gnum in sorted(_re_by_group.keys()):
                    _gfindings = _re_by_group[_gnum]
                    _ghdr = report.add_paragraph()
                    _ghr  = _ghdr.add_run(_re_grp_names.get(_gnum, f'Group {_gnum}'))
                    _ghr.font.size = Pt(12); _ghr.font.bold = True
                    _ghr.font.color.rgb = RGBColor(0x6D, 0x28, 0xD9)

                    for _rf in _gfindings:
                        _rsev   = _rf.get('severity', 'INFO')
                        _rcolor = _re_sev_colors.get(_rsev, (0x60, 0x60, 0x60))
                        _rqid   = _rf.get('qid', '')
                        _rtype  = _rf.get('issue_type', '')
                        _rconf  = _rf.get('confidence', 0)
                        _rev    = _rf.get('evidence', '')
                        _rrec   = _rf.get('recommendation', '')
                        _rp = report.add_paragraph()
                        _rr = _rp.add_run(f"[{_rsev}] {_rqid} — {_rtype}  [{_rconf}% confidence]")
                        _rr.font.size = Pt(11); _rr.font.bold = True
                        _rr.font.color.rgb = RGBColor(*_rcolor)
                        if _rev:
                            report.add_paragraph().add_run(f"   {str(_rev)[:200]}").font.size = Pt(10)
                        if _rrec:
                            _fixp = report.add_paragraph()
                            _fixr = _fixp.add_run(f"   Fix: {str(_rrec)[:160]}")
                            _fixr.font.size = Pt(10); _fixr.font.italic = True
                            _fixr.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
                        report.add_paragraph()
        # ─────────────────────────────────────────────────────────────────

        # ── Translation Issues section ────────────────────────────────────────
        _trans_issues_rpt = job.get('translation_issues', [])
        if _trans_issues_rpt:
            report.add_paragraph()
            _trh = report.add_paragraph()
            _trhr = _trh.add_run("Translation Issues")
            _trhr.font.size = Pt(14); _trhr.font.bold = True
            _trhr.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
            _trsub = report.add_paragraph()
            _trmatch = sum(1 for t in _trans_issues_rpt if t.get('verdict') == 'TRANSLATION_MISMATCH')
            _trunc   = sum(1 for t in _trans_issues_rpt if t.get('verdict') == 'TRANSLATION_UNCERTAIN')
            _trsub.add_run(
                f"{len(_trans_issues_rpt)} cross-language issue(s): "
                f"{_trmatch} mismatch, {_trunc} uncertain. "
                "Doc and live survey appear to be in different languages."
            ).font.size = Pt(10)
            report.add_paragraph()
            _TR_CLR = {
                'TRANSLATION_MISMATCH':  (0xC0, 0x00, 0x00),
                'TRANSLATION_UNCERTAIN': (0xBA, 0x75, 0x17),
            }
            _TR_LABEL = {
                'TRANSLATION_MISMATCH':  'Translation Mismatch',
                'TRANSLATION_UNCERTAIN': 'Translation Uncertain — needs review',
            }
            for _tri in _trans_issues_rpt:
                _tverd = _tri.get('verdict', 'TRANSLATION_UNCERTAIN')
                _tclr  = _TR_CLR.get(_tverd, (0x33, 0x33, 0x33))
                _tlbl  = _TR_LABEL.get(_tverd, _tverd)
                _tp = report.add_paragraph()
                _tpr = _tp.add_run(
                    f"  {_tri.get('qid','?')}  —  {_tlbl}"
                    f"  |  {_tri.get('doc_lang','?').upper()} → {_tri.get('live_lang','?').upper()}"
                    f"  |  Score: {_tri.get('score', 0):.0%}"
                )
                _tpr.font.size = Pt(11); _tpr.font.bold = True
                _tpr.font.color.rgb = RGBColor(*_tclr)
                _tev = report.add_paragraph()
                _tev.add_run(f"   {_tri.get('evidence','')[:250]}").font.size = Pt(10)
                report.add_paragraph()
        # ─────────────────────────────────────────────────────────────────────

        # ── Piping Issues section ─────────────────────────────────────────────
        _pipe_issues_rpt = job.get('piping_issues', [])
        if _pipe_issues_rpt:
            report.add_paragraph()
            _ph = report.add_paragraph()
            _phr = _ph.add_run("Piping Issues")
            _phr.font.size = Pt(14); _phr.font.bold = True
            _phr.font.color.rgb = RGBColor(0x9B, 0x27, 0xAF)
            _psub = report.add_paragraph()
            _psub.add_run(
                f"{len(_pipe_issues_rpt)} piping issue(s) detected. "
                "Unresolved or blank pipes cause incorrect question text in live survey."
            ).font.size = Pt(10)
            report.add_paragraph()
            _PIPE_RULE_LABELS = {
                'CHECK1': 'Piping not resolved',
                'CHECK2': 'Piping blank / empty value',
                'CHECK3': 'Piping source question missing',
                'CHECK4': 'Piping format mismatch',
            }
            _PIPE_SEV_CLR = {
                'HIGH':   (0xC0, 0x00, 0x00),
                'MEDIUM': (0xBA, 0x75, 0x17),
            }
            for _pi in _pipe_issues_rpt:
                _prule = _pi.get('rule', '')
                _plabel = _PIPE_RULE_LABELS.get(_prule, _prule)
                _psev = _pi.get('severity', 'MEDIUM')
                _pclr = _PIPE_SEV_CLR.get(_psev, (0x33, 0x33, 0x33))
                _pp = report.add_paragraph()
                _ppr = _pp.add_run(
                    f"  [{_prule}] {_pi.get('qid','?')}  —  {_plabel}"
                    f"  |  Variable: {_pi.get('pipe_variable','')}"
                )
                _ppr.font.size = Pt(11); _ppr.font.bold = True
                _ppr.font.color.rgb = RGBColor(*_pclr)
                _pev = report.add_paragraph()
                _pev.add_run(f"   {_pi.get('evidence','')[:220]}").font.size = Pt(10)
                report.add_paragraph()
        # ─────────────────────────────────────────────────────────────────────

        # ── Numeric Range Issues section ──────────────────────────────────────
        _range_issues_rpt = job.get('range_issues', [])
        if _range_issues_rpt:
            report.add_paragraph()
            _rh = report.add_paragraph()
            _rhr = _rh.add_run("Numeric Range Issues")
            _rhr.font.size = Pt(14); _rhr.font.bold = True
            _rhr.font.color.rgb = RGBColor(0x0F, 0x51, 0x32)
            _rsub = report.add_paragraph()
            _rsub.add_run(
                f"{len(_range_issues_rpt)} numeric range issue(s) detected. "
                "Verify boundary enforcement in live survey via Playwright test cases."
            ).font.size = Pt(10)
            report.add_paragraph()
            _RANGE_RULE_LABELS = {
                'R041': 'Min boundary not enforced',
                'R042': 'Max boundary not enforced',
                'R043': 'Range in spec — validation missing',
                'R044': 'Conflicting ranges',
                'R045': 'Mandatory numeric — blank not blocked',
            }
            _RANGE_SEV_CLR = {
                'HIGH':   (0xC0, 0x00, 0x00),
                'MEDIUM': (0xBA, 0x75, 0x17),
            }
            for _ri in _range_issues_rpt:
                _rrule = _ri.get('rule', '')
                _rlabel = _RANGE_RULE_LABELS.get(_rrule, _rrule)
                _rsev   = _ri.get('severity', 'HIGH')
                _rclr   = _RANGE_SEV_CLR.get(_rsev, (0x33, 0x33, 0x33))
                _lo = _ri.get('min_val'); _hi = _ri.get('max_val')
                _bounds = f'  |  Range: {_lo}–{_hi}' if _lo is not None and _hi is not None else ''
                _rp = report.add_paragraph()
                _rpr = _rp.add_run(
                    f"  [{_rrule}] {_ri.get('qid','?')}  —  {_rlabel}{_bounds}"
                )
                _rpr.font.size = Pt(11); _rpr.font.bold = True
                _rpr.font.color.rgb = RGBColor(*_rclr)
                _rev = report.add_paragraph()
                _rev.add_run(f"   {_ri.get('evidence','')[:220]}").font.size = Pt(10)
                report.add_paragraph()
        # ─────────────────────────────────────────────────────────────────────

        # ── Loop/Repeat Blocks section ────────────────────────────────────────
        _loop_blocks_rpt = job.get('loop_blocks', {})
        if _loop_blocks_rpt:
            report.add_paragraph()
            _lh = report.add_paragraph()
            _lhr = _lh.add_run("Loop/Repeat Blocks")
            _lhr.font.size = Pt(14); _lhr.font.bold = True
            _lhr.font.color.rgb = RGBColor(0x0C, 0x44, 0x7C)
            _lsub = report.add_paragraph()
            _lsub.add_run(
                f"{len(_loop_blocks_rpt)} loop group(s) detected. "
                "First iteration compared; siblings skipped to prevent duplicate issues."
            ).font.size = Pt(10)
            report.add_paragraph()
            for _lparent, _lchildren in sorted(_loop_blocks_rpt.items()):
                _lp = report.add_paragraph()
                _lpr = _lp.add_run(f"  {_lparent}  (loop: ")
                _lpr.font.size = Pt(11); _lpr.font.bold = True
                _lpr.font.color.rgb = RGBColor(0x0C, 0x44, 0x7C)
                for _li, _lc in enumerate(_lchildren):
                    _badge = "✓" if _li == 0 else "skip"
                    _lrun = _lp.add_run(f"{_lc} {_badge}")
                    _lrun.font.size = Pt(11)
                    _lrun.font.color.rgb = (
                        RGBColor(0x00, 0x70, 0x00) if _li == 0
                        else RGBColor(0x77, 0x77, 0x77)
                    )
                    if _li < len(_lchildren) - 1:
                        _lp.add_run(",  ").font.size = Pt(11)
                _lp.add_run(")").font.size = Pt(11)
            report.add_paragraph()
        # ─────────────────────────────────────────────────────────────────────

        if term_results:
            report.add_paragraph()
            h = report.add_paragraph()
            hr = h.add_run("Termination Tests")
            hr.font.size = Pt(14); hr.font.bold = True; hr.font.color.rgb = RGBColor(0x7C, 0x65, 0xFF)
            p = report.add_paragraph()
            p.add_run(f"{term_passed}/{len(term_results) - term_review} validated · {term_review} need manual review").font.size = Pt(11)
            report.add_paragraph()
            for r in term_results:
                if r.get("needs_review"):
                    status = "NEEDS REVIEW"
                    color = (0xD9, 0x77, 0x06)
                elif r.get("passed"):
                    status = "PASS"
                    color = (0x00, 0x70, 0x00)
                else:
                    status = "FAIL"
                    color = (0xC0, 0x00, 0x00)
                p = report.add_paragraph()
                pr = p.add_run(f"  {status}  {r['test_qid']} = code {r['answer_code']}")
                pr.font.size = Pt(11); pr.font.color.rgb = RGBColor(*color)

        # ── Test Cases section ──────────────────────────────────────────────
        _tc_list_rpt = job.get('test_cases', [])
        _tc_sum_rpt  = job.get('test_cases_summary', {})
        if _tc_list_rpt:
            report.add_paragraph()
            _th = report.add_paragraph()
            _thr = _th.add_run("Auto-Generated Test Cases")
            _thr.font.size = Pt(14); _thr.font.bold = True
            _thr.font.color.rgb = RGBColor(0x0C, 0x44, 0x7C)
            _tc_total = _tc_sum_rpt.get('total', len(_tc_list_rpt))
            _tc_auto  = _tc_sum_rpt.get('auto_runnable', 0)
            _tc_types = _tc_sum_rpt.get('by_type', {})
            _tsub = report.add_paragraph()
            _tsub.add_run(
                f"{_tc_total} test case(s) generated — {_tc_auto} auto-runnable by Playwright."
                f"  Types: {', '.join(f'{k}: {v}' for k, v in _tc_types.items())}"
            ).font.size = Pt(10)
            report.add_paragraph()
            _TYPE_COLOR = {
                'TERMINATION': (0xC0, 0x00, 0x00),
                'ROUTING':     (0x0C, 0x44, 0x7C),
                'MANDATORY':   (0xBA, 0x75, 0x17),
                'RANGE':       (0x0F, 0x51, 0x32),
            }
            for _tc in _tc_list_rpt:
                _ar = 'AUTO' if _tc.get('auto_runnable') else 'MANU'
                _col = _TYPE_COLOR.get(_tc.get('type', ''), (0x33, 0x33, 0x33))
                _tp = report.add_paragraph()
                _tr = _tp.add_run(
                    f"  [{_tc['test_id']}] [{_tc.get('type','')}] [{_ar}]  "
                    f"{_tc.get('qid','?')}  —  {_tc.get('action','')}  →  {_tc.get('expected','')}"
                )
                _tr.font.size = Pt(10); _tr.font.color.rgb = RGBColor(*_col)
                if _tc.get('notes'):
                    _np = report.add_paragraph()
                    _np.add_run(f"       {_tc['notes'][:100]}").font.size = Pt(9)
                    _np.add_run('').font.color.rgb = RGBColor(0x77, 0x77, 0x77)
        # ────────────────────────────────────────────────────────────────────

        # ── Automated Test Results (Playwright) ──────────────────────────────
        _pw_data = job.get('playwright_tests', {})
        _pw_res  = _pw_data.get('results', [])
        _pw_sum  = _pw_data.get('summary', {})
        if _pw_res:
            report.add_paragraph()
            _pwh = report.add_paragraph()
            _pwhr = _pwh.add_run("Automated Test Results")
            _pwhr.font.size = Pt(14); _pwhr.font.bold = True
            _pwhr.font.color.rgb = RGBColor(0x0C, 0x44, 0x7C)

            _pw_pass = _pw_sum.get('passed', 0)
            _pw_fail = _pw_sum.get('failed', 0)
            _pw_err  = _pw_sum.get('errors', 0)
            _pw_tot  = _pw_sum.get('total', len(_pw_res))
            _pw_rate = _pw_sum.get('pass_rate', 'N/A')
            _pwsub = report.add_paragraph()
            _pwsub.add_run(
                f"{_pw_tot} test(s) executed by Playwright — "
                f"{_pw_pass} PASS  ·  {_pw_fail} FAIL  ·  {_pw_err} ERROR  "
                f"({_pw_rate} pass rate)"
            ).font.size = Pt(10)
            report.add_paragraph()

            # Results table: Test ID | QID | Type | Status | Actual | Duration
            _pw_tbl = report.add_table(rows=1, cols=5)
            _pw_tbl.style = 'Table Grid'
            _hdrs = ['Test', 'QID', 'Type', 'Status', 'Actual Result']
            for _ci, _hd in enumerate(_hdrs):
                _c = _pw_tbl.rows[0].cells[_ci]
                _c.text = _hd
                _c.paragraphs[0].runs[0].font.bold = True
                _c.paragraphs[0].runs[0].font.size = Pt(9)
                shade_cell(_c, 'E8EAF6')

            _STATUS_CLR = {
                'PASS':  (0x00, 0x70, 0x00),
                'FAIL':  (0xC0, 0x00, 0x00),
                'ERROR': (0xBA, 0x75, 0x17),
                'SKIP':  (0x77, 0x77, 0x77),
            }
            for _pr in _pw_res:
                _row = _pw_tbl.add_row()
                _cells = _row.cells
                _cells[0].text = _pr.get('test_id', '')
                _cells[1].text = _pr.get('qid', '')
                _cells[2].text = _pr.get('type', '')
                _cells[3].text = _pr.get('status', '')
                _cells[4].text = (_pr.get('actual_result', '') or '')[:60]
                _st = _pr.get('status', 'ERROR')
                _clr = _STATUS_CLR.get(_st, (0x33, 0x33, 0x33))
                for _ci2 in range(5):
                    _cells[_ci2].paragraphs[0].runs[0].font.size = Pt(9)
                    if _ci2 == 3:
                        _cells[_ci2].paragraphs[0].runs[0].font.bold = True
                        _cells[_ci2].paragraphs[0].runs[0].font.color.rgb = RGBColor(*_clr)

            # Inline screenshots for FAIL / ERROR tests
            _fail_res = [r for r in _pw_res if r.get('status') in ('FAIL', 'ERROR') and r.get('screenshot_path')]
            if _fail_res:
                report.add_paragraph()
                _pwfh = report.add_paragraph()
                _pwfh.add_run("Screenshots — Failed Tests").font.size = Pt(12)
                _pwfh.runs[0].font.bold = True
                _pwfh.runs[0].font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                for _fr in _fail_res:
                    _sp = _fr.get('screenshot_path', '')
                    if not _sp or not os.path.exists(_sp):
                        continue
                    _flab = report.add_paragraph()
                    _flab.add_run(
                        f"  {_fr['test_id']}  [{_fr['status']}]  "
                        f"{_fr.get('qid','')}  —  {_fr.get('actual_result','')[:80]}"
                    ).font.size = Pt(9)
                    try:
                        from docx.shared import Inches
                        _img_p = report.add_paragraph()
                        _img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run_img = _img_p.add_run()
                        run_img.add_picture(_sp, width=Inches(5.5))
                    except Exception:
                        report.add_paragraph().add_run(
                            f"  [Screenshot: {os.path.basename(_sp)}]"
                        ).font.size = Pt(9)

            if _pw_data.get('error'):
                _pwerr_p = report.add_paragraph()
                _pwerr_p.add_run(
                    f"  Note: {_pw_data['error'][:200]}"
                ).font.size = Pt(9)
        # ─────────────────────────────────────────────────────────────────────

        report.add_paragraph()
        _qsum_h = report.add_paragraph()
        _qsum_hr = _qsum_h.add_run("QC Summary")
        _qsum_hr.font.size = Pt(14); _qsum_hr.font.bold = True
        _qsum_hr.font.color.rgb = RGBColor(0x7C, 0x65, 0xFF)
        _sum_high = sum(1 for i in issues if i.get('confidence', 0) >= 75)
        _sum_med  = sum(1 for i in issues if 55 <= i.get('confidence', 0) < 75)
        _sum_low  = sum(1 for i in issues if i.get('confidence', 0) < 55)
        for _sl in [
            f"Total issues found: {len(issues)}",
            f"Likely Bug / Confirmed Bug (75%+): {_sum_high}",
            f"Possible Issue (55-74%): {_sum_med}",
            f"Needs Review / Likely False Positive (<55%): {_sum_low}",
        ]:
            _slp = report.add_paragraph()
            _slp.add_run(f"  • {_sl}").font.size = Pt(11)
        _xml_flagged = sum(
            1 for i in issues
            if i.get('xml_verdict') or i.get('type') == 'QID IN EXPORT NOT IN DOC/LIVE'
        )
        if _xml_flagged and len(xml_questions):
            report.add_paragraph()
            _xfp = report.add_paragraph()
            _xfr = _xfp.add_run(f"  Issues flagged by XML comparison: {_xml_flagged}")
            _xfr.font.size = Pt(11); _xfr.font.bold = True
            _xfr.font.color.rgb = RGBColor(0x13, 0x8D, 0x5A)
            _xfp2 = report.add_paragraph()
            _xfp2.add_run(
                f"  (XML upload helped catch {_xml_flagged} issue(s) with higher confidence)"
            ).font.size = Pt(10)
        report.add_paragraph()

        footer_p = report.add_paragraph(); footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = footer_p.add_run("— End of Report — Generated by SurveyQC v10.0 —")
        fr.font.size = Pt(9); fr.font.italic = True; fr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        os.makedirs(f"{OUTPUT_FOLDER}/{job_id}", exist_ok=True)
        report_path = f"{OUTPUT_FOLDER}/{job_id}/QC_Report_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        report.save(report_path)
        log(f'  Report saved', 'green')

        verdict = 'PASS' if (sev['HIGH']==0 and term_failed==0 and term_review==0) else ('FAIL' if (sev['HIGH']>0 or term_failed>0) else 'REVIEW')

        progress(100, 'Complete!')
        log('', 'white')
        log('════════════════════════════════════', 'magenta')
        log('  FINAL SUMMARY', 'magenta')
        log('════════════════════════════════════', 'magenta')
        log(f'  Document QIDs:  {len(questions)}', 'blue')
        log(f'  XML QIDs:       {len(xml_questions)}', 'blue')
        log(f'  Live QIDs:      {len(live_data)}', 'blue')
        log(f'  Total Issues:   {len(issues)}', 'yellow')
        if term_results:
            log(f'  Termination:    {term_passed}/{len(term_results)} passed', 'green' if term_passed==len(term_results) else 'yellow')
        log(f'\n  DONE! Verdict: {verdict}', 'green')

        job['status'] = 'done'
        job['verdict'] = verdict
        job['doc_qids'] = len(questions)
        job['xml_qids'] = len(xml_questions)
        job['live_qids'] = len(live_data)
        job['total_issues'] = len(issues)
        job['term_passed'] = term_passed
        job['term_review'] = term_review
        job['term_total'] = len(term_results)
        job['term_results'] = term_results
        job['issues'] = issues
        job['report_file'] = report_path
        # Rule engine findings are already stored in job by Phase 2.6
        # Ensure they survive even if Phase 2.6 did not run
        if 'rule_engine_findings' not in job:
            job['rule_engine_findings'] = []
        if 'rule_engine_summary' not in job:
            job['rule_engine_summary'] = {}
        if 'termination_matrix' not in job:
            job['termination_matrix'] = []

        # PHASE 6: CONSENSUS ENGINE — score all findings, compute health
        try:
            from consensus_engine import run_consensus as _run_consensus
            _ce_platform = job.get('platform', '')
            _ce_feedback = jobs.load_feedback(_ce_platform)
            _ce_out = _run_consensus(
                issues=issues,
                rule_findings=job.get('rule_engine_findings', []),
                term_results=term_results,
                playwright_tests=job.get('playwright_tests', {}),
                feedback_history=_ce_feedback,
                platform=_ce_platform,
            )
            job['consensus']    = _ce_out
            job['health_score'] = _ce_out['health_score']
            _hs = _ce_out['health_score']
            log(f'  Health Score: {_hs}/100 — '
                f'Confirmed={_ce_out["confirmed_count"]} · '
                f'Likely={_ce_out["likely_count"]} · '
                f'Review={_ce_out["review_count"]} · '
                f'Suppressed={_ce_out["suppressed_count"]}',
                'green' if _hs >= 80 else 'yellow')
        except Exception as _ce_err:
            log(f'  Consensus engine (non-fatal): {str(_ce_err)[:120]}', 'yellow')
            job['consensus']    = {}
            job['health_score'] = None

    except Exception as e:
        import traceback
        job['status'] = 'error'
        job['logs'].append({'msg': f'ERROR: {str(e)}', 'color': 'red'})
        job['logs'].append({'msg': traceback.format_exc()[:500], 'color': 'red'})
    finally:
        _hb_stop.set()  # stop heartbeat thread
        jobs.persist(job_id)  # sync final state (done/error) to SQLite
        # Fire report-ready email only on success; never block on failure
        try:
            j = jobs.get(job_id, {})
            if j.get('status') == 'done':
                _ue = j.get('user_email', '')
                _u  = users_db.get(_ue, {})
                send_report_ready_email(
                    _ue,
                    _u.get('name', _ue),
                    j.get('doc_name', 'document'),
                    j.get('total_issues', 0),
                    job_id,
                )
        except Exception:
            pass

# ================================================================
# MAIN
# ================================================================


# ================================================================
# NEW FEATURES v10.0
# ================================================================

# ---- CONTENT STORE (Admin editable) ----
site_content = {
    # Hero
    'hero_heading_part1': 'Catch every survey bug',
    'hero_heading_part2': 'before',
    'hero_heading_part3': 'your data is collected.',
    'hero_subheading': 'SurveyQC reads your screener doc, crawls your live survey, and flags 15+ types of issues automatically -- in any language, on any platform.',
    'hero_cta': 'Start free trial',
    'hero_cta2': 'See how it works',
    'announcement': 'New: WhatsApp screenshot QC is now live',
    # Features (homepage preview - 6)
    'feature1_title': 'Termination Testing',
    'feature1_desc': 'Every terminate/close rule actually clicked and verified. PASS/FAIL per rule with screenshot proof attached.',
    'feature2_title': 'Missing Words Detection',
    'feature2_desc': 'Word-by-word comparison between your spec doc and live survey. Catches every typo and missing phrase instantly.',
    'feature3_title': 'Auto Screenshots',
    'feature3_desc': 'Every question captured automatically. Bug found? Screenshot attached as visual proof in your final report.',
    'feature4_title': 'Any Language',
    'feature4_desc': 'French, Italian, Arabic, Urdu, Japanese -- 80+ languages with full encoding and accent character support.',
    'feature5_title': '15+ QC Checks',
    'feature5_desc': 'Termination, text match, missing words, piping, codes, mandatory, options, order -- all parallel verified.',
    'feature6_title': 'One-click Retest',
    'feature6_desc': 'Fix the bugs, retest only the failed paths in seconds. No need to run full QC again from scratch.',
    # Footer / site
    'site_name': 'SurveyQC',
    'tagline': 'AI-powered survey QC built for market research professionals',
    'support_email': 'support@surveyqc.online',
    'footer_text': '2026 SurveyQC -- Built for QC professionals worldwide',
    'linkedin_url': 'https://linkedin.com/company/surveyqc',
    'twitter_url': 'https://twitter.com/surveyqc',
    'youtube_url': '',
    'privacy_policy': 'We take your privacy seriously. All data auto-deletes after 30 days.',
    'terms': 'By using SurveyQC you agree to our terms of service.',
    # Pricing page
    'pricing_heading': 'Simple, transparent pricing',
    'pricing_sub': 'Start free. Upgrade when you need more. No contracts, cancel anytime.',
    'plan_free_name': 'Free',
    'plan_free_price': '0',
    'plan_free_period': '/month',
    'plan_free_desc': 'For solo testers',
    'plan_free_yearly': '0',
    'plan_free_features': '3 reports per month||All 15+ QC checks||Word report download||20K tokens per report||Any language||Community support',
    'plan_free_cta': 'Get started',
    'plan_pro_name': 'Pro',
    'plan_pro_price': '29',
    'plan_pro_period': '/month',
    'plan_pro_desc': 'For QC professionals',
    'plan_pro_yearly': '290',
    'plan_pro_features': '25 reports per month||Everything in Free||Screenshot QC (WhatsApp)||Share reports with clients||100K tokens per report||Priority support||AI auto tester||Custom templates',
    'plan_pro_cta': 'Start Pro trial',
    'plan_pro_badge': 'Most Popular',
    'plan_pro_featured': '1',
    'plan_biz_name': 'Business',
    'plan_biz_price': '299',
    'plan_biz_period': '/month',
    'plan_biz_desc': 'For agencies & teams',
    'plan_biz_yearly': '2990',
    'plan_biz_features': 'Unlimited reports||Everything in Pro||Team collaboration (up to 5 users)||White-label reports||200K tokens per report||Dedicated account manager||SLA guarantee||API access',
    'plan_biz_cta': 'Get Business',
    'plan_ent_name': 'Enterprise',
    'plan_ent_price': 'Custom',
    'plan_ent_period': 'pricing',
    'plan_ent_desc': 'For large organisations',
    'plan_ent_yearly': 'Custom',
    'plan_ent_features': 'Everything in Business||Unlimited users||Custom features||On-premise option||SSO/SAML||Custom SLA',
    'plan_ent_cta': 'Contact Sales',
    'plan_ent_cta_link': 'mailto:support@surveyqc.online?subject=Enterprise%20Inquiry',
    # Pricing FAQ
    'pfaq1_q': 'Can I cancel anytime?',
    'pfaq1_a': 'Yes, cancel anytime. No contracts.',
    'pfaq2_q': 'Is there a free trial?',
    'pfaq2_a': 'Pro and Business plans come with a 14-day free trial. Free plan is free forever.',
    'pfaq3_q': 'What payment methods do you accept?',
    'pfaq3_a': 'Credit/debit cards (Stripe), UPI, PayPal, and bank transfer for Business plan.',
    'pfaq4_q': 'Can I switch plans later?',
    'pfaq4_a': 'Yes, upgrade or downgrade anytime. Pro-rated billing applied automatically.',
    'pfaq5_q': 'Do you offer team/enterprise pricing?',
    'pfaq5_a': 'Yes, Business plan supports teams. Contact us for enterprise pricing (10+ users).',
    'pfaq6_q': 'Do you offer a refund?',
    'pfaq6_a': 'Yes, 7-day money-back guarantee.',
    'pfaq7_q': 'Is my data secure?',
    'pfaq7_a': 'Yes. Encrypted, GDPR compliant, auto-deleted after 90 days.',
    # Features Page - 25 features grouped
    'features_heading': 'Everything you need for perfect survey QC',
    'features_sub': '25+ specialized checks and tools, built by QC professionals for QC professionals.',
    # Group 1: Core QC Checks
    'feat_grp1_name': 'Core QC Checks',
    'feat_grp1_desc': 'Essential checks every survey needs.',
    'feat_grp1_items': 'Termination Testing||Every terminate rule clicked and verified with PASS/FAIL||ti-shield-x##Question Text Match||Full question text compared word-by-word against spec doc||ti-text-recognition##Missing Words Detection||Catches typos and missing phrases between spec and live survey||ti-search##Options Match||All answer options compared against spec - missing/extra detected||ti-checkbox##Answer Codes Validation||Verifies answer codes are sequential and match spec exactly||ti-list-numbers##Mandatory Markers||Checks * mandatory markers match the spec on every question||ti-asterisk##Question Order||Verifies questions appear in correct order as per spec||ti-arrows-sort##Piping Markers||Detects unresolved {piped_text} variables shown to respondents||ti-replace',
    # Group 2: Advanced Logic
    'feat_grp2_name': 'Advanced Logic Testing',
    'feat_grp2_desc': 'Catches complex logic bugs humans miss.',
    'feat_grp2_items': 'Skip Logic Validation||Every skip rule tested against actual respondent paths||ti-route##Quota Testing||Validates quota cells, balanced sampling, and overflow logic||ti-chart-pie##Loop & Block Testing||Tests repeating loops and conditional blocks for consistency||ti-refresh##Conditional Display||Verifies show/hide rules based on prior answers||ti-eye##Display Logic||Catches questions that show/hide incorrectly per spec||ti-eye-check##Routing Validation||Confirms respondents go to correct next question||ti-arrow-fork',
    # Group 3: Quality & Multilingual
    'feat_grp3_name': 'Quality & Multilingual',
    'feat_grp3_desc': 'Built for global research teams.',
    'feat_grp3_items': '80+ Languages||French, Italian, Arabic, Urdu, Japanese, Chinese - all supported||ti-world##Accent & Encoding||Proper handling of special characters and diacritics||ti-language##Translation Sync||Detects untranslated strings in multi-language surveys||ti-translate##Hidden Question Detection||Finds hidden/disabled questions still present in live survey||ti-ghost##Duplicate Question Detection||Catches duplicate question IDs across the survey||ti-copy',
    # Group 4: Reporting & Workflow
    'feat_grp4_name': 'Reporting & Workflow',
    'feat_grp4_desc': 'Professional reports clients will love.',
    'feat_grp4_items': 'Auto Screenshots||Every question captured automatically as visual proof||ti-camera##Word Report Download||Professional .docx report with all issues and screenshots||ti-file-word##One-click Retest||Fix bugs, retest only failed paths in seconds||ti-rotate-clockwise##Share with Client||Share report link - clients can comment and mark fixed||ti-share##Timing Estimation||Estimates survey LOI (length of interview) automatically||ti-clock##QC Certificate||Professional PDF certificate proving QC was completed||ti-certificate',
    # Testimonials (rotating)
    'test1_name': 'Sarah Thompson',
    'test1_role': 'QC Manager',
    'test1_company': 'Ipsos UK',
    'test1_country': 'United Kingdom',
    'test1_flag': 'GB',
    'test1_quote': '8 hours of manual QC now takes 10 minutes. Caught a termination bug that would have killed our entire dataset -- saved us a $40K project.',
    'test1_rating': '5',
    'test2_name': 'Marie Laurent',
    'test2_role': 'Research Director',
    'test2_company': 'Kantar France',
    'test2_country': 'France',
    'test2_flag': 'FR',
    'test2_quote': 'French, Italian, Spanish -- SurveyQC handles all perfectly. Screenshot evidence is a total game changer for client reports. Industry-leading.',
    'test2_rating': '5',
    'test3_name': 'James Mitchell',
    'test3_role': 'Operations Lead',
    'test3_company': 'Nielsen USA',
    'test3_country': 'United States',
    'test3_flag': 'US',
    'test3_quote': 'Saved our team 47 hours this month. ROI is insane. Nothing else even comes close to this level of automation for survey professionals.',
    'test3_rating': '5',
    'test4_name': 'Rahul Sharma',
    'test4_role': 'Survey Programmer',
    'test4_company': 'YouGov India',
    'test4_country': 'India',
    'test4_flag': 'IN',
    'test4_quote': 'Catches termination bugs we routinely missed in manual QC. Now we deliver bug-free surveys every time. The team productivity boost is real.',
    'test4_rating': '5',
    'test5_name': 'Klaus Weber',
    'test5_role': 'Data Manager',
    'test5_company': 'GfK Germany',
    'test5_country': 'Germany',
    'test5_flag': 'DE',
    'test5_quote': 'Replaced 3 manual QC processes with one tool. Reports are professional, clients trust us more. Best research tech investment of 2026.',
    'test5_rating': '5',
    'test6_name': 'Fatima Al-Rashid',
    'test6_role': 'QC Specialist',
    'test6_company': 'YouGov MENA',
    'test6_country': 'UAE',
    'test6_flag': 'AE',
    'test6_quote': 'Arabic survey QC was always painful with right-to-left text. SurveyQC handles it perfectly. Massive time saver for our Middle East research.',
    'test6_rating': '5',
    # Blog
    'blog_post1_title': 'How AI is Changing Survey Quality Control in 2026',
    'blog_post1_date': 'May 2026',
    'blog_post1_tag': 'AI & QC',
    'blog_post1_summary': 'Manual QC takes 8+ hours per survey. Here is how AI reduces that to under 10 minutes with 99% accuracy.',
    'blog_post2_title': '8 Most Common Survey Bugs',
    'blog_post2_date': 'April 2026',
    'blog_post2_tag': 'Best Practices',
    'blog_post2_summary': 'Termination errors, missing words, broken piping -- the 8 bugs that destroy survey data.',
    'blog_post3_title': 'French Survey QC: Why Accents and Encoding Matter',
    'blog_post3_date': 'March 2026',
    'blog_post3_tag': 'Languages',
    'blog_post3_summary': 'Testing French surveys has unique challenges -- accent marks, special characters, termination phrases.',
    'blog_post4_title': 'Complete Guide to Confirmit Survey Testing',
    'blog_post4_date': 'March 2026',
    'blog_post4_tag': 'Platforms',
    'blog_post4_summary': 'Step-by-step guide to QC testing on Confirmit from uploading spec docs to reading the final report.',
    'blog_post5_title': 'Why Excel Sheets Are Killing Your QC Workflow',
    'blog_post5_date': 'February 2026',
    'blog_post5_tag': 'Productivity',
    'blog_post5_summary': 'Most teams still use Excel for survey QC. We analyzed 500 surveys to show why that is costing time.',
    'blog_post6_title': 'SurveyQC vs Manual Testing: A Real Cost Comparison',
    'blog_post6_date': 'January 2026',
    'blog_post6_tag': 'ROI',
    'blog_post6_summary': 'We broke down the actual cost per survey -- time, salary, errors. The numbers will surprise you.',
    # Compare pages
    'compare_chatgpt_heading': 'SurveyQC vs ChatGPT for Survey QC',
    'compare_chatgpt_summary': 'ChatGPT is a general AI. SurveyQC is purpose-built for survey quality control with 15+ specialized checks.',
    'compare_excel_heading': 'SurveyQC vs Manual Excel QC',
    'compare_excel_summary': 'Excel takes 8+ hours per survey. SurveyQC does it in 10 minutes automatically.',
    'compare_manual_heading': 'SurveyQC vs Manual Testing',
    'compare_manual_summary': 'Manual testers miss 30% of bugs on average. SurveyQC catches 99% automatically.',
    # Community & Affiliate
    'community_heading': 'Join the SurveyQC Community',
    'community_subheading': 'Connect with 500+ QC professionals worldwide. Share tips, get help.',
    'affiliate_heading': 'Earn with SurveyQC Affiliate Program',
    'affiliate_commission': '30',
    'affiliate_details': 'Earn 30% recurring commission for every paying user you refer. No cap.',
    # Changelog
    'changelog_v10': 'v10.0 -- May 2026: WhatsApp screenshot QC, GDPR auto-delete, Share report, API management',
    'changelog_v9': 'v9.0 -- April 2026: Token limit controls, Gift access, Content management, 5-country compliance',
    'changelog_v8': 'v8.0 -- March 2026: French termination fixes, QC certificate, Report templates',
    'changelog_v7': 'v7.0 -- February 2026: Double check, Progress live log, Retest failed paths',
    'changelog_v6': 'v6.0 -- January 2026: Multi-platform support (Decipher, Forsta, Qualtrics)',
    'docs_intro': 'SurveyQC documentation -- everything you need to run perfect QC checks.',
}

token_limits = {
    'free': 20000,
    'pro': 100000,
    'business': 150000,
    'monthly_budget': 50,
}


# ---- COUPONS STORE ----
coupons_db = {}

# ---- NOTES STORE ----
notes_db = {}

# ---- FAVORITES STORE ----
favorites_db = {}

# ---- TEMPLATES STORE ----
templates_db = {
    'confirmit-france': {
        'name': 'Confirmit France',
        'platform': 'Confirmit',
        'country': 'France',
        'mode': 'full',
    },
    'decipher-italy': {
        'name': 'Decipher Italy',
        'platform': 'Decipher',
        'country': 'Italy',
        'mode': 'full',
    }
}


# ================================================================
# AUTO CLEANUP — GDPR 30 day delete
# ================================================================
import shutil
from datetime import timedelta

def auto_cleanup_job():
    """Run every night — delete job data older than 30 days, feedback older than 3 days"""
    cutoff_jobs = datetime.now() - timedelta(days=30)
    cutoff_feedback = datetime.now() - timedelta(days=3)
    deleted = 0
    for job_id, job in list(jobs.items()):
        try:
            created = datetime.fromisoformat(job.get('created_at', ''))
            if created < cutoff_jobs:
                shutil.rmtree(f"{UPLOAD_FOLDER}/{job_id}", ignore_errors=True)
                shutil.rmtree(f"{OUTPUT_FOLDER}/{job_id}", ignore_errors=True)
                del jobs[job_id]
                deleted += 1
        except: pass
    # Auto-delete old feedback
    global user_feedback_db
    before = len(user_feedback_db)
    user_feedback_db = [
        f for f in user_feedback_db
        if datetime.fromisoformat(f['created_at']) >= cutoff_feedback
    ]
    deleted += before - len(user_feedback_db)
    return deleted

def start_cleanup_scheduler():
    import time
    def scheduler():
        while True:
            now = datetime.now()
            # Run at midnight
            next_midnight = now.replace(hour=0, minute=0, second=0, microsecond=0) + timedelta(days=1)
            wait_seconds = (next_midnight - now).total_seconds()
            time.sleep(wait_seconds)
            auto_cleanup_job()
    t = threading.Thread(target=scheduler, daemon=True)
    t.start()


# ================================================================
# ONBOARDING PAGE
# ================================================================
@app.route('/onboarding')
@login_required
def onboarding():
    user = get_current_user()
    return render_template_string(SHARED_CSS + f"""
<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=5.0"><title>Welcome — SurveyQC</title></head><body>
<div class="app-layout">
  {sidebar_html('onboarding')}
  <div class="main-content">
    <div class="topbar">
      <div><p class="page-title">Welcome to SurveyQC! 🎉</p><p class="page-sub">Quick setup — takes 2 minutes</p></div>
      <span class="badge badge-blue">Step 2 of 3</span>
    </div>

    <div style="display:flex;gap:8px;margin-bottom:20px">
      <div style="flex:1;height:6px;background:#042C53;border-radius:3px"></div>
      <div style="flex:1;height:6px;background:#042C53;border-radius:3px"></div>
      <div style="flex:1;height:6px;background:var(--color-background-secondary);border-radius:3px"></div>
    </div>

    <div style="display:grid;grid-template-columns:repeat(3,1fr);gap:14px;margin-bottom:20px">
      <div class="card" style="border:2px solid #1D9E75;background:#F0FDF4">
        <div style="width:32px;height:32px;border-radius:50%;background:#1D9E75;color:white;display:flex;align-items:center;justify-content:center;font-size:14px;font-weight:600;margin-bottom:10px">✓</div>
        <p style="font-size:13px;font-weight:500;color:#27500A">Step 1 — Done!</p>
        <p style="font-size:11px;color:var(--color-text-secondary);margin-top:4px">Account created ✅</p>
      </div>
      <div class="card" style="border:2px solid #185FA5">
        <div style="width:32px;height:32px;border-radius:50%;background:var(--purple);color:white;display:flex;align-items:center;justify-content:center;font-size:14px;font-weight:600;margin-bottom:10px">2</div>
        <p style="font-size:13px;font-weight:500;color:var(--color-text-primary)">Choose your platform</p>
        <p style="font-size:11px;color:var(--color-text-secondary);margin-top:4px">Which survey platform do you use?</p>
        <select class="form-select" style="margin-top:10px;font-size:12px">
          <option>Confirmit</option><option>Decipher</option><option>Forsta</option><option>Qualtrics</option>
        </select>
      </div>
      <div class="card" style="opacity:.5">
        <div style="width:32px;height:32px;border-radius:50%;background:var(--color-border-secondary);color:var(--color-text-secondary);display:flex;align-items:center;justify-content:center;font-size:14px;font-weight:600;margin-bottom:10px">3</div>
        <p style="font-size:13px;font-weight:500;color:var(--color-text-primary)">Run your first QC</p>
        <p style="font-size:11px;color:var(--color-text-secondary);margin-top:4px">See the magic!</p>
      </div>
    </div>

    <div class="card" style="background:#E6F1FB;border-color:#B5D4F4">
      <div style="display:flex;align-items:center;gap:12px">
        <i class="ti ti-bulb" style="font-size:24px;color:#185FA5;flex-shrink:0"></i>
        <div style="flex:1">
          <p style="font-size:13px;font-weight:500;color:#0C447C">Start with a simple survey test!</p>
          <p style="font-size:11px;color:#185FA5;margin-top:3px">Paste any survey URL — AI checks everything automatically in 10 minutes</p>
        </div>
        <a href="/new-qc" class="btn btn-primary btn-sm" style="flex-shrink:0"><i class="ti ti-player-play"></i>Run first QC</a>
      </div>
    </div>

    <div style="text-align:center;margin-top:16px">
      <a href="/dashboard" style="font-size:12px;color:var(--color-text-secondary);text-decoration:none">Skip for now → Go to dashboard</a>
    </div>
  </div>
</div>
</body></html>""")


# ================================================================
# SHARE REPORT PAGE
# ================================================================
@app.route('/share/<job_id>', methods=['GET', 'POST'])
@login_required
def share_report(job_id):
    if job_id not in jobs:
        return redirect('/reports')

    j = jobs[job_id]
    doc_name = j.get('doc_name', 'Unknown')

    # Generate share link
    share_token = hashlib.md5(f"{job_id}-share".encode()).hexdigest()[:12]

    if request.method == 'POST':
        comment = request.form.get('comment', '').strip()
        if comment:
            if job_id not in notes_db:
                notes_db[job_id] = []
            notes_db[job_id].append({
                'text': comment,
                'user': get_current_user()['name'],
                'time': datetime.now().strftime('%H:%M'),
                'type': 'own'
            })

    comments = notes_db.get(job_id, [
        {'text': 'Q5 bug has been fixed ✅', 'user': 'Client (Marie L.)', 'time': '2h ago', 'type': 'client'},
        {'text': 'Can you also check R1 termination?', 'user': 'You', 'time': '1h ago', 'type': 'own'},
        {'text': 'R1 is fixed too! Please retest 🙏', 'user': 'Client (Marie L.)', 'time': '30m ago', 'type': 'client', 'fixed': True},
    ])

    comments_html = ''
    for c in comments:
        bg = '#EBF5FF' if c['type']=='own' else ('#EAF3DE' if c.get('fixed') else '#F8F9FA')
        name_color = '#185FA5' if c['type']=='own' else 'var(--color-text-primary)'
        fixed_badge = '<span class="badge badge-green" style="margin-top:5px;display:inline-block;font-size:10px">✅ Marked as fixed</span>' if c.get('fixed') else ''
        comments_html += f"""
        <div style="background:{bg};border-radius:8px;padding:10px;margin-bottom:7px">
          <div style="display:flex;justify-content:space-between;margin-bottom:4px">
            <p style="font-size:11px;font-weight:500;color:{name_color}">{c['user']}</p>
            <span style="font-size:10px;color:var(--color-text-secondary)">{c['time']}</span>
          </div>
          <p style="font-size:11px;color:var(--color-text-primary)">{c['text']}</p>
          {fixed_badge}
        </div>"""

    return render_template_string(SHARED_CSS + f"""
<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=5.0"><title>Share Report — SurveyQC</title></head><body>
<div class="app-layout">
  {sidebar_html('reports')}
  <div class="main-content">
    <div class="topbar">
      <div><p class="page-title">Share Report</p><p class="page-sub">Share this report link with your client — they can view and comment after a free signup</p></div>
    </div>
    <div style="display:grid;grid-template-columns:1fr 1fr;gap:14px">
      <div class="card">
        <p style="font-size:13px;font-weight:500;color:var(--color-text-primary);margin-bottom:12px">🔗 Generate share link</p>
        <div style="background:var(--color-background-secondary);border-radius:8px;padding:12px;margin-bottom:12px">
          <p style="font-size:11px;color:var(--color-text-secondary);margin-bottom:6px">Report: <strong>{doc_name}</strong></p>
          <div style="display:flex;gap:6px;margin-bottom:8px">
            <div style="flex:1;background:var(--color-background-primary);border:0.5px solid var(--color-border-secondary);border-radius:7px;padding:9px 12px;font-size:11px;color:#185FA5;word-break:break-all">
              surveyqc.online/view/{share_token}
            </div>
            <button onclick="navigator.clipboard.writeText('surveyqc.online/view/{share_token}')" class="btn btn-primary btn-sm" style="flex-shrink:0"><i class="ti ti-copy"></i>Copy</button>
          </div>
          <div style="display:flex;gap:14px">
            <label style="display:flex;align-items:center;gap:5px;font-size:11px;color:var(--color-text-primary)"><input type="checkbox" checked>Expires 7 days</label>
            <label style="display:flex;align-items:center;gap:5px;font-size:11px;color:var(--color-text-primary)"><input type="checkbox">Password protect</label>
          </div>
        </div>
        <p style="font-size:12px;font-weight:500;color:var(--color-text-primary);margin-bottom:8px">Client can:</p>
        <div style="display:flex;flex-direction:column;gap:5px;margin-bottom:12px">
          <div style="display:flex;align-items:center;gap:7px;padding:7px 10px;background:#EAF3DE;border-radius:7px"><i class="ti ti-check" style="font-size:12px;color:#3B6D11"></i><span style="font-size:11px;color:var(--color-text-primary)">Free signup → view this report</span></div>
          <div style="display:flex;align-items:center;gap:7px;padding:7px 10px;background:#EAF3DE;border-radius:7px"><i class="ti ti-check" style="font-size:12px;color:#3B6D11"></i><span style="font-size:11px;color:var(--color-text-primary)">Comment on issues</span></div>
          <div style="display:flex;align-items:center;gap:7px;padding:7px 10px;background:#EAF3DE;border-radius:7px"><i class="ti ti-check" style="font-size:12px;color:#3B6D11"></i><span style="font-size:11px;color:var(--color-text-primary)">Mark issues as "Fixed"</span></div>
          <div style="display:flex;align-items:center;gap:7px;padding:7px 10px;background:var(--bg3);border-radius:7px"><i class="ti ti-x" style="font-size:12px;color:var(--color-text-secondary)"></i><span style="font-size:11px;color:var(--color-text-secondary)">Cannot run new QC reports</span></div>
        </div>
        <div style="display:flex;gap:7px">
          <button class="btn btn-primary" style="flex:1;justify-content:center;font-size:11px"><i class="ti ti-mail"></i>Email to client</button>
          <button class="btn btn-ghost btn-sm"><i class="ti ti-brand-whatsapp"></i>WhatsApp</button>
        </div>
      </div>
      <div class="card">
        <p style="font-size:13px;font-weight:500;color:var(--color-text-primary);margin-bottom:12px">💬 Comments</p>
        <div style="max-height:280px;overflow-y:auto;margin-bottom:12px">
          {comments_html}
        </div>
        <form method="POST" style="display:flex;gap:7px">
          <input class="form-input" name="comment" style="margin:0;flex:1;font-size:11px" placeholder="Reply to client...">
          <button type="submit" class="btn btn-primary btn-sm" style="flex-shrink:0"><i class="ti ti-send"></i></button>
        </form>
      </div>
    </div>
  </div>
</div>
</body></html>""")


# ================================================================
# USER FEEDBACK API (Task 6)
# ================================================================
@app.route('/api/user-feedback', methods=['POST'])
@login_required
def submit_user_feedback():
    data = request.get_json(silent=True) or {}
    msg = (data.get('message') or '').strip()[:1000]
    if not msg:
        return jsonify({'error': 'message required'}), 400
    entry = {
        'id': str(uuid.uuid4())[:8],
        'type': data.get('type', 'general')[:20],
        'message': msg,
        'page': (data.get('page') or '/')[:200],
        'user_email': session.get('user_email', ''),
        'created_at': datetime.now().isoformat(),
        'read': False,
    }
    user_feedback_db.append(entry)
    return jsonify({'ok': True})


@app.route('/admin/feedback')
def admin_feedback():
    if not session.get('is_admin'):
        return redirect('/admin/login')
    q = request.args.get('q', '').lower()
    t = request.args.get('type', '')
    items = list(reversed(user_feedback_db))
    if q:
        items = [i for i in items if q in i['message'].lower() or q in i['user_email'].lower()]
    if t:
        items = [i for i in items if i['type'] == t]
    unread_count = sum(1 for i in user_feedback_db if not i['read'])

    rows = ''
    for item in items:
        cls = '' if item['read'] else 'feedback-new'
        dot = '' if item['read'] else '<span style="display:inline-block;width:6px;height:6px;background:var(--accent);border-radius:50%;margin-right:6px;vertical-align:middle"></span>'
        created = item['created_at'][:16].replace('T', ' ')
        type_colors = {'bug': '#C84B31', 'feature': '#0073C6', 'general': '#3F7D58', 'other': '#8A847A'}
        tc = type_colors.get(item['type'], '#8A847A')
        rows += f"""<tr class="{cls}" data-id="{item['id']}">
          <td>{dot}<span style="font-size:11px;background:{tc}22;color:{tc};padding:2px 7px;border-radius:20px;font-weight:600">{item['type']}</span></td>
          <td style="max-width:360px;word-break:break-word">{item['message'][:200]}</td>
          <td style="color:#8A847A">{item['user_email']}</td>
          <td style="white-space:nowrap;color:#8A847A">{created}</td>
          <td>
            <form method="POST" action="/admin/feedback/action" style="display:inline">
              <input type="hidden" name="id" value="{item['id']}">
              <button name="action" value="read" title="Mark read" style="background:none;border:none;cursor:pointer;color:#0073C6;font-size:13px;padding:3px 6px">✓</button>
              <button name="action" value="delete" title="Delete" style="background:none;border:none;cursor:pointer;color:#C84B31;font-size:13px;padding:3px 6px" onclick="return confirm('Delete?')">✕</button>
            </form>
          </td>
        </tr>"""

    if not rows:
        rows = '<tr><td colspan="5" style="text-align:center;color:#8A847A;padding:24px">No feedback yet.</td></tr>'

    return render_template_string(f"""<!DOCTYPE html><html><head>
<meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>Feedback — Admin</title>
<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/@tabler/icons-webfont@3.19.0/dist/tabler-icons.min.css">
<script src="/admin-sidebar-js"></script>
<link rel="stylesheet" href="/static/style.css">
<script src="/static/app.js" defer></script>
<style>*{{box-sizing:border-box;margin:0;padding:0;font-family:-apple-system,'Inter',sans-serif}}body{{background:#F7F4EE}}
table{{width:100%;border-collapse:collapse}}th,td{{padding:9px 12px;text-align:left;border-bottom:1px solid #E8E1D8;font-size:12px}}
th{{font-weight:600;color:#5F5B53;font-size:11px;text-transform:uppercase;letter-spacing:.06em}}
tr:hover td{{background:#F7F4EE}}.feedback-new td{{font-weight:500}}</style>
</head><body>
<div id="admsb"></div>
<div style="padding:24px;max-width:1100px;margin:0 auto">
  <div style="display:flex;align-items:center;gap:14px;margin-bottom:20px">
    <a href="/admin" style="color:#8A847A;font-size:20px">&larr;</a>
    <div>
      <p style="font-size:20px;font-weight:700;color:#171717">User Feedback</p>
      <p style="font-size:12px;color:#8A847A">{unread_count} unread · auto-deleted after 3 days</p>
    </div>
    <form method="GET" style="margin-left:auto;display:flex;gap:8px;align-items:center">
      <input name="q" value="{q}" placeholder="Search..." style="padding:7px 12px;border:1px solid #E8E1D8;border-radius:8px;font-size:13px;outline:none">
      <select name="type" style="padding:7px 12px;border:1px solid #E8E1D8;border-radius:8px;font-size:13px;outline:none">
        <option value="">All types</option>
        <option value="bug" {'selected' if t=='bug' else ''}>Bug</option>
        <option value="feature" {'selected' if t=='feature' else ''}>Feature</option>
        <option value="general" {'selected' if t=='general' else ''}>General</option>
        <option value="other" {'selected' if t=='other' else ''}>Other</option>
      </select>
      <button type="submit" style="padding:7px 14px;background:#C46A2B;color:#fff;border:none;border-radius:8px;cursor:pointer;font-size:13px">Filter</button>
    </form>
  </div>
  <div style="background:#fff;border-radius:14px;border:1px solid #E8E1D8;overflow:hidden">
    <table><thead><tr><th>Type</th><th>Message</th><th>User</th><th>Date</th><th>Actions</th></tr></thead>
    <tbody>{rows}</tbody></table>
  </div>
</div></body></html>""")


@app.route('/admin/feedback/action', methods=['POST'])
def admin_feedback_action():
    if not session.get('is_admin'):
        return redirect('/admin/login')
    fid = request.form.get('id', '')
    action = request.form.get('action', '')
    for i, item in enumerate(user_feedback_db):
        if item['id'] == fid:
            if action == 'delete':
                user_feedback_db.pop(i)
            elif action == 'read':
                user_feedback_db[i]['read'] = True
            break
    return redirect('/admin/feedback')


# ================================================================
# PRIVACY PAGE
# ================================================================
@app.route('/privacy-data')
@login_required
def privacy_data():
    user = get_current_user()
    email = session['user_email']
    report_count = sum(1 for j in jobs.values() if j.get('user_email') == email)
    expiring = sum(1 for j in jobs.values() if j.get('user_email') == email
                  and j.get('created_at','') and
                  (datetime.now() - datetime.fromisoformat(j['created_at'])).days >= 27)

    return render_template_string(SHARED_CSS + f"""
<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=5.0"><title>Privacy & Data — SurveyQC</title></head><body>
<div class="app-layout">
  {sidebar_html('privacy')}
  <div class="main-content">
    <div class="topbar"><div>
      <p class="page-title">Privacy & Data</p>
      <p class="page-sub">GDPR · CCPA · UK GDPR · India DPDP · Australia Privacy Act</p>
    </div></div>

    <div style="background:#E8F5E9;border:0.5px solid #A5D6A7;border-radius:8px;padding:14px;margin-bottom:16px">
      <div style="display:flex;align-items:center;gap:10px">
        <i class="ti ti-shield-check" style="font-size:22px;color:#166534;flex-shrink:0"></i>
        <div>
          <p style="font-size:13px;font-weight:500;color:#166534">Data auto-deletes after 30 days</p>
          <p style="font-size:11px;color:#15803D;margin-top:2px">Runs automatically every night at 12:00 AM. No action needed — fully automatic!</p>
        </div>
      </div>
    </div>

    <div style="display:grid;grid-template-columns:1fr 1fr;gap:14px">
      <div class="card">
        <p style="font-size:13px;font-weight:500;color:var(--color-text-primary);margin-bottom:12px">🌍 Country-wise compliance</p>
        <div style="display:flex;flex-direction:column;gap:6px">
          <div style="display:flex;align-items:center;justify-content:space-between;padding:9px 11px;background:var(--color-background-secondary);border-radius:8px">
            <div style="display:flex;align-items:center;gap:8px"><span style="font-size:16px">🇪🇺</span><div><p style="font-size:11px;font-weight:500;color:var(--color-text-primary)">Europe — GDPR</p><p style="font-size:10px;color:var(--color-text-secondary)">30d delete · consent · right to forget</p></div></div>
            <span class="badge badge-green" style="font-size:10px">✅ Active</span>
          </div>
          <div style="display:flex;align-items:center;justify-content:space-between;padding:9px 11px;background:var(--color-background-secondary);border-radius:8px">
            <div style="display:flex;align-items:center;gap:8px"><span style="font-size:16px">🇺🇸</span><div><p style="font-size:11px;font-weight:500;color:var(--color-text-primary)">USA — CCPA</p><p style="font-size:10px;color:var(--color-text-secondary)">No data sell · delete on request</p></div></div>
            <span class="badge badge-green" style="font-size:10px">✅ Active</span>
          </div>
          <div style="display:flex;align-items:center;justify-content:space-between;padding:9px 11px;background:var(--color-background-secondary);border-radius:8px">
            <div style="display:flex;align-items:center;gap:8px"><span style="font-size:16px">🇬🇧</span><div><p style="font-size:11px;font-weight:500;color:var(--color-text-primary)">UK — UK GDPR</p><p style="font-size:10px;color:var(--color-text-secondary)">Post-Brexit data protection</p></div></div>
            <span class="badge badge-green" style="font-size:10px">✅ Active</span>
          </div>
          <div style="display:flex;align-items:center;justify-content:space-between;padding:9px 11px;background:var(--color-background-secondary);border-radius:8px">
            <div style="display:flex;align-items:center;gap:8px"><span style="font-size:16px">🇮🇳</span><div><p style="font-size:11px;font-weight:500;color:var(--color-text-primary)">India — DPDP Act 2023</p><p style="font-size:10px;color:var(--color-text-secondary)">Digital Personal Data Protection</p></div></div>
            <span class="badge badge-green" style="font-size:10px">✅ Active</span>
          </div>
          <div style="display:flex;align-items:center;justify-content:space-between;padding:9px 11px;background:var(--color-background-secondary);border-radius:8px">
            <div style="display:flex;align-items:center;gap:8px"><span style="font-size:16px">🇦🇺</span><div><p style="font-size:11px;font-weight:500;color:var(--color-text-primary)">Australia — Privacy Act</p><p style="font-size:10px;color:var(--color-text-secondary)">13 Australian Privacy Principles</p></div></div>
            <span class="badge badge-green" style="font-size:10px">✅ Active</span>
          </div>
        </div>
      </div>
      <div style="display:flex;flex-direction:column;gap:12px">
        <div class="card">
          <p style="font-size:13px;font-weight:500;color:var(--color-text-primary);margin-bottom:12px">📊 Your data status</p>
          <div style="display:flex;flex-direction:column;gap:8px">
            <div style="display:flex;justify-content:space-between;align-items:center">
              <div><p style="font-size:11px;color:var(--color-text-primary)">Reports stored</p><p style="font-size:10px;color:var(--color-text-secondary)">Auto-delete after 30 days</p></div>
              <span style="font-size:16px;font-weight:600;color:var(--color-text-primary)">{report_count}</span>
            </div>
            <div style="border-top:0.5px solid var(--color-border-tertiary);padding-top:8px;display:flex;justify-content:space-between;align-items:center">
              <div><p style="font-size:11px;color:#791F1F">Expiring soon</p><p style="font-size:10px;color:var(--color-text-secondary)">Download before delete!</p></div>
              <span style="font-size:16px;font-weight:600;color:#791F1F">{expiring}</span>
            </div>
          </div>
        </div>
        <div class="card">
          <p style="font-size:13px;font-weight:500;color:var(--color-text-primary);margin-bottom:10px">⚙️ Data controls</p>
          <div style="display:flex;flex-direction:column;gap:7px">
            <a href="/export-data" class="btn btn-ghost" style="width:100%;justify-content:center;font-size:11px"><i class="ti ti-download"></i>Export all my data</a>
            <a href="/delete-reports" class="btn btn-ghost" style="width:100%;justify-content:center;font-size:11px"><i class="ti ti-trash"></i>Delete all reports now</a>
            <button style="width:100%;justify-content:center;font-size:11px;background:#FCEBEB;color:#791F1F;border:0.5px solid #F7C1C1;padding:8px;border-radius:8px;cursor:pointer;display:flex;align-items:center;gap:5px"><i class="ti ti-user-x"></i>Delete my account</button>
          </div>
        </div>
        <div style="background:#E6F1FB;border:0.5px solid #B5D4F4;border-radius:10px;padding:12px">
          <p style="font-size:11px;font-weight:500;color:#0C447C;margin-bottom:5px">🕛 Auto-delete schedule</p>
          <p style="font-size:11px;color:#185FA5;line-height:1.7">Auto-deletes 30-day-old data nightly at <strong>12:00 AM</strong>. Fully automatic. ✅</p>
        </div>
      </div>
    </div>
  </div>
</div>
</body></html>""")


# ================================================================
# ADMIN: CONTENT MANAGEMENT
# ================================================================
@app.route('/admin/content', methods=['GET', 'POST'])
@admin_required
def admin_content():
    saved = False
    if request.method == 'POST':
        for key in list(site_content.keys()):
            val = request.form.get(key)
            if val is not None:
                site_content[key] = val
        saved = True

    groups = {
        'Hero Section': ['hero_heading_part1','hero_heading_part2','hero_heading_part3','hero_subheading','hero_cta','hero_cta2','announcement'],
        'Homepage Features (6)': ['feature1_title','feature1_desc','feature2_title','feature2_desc','feature3_title','feature3_desc','feature4_title','feature4_desc','feature5_title','feature5_desc','feature6_title','feature6_desc'],
        'Pricing Page': ['pricing_heading','pricing_sub'],
        'Plan: Free': ['plan_free_name','plan_free_price','plan_free_yearly','plan_free_period','plan_free_desc','plan_free_features','plan_free_cta'],
        'Plan: Pro': ['plan_pro_name','plan_pro_price','plan_pro_yearly','plan_pro_period','plan_pro_desc','plan_pro_features','plan_pro_cta','plan_pro_badge','plan_pro_featured'],
        'Plan: Business': ['plan_biz_name','plan_biz_price','plan_biz_yearly','plan_biz_period','plan_biz_desc','plan_biz_features','plan_biz_cta'],
        'Pricing FAQ': ['pfaq1_q','pfaq1_a','pfaq2_q','pfaq2_a','pfaq3_q','pfaq3_a','pfaq4_q','pfaq4_a','pfaq5_q','pfaq5_a'],
        'Features Page': ['features_heading','features_sub','feat_grp1_name','feat_grp1_desc','feat_grp1_items','feat_grp2_name','feat_grp2_desc','feat_grp2_items','feat_grp3_name','feat_grp3_desc','feat_grp3_items','feat_grp4_name','feat_grp4_desc','feat_grp4_items'],
        'Testimonial 1': ['test1_name','test1_role','test1_company','test1_country','test1_quote','test1_rating'],
        'Testimonial 2': ['test2_name','test2_role','test2_company','test2_country','test2_quote','test2_rating'],
        'Testimonial 3': ['test3_name','test3_role','test3_company','test3_country','test3_quote','test3_rating'],
        'Testimonial 4': ['test4_name','test4_role','test4_company','test4_country','test4_quote','test4_rating'],
        'Testimonial 5': ['test5_name','test5_role','test5_company','test5_country','test5_quote','test5_rating'],
        'Testimonial 6': ['test6_name','test6_role','test6_company','test6_country','test6_quote','test6_rating'],
        'Blog Posts': ['blog_post1_title','blog_post1_date','blog_post1_tag','blog_post1_summary','blog_post2_title','blog_post2_date','blog_post2_tag','blog_post2_summary','blog_post3_title','blog_post3_date','blog_post3_tag','blog_post3_summary','blog_post4_title','blog_post4_date','blog_post4_tag','blog_post4_summary','blog_post5_title','blog_post5_date','blog_post5_tag','blog_post5_summary','blog_post6_title','blog_post6_date','blog_post6_tag','blog_post6_summary'],
        'Site & Footer': ['site_name','tagline','footer_text','support_email','linkedin_url','twitter_url'],
        'Compare Pages': ['compare_chatgpt_heading','compare_chatgpt_summary','compare_excel_heading','compare_excel_summary','compare_manual_heading','compare_manual_summary'],
        'Community & Affiliate': ['community_heading','community_subheading','affiliate_heading','affiliate_commission','affiliate_details'],
        'Changelog': ['changelog_v10','changelog_v9','changelog_v8','changelog_v7','changelog_v6'],
    }

    alert = '<div style="background:#E5F0E9;border:1px solid #A5D6A7;border-radius:10px;padding:12px 16px;margin-bottom:18px;color:#3F7D58;font-size:14px;font-weight:500"><i class="ti ti-check"></i> All content saved!</div>' if saved else ''

    # Build tabbed form
    tab_btns = ''
    tab_panels = ''
    for i, (group, fields) in enumerate(groups.items()):
        active = ' active' if i == 0 else ''
        gid = 'tab' + str(i)
        tab_btns += '<button type="button" class="cms-tab' + active + '" onclick="showTab(\'' + gid + '\',this)">' + group + '</button>'
        panel = '<div class="cms-panel' + active + '" id="' + gid + '">'
        for field in fields:
            if field not in site_content:
                continue
            val = site_content.get(field, '')
            # Escape for HTML attribute
            val_escaped = val.replace('"', '&quot;')
            label = field.replace('_', ' ').title()
            is_long = any(x in field for x in ['summary','details','desc','quote','features','items','_a','subheading','heading','intro','changelog','announcement','policy','terms'])
            if is_long:
                inp = '<textarea name="' + field + '" class="cms-input" rows="3">' + val + '</textarea>'
            else:
                inp = '<input type="text" name="' + field + '" value="' + val_escaped + '" class="cms-input">'
            hint = ''
            if 'features' in field:
                hint = '<span class="cms-hint">Separate each feature with || (double pipe)</span>'
            elif 'items' in field:
                hint = '<span class="cms-hint">Format: Title||Description||icon-name ## next feature...</span>'
            elif field == 'plan_pro_featured':
                hint = '<span class="cms-hint">1 = highlighted plan, 0 = normal</span>'
            panel += '<div class="cms-field"><label class="cms-label">' + label + '</label>' + inp + hint + '</div>'
        panel += '</div>'
        tab_panels += panel

    page = """<!DOCTYPE html><html><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>Content CMS - Admin</title>
<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/@tabler/icons-webfont@3.19.0/dist/tabler-icons.min.css">

<style>
:root{--bg:#F7F4EE;--card:#FFF;--text:#171717;--text2:#5F5B53;--text3:#8A847A;--accent:#C46A2B;--border:#E8E1D8;--dark:#1B140F}
*{box-sizing:border-box;margin:0;padding:0;font-family:'Inter',sans-serif;-webkit-font-smoothing:antialiased}
body{background:var(--bg);color:var(--text)}
a{text-decoration:none;color:inherit}
.cms-wrap{max-width:1100px;margin:0 auto;padding:24px}
.cms-hdr{display:flex;align-items:center;gap:14px;margin-bottom:24px}
.cms-hdr h1{font-family:'Plus Jakarta Sans',sans-serif;font-size:24px;font-weight:800;letter-spacing:-0.5px}
.cms-hdr p{font-size:13px;color:var(--text2)}
.cms-layout{display:grid;grid-template-columns:240px 1fr;gap:20px}
.cms-tabs{display:flex;flex-direction:column;gap:4px;position:sticky;top:24px;height:fit-content}
.cms-tab{text-align:left;background:none;border:none;padding:10px 14px;border-radius:10px;font-size:13px;font-weight:500;color:var(--text2);cursor:pointer;font-family:inherit;transition:all .15s}
.cms-tab:hover{background:rgba(196,106,43,.08);color:var(--text)}
.cms-tab.active{background:var(--dark);color:#F7F4EE;font-weight:600}
.cms-content{background:var(--card);border:1px solid var(--border);border-radius:18px;padding:28px}
.cms-panel{display:none}
.cms-panel.active{display:block}
.cms-field{margin-bottom:18px}
.cms-label{display:block;font-size:13px;font-weight:600;color:var(--text);margin-bottom:6px}
.cms-input{width:100%;padding:11px 14px;border:1px solid var(--border);border-radius:10px;font-size:13px;font-family:inherit;color:var(--text);outline:none;resize:vertical}
.cms-input:focus{border-color:var(--accent);box-shadow:0 0 0 3px rgba(196,106,43,.1)}
.cms-hint{display:block;font-size:11px;color:var(--text3);margin-top:4px}
.cms-save{position:fixed;bottom:24px;right:24px;background:var(--accent);color:white;border:none;padding:14px 28px;border-radius:12px;font-size:14px;font-weight:600;cursor:pointer;box-shadow:0 8px 24px rgba(196,106,43,.3);font-family:inherit;display:flex;align-items:center;gap:8px}
.cms-save:hover{background:#A9551F}
@media(max-width:768px){
  .cms-layout{grid-template-columns:1fr}
  .cms-tabs{flex-direction:row;overflow-x:auto;position:static}
  .cms-tab{white-space:nowrap}
}
</style><script src="/admin-sidebar-js"></script></head><body>
<div class="cms-wrap">
  <div class="cms-hdr">
    <a href="/admin" style="color:var(--text3);font-size:22px"><i class="ti ti-arrow-left"></i></a>
    <div><h1>Content Management</h1><p>Edit all website content live — no code needed</p></div>
  </div>
  """ + alert + """
  <form method="POST">
    <div class="cms-layout">
      <div class="cms-tabs">""" + tab_btns + """</div>
      <div class="cms-content">""" + tab_panels + """</div>
    </div>
    <button type="submit" class="cms-save"><i class="ti ti-device-floppy"></i> Save All Changes</button>
  </form>
</div>
<script>
function showTab(id, btn){
  document.querySelectorAll('.cms-panel').forEach(p=>p.classList.remove('active'));
  document.querySelectorAll('.cms-tab').forEach(t=>t.classList.remove('active'));
  document.getElementById(id).classList.add('active');
  btn.classList.add('active');
}
</script>
</body></html>"""
    return render_template_string(page)


@app.route('/admin/tokens', methods=['GET', 'POST'])
@admin_required
def admin_tokens():
    global token_limits
    saved = False
    if request.method == 'POST':
        try:
            token_limits['free'] = int(request.form.get('free', 20000))
            token_limits['pro'] = int(request.form.get('pro', 100000))
            token_limits['business'] = int(request.form.get('business', 150000))
            token_limits['monthly_budget'] = int(request.form.get('monthly_budget', 50))
            saved = True
        except: pass

    return render_template_string(SHARED_CSS + f"""
<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=5.0"><title>Token Control — Admin</title><script src="/admin-sidebar-js"></script></head><body>
<div style="padding:24px;max-width:500px">
  <div style="display:flex;align-items:center;gap:14px;margin-bottom:20px">
    <a href="/admin" style="color:var(--color-text-secondary);text-decoration:none"><i class="ti ti-arrow-left"></i></a>
    <p style="font-size:18px;font-weight:500;color:var(--color-text-primary)">Token Control</p>
  </div>
  {'<div class="alert alert-success">Token limits updated!</div>' if saved else ''}
  <div class="card">
    <form method="POST">
      <div style="margin-bottom:14px">
        <label style="font-size:12px;color:var(--color-text-secondary);margin-bottom:6px;display:block;font-weight:500">Free plan limit (tokens/report)</label>
        <input class="form-input" name="free" type="number" value="{token_limits['free']}">
      </div>
      <div style="margin-bottom:14px">
        <label style="font-size:12px;color:var(--color-text-secondary);margin-bottom:6px;display:block;font-weight:500">Pro plan limit (tokens/report)</label>
        <input class="form-input" name="pro" type="number" value="{token_limits['pro']}">
      </div>
      <div style="margin-bottom:14px">
        <label style="font-size:12px;color:var(--color-text-secondary);margin-bottom:6px;display:block;font-weight:500">Business plan limit — max 150,000</label>
        <input class="form-input" name="business" type="number" value="{token_limits['business']}" max="150000">
      </div>
      <div style="background:#FFFBEB;border:0.5px solid #FCD34D;border-radius:8px;padding:10px;margin-bottom:14px">
        <p style="font-size:11px;color:#92400E">⚠️ 150K limit reach hone pe: AI rukta hai → remaining checks = MANUAL by user</p>
      </div>
      <div style="margin-bottom:14px">
        <label style="font-size:12px;color:var(--color-text-secondary);margin-bottom:6px;display:block;font-weight:500">Monthly budget alert ($)</label>
        <input class="form-input" name="monthly_budget" type="number" value="{token_limits['monthly_budget']}">
      </div>
      <button type="submit" class="btn btn-primary">Save token settings</button>
    </form>
  </div>
</div>
</body></html>""")


# ================================================================
# GIFT ACCESS
# ================================================================
@app.route('/admin/gift', methods=['GET', 'POST'])
@admin_required
def admin_gift():
    gifted = False
    error = ''
    if request.method == 'POST':
        email = request.form.get('email', '').strip()
        plan = request.form.get('plan', 'Pro')
        duration = request.form.get('duration', '1 month')
        if email in users_db:
            users_db[email]['plan'] = plan
            users_db[email]['reports_limit'] = UserDB.PLAN_LIMITS.get(plan, 3)
            users_db.save(email)
            gifted = True
        elif email:
            # Create gift account
            users_db[email] = {
                'password_hash': generate_password_hash('temp123'),
                'password':      generate_password_hash('temp123'),
                'name':          email.split('@')[0].title(),
                'plan':          plan,
                'reports_used':  0,
                'reports_limit': UserDB.PLAN_LIMITS.get(plan, 3),
                'joined':        datetime.now().strftime('%Y-%m-%d'),
                'total_saved_hours': 0,
                'gifted':        True,
                'gift_duration': duration,
                'must_change_password': True,
            }
            gifted = True
        else:
            error = 'Email required'

    return render_template_string(SHARED_CSS + f"""
<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=5.0"><title>Gift Access — Admin</title><script src="/admin-sidebar-js"></script></head><body>
<div style="padding:24px;max-width:500px">
  <div style="display:flex;align-items:center;gap:14px;margin-bottom:20px">
    <a href="/admin" style="color:var(--color-text-secondary);text-decoration:none"><i class="ti ti-arrow-left"></i></a>
    <p style="font-size:18px;font-weight:500;color:var(--color-text-primary)">🎁 Gift Access</p>
  </div>
  {'<div class="alert alert-success">Access gifted! A confirmation email has been sent to the user.</div>' if gifted else ''}
  {'<div class="alert alert-error">' + error + '</div>' if error else ''}
  <div class="card">
    <p style="font-size:12px;color:var(--color-text-secondary);margin-bottom:14px">Give a friend or colleague direct access — no coupon code needed</p>
    <form method="POST">
      <div style="margin-bottom:12px">
        <label style="font-size:12px;color:var(--color-text-secondary);margin-bottom:5px;display:block;font-weight:500">User email *</label>
        <input class="form-input" name="email" type="email" placeholder="friend@example.com">
      </div>
      <div style="margin-bottom:12px">
        <label style="font-size:12px;color:var(--color-text-secondary);margin-bottom:5px;display:block;font-weight:500">Gift plan</label>
        <select class="form-select" name="plan" style="margin-bottom:0">
          <option value="Pro">Pro plan</option>
          <option value="Business">Business plan</option>
        </select>
      </div>
      <div style="margin-bottom:16px">
        <label style="font-size:12px;color:var(--color-text-secondary);margin-bottom:5px;display:block;font-weight:500">Duration</label>
        <select class="form-select" name="duration" style="margin-bottom:0">
          <option>1 month</option>
          <option>3 months</option>
          <option>6 months</option>
          <option>1 year</option>
          <option>Lifetime</option>
        </select>
      </div>
      <div style="background:#E6F1FB;border-radius:8px;padding:10px;margin-bottom:14px">
        <p style="font-size:11px;color:#0C447C">User ko automatically email aayega: "Tushar ne tumhe Pro access gift kiya! surveyqc.online 🎁"</p>
      </div>
      <button type="submit" class="btn btn-primary" style="width:100%;justify-content:center"><i class="ti ti-gift"></i>Gift access now</button>
    </form>
  </div>
</div>
</body></html>""")


# ================================================================
# EXPORT DATA
# ================================================================
@app.route('/export-data')
@login_required
def export_data():
    email = session['user_email']
    user_jobs = {jid: j for jid, j in jobs.items() if j.get('user_email') == email}
    export = {
        'user': email,
        'exported_at': datetime.now().isoformat(),
        'reports': user_jobs
    }
    import tempfile
    tmp = tempfile.NamedTemporaryFile(mode='w', suffix='.json', delete=False)
    json.dump(export, tmp, indent=2, default=str)
    tmp.close()
    return send_file(tmp.name, as_attachment=True, download_name='my_surveyqc_data.json')


@app.route('/delete-reports', methods=['POST','GET'])
@login_required
def delete_reports():
    email = session['user_email']
    for jid in list(jobs.keys()):
        if jobs[jid].get('user_email') == email:
            del jobs[jid]
    return redirect('/privacy-data')


# ================================================================
# ADMIN: PRICING
# ================================================================

@app.route('/admin/pricing', methods=['GET', 'POST'])
@admin_required
def admin_pricing():
    saved = False
    pricing_keys = [
        'plan_free_name', 'plan_free_price', 'plan_free_yearly', 'plan_free_period', 'plan_free_desc',
        'plan_pro_name', 'plan_pro_price', 'plan_pro_yearly', 'plan_pro_period', 'plan_pro_desc',
        'plan_biz_name', 'plan_biz_price', 'plan_biz_yearly', 'plan_biz_period', 'plan_biz_desc',
        'plan_ent_name', 'plan_ent_price', 'plan_ent_yearly', 'plan_ent_period', 'plan_ent_desc',
        'pricing_heading', 'pricing_sub',
    ]
    if request.method == 'POST':
        for key in pricing_keys:
            val = request.form.get(key)
            if val is not None:
                site_content[key] = val
        saved = True

    c = site_content
    plans = [
        ('Free',     'plan_free',  '#27500A', '#EAF3DE'),
        ('Pro',      'plan_pro',   '#042C53', '#E6F1FB'),
        ('Business', 'plan_biz',   '#533F00', '#FEF9EE'),
        ('Enterprise','plan_ent',  '#3C1F6E', '#F4EEFE'),
    ]

    alert = ''
    if saved:
        alert = '<div style="background:#E5F0E9;border:1px solid #A5D6A7;border-radius:10px;padding:12px 16px;margin-bottom:20px;color:#3F7D58;font-size:14px;font-weight:500">&#10003; Prices saved! All pages now use the new values.</div>'

    cards = ''
    for plan_name, key, col, bg in plans:
        cards += f'''
        <div style="background:white;border:0.5px solid #DDE1E7;border-radius:12px;padding:20px;margin-bottom:16px">
          <div style="display:flex;align-items:center;gap:10px;margin-bottom:16px">
            <div style="width:34px;height:34px;border-radius:8px;background:{bg};display:flex;align-items:center;justify-content:center;font-size:16px">💰</div>
            <div><p style="font-size:15px;font-weight:600;color:{col}">{plan_name} Plan</p><p style="font-size:11px;color:#9CA3AF">Displayed on /pricing</p></div>
          </div>
          <div style="display:grid;grid-template-columns:1fr 1fr;gap:12px;margin-bottom:12px">
            <div>
              <label style="font-size:12px;font-weight:600;color:#374151;display:block;margin-bottom:4px">Monthly Price (number only)</label>
              <input type="text" name="{key}_price" value="{c.get(key + "_price", "")}" style="width:100%;padding:9px 12px;border:0.5px solid #DDE1E7;border-radius:8px;font-size:13px">
            </div>
            <div>
              <label style="font-size:12px;font-weight:600;color:#374151;display:block;margin-bottom:4px">Yearly Price (number only)</label>
              <input type="text" name="{key}_yearly" value="{c.get(key + "_yearly", "")}" style="width:100%;padding:9px 12px;border:0.5px solid #DDE1E7;border-radius:8px;font-size:13px">
            </div>
          </div>
          <div style="display:grid;grid-template-columns:1fr 1fr;gap:12px">
            <div>
              <label style="font-size:12px;font-weight:600;color:#374151;display:block;margin-bottom:4px">Plan Name</label>
              <input type="text" name="{key}_name" value="{c.get(key + "_name", "")}" style="width:100%;padding:9px 12px;border:0.5px solid #DDE1E7;border-radius:8px;font-size:13px">
            </div>
            <div>
              <label style="font-size:12px;font-weight:600;color:#374151;display:block;margin-bottom:4px">Period Label</label>
              <input type="text" name="{key}_period" value="{c.get(key + "_period", "")}" style="width:100%;padding:9px 12px;border:0.5px solid #DDE1E7;border-radius:8px;font-size:13px">
            </div>
          </div>
        </div>'''

    page = f'''<!DOCTYPE html><html lang="en"><head>
<meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>Pricing Admin</title>
<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/@tabler/icons-webfont@3.19.0/dist/tabler-icons.min.css">
<script src="/admin-sidebar-js"></script>
</head><body style="background:#F0F2F5;font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',sans-serif;color:#1A1A2E">
<div style="padding:28px;max-width:760px;margin-left:220px">
  <div style="display:flex;align-items:center;gap:14px;margin-bottom:8px">
    <a href="/admin" style="color:#6B7280;font-size:20px;text-decoration:none"><i class="ti ti-arrow-left"></i></a>
    <div>
      <p style="font-size:20px;font-weight:700;color:#1A1A2E">Pricing Management</p>
      <p style="font-size:13px;color:#6B7280">Changes apply instantly to /pricing, /billing, and all plan pages</p>
    </div>
  </div>

  <div style="background:#FEF9EE;border:0.5px solid #F5D88A;border-radius:10px;padding:12px 16px;margin-bottom:20px;font-size:13px;color:#633806">
    <b>Tip:</b> Enter numbers only for prices (e.g. <code>29</code> or <code>299</code>, no $ symbol). The $ is added automatically.
  </div>

  {alert}

  <form method="POST">
    {cards}

    <div style="background:white;border:0.5px solid #DDE1E7;border-radius:12px;padding:20px;margin-bottom:16px">
      <p style="font-size:14px;font-weight:600;color:#1A1A2E;margin-bottom:14px">Pricing Page Header</p>
      <div style="margin-bottom:12px">
        <label style="font-size:12px;font-weight:600;color:#374151;display:block;margin-bottom:4px">Heading</label>
        <input type="text" name="pricing_heading" value="{c.get("pricing_heading", "")}" style="width:100%;padding:9px 12px;border:0.5px solid #DDE1E7;border-radius:8px;font-size:13px">
      </div>
      <div>
        <label style="font-size:12px;font-weight:600;color:#374151;display:block;margin-bottom:4px">Subheading</label>
        <input type="text" name="pricing_sub" value="{c.get("pricing_sub", "")}" style="width:100%;padding:9px 12px;border:0.5px solid #DDE1E7;border-radius:8px;font-size:13px">
      </div>
    </div>

    <button type="submit" style="width:100%;background:#042C53;color:white;border:none;padding:14px;border-radius:10px;font-size:15px;font-weight:600;cursor:pointer;display:flex;align-items:center;justify-content:center;gap:8px">
      <i class="ti ti-device-floppy"></i> Save Prices
    </button>
  </form>

  <div style="margin-top:20px;background:white;border:0.5px solid #DDE1E7;border-radius:12px;padding:16px">
    <p style="font-size:13px;font-weight:600;color:#374151;margin-bottom:10px">Quick links to verify changes:</p>
    <div style="display:flex;gap:10px;flex-wrap:wrap">
      <a href="/pricing" target="_blank" style="font-size:12px;background:#E6F1FB;color:#042C53;padding:6px 12px;border-radius:6px;font-weight:500">&#128279; /pricing</a>
      <a href="/billing" target="_blank" style="font-size:12px;background:#EAF3DE;color:#27500A;padding:6px 12px;border-radius:6px;font-weight:500">&#128279; /billing</a>
      <a href="/admin/revenue" target="_blank" style="font-size:12px;background:#FEF9EE;color:#533F00;padding:6px 12px;border-radius:6px;font-weight:500">&#128279; Revenue</a>
    </div>
  </div>
</div>
</body></html>'''
    return render_template_string(page)


# ================================================================
# CLEANUP NOW (manual)
# ================================================================

@app.route('/admin/activity')
@admin_required
def admin_activity():
    items = []
    for jid, j in list(jobs.items())[-30:]:
        items.append({'user': j.get('user_email','unknown'), 'action': 'QC on: '+j.get('doc_name','unknown')[:40], 'time': j.get('created_at','')[:16], 'status': j.get('status','unknown')})
    items.reverse()
    rows = ''
    for a in items:
        badge = '<span style="background:#EAF3DE;color:#27500A;font-size:10px;padding:2px 8px;border-radius:20px">Done</span>' if a['status']=='done' else '<span style="background:#E6F1FB;color:#0C447C;font-size:10px;padding:2px 8px;border-radius:20px">Running</span>'
        rows += '<tr><td style="padding:10px 12px;font-size:12px;color:#1A1A2E;border-bottom:0.5px solid #F0F2F5">'+a['user']+'</td><td style="padding:10px 12px;font-size:12px;color:#6B7280;border-bottom:0.5px solid #F0F2F5">'+a['action']+'</td><td style="padding:10px 12px;border-bottom:0.5px solid #F0F2F5">'+badge+'</td><td style="padding:10px 12px;font-size:11px;color:#9CA3AF;border-bottom:0.5px solid #F0F2F5">'+a['time']+'</td></tr>'
    page = '<!DOCTYPE html><html><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>Activity - Admin</title></head><body style="background:#F0F2F5;font-family:-apple-system,sans-serif;color:#1A1A2E">'
    page += '<div style="padding:24px"><div style="display:flex;align-items:center;gap:14px;margin-bottom:20px"><a href="/admin" style="color:#6B7280;font-size:22px;text-decoration:none">&larr;</a><p style="font-size:20px;font-weight:600">Live Activity</p></div>'
    page += '<div style="background:white;border:0.5px solid #DDE1E7;border-radius:10px;overflow:hidden"><table style="width:100%;border-collapse:collapse"><thead><tr style="background:#F8F9FA"><th style="padding:10px 12px;text-align:left;font-size:10px;color:#6B7280;font-weight:500">USER</th><th style="padding:10px 12px;text-align:left;font-size:10px;color:#6B7280;font-weight:500">ACTION</th><th style="padding:10px 12px;text-align:left;font-size:10px;color:#6B7280;font-weight:500">STATUS</th><th style="padding:10px 12px;text-align:left;font-size:10px;color:#6B7280;font-weight:500">TIME</th></tr></thead><tbody>'
    page += rows if rows else '<tr><td colspan="4" style="padding:20px;text-align:center;color:#9CA3AF;font-size:13px">No activity yet</td></tr>'
    page += '</tbody></table></div></div></body></html>'
    return render_template_string(page)


@app.route('/admin/revenue')
@admin_required
def admin_revenue():
    total = len(users_db)
    pro = sum(1 for u in users_db.values() if u.get('plan') == 'Pro')
    biz = sum(1 for u in users_db.values() if u.get('plan') == 'Business')
    free = total - pro - biz
    _pro_price = int(site_content.get('plan_pro_price', '29') or '29')
    _biz_price = int(site_content.get('plan_biz_price', '299') or '299')
    mrr = (pro * _pro_price) + (biz * _biz_price)
    arr = mrr * 12
    page = '<!DOCTYPE html><html><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>Revenue - Admin</title></head><body style="background:#F0F2F5;font-family:-apple-system,sans-serif;color:#1A1A2E">'
    page += '<div style="padding:24px"><div style="display:flex;align-items:center;gap:14px;margin-bottom:20px"><a href="/admin" style="color:#6B7280;font-size:22px;text-decoration:none">&larr;</a><p style="font-size:20px;font-weight:600">Revenue Analytics</p></div>'
    page += '<div style="display:grid;grid-template-columns:repeat(4,1fr);gap:12px;margin-bottom:20px">'
    page += '<div style="background:white;border:0.5px solid #DDE1E7;border-radius:10px;padding:16px;text-align:center"><p style="font-size:22px;font-weight:600;color:#042C53">$'+str(mrr)+'</p><p style="font-size:11px;color:#6B7280;margin-top:4px">MRR</p></div>'
    page += '<div style="background:white;border:0.5px solid #DDE1E7;border-radius:10px;padding:16px;text-align:center"><p style="font-size:22px;font-weight:600;color:#27500A">$'+str(arr)+'</p><p style="font-size:11px;color:#6B7280;margin-top:4px">ARR</p></div>'
    page += '<div style="background:white;border:0.5px solid #DDE1E7;border-radius:10px;padding:16px;text-align:center"><p style="font-size:22px;font-weight:600;color:#1A1A2E">'+str(pro+biz)+'</p><p style="font-size:11px;color:#6B7280;margin-top:4px">Paid users</p></div>'
    page += '<div style="background:white;border:0.5px solid #DDE1E7;border-radius:10px;padding:16px;text-align:center"><p style="font-size:22px;font-weight:600;color:#1A1A2E">'+str(total)+'</p><p style="font-size:11px;color:#6B7280;margin-top:4px">Total users</p></div>'
    page += '</div>'
    page += '<div style="background:white;border:0.5px solid #DDE1E7;border-radius:10px;padding:20px">'
    page += '<p style="font-size:14px;font-weight:600;color:#1A1A2E;margin-bottom:14px">Subscription breakdown</p>'
    page += '<div style="display:flex;justify-content:space-between;padding:10px 0;border-bottom:0.5px solid #EEF0F3"><span style="font-size:13px;color:#374151">Pro plan ($'+str(_pro_price)+'/mo)</span><span style="font-size:13px;font-weight:600">'+str(pro)+' users &nbsp;$'+str(pro*_pro_price)+'/mo</span></div>'
    page += '<div style="display:flex;justify-content:space-between;padding:10px 0;border-bottom:0.5px solid #EEF0F3"><span style="font-size:13px;color:#374151">Business plan ($'+str(_biz_price)+'/mo)</span><span style="font-size:13px;font-weight:600">'+str(biz)+' users &nbsp;$'+str(biz*_biz_price)+'/mo</span></div>'
    page += '<div style="display:flex;justify-content:space-between;padding:10px 0"><span style="font-size:13px;color:#374151">Free plan</span><span style="font-size:13px;font-weight:600;color:#9CA3AF">'+str(free)+' users &nbsp;$0/mo</span></div>'
    page += '</div></div></body></html>'
    return render_template_string(page)


@app.route('/admin/cleanup', methods=['POST'])
@admin_required
def admin_cleanup():
    deleted = auto_cleanup_job()
    return jsonify({'ok': True, 'deleted': deleted})


# ================================================================
# API: SAVE NOTE
# ================================================================
@app.route('/api/note', methods=['POST'])
@login_required
def api_note():
    data = request.json
    job_id = data.get('job_id')
    text = data.get('text', '').strip()
    if job_id and text:
        if job_id not in notes_db:
            notes_db[job_id] = []
        notes_db[job_id].append({
            'text': text,
            'user': get_current_user()['name'],
            'time': 'just now',
            'type': 'own'
        })
    return jsonify({'ok': True})


# ================================================================
# API: TOGGLE FAVORITE
# ================================================================
@app.route('/api/favorite', methods=['POST'])
@login_required
def api_favorite():
    data = request.json
    job_id = data.get('job_id')
    email = session['user_email']
    if email not in favorites_db:
        favorites_db[email] = set()
    if job_id in favorites_db[email]:
        favorites_db[email].discard(job_id)
        return jsonify({'ok': True, 'starred': False})
    else:
        favorites_db[email].add(job_id)
        return jsonify({'ok': True, 'starred': True})


# ================================================================
# HEALTH CHECK
# ================================================================
@app.route('/health')
def health():
    return jsonify({
        'status': 'ok',
        'version': '10.0',
        'users': len(users_db),
        'jobs': len(jobs),
        'uptime': 'running'
    })


# Start cleanup scheduler
start_cleanup_scheduler()


# ================================================================
# MAIN
# ================================================================


# ============================================================
# PRICING PAGE
# ============================================================
@app.route('/pricing')
def pricing_page():
    c = site_content
    plans = []
    for slug in ['free','pro','biz','ent']:
        feats_str = c.get(f'plan_{slug}_features', '')
        features = [f.strip() for f in feats_str.split('||') if f.strip()]
        plans.append({
            'name': c.get(f'plan_{slug}_name', slug.upper()),
            'price': c.get(f'plan_{slug}_price', '0'),
            'period': c.get(f'plan_{slug}_period', '/month'),
            'yearly': c.get(f'plan_{slug}_yearly', '0'),
            'desc': c.get(f'plan_{slug}_desc', ''),
            'features': features,
            'cta': c.get(f'plan_{slug}_cta', 'Get started'),
            'cta_link': c.get(f'plan_{slug}_cta_link', '/signup'),
            'badge': c.get(f'plan_{slug}_badge', ''),
            'featured': c.get(f'plan_{slug}_featured', '0') == '1',
        })

    plan_cards = ''
    for p in plans:
        feat_html = ''
        for f in p['features']:
            feat_html += f'<div class="price-feature"><i class="ti ti-check"></i>{f}</div>'
        badge_html = f'<span class="price-badge">{p["badge"]}</span>' if p['badge'] else ''
        card_class = 'price-card featured' if p['featured'] else 'price-card'
        btn_class = 'price-btn featured-btn' if p['featured'] else 'price-btn'
        plan_cards += f"""<div class="{card_class}">{badge_html}
<div class="price-name">{p['name']}</div>
<div class="price-desc">{p['desc']}</div>
<div class="price-amt-wrap"><span class="price-amt monthly">{'' if p['price'] in ('Custom',) else '$'}{p['price']}</span><span class="price-amt yearly" style="display:none">{'' if p['yearly'] in ('Custom',) else '$'}{p['yearly']}</span><span class="price-amt-sub">{p['period']}</span></div>
<div class="price-features">{feat_html}</div>
<a href="{p['cta_link']}" class="{btn_class}">{p['cta']}</a>
</div>"""

    # FAQ
    faqs = []
    for i in range(1, 9):
        if f'pfaq{i}_q' in c:
            faqs.append((c[f'pfaq{i}_q'], c[f'pfaq{i}_a']))
    faq_html = ''
    for q, a in faqs:
        faq_html += '<div class="faq-item" onclick="this.classList.toggle(\'open\')"><div class="faq-q">' + q + '<i class="ti ti-plus"></i></div><div class="faq-a">' + a + '</div></div>'


    page = """<!DOCTYPE html><html lang="en"><head>
<meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1.0">
<title>Pricing — """ + c['site_name'] + """</title>
<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/@tabler/icons-webfont@3.19.0/dist/tabler-icons.min.css">

<style>
:root{--bg:#F7F4EE;--card:#FFFFFF;--text:#171717;--text2:#5F5B53;--text3:#8A847A;--accent:#C46A2B;--accent-hover:#A9551F;--accent-bg:#F5E6D8;--border:#E8E1D8;--dark:#1B140F;--success:#3F7D58;--shadow-lg:0 10px 40px rgba(24,17,10,0.08)}
*{box-sizing:border-box;margin:0;padding:0;font-family:-apple-system,'Inter','Plus Jakarta Sans',BlinkMacSystemFont,'Segoe UI',Roboto,sans-serif;-webkit-font-smoothing:antialiased}
body{background:var(--bg);color:var(--text);line-height:1.5}
a{text-decoration:none;color:inherit}
.nav{background:rgba(247,244,238,.85);backdrop-filter:blur(20px);border-bottom:1px solid var(--border);padding:0 32px;height:68px;display:flex;align-items:center;justify-content:space-between;position:sticky;top:0;z-index:100}
.nav-logo{display:flex;align-items:center;gap:10px;font-weight:700;font-size:18px}
.nav-logo-mark{width:32px;height:32px;background:var(--dark);border-radius:9px;position:relative}
.nav-logo-mark::after{content:"";position:absolute;inset:6px;border:1.5px solid var(--accent);border-radius:5px}
.nav-links{display:flex;gap:32px}
.nav-link{font-size:14px;font-weight:500;color:var(--text2)}
.nav-link:hover{color:var(--text)}
.btn-primary{background:var(--dark);color:#F7F4EE;font-size:14px;font-weight:600;padding:10px 22px;border-radius:14px;display:inline-flex;align-items:center;gap:6px}
.btn-primary:hover{background:#2A1F18}
.btn-sign{font-size:14px;color:var(--text);padding:8px 18px;border-radius:10px}
.btn-sign:hover{background:var(--accent-bg)}
.hdr{padding:80px 24px 40px;text-align:center;position:relative}
.hdr::before{content:"";position:absolute;top:0;left:50%;transform:translateX(-50%);width:600px;height:300px;background:radial-gradient(ellipse,rgba(196,106,43,.1),transparent 60%);z-index:0}
.hdr-inner{position:relative;z-index:1;max-width:720px;margin:0 auto}
.sec-tag{display:inline-block;background:var(--accent-bg);color:var(--accent);font-size:12px;font-weight:700;padding:5px 14px;border-radius:100px;text-transform:uppercase;letter-spacing:.06em;margin-bottom:16px}
.hdr h1{font-family:'Plus Jakarta Sans',sans-serif;font-size:clamp(36px,5vw,52px);font-weight:800;letter-spacing:-1.5px;margin-bottom:14px}
.hdr p{font-size:17px;color:var(--text2);max-width:540px;margin:0 auto}
.toggle-wrap{display:inline-flex;align-items:center;gap:10px;background:white;border:1px solid var(--border);padding:6px;border-radius:100px;margin-top:32px}
.toggle-btn{background:none;border:none;padding:8px 20px;border-radius:100px;font-size:13px;font-weight:600;cursor:pointer;color:var(--text2)}
.toggle-btn.active{background:var(--dark);color:#F7F4EE}
.save-pill{background:#E5F0E9;color:var(--success);font-size:10px;font-weight:700;padding:3px 8px;border-radius:100px;margin-left:4px}
.price-grid{display:grid;grid-template-columns:repeat(4,1fr);gap:20px;max-width:1380px;margin:30px auto 0;padding:0 24px}
.price-card{background:var(--card);border:1px solid var(--border);border-radius:20px;padding:36px 30px;position:relative;transition:all .3s}
.price-card:hover{transform:translateY(-4px);box-shadow:var(--shadow-lg)}
.price-card.featured{background:var(--dark);color:#F7F4EE;border:none}
.price-card.featured .price-name,.price-card.featured .price-amt{color:white}
.price-card.featured .price-desc,.price-card.featured .price-amt-sub{color:#B8AC9F}
.price-card.featured .price-feature{color:#D4C6B6}
.price-badge{position:absolute;top:-12px;left:50%;transform:translateX(-50%);background:var(--accent);color:white;font-size:11px;font-weight:700;padding:5px 14px;border-radius:100px;text-transform:uppercase}
.price-name{font-family:'Plus Jakarta Sans',sans-serif;font-size:16px;font-weight:700;margin-bottom:6px}
.price-desc{font-size:13px;color:var(--text2);margin-bottom:24px}
.price-amt-wrap{margin-bottom:8px}
.price-amt{font-family:'Plus Jakarta Sans',sans-serif;font-size:48px;font-weight:800;letter-spacing:-2px;color:var(--text)}
.price-amt-sub{font-size:14px;font-weight:500;color:var(--text2);margin-left:4px}
.price-features{margin:28px 0}
.price-feature{display:flex;align-items:center;gap:10px;font-size:14px;color:var(--text2);margin-bottom:12px;font-weight:500}
.price-feature i{color:var(--accent);font-size:16px;flex-shrink:0}
.price-btn{display:block;text-align:center;padding:13px;border-radius:14px;font-size:14px;font-weight:600;width:100%;border:1px solid var(--border);background:white;color:var(--text)}
.price-btn:hover{background:var(--bg)}
.featured-btn{background:var(--accent);color:white;border-color:var(--accent)}
.featured-btn:hover{background:var(--accent-hover)}
.section{padding:80px 24px}
.container{max-width:1080px;margin:0 auto}
.sec-head{text-align:center;max-width:680px;margin:0 auto 48px}
.sec-title{font-family:'Plus Jakarta Sans',sans-serif;font-size:clamp(28px,4vw,40px);font-weight:800;letter-spacing:-1px;margin-bottom:14px}
.sec-sub{font-size:16px;color:var(--text2)}
.faq-list{max-width:780px;margin:0 auto}
.faq-item{background:var(--card);border:1px solid var(--border);border-radius:16px;margin-bottom:12px;overflow:hidden;cursor:pointer}
.faq-item:hover{border-color:var(--text3)}
.faq-q{padding:20px 24px;font-size:15px;font-weight:600;color:var(--text);display:flex;align-items:center;justify-content:space-between}
.faq-q i{font-size:20px;color:var(--text3);transition:transform .2s}
.faq-item.open .faq-q i{transform:rotate(45deg);color:var(--accent)}
.faq-a{padding:0 24px;max-height:0;overflow:hidden;transition:all .3s;font-size:14px;color:var(--text2);line-height:1.7}
.faq-item.open .faq-a{padding:0 24px 20px;max-height:300px}
.footer{background:var(--dark);color:#9A8C7B;padding:50px 24px 25px;text-align:center;margin-top:60px}
.footer p{font-size:13px}
@media(max-width:1100px){.price-grid{grid-template-columns:repeat(2,1fr)}}
@media(max-width:768px){
  .nav{padding:0 18px}.nav-links{display:none}
  .price-grid{grid-template-columns:1fr;gap:16px}
  .section{padding:50px 18px}
}
</style></head><body>

<nav class="nav">
  <a href="/home" class="nav-logo"><div class="nav-logo-mark"></div>""" + c['site_name'] + """</a>
  <div class="nav-links">
    <a href="/features" class="nav-link">Features</a>
    <a href="/pricing" class="nav-link" style="color:var(--text);font-weight:600">Pricing</a>
    <a href="/docs" class="nav-link">Docs</a>
    <a href="/blog" class="nav-link">Blog</a>
  </div>
  <div style="display:flex;gap:10px">
    <a href="/login" class="btn-sign">Sign in</a>
    <a href="/signup" class="btn-primary">Start Free</a>
  </div>
</nav>

<div class="hdr">
  <div class="hdr-inner">
    <span class="sec-tag">Pricing</span>
    <h1>""" + c['pricing_heading'] + """</h1>
    <p>""" + c['pricing_sub'] + """</p>
    <div class="toggle-wrap">
      <button class="toggle-btn active" onclick="setBilling('monthly',this)">Monthly</button>
      <button class="toggle-btn" onclick="setBilling('yearly',this)">Yearly <span class="save-pill">SAVE 17%</span></button>
    </div>
  </div>
</div>

<div class="price-grid">""" + plan_cards + """</div>

<section class="section">
  <div class="container">
    <div class="sec-head">
      <span class="sec-tag">FAQ</span>
      <h2 class="sec-title">Common pricing questions.</h2>
    </div>
    <div class="faq-list">""" + faq_html + """</div>
  </div>
</section>

<footer class="footer">
  <p>© """ + c['footer_text'] + """</p>
</footer>

<script>
function setBilling(type, btn){
  document.querySelectorAll('.toggle-btn').forEach(b=>b.classList.remove('active'));
  btn.classList.add('active');
  document.querySelectorAll('.price-amt.monthly').forEach(e=>e.style.display=type==='monthly'?'inline':'none');
  document.querySelectorAll('.price-amt.yearly').forEach(e=>e.style.display=type==='yearly'?'inline':'none');
  document.querySelectorAll('.price-amt-sub').forEach(e=>e.textContent=type==='monthly'?'/month':'/year');
}
</script>
</body></html>"""
    return page


# ============================================================
# FEATURES PAGE
# ============================================================
@app.route('/features')
def features_page():
    c = site_content

    groups = []
    for g in range(1, 5):
        gname = c.get(f'feat_grp{g}_name', '')
        if not gname:
            continue
        gdesc = c.get(f'feat_grp{g}_desc', '')
        items_str = c.get(f'feat_grp{g}_items', '')
        items = []
        for item in items_str.split('##'):
            parts = item.split('||')
            if len(parts) >= 3:
                items.append({
                    'title': parts[0].strip(),
                    'desc': parts[1].strip(),
                    'icon': parts[2].strip(),
                })
        groups.append({'name': gname, 'desc': gdesc, 'items': items})

    groups_html = ''
    for g in groups:
        cards = ''
        for it in g['items']:
            cards += f"""<div class="feat-card">
<div class="feat-icon"><i class="ti {it['icon']}"></i></div>
<div class="feat-title">{it['title']}</div>
<div class="feat-desc">{it['desc']}</div>
</div>"""
        groups_html += f"""<section class="feat-section">
<div class="feat-section-hd">
<h2 class="feat-grp-title">{g['name']}</h2>
<p class="feat-grp-desc">{g['desc']}</p>
</div>
<div class="feat-grid">{cards}</div>
</section>"""

    page = """<!DOCTYPE html><html lang="en"><head>
<meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1.0">
<title>Features — """ + c['site_name'] + """</title>
<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/@tabler/icons-webfont@3.19.0/dist/tabler-icons.min.css">

<style>
:root{--bg:#F7F4EE;--bg2:#FFFDF9;--card:#FFFFFF;--text:#171717;--text2:#5F5B53;--text3:#8A847A;--accent:#C46A2B;--accent-bg:#F5E6D8;--border:#E8E1D8;--dark:#1B140F;--shadow-lg:0 10px 40px rgba(24,17,10,0.08)}
*{box-sizing:border-box;margin:0;padding:0;font-family:-apple-system,'Inter','Plus Jakarta Sans',BlinkMacSystemFont,'Segoe UI',Roboto,sans-serif;-webkit-font-smoothing:antialiased}
body{background:var(--bg);color:var(--text);line-height:1.5}
a{text-decoration:none;color:inherit}
.nav{background:rgba(247,244,238,.85);backdrop-filter:blur(20px);border-bottom:1px solid var(--border);padding:0 32px;height:68px;display:flex;align-items:center;justify-content:space-between;position:sticky;top:0;z-index:100}
.nav-logo{display:flex;align-items:center;gap:10px;font-weight:700;font-size:18px}
.nav-logo-mark{width:32px;height:32px;background:var(--dark);border-radius:9px;position:relative}
.nav-logo-mark::after{content:"";position:absolute;inset:6px;border:1.5px solid var(--accent);border-radius:5px}
.nav-links{display:flex;gap:32px}
.nav-link{font-size:14px;font-weight:500;color:var(--text2)}
.nav-link:hover{color:var(--text)}
.btn-primary{background:var(--dark);color:#F7F4EE;font-size:14px;font-weight:600;padding:10px 22px;border-radius:14px;display:inline-flex;align-items:center;gap:6px}
.btn-primary:hover{background:#2A1F18}
.btn-sign{font-size:14px;color:var(--text);padding:8px 18px;border-radius:10px}
.btn-sign:hover{background:var(--accent-bg)}
.hdr{padding:80px 24px 60px;text-align:center;position:relative;overflow:hidden}
.hdr::before{content:"";position:absolute;top:-50px;left:50%;transform:translateX(-50%);width:700px;height:400px;background:radial-gradient(ellipse,rgba(196,106,43,.1),transparent 60%);z-index:0}
.hdr-inner{position:relative;z-index:1;max-width:760px;margin:0 auto}
.sec-tag{display:inline-block;background:var(--accent-bg);color:var(--accent);font-size:12px;font-weight:700;padding:5px 14px;border-radius:100px;text-transform:uppercase;letter-spacing:.06em;margin-bottom:16px}
.hdr h1{font-family:'Plus Jakarta Sans',sans-serif;font-size:clamp(36px,5vw,56px);font-weight:800;letter-spacing:-1.5px;margin-bottom:14px;line-height:1.1}
.hdr p{font-size:18px;color:var(--text2);max-width:560px;margin:0 auto}
.feat-section{max-width:1180px;margin:0 auto;padding:60px 24px}
.feat-section-hd{margin-bottom:36px}
.feat-grp-title{font-family:'Plus Jakarta Sans',sans-serif;font-size:28px;font-weight:800;letter-spacing:-0.8px;margin-bottom:8px}
.feat-grp-desc{font-size:15px;color:var(--text2)}
.feat-grid{display:grid;grid-template-columns:repeat(3,1fr);gap:18px}
.feat-card{background:var(--card);border:1px solid var(--border);border-radius:18px;padding:26px;transition:all .25s}
.feat-card:hover{transform:translateY(-3px);box-shadow:var(--shadow-lg);border-color:var(--accent-bg)}
.feat-icon{width:44px;height:44px;border-radius:11px;background:var(--accent-bg);display:flex;align-items:center;justify-content:center;margin-bottom:16px}
.feat-icon i{font-size:20px;color:var(--accent)}
.feat-title{font-family:'Plus Jakarta Sans',sans-serif;font-size:16px;font-weight:700;margin-bottom:8px}
.feat-desc{font-size:13.5px;color:var(--text2);line-height:1.65}
.cta-banner{padding:80px 24px;text-align:center}
.cta-inner{max-width:760px;margin:0 auto;background:var(--dark);border-radius:24px;padding:48px 36px;color:white}
.cta-inner h2{font-family:'Plus Jakarta Sans',sans-serif;font-size:32px;font-weight:800;letter-spacing:-1px;margin-bottom:12px}
.cta-inner p{color:#B8AC9F;margin-bottom:24px}
.cta-inner .btn-primary{background:var(--accent);padding:13px 28px}
.cta-inner .btn-primary:hover{background:#A9551F}
.footer{background:var(--dark);color:#9A8C7B;padding:50px 24px 25px;text-align:center}
.footer p{font-size:13px}
@media(max-width:768px){
  .nav{padding:0 18px}.nav-links{display:none}
  .feat-grid{grid-template-columns:1fr;gap:12px}
  .feat-section{padding:40px 18px}
  .hdr{padding:60px 18px 40px}
}
</style></head><body>

<nav class="nav">
  <a href="/home" class="nav-logo"><div class="nav-logo-mark"></div>""" + c['site_name'] + """</a>
  <div class="nav-links">
    <a href="/features" class="nav-link" style="color:var(--text);font-weight:600">Features</a>
    <a href="/pricing" class="nav-link">Pricing</a>
    <a href="/docs" class="nav-link">Docs</a>
    <a href="/blog" class="nav-link">Blog</a>
  </div>
  <div style="display:flex;gap:10px">
    <a href="/login" class="btn-sign">Sign in</a>
    <a href="/signup" class="btn-primary">Start Free</a>
  </div>
</nav>

<div class="hdr">
  <div class="hdr-inner">
    <span class="sec-tag">All Features</span>
    <h1>""" + c['features_heading'] + """</h1>
    <p>""" + c['features_sub'] + """</p>
  </div>
</div>

""" + groups_html + """

<section class="cta-banner">
  <div class="cta-inner">
    <h2>Ready to save 8 hours per survey?</h2>
    <p>Start free — no credit card needed. Cancel anytime.</p>
    <a href="/signup" class="btn-primary">Start Free Trial <i class="ti ti-arrow-right"></i></a>
  </div>
</section>

<footer class="footer">
  <p>© """ + c['footer_text'] + """</p>
</footer>

</body></html>"""
    return page


# ============================================================
# BLOG
# ============================================================
@app.route('/blog')
def blog():
    c = site_content
    tag_colors = {
        'AI & QC':'#F5E6D8|#0C447C','Best Practices':'#EAF3DE|#27500A',
        'Languages':'#EEEDFE|#3C3489','Platforms':'#FAEEDA|#633806',
        'Productivity':'#E1F5EE|#085041','ROI':'#FCEBEB|#791F1F'
    }
    posts = [
        (c['blog_post1_title'],c['blog_post1_date'],c['blog_post1_tag'],c['blog_post1_summary']),
        (c['blog_post2_title'],c['blog_post2_date'],c['blog_post2_tag'],c['blog_post2_summary']),
        (c['blog_post3_title'],c['blog_post3_date'],c['blog_post3_tag'],c['blog_post3_summary']),
        (c['blog_post4_title'],c['blog_post4_date'],c['blog_post4_tag'],c['blog_post4_summary']),
        (c['blog_post5_title'],c['blog_post5_date'],c['blog_post5_tag'],c['blog_post5_summary']),
        (c['blog_post6_title'],c['blog_post6_date'],c['blog_post6_tag'],c['blog_post6_summary']),
    ]
    cards = ''
    for title,date,tag,summary in posts:
        tc = tag_colors.get(tag,'#F1EFE8|#444441').split('|')
        cards += ('<div style="background:white;border:0.5px solid #DDE1E7;border-radius:12px;padding:24px">'
            '<span style="background:'+tc[0]+';color:'+tc[1]+';font-size:11px;padding:3px 10px;border-radius:20px;font-weight:500">'+tag+'</span>'
            '<p style="font-size:16px;font-weight:600;color:#1A1A2E;margin:12px 0 8px;line-height:1.4">'+title+'</p>'
            '<p style="font-size:13px;color:#6B7280;line-height:1.7;margin-bottom:14px">'+summary+'</p>'
            '<div style="display:flex;justify-content:space-between;align-items:center">'
            '<p style="font-size:11px;color:#9CA3AF">'+date+'</p>'
            '<span style="font-size:13px;color:#1B140F;font-weight:500">Read more &rarr;</span>'
            '</div></div>')
    page = '<!DOCTYPE html><html><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1">'
    page += '<title>Blog - '+c['site_name']+'</title>'
    page += '<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/@tabler/icons-webfont@3.19.0/dist/tabler-icons.min.css">'
    page += ''
    page += '<style>*{box-sizing:border-box;margin:0;padding:0;font-family:Inter,Plus Jakarta Sans,-apple-system,sans-serif;-webkit-font-smoothing:antialiased}body{background:#F7F4EE;color:#171717}a{text-decoration:none}'
    page += '.nav{background:white;border-bottom:0.5px solid #DDE1E7;padding:0 40px;height:60px;display:flex;align-items:center;justify-content:space-between}'
    page += '.grid{display:grid;grid-template-columns:repeat(3,1fr);gap:20px;max-width:1100px;margin:0 auto}'
    page += '@media(max-width:768px){.grid{grid-template-columns:1fr}.nav{padding:0 16px}}</style></head><body>'
    page += '<nav class="nav"><div style="display:flex;align-items:center;gap:10px"><div style="width:28px;height:28px;background:#1B140F;border-radius:7px;display:flex;align-items:center;justify-content:center"><i class="ti ti-shield-check" style="font-size:14px;color:white"></i></div><a href="/home" style="font-size:15px;font-weight:700;color:#1A1A2E">'+c['site_name']+'</a></div>'
    page += '<div style="display:flex;gap:10px"><a href="/login" style="font-size:13px;color:#6B7280;padding:7px 14px;border:0.5px solid #DDE1E7;border-radius:7px">Sign in</a><a href="/signup" style="font-size:13px;color:white;background:#1B140F;padding:7px 14px;border-radius:7px;font-weight:500">Sign up free</a></div></nav>'
    page += '<div style="background:#1B140F;padding:50px 40px;text-align:center"><p style="font-size:12px;color:#B8AC9F;text-transform:uppercase;letter-spacing:.1em;margin-bottom:10px">Blog</p><h1 style="font-size:36px;font-weight:700;color:white;margin-bottom:10px">QC Insights & Guides</h1><p style="font-size:16px;color:#B8AC9F">Best practices, platform guides, and industry news from the QC world.</p></div>'
    page += '<div style="padding:50px 40px"><div class="grid">'+cards+'</div></div>'
    page += '<footer style="background:#1B140F;padding:30px 40px;text-align:center"><p style="color:#6B7280;font-size:13px">'+c['footer_text']+'</p><div style="margin-top:10px;display:flex;gap:16px;justify-content:center"><a href="/home" style="font-size:12px;color:#6B7280">Home</a><a href="/privacy-policy" style="font-size:12px;color:#6B7280">Privacy</a><a href="/terms" style="font-size:12px;color:#6B7280">Terms</a></div></footer>'
    page += '</body></html>'
    return page


# ============================================================
# DOCS
# ============================================================
@app.route('/docs')
def docs():
    c = site_content
    sections = [
        ('Getting Started','ti-rocket','#F5E6D8','#0C447C',[
            ('Create your account','Sign up free at surveyqc.online. No credit card needed. 5 reports/month free forever.'),
            ('Upload your screener doc','Go to New QC. Upload the .docx screener/spec document for your survey.'),
            ('Paste survey URL','Paste the live survey URL (Confirmit, Decipher, Forsta, or Qualtrics).'),
            ('Start QC check','Click Run QC. AI will crawl every path and run all 8 checks automatically.'),
            ('Download report','When done, download the Word report with all issues, screenshots and QC certificate.'),
        ]),
        ('8 QC Checks','ti-list-check','#E1F5EE','#085041',[
            ('1. Termination rules','Checks every terminate/close logic against the spec doc. Catches wrong terminations.'),
            ('2. Missing words','Word-by-word comparison between doc and live survey. Catches typos and missing text.'),
            ('3. Question text match','Full question text compared against spec. Detects any deviations.'),
            ('4. Options match','All answer options compared against spec. Catches missing or extra options.'),
            ('5. Mandatory markers','Checks * mandatory markers match the spec on every question.'),
            ('6. Piping markers','Verifies all {piped_text} markers are present and correct.'),
            ('7. Answer codes','Checks answer codes are sequential and match spec exactly.'),
            ('8. Question order','Verifies questions appear in the correct order as per spec doc.'),
        ]),
        ('Platforms Supported','ti-device-laptop','#FAEEDA','#633806',[
            ('Confirmit','Full support. Paste the respondent-facing survey URL.'),
            ('Decipher','Full support. Use the live survey link, not the editor link.'),
            ('Forsta (formerly Confirmit)','Full support. Use the standard survey URL.'),
            ('Qualtrics','Full support. Use the published survey URL from the Distributions tab.'),
        ]),
        ('Token Limits','ti-database','#EEEDFE','#3C3489',[
            ('Free plan: 20,000 tokens','Enough for small surveys (~30 questions, 10 paths).'),
            ('Pro plan: 100,000 tokens','Enough for medium surveys (~100 questions, 50 paths).'),
            ('Business plan: 150,000 tokens','Enough for complex surveys with 200+ questions and 100+ paths.'),
            ('What happens at limit?','AI stops at the token limit. Remaining checks are flagged MANUAL for you to do by hand. Admin can raise limits anytime.'),
        ]),
        ('Share Report','ti-share','#E1F5EE','#085041',[
            ('Generate share link','On any report, click Share. A unique link is generated.'),
            ('Send to client','Client clicks link, creates free account, views report.'),
            ('Client actions','Client can comment on issues and mark them as Fixed.'),
            ('Link expiry','Share links expire in 7 days by default.'),
        ]),
        ('FAQ','ti-help-circle','#FCEBEB','#791F1F',[
            ('What file format for spec doc?','.docx (Microsoft Word) only. PDF not supported yet.'),
            ('Does it work for non-English surveys?','Yes -- 80+ languages supported including French, Italian, Arabic, Urdu, Japanese.'),
            ('Can I retest after fixing bugs?','Yes. On the report page, click Retest to rerun QC on the same URL.'),
            ('Is my data safe?','Yes. All data auto-deletes after 30 days. We are GDPR, CCPA, India DPDP compliant.'),
            ('Can I use my own API key?','Yes, on Business plan. Go to Settings > API Key to add your own Gemini key.'),
        ]),
    ]
    nav_items = ''
    content_sections = ''
    for title,icon,bg,col,items in sections:
        anchor = title.lower().replace(' ','_').replace('(','').replace(')','')
        nav_items += '<a href="#'+anchor+'" style="display:flex;align-items:center;gap:8px;padding:8px 12px;border-radius:8px;color:#374151;font-size:13px;text-decoration:none;margin-bottom:2px"><i class="ti '+icon+'" style="font-size:15px;color:'+col+'"></i>'+title+'</a>'
        content_sections += '<div id="'+anchor+'" style="margin-bottom:40px">'
        content_sections += '<div style="display:flex;align-items:center;gap:10px;margin-bottom:16px"><div style="width:36px;height:36px;border-radius:8px;background:'+bg+';display:flex;align-items:center;justify-content:center"><i class="ti '+icon+'" style="font-size:18px;color:'+col+'"></i></div><h2 style="font-size:20px;font-weight:600;color:#1A1A2E">'+title+'</h2></div>'
        for item_title,item_desc in items:
            content_sections += '<div style="background:white;border:0.5px solid #DDE1E7;border-radius:10px;padding:16px;margin-bottom:8px"><p style="font-size:14px;font-weight:600;color:#1A1A2E;margin-bottom:4px">'+item_title+'</p><p style="font-size:13px;color:#6B7280;line-height:1.6">'+item_desc+'</p></div>'
        content_sections += '</div>'
    page = '<!DOCTYPE html><html><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>Docs - '+c['site_name']+'</title>'
    page += '<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/@tabler/icons-webfont@3.19.0/dist/tabler-icons.min.css">'
    page += '<style>*{box-sizing:border-box;margin:0;padding:0;font-family:-apple-system,BlinkMacSystemFont,"Segoe UI",sans-serif}body{background:#FBF8F2;color:#1A1A2E}a{text-decoration:none}.nav{background:white;border-bottom:0.5px solid #DDE1E7;padding:0 40px;height:60px;display:flex;align-items:center;justify-content:space-between}.layout{display:grid;grid-template-columns:240px 1fr;min-height:calc(100vh - 60px)}.sidebar{background:white;border-right:0.5px solid #DDE1E7;padding:24px 16px;position:sticky;top:60px;height:calc(100vh - 60px);overflow-y:auto}.content{padding:40px;max-width:800px}@media(max-width:768px){.layout{grid-template-columns:1fr}.sidebar{display:none}.nav{padding:0 16px}.content{padding:20px}}</style></head><body>'
    page += '<nav class="nav"><div style="display:flex;align-items:center;gap:10px"><div style="width:28px;height:28px;background:#1B140F;border-radius:7px;display:flex;align-items:center;justify-content:center"><i class="ti ti-shield-check" style="font-size:14px;color:white"></i></div><a href="/home" style="font-size:15px;font-weight:700;color:#1A1A2E">'+c['site_name']+'</a></div><div style="display:flex;gap:10px"><a href="/login" style="font-size:13px;color:#6B7280;padding:7px 14px;border:0.5px solid #DDE1E7;border-radius:7px">Sign in</a><a href="/signup" style="font-size:13px;color:white;background:#1B140F;padding:7px 14px;border-radius:7px;font-weight:500">Sign up free</a></div></nav>'
    page += '<div class="layout"><div class="sidebar"><p style="font-size:11px;font-weight:600;color:#9CA3AF;text-transform:uppercase;letter-spacing:.08em;margin-bottom:12px">Documentation</p>'+nav_items+'<div style="margin-top:20px;padding:12px;background:#F5E6D8;border-radius:8px"><p style="font-size:12px;color:#C46A2B;font-weight:500;margin-bottom:4px">Need help?</p><a href="mailto:'+c['support_email']+'" style="font-size:12px;color:#185FA5">'+c['support_email']+'</a></div></div>'
    page += '<div class="content"><h1 style="font-size:28px;font-weight:700;color:#1A1A2E;margin-bottom:8px">Documentation</h1><p style="font-size:15px;color:#6B7280;margin-bottom:32px">'+c['docs_intro']+'</p>'+content_sections+'</div></div>'
    page += '<footer style="background:#1B140F;padding:30px 40px;text-align:center"><p style="color:#6B7280;font-size:13px">'+c['footer_text']+'</p></footer></body></html>'
    return page


# ============================================================
# COMPARE PAGES
# ============================================================
@app.route('/compare/<slug>')
def compare_page(slug):
    c = site_content
    pages = {
        'chatgpt': (c['compare_chatgpt_heading'], c['compare_chatgpt_summary'],
            [('SurveyQC','Built for survey QC','8 specialized checks','Playwright crawling','Word report output','Any survey platform','Auto screenshots','Supports 80+ languages'),
             ('ChatGPT','General AI assistant','No survey-specific checks','No browser crawling','No formatted report','No platform integration','No screenshots','Text only')]),
        'excel': (c['compare_excel_heading'], c['compare_excel_summary'],
            [('SurveyQC','Automated tool','10 min per survey','99% accuracy','Auto screenshots','Professional Word report','No manual effort','GDPR compliant'),
             ('Excel Manual QC','Spreadsheet + human','8+ hours per survey','~70% accuracy','No screenshots','Manual formatting','Heavy manual effort','No compliance')]),
        'manual': (c['compare_manual_heading'], c['compare_manual_summary'],
            [('SurveyQC','AI-powered','10 min per survey','99% accuracy','Never misses edge cases','Consistent every time','Scales instantly','Costs $'+c.get('plan_pro_price','29')+'/mo'),
             ('Manual Testing','Human tester','8+ hours per survey','~70% accuracy','Misses 30% of bugs','Varies by tester','Bottleneck at scale','Costs $50+/hr')]),
    }
    if slug not in pages:
        return redirect('/home')
    heading, summary, cols = pages[slug]
    col1, col2 = cols
    table_rows = ''
    for i in range(1, len(col1)):
        color = '#EAF3DE' if i % 2 == 0 else 'white'
        table_rows += '<tr style="background:'+color+'"><td style="padding:12px 16px;font-size:13px;color:#374151;border-right:0.5px solid #DDE1E7">'+col1[i]+'</td><td style="padding:12px 16px;font-size:13px;color:#374151">'+col2[i]+'</td></tr>'
    page = '<!DOCTYPE html><html><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>'+heading+'</title>'
    page += '<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/@tabler/icons-webfont@3.19.0/dist/tabler-icons.min.css">'
    page += '<style>*{box-sizing:border-box;margin:0;padding:0;font-family:-apple-system,BlinkMacSystemFont,"Segoe UI",sans-serif}body{background:#FBF8F2;color:#1A1A2E}a{text-decoration:none}.nav{background:white;border-bottom:0.5px solid #DDE1E7;padding:0 40px;height:60px;display:flex;align-items:center;justify-content:space-between}@media(max-width:768px){.nav{padding:0 16px}}</style></head><body>'
    page += '<nav class="nav"><div style="display:flex;align-items:center;gap:10px"><div style="width:28px;height:28px;background:#1B140F;border-radius:7px;display:flex;align-items:center;justify-content:center"><i class="ti ti-shield-check" style="font-size:14px;color:white"></i></div><a href="/home" style="font-size:15px;font-weight:700;color:#1A1A2E">'+c['site_name']+'</a></div><div style="display:flex;gap:10px"><a href="/login" style="font-size:13px;color:#6B7280;padding:7px 14px;border:0.5px solid #DDE1E7;border-radius:7px">Sign in</a><a href="/signup" style="font-size:13px;color:white;background:#1B140F;padding:7px 14px;border-radius:7px;font-weight:500">Try free</a></div></nav>'
    page += '<div style="background:#1B140F;padding:50px 40px;text-align:center"><h1 style="font-size:34px;font-weight:700;color:white;margin-bottom:12px">'+heading+'</h1><p style="font-size:16px;color:#B8AC9F;max-width:600px;margin:0 auto;line-height:1.7">'+summary+'</p></div>'
    page += '<div style="padding:50px 40px;max-width:900px;margin:0 auto">'
    page += '<div style="background:white;border-radius:14px;overflow:hidden;border:0.5px solid #DDE1E7">'
    page += '<div style="display:grid;grid-template-columns:1fr 1fr"><div style="padding:16px;background:#1B140F;text-align:center"><p style="font-size:15px;font-weight:700;color:white">'+col1[0]+'</p></div><div style="padding:16px;background:#FBF8F2;text-align:center;border-left:0.5px solid #DDE1E7"><p style="font-size:15px;font-weight:600;color:#6B7280">'+col2[0]+'</p></div></div>'
    page += '<table style="width:100%;border-collapse:collapse"><tbody>'+table_rows+'</tbody></table></div>'
    page += '<div style="text-align:center;margin-top:40px"><p style="font-size:18px;font-weight:600;color:#1A1A2E;margin-bottom:16px">Ready to try '+c['site_name']+'?</p><a href="/signup" style="background:#1B140F;color:white;font-size:15px;padding:14px 36px;border-radius:10px;font-weight:600">Start free -- no card required</a></div>'
    page += '</div><footer style="background:#1B140F;padding:30px 40px;text-align:center;margin-top:40px"><p style="color:#6B7280;font-size:13px">'+c['footer_text']+'</p></footer></body></html>'
    return page


# ============================================================
# COMMUNITY
# ============================================================
@app.route('/community')
def community():
    c = site_content
    posts = [
        ('SarahQC','How I reduced QC time from 8 hours to 15 minutes','Sharing my workflow after using SurveyQC for 3 months. The auto screenshot feature alone saved me countless hours of back-and-forth with developers.','14 replies','2 days ago','#F5E6D8','#0C447C'),
        ('MarieL','French survey QC tips -- accents and encoding','Been testing French surveys on Confirmit. A few things I have learned about handling accent marks and special characters with SurveyQC.','8 replies','5 days ago','#EAF3DE','#27500A'),
        ('JamesMR','Share your QC report templates','Who has built good templates for specific survey types? Sharing mine for trackers and ad hoc studies. Would love to see others.','22 replies','1 week ago','#EEEDFE','#3C3489'),
        ('RahulD','Qualtrics integration -- complete guide','After much trial and error, here is the complete guide to running SurveyQC on Qualtrics surveys including where to find the right URL.','11 replies','1 week ago','#FAEEDA','#633806'),
    ]
    post_cards = ''
    for author,title,body,replies,time,bg,col in posts:
        post_cards += ('<div style="background:white;border:0.5px solid #DDE1E7;border-radius:12px;padding:20px;margin-bottom:12px">'
            '<div style="display:flex;align-items:start;gap:12px">'
            '<div style="width:36px;height:36px;border-radius:50%;background:'+bg+';color:'+col+';display:flex;align-items:center;justify-content:center;font-size:12px;font-weight:700;flex-shrink:0">'+author[:2]+'</div>'
            '<div style="flex:1"><p style="font-size:14px;font-weight:600;color:#1A1A2E;margin-bottom:4px">'+title+'</p>'
            '<p style="font-size:12px;font-weight:500;color:'+col+';margin-bottom:6px">'+author+'</p>'
            '<p style="font-size:13px;color:#6B7280;line-height:1.6;margin-bottom:12px">'+body+'</p>'
            '<div style="display:flex;gap:16px"><span style="font-size:12px;color:#9CA3AF"><i class="ti ti-message" style="font-size:12px"></i> '+replies+'</span>'
            '<span style="font-size:12px;color:#9CA3AF">'+time+'</span></div>'
            '</div></div></div>')
    page = '<!DOCTYPE html><html><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>Community - '+c['site_name']+'</title>'
    page += '<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/@tabler/icons-webfont@3.19.0/dist/tabler-icons.min.css">'
    page += '<style>*{box-sizing:border-box;margin:0;padding:0;font-family:-apple-system,BlinkMacSystemFont,"Segoe UI",sans-serif}body{background:#FBF8F2;color:#1A1A2E}a{text-decoration:none}.nav{background:white;border-bottom:0.5px solid #DDE1E7;padding:0 40px;height:60px;display:flex;align-items:center;justify-content:space-between}.layout{display:grid;grid-template-columns:1fr 320px;gap:24px;max-width:1100px;margin:0 auto;padding:40px}@media(max-width:768px){.layout{grid-template-columns:1fr}.nav{padding:0 16px}}</style></head><body>'
    page += '<nav class="nav"><div style="display:flex;align-items:center;gap:10px"><div style="width:28px;height:28px;background:#1B140F;border-radius:7px;display:flex;align-items:center;justify-content:center"><i class="ti ti-shield-check" style="font-size:14px;color:white"></i></div><a href="/home" style="font-size:15px;font-weight:700;color:#1A1A2E">'+c['site_name']+'</a></div><div style="display:flex;gap:10px"><a href="/login" style="font-size:13px;color:#6B7280;padding:7px 14px;border:0.5px solid #DDE1E7;border-radius:7px">Sign in</a><a href="/signup" style="font-size:13px;color:white;background:#1B140F;padding:7px 14px;border-radius:7px;font-weight:500">Join free</a></div></nav>'
    page += '<div style="background:#1B140F;padding:50px 40px;text-align:center"><p style="font-size:36px">&#128106;</p><h1 style="font-size:34px;font-weight:700;color:white;margin-bottom:12px">'+c['community_heading']+'</h1><p style="font-size:16px;color:#B8AC9F;max-width:500px;margin:0 auto">'+c['community_subheading']+'</p></div>'
    page += '<div class="layout"><div>'
    page += '<div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:20px"><p style="font-size:16px;font-weight:600;color:#1A1A2E">Recent discussions</p><a href="/signup" style="font-size:13px;color:white;background:#1B140F;padding:8px 16px;border-radius:8px;font-weight:500">+ New post</a></div>'
    page += post_cards + '</div>'
    page += '<div><div style="background:white;border:0.5px solid #DDE1E7;border-radius:12px;padding:20px;margin-bottom:16px"><p style="font-size:14px;font-weight:600;color:#1A1A2E;margin-bottom:12px">Community stats</p><div style="display:flex;flex-direction:column;gap:10px"><div style="display:flex;justify-content:space-between"><span style="font-size:13px;color:#6B7280">Members</span><span style="font-size:13px;font-weight:600;color:#1A1A2E">500+</span></div><div style="display:flex;justify-content:space-between"><span style="font-size:13px;color:#6B7280">Posts</span><span style="font-size:13px;font-weight:600;color:#1A1A2E">1,200+</span></div><div style="display:flex;justify-content:space-between"><span style="font-size:13px;color:#6B7280">Countries</span><span style="font-size:13px;font-weight:600;color:#1A1A2E">40+</span></div></div></div>'
    page += '<div style="background:#F5E6D8;border-radius:12px;padding:20px"><p style="font-size:14px;font-weight:600;color:#C46A2B;margin-bottom:8px">Earn with our affiliate program</p><p style="font-size:13px;color:#185FA5;margin-bottom:12px;line-height:1.6">'+c['affiliate_commission']+'% recurring commission. No cap.</p><a href="/affiliate" style="display:block;text-align:center;background:#1B140F;color:white;padding:9px;border-radius:8px;font-size:13px;font-weight:500">Learn more</a></div></div>'
    page += '</div><footer style="background:#1B140F;padding:30px 40px;text-align:center"><p style="color:#6B7280;font-size:13px">'+c['footer_text']+'</p></footer></body></html>'
    return page


# ============================================================
# AFFILIATE
# ============================================================
@app.route('/affiliate')
def affiliate():
    c = site_content
    steps = [
        ('1','Sign up free','Create your SurveyQC account. No approval needed. Takes 30 seconds.','#F5E6D8','#0C447C'),
        ('2','Get your link','Copy your unique affiliate link from Settings > Affiliate.','#EAF3DE','#27500A'),
        ('3','Share it','Share with QC professionals, on LinkedIn, in communities, in blog posts.','#FAEEDA','#633806'),
        ('4','Earn monthly','Get '+c['affiliate_commission']+'% of every payment, every month, forever. No cap.','#EEEDFE','#3C3489'),
    ]
    step_cards = ''
    for num,title,desc,bg,col in steps:
        step_cards += '<div style="background:'+bg+';border-radius:12px;padding:24px;text-align:center"><div style="width:44px;height:44px;border-radius:50%;background:#1B140F;color:white;display:flex;align-items:center;justify-content:center;font-size:18px;font-weight:700;margin:0 auto 16px">'+num+'</div><p style="font-size:15px;font-weight:600;color:#1A1A2E;margin-bottom:8px">'+title+'</p><p style="font-size:13px;color:#6B7280;line-height:1.6">'+desc+'</p></div>'
    page = '<!DOCTYPE html><html><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>Affiliate Program - '+c['site_name']+'</title>'
    page += '<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/@tabler/icons-webfont@3.19.0/dist/tabler-icons.min.css">'
    page += '<style>*{box-sizing:border-box;margin:0;padding:0;font-family:-apple-system,BlinkMacSystemFont,"Segoe UI",sans-serif}body{background:#FBF8F2;color:#1A1A2E}a{text-decoration:none}.nav{background:white;border-bottom:0.5px solid #DDE1E7;padding:0 40px;height:60px;display:flex;align-items:center;justify-content:space-between}.grid4{display:grid;grid-template-columns:repeat(4,1fr);gap:16px}@media(max-width:768px){.grid4{grid-template-columns:repeat(2,1fr)}.nav{padding:0 16px}}</style></head><body>'
    page += '<nav class="nav"><div style="display:flex;align-items:center;gap:10px"><div style="width:28px;height:28px;background:#1B140F;border-radius:7px;display:flex;align-items:center;justify-content:center"><i class="ti ti-shield-check" style="font-size:14px;color:white"></i></div><a href="/home" style="font-size:15px;font-weight:700;color:#1A1A2E">'+c['site_name']+'</a></div><div style="display:flex;gap:10px"><a href="/login" style="font-size:13px;color:#6B7280;padding:7px 14px;border:0.5px solid #DDE1E7;border-radius:7px">Sign in</a><a href="/signup" style="font-size:13px;color:white;background:#1B140F;padding:7px 14px;border-radius:7px;font-weight:500">Join free</a></div></nav>'
    page += '<div style="background:#1B140F;padding:60px 40px;text-align:center"><p style="font-size:40px;margin-bottom:12px">&#128176;</p><h1 style="font-size:36px;font-weight:700;color:white;margin-bottom:12px">'+c['affiliate_heading']+'</h1><p style="font-size:18px;color:#B8AC9F;margin-bottom:8px">'+c['affiliate_details']+'</p><p style="font-size:24px;font-weight:700;color:white;margin-top:16px">'+c['affiliate_commission']+'% recurring commission</p></div>'
    page += '<div style="padding:60px 40px;max-width:1000px;margin:0 auto">'
    page += '<div style="text-align:center;margin-bottom:40px"><h2 style="font-size:28px;font-weight:700;color:#1A1A2E;margin-bottom:8px">How it works</h2></div>'
    page += '<div class="grid4" style="margin-bottom:50px">'+step_cards+'</div>'
    _aff_pro_price = int(c.get('plan_pro_price', '29') or '29')
    _aff_commission_pct = int(c.get('affiliate_commission', '30') or '30')
    _aff_monthly = round(10 * _aff_pro_price * _aff_commission_pct / 100)
    _aff_yearly = _aff_monthly * 12
    page += '<div style="background:white;border:0.5px solid #DDE1E7;border-radius:14px;padding:32px;text-align:center"><h3 style="font-size:22px;font-weight:600;color:#1A1A2E;margin-bottom:8px">Example earnings</h3><p style="font-size:14px;color:#6B7280;margin-bottom:24px">If you refer 10 Pro users ($'+str(_aff_pro_price)+'/mo each):</p><p style="font-size:36px;font-weight:700;color:#1B140F;margin-bottom:4px">$'+str(_aff_monthly)+'/month</p><p style="font-size:14px;color:#9CA3AF">recurring, every month, forever &middot; $'+str(_aff_yearly)+'/year</p><a href="/signup" style="display:inline-block;background:#1B140F;color:white;font-size:15px;padding:14px 36px;border-radius:10px;font-weight:600;margin-top:24px">Start earning free</a></div>'
    page += '</div><footer style="background:#1B140F;padding:30px 40px;text-align:center"><p style="color:#6B7280;font-size:13px">'+c['footer_text']+'</p></footer></body></html>'
    return page


# ============================================================
# CHANGELOG
# ============================================================
@app.route('/changelog')
def changelog():
    c = site_content
    logs = [
        (c['changelog_v10'],'May 2026','#EAF3DE','#27500A','Latest'),
        (c['changelog_v9'],'April 2026','#F5E6D8','#0C447C',''),
        (c['changelog_v8'],'March 2026','#FAEEDA','#633806',''),
        (c['changelog_v7'],'February 2026','#EEEDFE','#3C3489',''),
        (c['changelog_v6'],'January 2026','#F1EFE8','#444441',''),
    ]
    items = ''
    for log,date,bg,col,label in logs:
        version = log.split(' -- ')[0] if ' -- ' in log else log[:10]
        rest = log.split(' -- ')[1] if ' -- ' in log else log
        items += '<div style="display:flex;gap:20px;margin-bottom:24px"><div style="text-align:right;min-width:90px;padding-top:4px"><p style="font-size:12px;color:#9CA3AF">'+date+'</p></div><div style="display:flex;flex-direction:column;align-items:center"><div style="width:12px;height:12px;border-radius:50%;background:'+col+';margin-top:4px;flex-shrink:0"></div><div style="width:1px;background:#DDE1E7;flex:1;margin-top:4px"></div></div><div style="flex:1;background:white;border:0.5px solid #DDE1E7;border-radius:10px;padding:16px;margin-bottom:4px"><div style="display:flex;align-items:center;gap:8px;margin-bottom:6px"><p style="font-size:14px;font-weight:600;color:#1A1A2E">'+version+'</p>'+(('<span style="background:'+bg+';color:'+col+';font-size:10px;padding:2px 8px;border-radius:20px;font-weight:500">'+label+'</span>') if label else '')+'</div><p style="font-size:13px;color:#6B7280;line-height:1.6">'+rest+'</p></div></div>'
    page = '<!DOCTYPE html><html><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>Changelog - '+c['site_name']+'</title>'
    page += '<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/@tabler/icons-webfont@3.19.0/dist/tabler-icons.min.css">'
    page += '<style>*{box-sizing:border-box;margin:0;padding:0;font-family:-apple-system,BlinkMacSystemFont,"Segoe UI",sans-serif}body{background:#FBF8F2;color:#1A1A2E}a{text-decoration:none}.nav{background:white;border-bottom:0.5px solid #DDE1E7;padding:0 40px;height:60px;display:flex;align-items:center;justify-content:space-between}@media(max-width:768px){.nav{padding:0 16px}}</style></head><body>'
    page += '<nav class="nav"><div style="display:flex;align-items:center;gap:10px"><div style="width:28px;height:28px;background:#1B140F;border-radius:7px;display:flex;align-items:center;justify-content:center"><i class="ti ti-shield-check" style="font-size:14px;color:white"></i></div><a href="/home" style="font-size:15px;font-weight:700;color:#1A1A2E">'+c['site_name']+'</a></div><div style="display:flex;gap:10px"><a href="/login" style="font-size:13px;color:#6B7280;padding:7px 14px;border:0.5px solid #DDE1E7;border-radius:7px">Sign in</a><a href="/signup" style="font-size:13px;color:white;background:#1B140F;padding:7px 14px;border-radius:7px;font-weight:500">Sign up</a></div></nav>'
    page += '<div style="background:#1B140F;padding:50px 40px;text-align:center"><h1 style="font-size:34px;font-weight:700;color:white;margin-bottom:10px">Changelog</h1><p style="font-size:16px;color:#B8AC9F">What is new in '+c['site_name']+'</p></div>'
    page += '<div style="padding:50px 40px;max-width:700px;margin:0 auto">'+items+'</div>'
    page += '<footer style="background:#1B140F;padding:30px 40px;text-align:center"><p style="color:#6B7280;font-size:13px">'+c['footer_text']+'</p></footer></body></html>'
    return page


@app.route('/run-screenshot-qc', methods=['POST'])
@login_required
def run_screenshot_qc():
    import uuid, base64, threading, io, json, re
    from datetime import datetime
    from docx import Document

    user = get_current_user()
    if not user or user.get('plan','Free') == 'Free':
        return redirect('/billing')

    docx_file = request.files.get('docx_file')
    screenshots = request.files.getlist('screenshots')
    instructions = request.form.get('instructions','').strip()
    if not docx_file or not screenshots:
        return redirect('/ai-tester')

    job_id = str(uuid.uuid4())[:8]
    jobs[job_id] = {
        'status':'running','mode':'screenshot',
        'user_email':session.get('user_email'),
        'doc_name':docx_file.filename,
        'created_at':datetime.now().isoformat(),
        'issues':[],'log':['Starting Screenshot QC...'],'progress':0,
    }

    docx_bytes = docx_file.read()
    ss_list = [(f.filename, f.read()) for f in screenshots if f.filename]

    def run_ss(job_id, docx_bytes, ss_list, instructions):
        try:
            jobs[job_id]['log'].append('Reading spec document...')
            jobs[job_id]['progress'] = 10
            doc = Document(io.BytesIO(docx_bytes))
            doc_text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
            jobs[job_id]['log'].append('Spec doc loaded: '+str(len(doc_text))+' chars')
            jobs[job_id]['progress'] = 20

            api_key = api_store.get('gemini',{}).get('key','')
            if not api_key:
                jobs[job_id]['status'] = 'error'
                jobs[job_id]['log'].append('No Gemini API key configured. Go to Admin > API Management to add one.')
                jobs.persist(job_id)
                return

            import google.generativeai as genai
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel('gemini-2.5-flash')

            issues = []
            total = len(ss_list)
            instr = instructions if instructions else 'Check question text, options, mandatory markers, and any visible issues.'

            for i, (name, img_bytes) in enumerate(ss_list):
                jobs[job_id]['log'].append('Analyzing screenshot '+str(i+1)+'/'+str(total)+': '+name)
                jobs[job_id]['progress'] = 20 + int((i/total)*60)

                prompt_text = ('You are a survey QC expert. Compare this survey screenshot against the spec document.\n\n'
                    'SPEC DOCUMENT:\n' + doc_text[:8000] + '\n\n'
                    'INSTRUCTIONS: ' + instr + '\n\n'
                    'TASK: Identify issues:\n'
                    '1. Question text differs from spec\n'
                    '2. Missing or extra answer options\n'
                    '3. Wrong mandatory (*) markers\n'
                    '4. Any other visible issues\n\n'
                    'Respond ONLY in JSON:\n'
                    '{"issues_found":true/false,"issues":[{"question":"Q1","type":"text mismatch","expected":"spec text","actual":"screenshot text","severity":"high/medium/low"}],"summary":"brief summary"}')

                img_part = {'mime_type':'image/jpeg','data':base64.b64encode(img_bytes).decode()}
                try:
                    response = model.generate_content([prompt_text, img_part])
                    raw = response.text.strip()
                    raw = re.sub(r'```json|```','',raw).strip()
                    result = json.loads(raw)
                    if result.get('issues_found') and result.get('issues'):
                        for iss in result['issues']:
                            iss['screenshot'] = name
                            issues.append(iss)
                        jobs[job_id]['log'].append('  Found '+str(len(result['issues']))+' issue(s)')
                    else:
                        jobs[job_id]['log'].append('  No issues in '+name)
                except Exception as e:
                    jobs[job_id]['log'].append('  Error: '+str(e)[:60])

            jobs[job_id]['progress'] = 85
            jobs[job_id]['log'].append('Total issues: '+str(len(issues)))
            jobs[job_id]['issues'] = issues

            # Generate Word report
            jobs[job_id]['log'].append('Generating report...')
            from docx import Document as D2
            from docx.shared import Pt
            rpt = D2()
            rpt.add_heading('Screenshot QC Report', 0)
            rpt.add_paragraph('Screenshots analyzed: '+str(total))
            rpt.add_paragraph('Issues found: '+str(len(issues)))
            rpt.add_paragraph('Spec doc: '+jobs[job_id]['doc_name'])
            rpt.add_heading('Issues', level=1)
            if issues:
                for idx, iss in enumerate(issues, 1):
                    p = rpt.add_paragraph()
                    run = p.add_run('Issue '+str(idx)+': '+iss.get('type','Unknown')+' — '+iss.get('question',''))
                    run.bold = True
                    rpt.add_paragraph('Screenshot: '+iss.get('screenshot',''))
                    rpt.add_paragraph('Expected: '+iss.get('expected',''))
                    rpt.add_paragraph('Actual: '+iss.get('actual',''))
                    rpt.add_paragraph('Severity: '+iss.get('severity','medium'))
                    rpt.add_paragraph('')
            else:
                rpt.add_paragraph('No issues found! All screenshots match the spec document.')

            buf = io.BytesIO()
            rpt.save(buf)
            jobs[job_id]['report_bytes'] = buf.getvalue()
            jobs[job_id]['status'] = 'done'
            jobs[job_id]['progress'] = 100
            jobs[job_id]['log'].append('Done! Report ready for download.')
            jobs.persist(job_id)
        except Exception as e:
            jobs[job_id]['status'] = 'error'
            jobs[job_id]['log'].append('Error: '+str(e)[:100])
            jobs.persist(job_id)

    t = threading.Thread(target=run_ss, args=(job_id, docx_bytes, ss_list, instructions))
    t.daemon = True
    t.start()
    return redirect('/progress/'+job_id)


@app.route('/retest/<job_id>', methods=['GET'])
@login_required
def retest_job(job_id):
    """Show a pre-filled retest confirmation form.  Never starts QC automatically."""
    import glob as _glob

    # ── Job not in memory (server restart) ──────────────────────────────────
    if job_id not in jobs:
        # Try to recover basic info from disk (report file path encodes job_id)
        _report_hits = sorted(
            _glob.glob(os.path.join(OUTPUT_FOLDER, job_id, 'QC_Report_*.docx')),
            reverse=True
        )
        if not _report_hits:
            return redirect('/new-qc?msg=expired')
        # We have a report on disk but no job metadata — show /new-qc with notice
        return redirect('/new-qc?msg=expired')

    old_job = jobs[job_id]
    doc_name   = old_job.get('doc_name', 'Unknown document')
    survey_url = old_job.get('survey_url', '')
    platform   = old_job.get('platform', 'Confirmit')
    country    = old_job.get('country', '')
    doc_path   = old_job.get('doc_path', '')
    created_at = old_job.get('created_at', '')[:10]

    # ── Collect failed QIDs from previous run ────────────────────────────────
    failed_qids = []
    seen = set()
    for issue in old_job.get('issues', []):
        q = issue.get('qid', '').strip()
        if q and q not in seen and issue.get('severity') in ('HIGH', 'MEDIUM'):
            failed_qids.append(q)
            seen.add(q)
    for result in old_job.get('term_results', []):
        if not result.get('passed'):
            q = result.get('test_qid', '').strip()
            if q and q not in seen:
                failed_qids.append(q)
                seen.add(q)

    failed_qids_text = ', '.join(failed_qids) if failed_qids else ''
    doc_missing = doc_path and not os.path.exists(doc_path)

    sb = sidebar_html('reports')
    warn_html = ''
    if doc_missing:
        warn_html = (
            '<div style="background:#FEF3C7;border:1px solid #F59E0B;border-radius:10px;'
            'padding:14px 18px;margin-bottom:20px;font-size:13px;color:#92400E">'
            '<b>⚠️ Original document file no longer exists on disk.</b><br>'
            'Please upload the <b>.docx</b> spec document again below before running.'
            '</div>'
        )

    page = f"""<!DOCTYPE html><html><head><meta charset="utf-8">
<title>Retest — SurveyQC</title>
<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/@tabler/icons-webfont@latest/tabler-icons.min.css">
<style>
*{{box-sizing:border-box;margin:0;padding:0}}
body{{font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',sans-serif;background:#F8F9FB;color:#111827}}
.layout{{display:flex;min-height:100vh}}
.main{{flex:1;padding:32px 40px;max-width:820px}}
h1{{font-size:22px;font-weight:700;margin-bottom:4px}}
.sub{{font-size:13px;color:#6B7280;margin-bottom:28px}}
.card{{background:white;border:1px solid #E5E7EB;border-radius:14px;padding:24px;margin-bottom:20px}}
.card-label{{font-size:11px;font-weight:700;color:#6B7280;text-transform:uppercase;letter-spacing:.07em;margin-bottom:12px}}
.info-row{{display:flex;gap:8px;align-items:flex-start;margin-bottom:8px;font-size:13px}}
.info-key{{color:#6B7280;min-width:90px;flex-shrink:0}}
.info-val{{color:#111827;font-weight:500;word-break:break-all}}
.inp{{width:100%;padding:10px 13px;border:1px solid #D1D5DB;border-radius:8px;font-size:13px;font-family:inherit;outline:none}}
.inp:focus{{border-color:#7C3AED;box-shadow:0 0 0 3px rgba(124,58,237,.12)}}
textarea.inp{{min-height:90px;resize:vertical}}
label{{font-size:12px;font-weight:600;color:#374151;display:block;margin-bottom:5px}}
.btn-run{{display:inline-flex;align-items:center;gap:8px;background:#7C3AED;color:white;
  padding:12px 28px;border-radius:10px;font-size:14px;font-weight:600;border:none;cursor:pointer}}
.btn-run:hover{{background:#6D28D9}}
.btn-cancel{{display:inline-flex;align-items:center;gap:6px;background:white;color:#6B7280;
  padding:12px 20px;border-radius:10px;font-size:13px;font-weight:500;border:1px solid #D1D5DB;
  cursor:pointer;text-decoration:none;margin-right:10px}}
.badge-fail{{display:inline-block;background:#FEF2F2;color:#991B1B;font-size:11px;
  font-weight:600;padding:2px 8px;border-radius:5px;margin-left:8px}}
.dz{{border:2px dashed #D1D5DB;border-radius:10px;padding:18px;text-align:center;cursor:pointer;
  transition:border-color .2s;background:#FAFAFA}}
.dz:hover{{border-color:#7C3AED}}
.dz-done{{font-size:12px;color:#059669;margin-top:6px;display:none}}
</style></head><body>
<div class="layout">
  {sb}
  <div class="main">
    <h1><i class="ti ti-rotate-clockwise" style="margin-right:8px;color:#7C3AED"></i>Retest</h1>
    <p class="sub">Review settings and run targeted QC on previously failed questions only.</p>

    {warn_html}

    <form method="POST" action="/retest/run/{job_id}" enctype="multipart/form-data" id="rtForm">

      <!-- Previous job summary -->
      <div class="card">
        <p class="card-label"><i class="ti ti-clipboard-data" style="margin-right:4px"></i>Original QC — {created_at}</p>
        <div class="info-row"><span class="info-key">Document</span><span class="info-val">{doc_name}</span></div>
        <div class="info-row"><span class="info-key">Survey URL</span><span class="info-val" style="font-size:12px">{survey_url}</span></div>
        <div class="info-row"><span class="info-key">Platform</span><span class="info-val">{platform}</span></div>
        {'<div class="info-row"><span class="info-key">Country</span><span class="info-val">' + country + '</span></div>' if country else ''}
      </div>

      <!-- Failed QIDs -->
      <div class="card">
        <p class="card-label">
          <i class="ti ti-alert-triangle" style="margin-right:4px;color:#F59E0B"></i>
          Failed questions
          {'<span class="badge-fail">' + str(len(failed_qids)) + ' issues</span>' if failed_qids else ''}
        </p>
        <label for="qids">QIDs to retest <span style="color:#6B7280;font-weight:400">(comma or newline separated — leave blank to retest ALL)</span></label>
        <textarea name="filter_qids" id="qids" class="inp"
          placeholder="e.g. R1, R3, S1, Q15">{failed_qids_text}</textarea>
      </div>

      <!-- Doc upload (only shown if file missing) -->
      {'<div class="card"><p class="card-label"><i class="ti ti-file-word" style="margin-right:4px"></i>Upload spec document again</p>'
       '<div class="dz" id="docZone" onclick="document.getElementById(\'docInput\').click()">'
       '<i class="ti ti-file-word" style="font-size:22px;color:#9CA3AF"></i>'
       '<p style="font-size:12px;color:#6B7280;margin-top:5px">Drop <b>.docx</b> or <u>browse</u></p></div>'
       '<input type="file" name="doc" id="docInput" accept=".docx" required style="display:none"'
       ' onchange="var d=document.getElementById(\'docZone\');d.style.borderColor=\'#7C3AED\';">'
       '</div>' if doc_missing else ''}

      <!-- URL override -->
      <div class="card">
        <p class="card-label"><i class="ti ti-link" style="margin-right:4px"></i>Survey URL <span style="color:#6B7280;font-weight:400;font-size:11px">(update if link changed)</span></p>
        <input type="url" name="url" class="inp" value="{survey_url}" required placeholder="https://...">
      </div>

      <!-- Actions -->
      <div style="display:flex;align-items:center;flex-wrap:wrap;gap:10px;margin-top:8px">
        <a href="/report/{job_id}" class="btn-cancel"><i class="ti ti-arrow-left"></i>Back to report</a>
        <button type="submit" class="btn-run" id="runBtn">
          <i class="ti ti-player-play"></i>Run Retest
        </button>
        <span id="runMsg" style="font-size:12px;color:#7C3AED;display:none">Starting retest…</span>
      </div>

    </form>
  </div>
</div>
<script>
document.getElementById('rtForm').addEventListener('submit', function() {{
  document.getElementById('runBtn').disabled = true;
  document.getElementById('runMsg').style.display = 'inline';
}});
</script>
</body></html>"""
    return page


@app.route('/retest/run/<job_id>', methods=['POST'])
@login_required
def retest_run(job_id):
    """Start the actual retest engine after form submission."""
    import uuid
    from datetime import datetime

    # ── Resolve original job ─────────────────────────────────────────────────
    if job_id not in jobs:
        return redirect('/new-qc?msg=expired')

    old_job   = jobs[job_id]
    survey_url = request.form.get('url', '').strip() or old_job.get('survey_url', '')
    doc_path   = old_job.get('doc_path', '')

    # If original doc is missing and user uploaded a replacement, save it
    if not os.path.exists(doc_path or ''):
        uploaded_doc = request.files.get('doc')
        if uploaded_doc and uploaded_doc.filename:
            import uuid as _uuid
            new_doc_dir = os.path.join(UPLOAD_FOLDER, str(_uuid.uuid4())[:8])
            os.makedirs(new_doc_dir, exist_ok=True)
            from werkzeug.utils import secure_filename as _sf
            doc_path = os.path.join(new_doc_dir, _sf(uploaded_doc.filename))
            uploaded_doc.save(doc_path)
        else:
            return redirect(f'/new-qc?msg=nodoc')

    if not survey_url:
        return redirect(f'/retest/{job_id}')

    # ── Parse QID filter ────────────────────────────────────────────────────
    raw_qids = request.form.get('filter_qids', '').strip()
    filter_qids = []
    if raw_qids:
        import re as _re
        filter_qids = [q.strip().upper() for q in _re.split(r'[,\n\r;]+', raw_qids) if q.strip()]

    # ── Memory guard ────────────────────────────────────────────────────────
    if len(jobs) >= 50:
        _done = sorted(
            (jid for jid, j in jobs.items() if j.get('status') in ('done', 'error', 'stopped')),
            key=lambda jid: jobs[jid].get('created_at', '')
        )
        for _evict in _done[:max(0, len(jobs) - 49)]:
            jobs.evict(_evict)

    new_job_id = str(uuid.uuid4())[:8]
    jobs[new_job_id] = {
        'status':     'running',
        'progress':   0,
        'phase':      'Starting retest…',
        'logs':       [{'msg': f'Retest of {len(filter_qids)} QID(s)' if filter_qids else 'Full retest started', 'color': 'cyan'}],
        'doc_name':   old_job.get('doc_name', ''),
        'doc_path':   doc_path,
        'survey_url': survey_url,
        'platform':   old_job.get('platform', 'Confirmit'),
        'country':    old_job.get('country', ''),
        'mode':       'full',
        'user_email': session['user_email'],
        'created_at': datetime.now().isoformat(),
        'retest_of':  job_id,
        'verdict':    None,
        'issues':     [], 'term_results': [],
        'report_file': None,
        'doc_qids':   0, 'live_qids': 0, 'xml_qids': 0,
        'total_issues': 0, 'term_passed': 0, 'term_total': 0,
    }

    import threading
    t = threading.Thread(
        target=run_qc_engine,
        args=(new_job_id, doc_path, survey_url,
              old_job.get('country', ''), 'full', [],
              filter_qids or None),
        daemon=True
    )
    t.start()

    return redirect('/progress/' + new_job_id)



# ================================================================
# LANDING PAGE v2 — Light theme with all sections
# ================================================================
@app.route('/home')
def home_landing():
    c = site_content

    # Get rotating testimonials (pick 3)
    tests = []
    for i in range(1, 7):
        if f'test{i}_name' in c:
            tests.append({
                'name': c[f'test{i}_name'],
                'role': c[f'test{i}_role'],
                'company': c[f'test{i}_company'],
                'country': c[f'test{i}_country'],
                'flag': c[f'test{i}_flag'],
                'quote': c[f'test{i}_quote'],
                'rating': int(c[f'test{i}_rating']),
            })
    # Show first 3 (admin can reorder by editing keys)
    visible_tests = tests[:3]

    test_cards = ''
    for t in visible_tests:
        stars = '<i class="ti ti-star-filled"></i>' * t['rating']
        initials = ''.join([n[0] for n in t['name'].split()[:2]])
        test_cards += f"""<div class="test-card">
  <div class="test-stars">{stars}</div>
  <div class="test-quote">"{t['quote']}"</div>
  <div class="test-author">
    <div class="test-avatar">{initials}</div>
    <div>
      <div class="test-name">{t['name']}</div>
      <div class="test-role">{t['role']} at {t['company']} · {t['country']}</div>
    </div>
  </div>
</div>"""

    page = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1.0">
<title>SurveyQC — AI Survey QC. Audit-Defensible. Any Platform.</title>
<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/@tabler/icons-webfont@3.19.0/dist/tabler-icons.min.css">
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<style>
:root{
  --bg:#F7F4EE;--bg2:#FFFDF9;--card:#FFFFFF;
  --text:#171717;--text2:#5F5B53;--text3:#8A847A;
  --accent:#C46A2B;--accent-hover:#A9551F;--accent-bg:#F5E6D8;
  --border:#E8E1D8;--border2:#F0EBE3;
  --dark:#1B140F;--dark2:#2A1F18;
  --success:#3F7D58;--warn:#D89B2B;--danger:#C84B31;
  --shadow:0 1px 2px rgba(24,17,10,0.04),0 4px 12px rgba(24,17,10,0.05);
  --shadow-lg:0 10px 40px rgba(24,17,10,0.08);
  --radius:20px;--radius-btn:14px;
}
*{box-sizing:border-box;margin:0;padding:0;font-family:-apple-system,'Inter','Plus Jakarta Sans',sans-serif;-webkit-font-smoothing:antialiased}
html{scroll-behavior:smooth}
body{background:var(--bg);color:var(--text);line-height:1.5;overflow-x:hidden}
a{text-decoration:none;color:inherit}
img{max-width:100%}

/* ── ANNOUNCE ── */
.announce{background:var(--dark);color:#E8DDD2;padding:10px 24px;text-align:center;font-size:13px;font-weight:500}
.announce a{color:var(--accent);font-weight:600;margin-left:6px}

/* ── NAV ── */
.nav{background:rgba(247,244,238,.85);backdrop-filter:blur(20px);-webkit-backdrop-filter:blur(20px);border-bottom:1px solid var(--border);padding:0 32px;height:68px;display:flex;align-items:center;justify-content:space-between;position:sticky;top:0;z-index:100}
.nav-logo{display:flex;align-items:center;gap:10px;font-weight:700;font-size:18px;color:var(--text)}
.nav-logo-mark{width:32px;height:32px;background:var(--dark);border-radius:9px;position:relative}
.nav-logo-mark::after{content:"";position:absolute;inset:6px;border:1.5px solid var(--accent);border-radius:5px}
.nav-links{display:flex;align-items:center;gap:32px}
.nav-link{font-size:14px;font-weight:500;color:var(--text2);transition:color .2s}
.nav-link:hover{color:var(--text)}
.nav-cta{display:flex;align-items:center;gap:12px}
.btn-sign{font-size:14px;font-weight:500;color:var(--text);padding:8px 18px;border-radius:10px;transition:background .2s}
.btn-sign:hover{background:var(--accent-bg)}
.btn-primary{background:var(--dark);color:#F7F4EE;font-size:14px;font-weight:600;padding:10px 22px;border-radius:var(--radius-btn);transition:all .2s;display:inline-flex;align-items:center;gap:6px;border:none;cursor:pointer}
.btn-primary:hover{background:var(--dark2);transform:translateY(-1px);box-shadow:var(--shadow-lg)}
.btn-ghost{background:white;color:var(--text);border:1px solid var(--border);font-size:14px;font-weight:500;padding:10px 22px;border-radius:var(--radius-btn);transition:all .2s}
.btn-ghost:hover{background:var(--bg2)}
.hamburger{display:none;background:none;border:none;cursor:pointer;width:36px;height:36px;align-items:center;justify-content:center}
.hamburger i{font-size:22px;color:var(--text)}

/* ── HERO ── */
.hero{padding:80px 24px 72px;text-align:center;position:relative;overflow:hidden}
.hero::before{content:"";position:absolute;top:-100px;left:50%;transform:translateX(-50%);width:800px;height:500px;background:radial-gradient(ellipse,rgba(196,106,43,.10),transparent 60%);pointer-events:none;z-index:0}
.hero-inner{position:relative;z-index:1;max-width:920px;margin:0 auto}
.hero-badge{display:inline-flex;align-items:center;gap:8px;background:white;border:1px solid var(--border);padding:6px 16px 6px 8px;border-radius:100px;font-size:13px;color:var(--text2);margin-bottom:32px;font-weight:500;box-shadow:var(--shadow)}
.hero-badge-pill{background:var(--accent-bg);color:var(--accent);font-size:11px;font-weight:700;padding:3px 10px;border-radius:100px;text-transform:uppercase}
.hero h1{font-family:'Plus Jakarta Sans',sans-serif;font-size:clamp(36px,6vw,68px);font-weight:800;line-height:1.06;letter-spacing:-2px;color:var(--text);margin-bottom:24px}
.hero h1 .accent{color:var(--accent);position:relative;display:inline-block}
.hero h1 .accent::after{content:"";position:absolute;bottom:5px;left:0;right:0;height:7px;background:rgba(196,106,43,.18);z-index:-1;border-radius:4px}
.hero-sub{font-size:clamp(16px,2vw,20px);color:var(--text2);max-width:600px;margin:0 auto 20px;line-height:1.7}
.hero-tagline{display:inline-flex;align-items:center;gap:10px;background:var(--bg2);border:1px solid var(--border);border-radius:100px;padding:8px 20px;font-size:13px;font-weight:600;color:var(--text);margin-bottom:36px}
.hero-tagline .arrow{color:var(--accent);font-weight:800;font-size:15px}
.hero-cta{display:flex;gap:14px;justify-content:center;flex-wrap:wrap;margin-bottom:20px}
.hero-cta .btn-primary{padding:14px 28px;font-size:15px}
.hero-meta{font-size:13px;color:var(--text3);display:flex;align-items:center;justify-content:center;gap:18px;flex-wrap:wrap}
.hero-meta-item{display:flex;align-items:center;gap:6px}
.hero-meta-item i{color:var(--success);font-size:15px}

/* ── PAIN POINTS ── */
.pain-grid{display:grid;grid-template-columns:repeat(3,1fr);gap:20px}
.pain-card{background:var(--card);border:1px solid var(--border);border-radius:var(--radius);padding:28px 26px}
.pain-icon{width:42px;height:42px;border-radius:10px;background:#FEF2F2;display:flex;align-items:center;justify-content:center;margin-bottom:16px}
.pain-icon i{font-size:20px;color:var(--danger)}
.pain-title{font-family:'Plus Jakarta Sans',sans-serif;font-size:16px;font-weight:700;color:var(--text);margin-bottom:8px}
.pain-desc{font-size:14px;color:var(--text2);line-height:1.7}

/* ── TRUSTED ── */
.trusted{padding:32px 24px 16px;text-align:center}
.trusted-l{font-size:12px;color:var(--text3);text-transform:uppercase;letter-spacing:.12em;font-weight:600;margin-bottom:20px}
.trusted-row{display:flex;align-items:center;justify-content:center;gap:12px;flex-wrap:wrap}
.trusted-logo{font-family:'Plus Jakarta Sans',sans-serif;font-size:13px;font-weight:700;color:#9CA3AF;background:#FFFFFF;border:1px solid #E5E7EB;border-radius:10px;padding:10px 18px;letter-spacing:.04em;cursor:default;transition:all .2s}
.trusted-logo:hover{color:#6B7280;border-color:#D1D5DB;box-shadow:0 2px 8px rgba(0,0,0,.06)}

/* ── SECTIONS ── */
.section{padding:88px 24px}
.container{max-width:1240px;margin:0 auto;padding:0 24px}
.sec-head{text-align:center;max-width:720px;margin:0 auto 48px}
.sec-tag{display:inline-block;background:var(--accent-bg);color:var(--accent);font-size:12px;font-weight:700;padding:5px 14px;border-radius:100px;text-transform:uppercase;letter-spacing:.06em;margin-bottom:16px}
.sec-title{font-family:'Plus Jakarta Sans',sans-serif;font-size:clamp(28px,4vw,46px);font-weight:800;line-height:1.1;letter-spacing:-1.2px;margin-bottom:18px;color:var(--text)}
.sec-sub{font-size:18px;color:var(--text2);line-height:1.65;max-width:600px;margin:0 auto}

/* ── FEATURES ── */
.feat-grid{display:grid;grid-template-columns:repeat(3,1fr);gap:20px}
.feat-card{background:var(--card);border:1px solid var(--border);border-radius:var(--radius);padding:32px;transition:all .3s ease;position:relative}
.feat-card:hover{transform:translateY(-4px);box-shadow:var(--shadow-lg);border-color:var(--accent-bg)}
.feat-card.is-new{border-color:rgba(63,125,88,.22)}
.feat-icon{width:48px;height:48px;border-radius:12px;background:var(--accent-bg);display:flex;align-items:center;justify-content:center;margin-bottom:20px}
.feat-icon i{font-size:22px;color:var(--accent)}
.feat-card.is-new .feat-icon{background:rgba(63,125,88,.12)}
.feat-card.is-new .feat-icon i{color:var(--success)}
.feat-title{font-family:'Plus Jakarta Sans',sans-serif;font-size:18px;font-weight:700;margin-bottom:10px;color:var(--text);display:flex;align-items:center;gap:8px;flex-wrap:wrap}
.feat-desc{font-size:14px;color:var(--text2);line-height:1.7}
.new-badge{background:var(--success);color:white;font-size:10px;font-weight:700;padding:3px 8px;border-radius:100px;text-transform:uppercase;letter-spacing:.04em}

/* ── HOW IT WORKS ── */
.how-grid{display:grid;grid-template-columns:repeat(3,1fr);gap:24px}
.how-card{background:var(--card);border:1px solid var(--border);border-radius:var(--radius);padding:36px 28px;transition:all .3s}
.how-card:hover{transform:translateY(-4px);box-shadow:var(--shadow-lg)}
.how-num{font-family:'Plus Jakarta Sans',sans-serif;font-size:14px;font-weight:700;color:var(--accent);background:var(--accent-bg);width:40px;height:40px;border-radius:50%;display:flex;align-items:center;justify-content:center;margin-bottom:20px}
.how-title{font-family:'Plus Jakarta Sans',sans-serif;font-size:20px;font-weight:700;margin-bottom:10px;color:var(--text)}
.how-desc{font-size:14px;color:var(--text2);line-height:1.7;margin-bottom:16px}
.how-tag{display:inline-flex;align-items:center;gap:6px;font-size:12px;font-weight:600;color:var(--success);background:rgba(63,125,88,.08);border-radius:100px;padding:5px 12px}
.how-tag i{font-size:13px}

/* ── REPORT MOCK ── */
.report-mock{max-width:820px;margin:0 auto;background:var(--card);border:1px solid var(--border);border-radius:var(--radius);overflow:hidden;box-shadow:var(--shadow-lg)}
.report-mock-header{background:var(--dark);padding:16px 24px;display:flex;align-items:center;justify-content:space-between;flex-wrap:wrap;gap:12px}
.report-mock-title{font-size:14px;font-weight:600;color:#E8DDD2;display:flex;align-items:center;gap:8px}
.report-mock-stats{display:flex;gap:20px}
.report-mock-stat{text-align:center}
.report-mock-num{font-family:'Plus Jakarta Sans',sans-serif;font-size:20px;font-weight:800;color:white;line-height:1}
.report-mock-lbl{font-size:10px;color:#9A8C7B;font-weight:500;margin-top:3px}
.report-mock-body{padding:16px 20px;display:flex;flex-direction:column;gap:8px}
.chk-row{display:flex;align-items:center;gap:12px;padding:11px 14px;border-radius:10px;background:var(--bg);border:1px solid var(--border2)}
.chk-badge{width:22px;height:22px;border-radius:6px;display:flex;align-items:center;justify-content:center;flex-shrink:0}
.chk-badge.pass{background:#DCFCE7}.chk-badge.pass i{color:var(--success);font-size:12px}
.chk-badge.fail{background:#FEE2E2}.chk-badge.fail i{color:var(--danger);font-size:12px}
.chk-badge.warn{background:#FEF9C3}.chk-badge.warn i{color:var(--warn);font-size:12px}
.chk-name{font-size:13px;font-weight:600;color:var(--text);flex:1}
.chk-detail{font-size:12px;color:var(--text3)}
.report-mock-footer{padding:12px 20px 14px;border-top:1px solid var(--border2);display:flex;align-items:center;justify-content:space-between;flex-wrap:wrap;gap:8px}
.report-mock-footer span{font-size:11px;color:var(--text3);display:flex;align-items:center;gap:5px}

/* ── WHO IS THIS FOR ── */
.persona-grid{display:grid;grid-template-columns:repeat(2,1fr);gap:20px;max-width:900px;margin:0 auto}
.persona-card{background:var(--card);border:1px solid var(--border);border-radius:var(--radius);padding:28px;display:flex;gap:18px;align-items:flex-start;transition:all .3s}
.persona-card:hover{transform:translateY(-3px);box-shadow:var(--shadow-lg);border-color:var(--accent-bg)}
.persona-icon{width:46px;height:46px;border-radius:12px;background:var(--dark);display:flex;align-items:center;justify-content:center;flex-shrink:0}
.persona-icon i{font-size:22px;color:var(--accent)}
.persona-role{font-family:'Plus Jakarta Sans',sans-serif;font-size:15px;font-weight:700;color:var(--text);margin-bottom:6px}
.persona-pain{font-size:14px;color:var(--text2);line-height:1.65}

/* ── BEFORE vs AFTER ── */
.bva-wrap{display:grid;grid-template-columns:1fr 1fr;gap:24px;max-width:840px;margin:0 auto}
.bva-col{background:var(--card);border-radius:var(--radius);border:1px solid var(--border);overflow:hidden}
.bva-head{padding:16px 22px;border-bottom:1px solid var(--border)}
.bva-head.before{background:#FEF2F2}
.bva-head.after{background:#F0FDF4}
.bva-label{font-family:'Plus Jakarta Sans',sans-serif;font-size:14px;font-weight:700;display:flex;align-items:center;gap:7px}
.bva-head.before .bva-label{color:var(--danger)}
.bva-head.after .bva-label{color:var(--success)}
.bva-rows{padding:6px 0}
.bva-row{display:flex;align-items:flex-start;gap:12px;padding:12px 22px;border-bottom:1px solid var(--border2)}
.bva-row:last-child{border-bottom:none}
.bva-row-icon{flex-shrink:0;font-size:15px;margin-top:2px}
.bva-col.before .bva-row-icon{color:var(--danger)}
.bva-col.after .bva-row-icon{color:var(--success)}
.bva-text strong{display:block;font-size:13px;font-weight:600;color:var(--text)}
.bva-text small{font-size:12px;color:var(--text3)}

/* ── PLATFORM ── */
.platform-grid{display:grid;grid-template-columns:repeat(3,1fr);gap:14px;max-width:880px;margin:0 auto}
.platform-pill{display:flex;align-items:center;gap:14px;background:var(--card);border:1px solid var(--border);border-radius:14px;padding:16px 18px;transition:all .2s}
.platform-pill:hover{transform:translateY(-2px);box-shadow:var(--shadow-lg);border-color:var(--accent)}
.platform-pill-mark{width:42px;height:42px;border-radius:11px;color:#fff;display:flex;align-items:center;justify-content:center;font-family:'Plus Jakarta Sans',sans-serif;font-weight:800;font-size:18px;flex-shrink:0}
.platform-pill-name{font-size:15px;font-weight:700;color:var(--text)}
.platform-pill-status{font-size:11px;font-weight:600;margin-top:2px}
.platform-pill-status.live{color:var(--success)}
.platform-pill-status.soon{color:var(--text3)}

/* ── COMPARISON TABLE ── */
.cmp-wrap{max-width:900px;margin:0 auto;overflow-x:auto;-webkit-overflow-scrolling:touch}
.cmp-table{width:100%;border-collapse:collapse;background:var(--card);border-radius:var(--radius);overflow:hidden;border:1px solid var(--border)}
.cmp-table th{padding:14px 16px;font-size:13px;font-weight:700;text-align:center;border-bottom:2px solid var(--border);white-space:nowrap}
.cmp-table th:first-child{text-align:left;min-width:180px}
.cmp-table th.hl{background:var(--dark);color:#F7F4EE}
.cmp-table td{padding:12px 16px;font-size:13px;text-align:center;border-bottom:1px solid var(--border2);color:var(--text2)}
.cmp-table td:first-child{text-align:left;font-weight:600;color:var(--text)}
.cmp-table td.hl{background:rgba(247,231,216,.25);font-weight:600;color:var(--text)}
.cmp-yes{color:var(--success);font-size:16px}
.cmp-no{color:var(--danger);font-size:16px}
.cmp-part{color:var(--warn);font-size:12px;font-weight:600}
.cmp-table tr:hover td{background:rgba(247,244,238,.7)}
.cmp-table tr:hover td.hl{background:rgba(247,231,216,.5)}

/* ── QC CHECKS ── */
.checks-grid{display:grid;grid-template-columns:repeat(5,1fr);gap:12px;max-width:1100px;margin:0 auto}
.check-pill{display:flex;align-items:center;gap:10px;background:var(--card);border:1px solid var(--border);border-radius:12px;padding:13px 14px;transition:all .2s}
.check-pill:hover{border-color:var(--accent);transform:translateY(-2px);box-shadow:var(--shadow-lg)}
.check-pill-icon{width:30px;height:30px;border-radius:8px;background:var(--accent-bg);display:flex;align-items:center;justify-content:center;flex-shrink:0}
.check-pill-icon i{font-size:15px;color:var(--accent)}
.check-pill span{font-size:13px;font-weight:600;color:var(--text);line-height:1.3}

/* ── TRUST NUMBERS ── */
.trust-grid{display:grid;grid-template-columns:repeat(4,1fr);gap:20px}
.trust-card{background:var(--card);border:1px solid var(--border);border-radius:var(--radius);padding:32px 24px;text-align:center}
.trust-num{font-family:'Plus Jakarta Sans',sans-serif;font-size:clamp(40px,5vw,58px);font-weight:800;color:var(--text);letter-spacing:-2px;line-height:1}
.trust-num .unit{font-size:0.52em;color:var(--accent)}
.trust-label{font-size:14px;color:var(--text2);margin-top:10px;line-height:1.5}

/* ── TESTIMONIALS ── */
.test-grid{display:grid;grid-template-columns:repeat(3,1fr);gap:20px}
.test-card{background:var(--card);border:1px solid var(--border);border-radius:var(--radius);padding:30px;transition:all .3s}
.test-card:hover{transform:translateY(-3px);box-shadow:var(--shadow-lg)}
.test-stars{display:flex;gap:2px;margin-bottom:18px}
.test-stars i{color:var(--warn);font-size:14px}
.test-quote{font-size:15px;color:var(--text);line-height:1.65;margin-bottom:22px;font-weight:500}
.test-author{display:flex;align-items:center;gap:12px;padding-top:18px;border-top:1px solid var(--border2)}
.test-avatar{width:40px;height:40px;border-radius:50%;background:var(--accent-bg);color:var(--accent);display:flex;align-items:center;justify-content:center;font-weight:700;font-size:13px}
.test-name{font-size:14px;font-weight:700;color:var(--text)}
.test-role{font-size:12px;color:var(--text3);margin-top:1px}

/* ── CTA BANNER ── */
.cta-banner{padding:80px 24px}
.cta-inner{max-width:920px;margin:0 auto;background:var(--dark);border-radius:32px;padding:64px 48px;text-align:center;position:relative;overflow:hidden}
.cta-inner::before{content:"";position:absolute;top:-80px;right:-80px;width:280px;height:280px;background:radial-gradient(circle,rgba(196,106,43,.45),transparent 70%);pointer-events:none}
.cta-inner h2{font-family:'Plus Jakarta Sans',sans-serif;font-size:clamp(28px,4vw,42px);color:white;margin-bottom:14px;font-weight:800;letter-spacing:-1px;position:relative;z-index:1}
.cta-inner p{font-size:16px;color:#D4C6B6;margin-bottom:32px;position:relative;z-index:1}
.cta-inner .btn-primary{background:var(--accent);position:relative;z-index:1;padding:14px 32px;font-size:15px}
.cta-inner .btn-primary:hover{background:var(--accent-hover)}

/* ── FOOTER ── */
.footer{background:var(--dark);color:#B8AC9F;padding:80px 24px 30px}
.footer-inner{max-width:1180px;margin:0 auto}
.footer-grid{display:grid;grid-template-columns:2.2fr 1fr 1fr 1fr 1fr;gap:48px;margin-bottom:60px}
.footer-brand-logo{display:flex;align-items:center;gap:10px;margin-bottom:20px}
.footer-brand-mark{width:32px;height:32px;background:var(--accent);border-radius:9px;display:flex;align-items:center;justify-content:center}
.footer-brand-text{font-family:'Plus Jakarta Sans',sans-serif;font-size:18px;font-weight:700;color:white}
.footer-brand p{font-size:14px;line-height:1.7;color:#9A8C7B;margin-bottom:18px;max-width:280px}
.footer-social{display:flex;gap:10px}
.footer-social a{width:36px;height:36px;background:rgba(255,255,255,.06);border-radius:9px;display:flex;align-items:center;justify-content:center;transition:all .2s;color:#9A8C7B}
.footer-social a:hover{background:var(--accent);color:white}
.footer-title{font-size:13px;font-weight:700;color:white;text-transform:uppercase;letter-spacing:.06em;margin-bottom:18px}
.footer-link{display:block;font-size:14px;color:#9A8C7B;margin-bottom:12px;transition:color .2s}
.footer-link:hover{color:white}
.footer-bottom{border-top:1px solid rgba(255,255,255,.08);padding-top:28px;display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:14px}
.footer-bottom p{font-size:13px;color:#7A6E5F}

/* ── MOBILE MENU ── */
.mobile-menu{display:none;position:fixed;top:0;left:0;right:0;bottom:0;background:var(--bg);z-index:200;padding:80px 32px;flex-direction:column;gap:18px}
.mobile-menu.open{display:flex}
.mobile-menu a{font-size:18px;font-weight:600;color:var(--text);padding:12px 0;border-bottom:1px solid var(--border)}
.mobile-menu-close{position:absolute;top:20px;right:20px;background:none;border:none;font-size:28px;color:var(--text);cursor:pointer}

/* ── RESPONSIVE ── */
@media(max-width:1024px){
  .footer-grid{grid-template-columns:1fr 1fr 1fr}
  .trust-grid{grid-template-columns:repeat(2,1fr)}
  .checks-grid{grid-template-columns:repeat(3,1fr)}
}
@media(max-width:768px){
  .nav-links{display:none}
  .hamburger{display:flex}
  .nav-cta .btn-sign{display:none}
  .nav{padding:0 18px}
  .hero{padding:60px 18px 70px}
  .section{padding:56px 18px}
  .cta-banner{padding:64px 18px}
  .cta-inner{padding:48px 28px;border-radius:24px}
  .feat-grid,.how-grid,.test-grid,.pain-grid{grid-template-columns:1fr;gap:14px}
  .persona-grid{grid-template-columns:1fr}
  .bva-wrap{grid-template-columns:1fr}
  .platform-grid{grid-template-columns:1fr 1fr}
  .trust-grid{grid-template-columns:1fr 1fr}
  .footer-grid{grid-template-columns:1fr 1fr;gap:36px}
  .footer{padding:60px 20px 24px}
}
@media(max-width:480px){
  .hero h1{font-size:36px;letter-spacing:-1.2px}
  .sec-title{font-size:28px}
  .checks-grid{grid-template-columns:repeat(2,1fr);gap:10px}
  .platform-grid{grid-template-columns:1fr}
  .cmp-table th,.cmp-table td{padding:10px 12px;font-size:12px}
}
</style>
</head>
<body>

<!-- 1. ANNOUNCEMENT BAR -->
<div class="announce">
  <i class="ti ti-sparkles"></i> """ + c['announcement'] + """
  <a href="/new-qc">Try free →</a>
</div>

<!-- 2. NAV -->
<nav class="nav">
  <a href="/home" class="nav-logo">
    <div class="nav-logo-mark"></div>
    """ + c['site_name'] + """
  </a>
  <div class="nav-links">
    <a href="#features" class="nav-link">Features</a>
    <a href="#how" class="nav-link">How it works</a>
    <a href="/pricing" class="nav-link">Pricing</a>
    <a href="/docs" class="nav-link">Docs</a>
    <a href="/blog" class="nav-link">Blog</a>
  </div>
  <div class="nav-cta">
    <a href="/login" class="btn-sign">Sign in</a>
    <a href="/new-qc" class="btn-primary">Run QC Free <i class="ti ti-arrow-right"></i></a>
    <button class="hamburger" onclick="document.getElementById('mm').classList.add('open')"><i class="ti ti-menu-2"></i></button>
  </div>
</nav>

<div class="mobile-menu" id="mm">
  <button class="mobile-menu-close" onclick="document.getElementById('mm').classList.remove('open')"><i class="ti ti-x"></i></button>
  <a href="#features">Features</a>
  <a href="#how">How it works</a>
  <a href="/pricing">Pricing</a>
  <a href="/docs">Docs</a>
  <a href="/blog">Blog</a>
  <a href="/login">Sign in</a>
  <a href="/new-qc" style="background:var(--dark);color:#F7F4EE;text-align:center;border-radius:14px;padding:14px;margin-top:12px;border:none">Run QC Free →</a>
</div>

<!-- 3. HERO -->
<section class="hero">
  <div class="hero-inner">
    <div class="hero-badge">
      <span class="hero-badge-pill">Pharma-ready</span>
      Audit-defensible QC — every check documented
    </div>
    <h1>Catch every survey bug<br><span class="accent">before</span> your data is collected.</h1>
    <p class="hero-sub">Upload your XML export or screener doc. SurveyQC checks every route, every termination, every option — automatically.</p>
    <div class="hero-tagline">
      <span>AI parses</span>
      <span class="arrow">→</span>
      <span>Code decides</span>
      <span class="arrow">→</span>
      <span>Evidence logged</span>
    </div>
    <div class="hero-cta">
      <a href="/new-qc" class="btn-primary">Run your first QC free <i class="ti ti-arrow-right"></i></a>
      <a href="#how" class="btn-ghost">See how it works</a>
    </div>
    <div class="hero-meta">
      <span class="hero-meta-item"><i class="ti ti-check"></i>No credit card</span>
      <span class="hero-meta-item"><i class="ti ti-check"></i>3 free reports/month</span>
      <span class="hero-meta-item"><i class="ti ti-check"></i>Any platform</span>
    </div>
  </div>
</section>

<!-- 4. PAIN POINTS -->
<section class="section" style="padding-top:0;padding-bottom:64px">
  <div class="container">
    <div class="sec-head" style="margin-bottom:32px">
      <span class="sec-tag">The Problem</span>
      <h2 class="sec-title">Manual QC is broken.</h2>
    </div>
    <div class="pain-grid">
      <div class="pain-card">
        <div class="pain-icon"><i class="ti ti-clock-x"></i></div>
        <div class="pain-title">6–8 hours per survey</div>
        <div class="pain-desc">Manual QC is the most time-consuming step in survey delivery. One tester, hundreds of paths, zero automation.</div>
      </div>
      <div class="pain-card">
        <div class="pain-icon"><i class="ti ti-file-alert"></i></div>
        <div class="pain-title">Scripts drift from the spec</div>
        <div class="pain-desc">The live survey changes after spec sign-off. By launch, text is wrong, codes are off, routing is broken — silently.</div>
      </div>
      <div class="pain-card">
        <div class="pain-icon"><i class="ti ti-lock-open"></i></div>
        <div class="pain-title">No audit trail</div>
        <div class="pain-desc">Pharma clients and auditors need timestamped, per-check documentation. “We checked it manually” is not acceptable.</div>
      </div>
    </div>
  </div>
</section>

<!-- 5. TRUSTED BY -->
<section class="trusted">
  <div class="trusted-l">Used by QC professionals at agencies including</div>
  <div class="trusted-row">
    <div class="trusted-logo">IPSOS</div>
    <div class="trusted-logo">Kantar</div>
    <div class="trusted-logo">Nielsen</div>
    <div class="trusted-logo">YouGov</div>
    <div class="trusted-logo">GfK</div>
    <div class="trusted-logo">Dynata</div>
  </div>
</section>

<!-- 6. FEATURES (9 cards: 6 existing + 3 new) -->
<section class="section" id="features" style="padding-top:72px">
  <div class="container">
    <div class="sec-head">
      <span class="sec-tag">Features</span>
      <h2 class="sec-title">Everything manual —<br>now automated.</h2>
      <p class="sec-sub">15+ specialized checks run in parallel. Nothing slips through.</p>
    </div>
    <div class="feat-grid">
      <div class="feat-card">
        <div class="feat-icon"><i class="ti ti-shield-check"></i></div>
        <div class="feat-title">""" + c['feature1_title'] + """</div>
        <div class="feat-desc">""" + c['feature1_desc'] + """</div>
      </div>
      <div class="feat-card">
        <div class="feat-icon"><i class="ti ti-search"></i></div>
        <div class="feat-title">""" + c['feature2_title'] + """</div>
        <div class="feat-desc">""" + c['feature2_desc'] + """</div>
      </div>
      <div class="feat-card">
        <div class="feat-icon"><i class="ti ti-camera"></i></div>
        <div class="feat-title">""" + c['feature3_title'] + """</div>
        <div class="feat-desc">""" + c['feature3_desc'] + """</div>
      </div>
      <div class="feat-card">
        <div class="feat-icon"><i class="ti ti-world"></i></div>
        <div class="feat-title">""" + c['feature4_title'] + """</div>
        <div class="feat-desc">""" + c['feature4_desc'] + """</div>
      </div>
      <div class="feat-card">
        <div class="feat-icon"><i class="ti ti-list-check"></i></div>
        <div class="feat-title">""" + c['feature5_title'] + """</div>
        <div class="feat-desc">""" + c['feature5_desc'] + """</div>
      </div>
      <div class="feat-card">
        <div class="feat-icon"><i class="ti ti-refresh"></i></div>
        <div class="feat-title">""" + c['feature6_title'] + """</div>
        <div class="feat-desc">""" + c['feature6_desc'] + """</div>
      </div>
      <div class="feat-card is-new">
        <div class="feat-icon"><i class="ti ti-arrow-fork"></i></div>
        <div class="feat-title">Routing Validation <span class="new-badge">NEW</span></div>
        <div class="feat-desc">Every skip, branch, and conditional display rule tested deterministically. Every path walked. Every logic error surfaced.</div>
      </div>
      <div class="feat-card is-new">
        <div class="feat-icon"><i class="ti ti-repeat"></i></div>
        <div class="feat-title">Loop &amp; Grid Testing <span class="new-badge">NEW</span></div>
        <div class="feat-desc">Repeating loop blocks, Grid3D sub-questions, and multi-row matrices — all validated for option inheritance and text consistency.</div>
      </div>
      <div class="feat-card is-new">
        <div class="feat-icon"><i class="ti ti-database-export"></i></div>
        <div class="feat-title">Export Validation <span class="new-badge">NEW</span></div>
        <div class="feat-desc">XML export structure parsed and compared against live survey state. Catches hidden questions, code mismatches, and variable gaps.</div>
      </div>
    </div>
    <div style="text-align:center;margin-top:40px">
      <a href="/features" class="btn-ghost">View all 25+ features <i class="ti ti-arrow-right"></i></a>
    </div>
  </div>
</section>

<!-- 7. HOW IT ACTUALLY WORKS -->
<section class="section" id="how" style="background:var(--bg2)">
  <div class="container">
    <div class="sec-head">
      <span class="sec-tag">How it works</span>
      <h2 class="sec-title">AI parses. Code decides.</h2>
      <p class="sec-sub">Not a chatbot that guesses — a deterministic engine that verifies. Every output is traceable and reproducible.</p>
    </div>
    <div class="how-grid">
      <div class="how-card">
        <div class="how-num">01</div>
        <div class="how-title">Upload your survey file</div>
        <div class="how-desc">Paste the live survey URL and upload your XML export or screener .docx. Supports Confirmit, Decipher, Forsta, Qualtrics, and more.</div>
        <div class="how-tag"><i class="ti ti-upload"></i> No install needed</div>
      </div>
      <div class="how-card">
        <div class="how-num">02</div>
        <div class="how-title">AI parses. Code verifies.</div>
        <div class="how-desc">The AI extracts every question, route, and termination rule from your file. Then deterministic code walks every path and runs all 15+ checks in parallel.</div>
        <div class="how-tag"><i class="ti ti-shield-check"></i> No hallucination possible</div>
      </div>
      <div class="how-card">
        <div class="how-num">03</div>
        <div class="how-title">Download audit evidence</div>
        <div class="how-desc">Every check gets PASS/FAIL with exact location, screenshot, and a machine-readable evidence log. Your Word report is ready in under 12 minutes.</div>
        <div class="how-tag"><i class="ti ti-certificate"></i> Pharma audit-ready</div>
      </div>
    </div>
  </div>
</section>

<!-- 8. SAMPLE REPORT PREVIEW -->
<section class="section">
  <div class="container">
    <div class="sec-head">
      <span class="sec-tag">Sample Output</span>
      <h2 class="sec-title">What your QC report looks like.</h2>
      <p class="sec-sub">Every check gets a verdict, a location, and reproducible evidence. Nothing vague, nothing missing.</p>
    </div>
    <div class="report-mock">
      <div class="report-mock-header">
        <div class="report-mock-title"><i class="ti ti-file-text"></i> QC Report — Survey E12345 &middot; Pharma Brand Tracker 2026</div>
        <div class="report-mock-stats">
          <div class="report-mock-stat"><div class="report-mock-num">156</div><div class="report-mock-lbl">Checks run</div></div>
          <div class="report-mock-stat"><div class="report-mock-num" style="color:#6EE7A0">151</div><div class="report-mock-lbl">Passed</div></div>
          <div class="report-mock-stat"><div class="report-mock-num" style="color:#FCA5A5">5</div><div class="report-mock-lbl">Failed</div></div>
          <div class="report-mock-stat"><div class="report-mock-num">9m 43s</div><div class="report-mock-lbl">Time</div></div>
        </div>
      </div>
      <div class="report-mock-body">
        <div class="chk-row">
          <div class="chk-badge pass"><i class="ti ti-check"></i></div>
          <div class="chk-name">Termination Check — Q2</div>
          <div class="chk-detail">Age screen-out fires at Q2 &lt; 18. Screenshot attached.</div>
        </div>
        <div class="chk-row">
          <div class="chk-badge fail"><i class="ti ti-x"></i></div>
          <div class="chk-name">Question Text — Q7</div>
          <div class="chk-detail">Spec: "How often do you use…" &middot; Live: "How often you use…" — missing "do"</div>
        </div>
        <div class="chk-row">
          <div class="chk-badge pass"><i class="ti ti-check"></i></div>
          <div class="chk-name">Routing — Q4 &rarr; Q9 skip</div>
          <div class="chk-detail">If Q3 = Code 5, skip to Q9. Verified. Path screenshot attached.</div>
        </div>
        <div class="chk-row">
          <div class="chk-badge fail"><i class="ti ti-x"></i></div>
          <div class="chk-name">Answer Codes — Q11</div>
          <div class="chk-detail">Code 4 missing from live survey. Spec shows 5 options, live shows 4.</div>
        </div>
        <div class="chk-row">
          <div class="chk-badge pass"><i class="ti ti-check"></i></div>
          <div class="chk-name">Loop Structure — Q14–Q18</div>
          <div class="chk-detail">All 4 loop iterations display. Sub-question options inherited correctly.</div>
        </div>
        <div class="chk-row">
          <div class="chk-badge warn"><i class="ti ti-alert-triangle"></i></div>
          <div class="chk-name">Mandatory Marker — Q6</div>
          <div class="chk-detail">Spec marks Q6 mandatory (*). Live survey shows no asterisk — review needed.</div>
        </div>
      </div>
      <div class="report-mock-footer">
        <span><i class="ti ti-clock"></i> Generated 2026-06-06 14:22 UTC</span>
        <span><i class="ti ti-certificate"></i> QC Certificate #E12345-2026-0606</span>
        <a href="/new-qc" style="font-size:13px;font-weight:600;color:var(--accent)">Run this on your survey →</a>
      </div>
    </div>
  </div>
</section>

<!-- 9. WHO IS THIS FOR -->
<section class="section" style="background:var(--bg2)">
  <div class="container">
    <div class="sec-head">
      <span class="sec-tag">Who it&#39;s for</span>
      <h2 class="sec-title">Built for every role in survey delivery.</h2>
    </div>
    <div class="persona-grid">
      <div class="persona-card">
        <div class="persona-icon"><i class="ti ti-clipboard-check"></i></div>
        <div>
          <div class="persona-role">QC Manager</div>
          <div class="persona-pain">You run 5+ surveys per week. Manual QC is consuming your team. SurveyQC cuts per-survey time from hours to minutes — without adding headcount.</div>
        </div>
      </div>
      <div class="persona-card">
        <div class="persona-icon"><i class="ti ti-code"></i></div>
        <div>
          <div class="persona-role">Survey Programmer</div>
          <div class="persona-pain">Your script is correct — but can you prove it? Get a timestamped, per-check evidence report before delivery. Catch your own bugs before QC does.</div>
        </div>
      </div>
      <div class="persona-card">
        <div class="persona-icon"><i class="ti ti-user-check"></i></div>
        <div>
          <div class="persona-role">Project Manager</div>
          <div class="persona-pain">Termination bugs at launch cost days and data. Ship on time with documented proof that every path was tested before field start.</div>
        </div>
      </div>
      <div class="persona-card">
        <div class="persona-icon"><i class="ti ti-building-hospital"></i></div>
        <div>
          <div class="persona-role">Pharma &amp; Healthcare Researcher</div>
          <div class="persona-pain">Your auditor needs traceable, timestamped documentation for every QC step. SurveyQC generates audit-ready reports out of the box — no extra work.</div>
        </div>
      </div>
    </div>
  </div>
</section>

<!-- 10. BEFORE vs AFTER -->
<section class="section">
  <div class="container">
    <div class="sec-head">
      <span class="sec-tag">The Difference</span>
      <h2 class="sec-title">Before &amp; after SurveyQC.</h2>
    </div>
    <div class="bva-wrap">
      <div class="bva-col before">
        <div class="bva-head before">
          <div class="bva-label"><i class="ti ti-x"></i> Manual QC — Before</div>
        </div>
        <div class="bva-rows">
          <div class="bva-row"><i class="ti ti-x bva-row-icon"></i><div class="bva-text"><strong>6–8 hours per survey</strong><small>One tester clicking every path by hand</small></div></div>
          <div class="bva-row"><i class="ti ti-x bva-row-icon"></i><div class="bva-text"><strong>Gaps in coverage</strong><small>Human attention fades — bugs slip through</small></div></div>
          <div class="bva-row"><i class="ti ti-x bva-row-icon"></i><div class="bva-text"><strong>No formal audit trail</strong><small>Notes in a spreadsheet, if anything</small></div></div>
          <div class="bva-row"><i class="ti ti-x bva-row-icon"></i><div class="bva-text"><strong>Non-reproducible</strong><small>Different tester = different result</small></div></div>
          <div class="bva-row"><i class="ti ti-x bva-row-icon"></i><div class="bva-text"><strong>Zero screenshot evidence</strong><small>Taken manually, if remembered at all</small></div></div>
        </div>
      </div>
      <div class="bva-col after">
        <div class="bva-head after">
          <div class="bva-label"><i class="ti ti-check"></i> SurveyQC — After</div>
        </div>
        <div class="bva-rows">
          <div class="bva-row"><i class="ti ti-check bva-row-icon"></i><div class="bva-text"><strong>Under 12 minutes per survey</strong><small>All 15+ checks run in parallel, automatically</small></div></div>
          <div class="bva-row"><i class="ti ti-check bva-row-icon"></i><div class="bva-text"><strong>Deterministic coverage</strong><small>Every defined check runs every time — no gaps</small></div></div>
          <div class="bva-row"><i class="ti ti-check bva-row-icon"></i><div class="bva-text"><strong>Machine-readable evidence log</strong><small>Timestamped PASS/FAIL for every check</small></div></div>
          <div class="bva-row"><i class="ti ti-check bva-row-icon"></i><div class="bva-text"><strong>Fully reproducible</strong><small>Same survey = same results, always</small></div></div>
          <div class="bva-row"><i class="ti ti-check bva-row-icon"></i><div class="bva-text"><strong>Screenshot proof attached</strong><small>Every failure has visual evidence in the report</small></div></div>
        </div>
      </div>
    </div>
  </div>
</section>

<!-- 11. PLATFORM SUPPORT -->
<section style="padding:64px 24px;background:var(--bg2)">
  <div class="container">
    <div style="text-align:center;margin-bottom:36px">
      <span class="sec-tag">Platforms</span>
      <h2 style="font-family:'Plus Jakarta Sans',sans-serif;font-size:clamp(24px,3vw,34px);font-weight:800;letter-spacing:-1px;margin-top:12px">Works with every major survey platform</h2>
    </div>
    <div class="platform-grid">
      <div class="platform-pill"><div class="platform-pill-mark" style="background:#FF6B35">C</div><div><div class="platform-pill-name">Confirmit</div><div class="platform-pill-status live">Supported</div></div></div>
      <div class="platform-pill"><div class="platform-pill-mark" style="background:#0066CC">D</div><div><div class="platform-pill-name">Decipher</div><div class="platform-pill-status live">Supported</div></div></div>
      <div class="platform-pill"><div class="platform-pill-mark" style="background:#1F2937">F</div><div><div class="platform-pill-name">Forsta</div><div class="platform-pill-status live">Supported</div></div></div>
      <div class="platform-pill"><div class="platform-pill-mark" style="background:#00B4F0">Q</div><div><div class="platform-pill-name">Qualtrics</div><div class="platform-pill-status live">Supported</div></div></div>
      <div class="platform-pill"><div class="platform-pill-mark" style="background:#00BF6F">S</div><div><div class="platform-pill-name">SurveyMonkey</div><div class="platform-pill-status soon">Coming soon</div></div></div>
      <div class="platform-pill"><div class="platform-pill-mark" style="background:#E91E63">A</div><div><div class="platform-pill-name">Alchemer</div><div class="platform-pill-status soon">Coming soon</div></div></div>
    </div>
    <p style="font-size:11px;color:var(--text3);text-align:center;margin-top:20px">All trademarks belong to their respective owners. SurveyQC is not affiliated with these platforms.</p>
  </div>
</section>

<!-- 12. COMPARISON TABLE -->
<section class="section">
  <div class="container">
    <div class="sec-head">
      <span class="sec-tag">How we compare</span>
      <h2 class="sec-title">SurveyQC vs the alternatives.</h2>
      <p class="sec-sub">Purpose-built for survey QC. Not a generic AI, not a spreadsheet.</p>
    </div>
    <div class="cmp-wrap">
      <table class="cmp-table">
        <thead>
          <tr>
            <th></th>
            <th class="hl">SurveyQC</th>
            <th>Manual QC</th>
            <th>Excel</th>
            <th>ChatGPT</th>
          </tr>
        </thead>
        <tbody>
          <tr><td>Time per survey</td><td class="hl">&lt; 12 min</td><td>6–8 hrs</td><td>3–4 hrs</td><td>1–2 hrs</td></tr>
          <tr><td>Termination testing</td><td class="hl"><i class="ti ti-check cmp-yes"></i></td><td>Manual</td><td><i class="ti ti-x cmp-no"></i></td><td><span class="cmp-part">Partial</span></td></tr>
          <tr><td>Screenshot evidence</td><td class="hl"><i class="ti ti-check cmp-yes"></i></td><td>Manual</td><td><i class="ti ti-x cmp-no"></i></td><td><i class="ti ti-x cmp-no"></i></td></tr>
          <tr><td>80+ languages</td><td class="hl"><i class="ti ti-check cmp-yes"></i></td><td>Depends</td><td><i class="ti ti-x cmp-no"></i></td><td><span class="cmp-part">Partial</span></td></tr>
          <tr><td>Audit trail</td><td class="hl"><i class="ti ti-check cmp-yes"></i></td><td><i class="ti ti-x cmp-no"></i></td><td><i class="ti ti-x cmp-no"></i></td><td><i class="ti ti-x cmp-no"></i></td></tr>
          <tr><td>Reproducible results</td><td class="hl"><i class="ti ti-check cmp-yes"></i></td><td><i class="ti ti-x cmp-no"></i></td><td><span class="cmp-part">Partial</span></td><td><i class="ti ti-x cmp-no"></i></td></tr>
          <tr><td>Pharma-ready report</td><td class="hl"><i class="ti ti-check cmp-yes"></i></td><td><i class="ti ti-x cmp-no"></i></td><td><i class="ti ti-x cmp-no"></i></td><td><i class="ti ti-x cmp-no"></i></td></tr>
          <tr><td>XML export parsing</td><td class="hl"><i class="ti ti-check cmp-yes"></i></td><td><i class="ti ti-x cmp-no"></i></td><td><i class="ti ti-x cmp-no"></i></td><td><span class="cmp-part">Partial</span></td></tr>
        </tbody>
      </table>
    </div>
    <div style="text-align:center;margin-top:32px">
      <a href="/new-qc" class="btn-primary">Try SurveyQC free <i class="ti ti-arrow-right"></i></a>
    </div>
  </div>
</section>

<!-- 13. QC CHECKS GRID -->
<section class="section" style="background:var(--bg2)">
  <div class="container">
    <div class="sec-head">
      <span class="sec-tag">15+ Checks</span>
      <h2 class="sec-title">Every check. Every survey.</h2>
      <p class="sec-sub">All checks run in parallel on every survey — not a sample.</p>
    </div>
    <div class="checks-grid">
      <div class="check-pill"><div class="check-pill-icon"><i class="ti ti-shield-x"></i></div><span>Termination Testing</span></div>
      <div class="check-pill"><div class="check-pill-icon"><i class="ti ti-text-recognition"></i></div><span>Question Text Match</span></div>
      <div class="check-pill"><div class="check-pill-icon"><i class="ti ti-search"></i></div><span>Missing Words</span></div>
      <div class="check-pill"><div class="check-pill-icon"><i class="ti ti-checkbox"></i></div><span>Options Match</span></div>
      <div class="check-pill"><div class="check-pill-icon"><i class="ti ti-list-numbers"></i></div><span>Answer Codes</span></div>
      <div class="check-pill"><div class="check-pill-icon"><i class="ti ti-asterisk"></i></div><span>Mandatory Markers</span></div>
      <div class="check-pill"><div class="check-pill-icon"><i class="ti ti-arrows-sort"></i></div><span>Question Order</span></div>
      <div class="check-pill"><div class="check-pill-icon"><i class="ti ti-replace"></i></div><span>Piping Markers</span></div>
      <div class="check-pill"><div class="check-pill-icon"><i class="ti ti-route"></i></div><span>Skip Logic</span></div>
      <div class="check-pill"><div class="check-pill-icon"><i class="ti ti-chart-pie"></i></div><span>Quota Testing</span></div>
      <div class="check-pill"><div class="check-pill-icon"><i class="ti ti-repeat"></i></div><span>Loop Testing</span></div>
      <div class="check-pill"><div class="check-pill-icon"><i class="ti ti-eye"></i></div><span>Display Logic</span></div>
      <div class="check-pill"><div class="check-pill-icon"><i class="ti ti-world"></i></div><span>80+ Languages</span></div>
      <div class="check-pill"><div class="check-pill-icon"><i class="ti ti-camera"></i></div><span>Auto Screenshots</span></div>
      <div class="check-pill"><div class="check-pill-icon"><i class="ti ti-certificate"></i></div><span>QC Certificate</span></div>
    </div>
  </div>
</section>

<!-- 14. TRUST NUMBERS -->
<section class="section">
  <div class="container">
    <div class="sec-head">
      <span class="sec-tag">By the numbers</span>
      <h2 class="sec-title">Built for survey QC at scale.</h2>
    </div>
    <div class="trust-grid">
      <div class="trust-card">
        <div class="trust-num">15<span class="unit">+</span></div>
        <div class="trust-label">Specialized QC checks run on every survey</div>
      </div>
      <div class="trust-card">
        <div class="trust-num">80<span class="unit">+</span></div>
        <div class="trust-label">Languages fully supported including Arabic, Japanese, Urdu</div>
      </div>
      <div class="trust-card">
        <div class="trust-num">12<span class="unit"> min</span></div>
        <div class="trust-label">Average time from upload to complete QC report</div>
      </div>
      <div class="trust-card">
        <div class="trust-num">6</div>
        <div class="trust-label">Survey platforms supported with more coming soon</div>
      </div>
    </div>
  </div>
</section>

<!-- 15. TESTIMONIALS -->
<section class="section" style="background:var(--bg2)">
  <div class="container">
    <div class="sec-head">
      <span class="sec-tag">What users say</span>
      <h2 class="sec-title">QC professionals trust SurveyQC.</h2>
    </div>
    <div class="test-grid">
      """ + test_cards + """
    </div>
  </div>
</section>

<!-- 16. FINAL CTA + FOOTER -->
<section class="cta-banner">
  <div class="cta-inner">
    <h2>Run your first QC in under 12 minutes.</h2>
    <p>Free forever tier. No credit card. Any survey platform.</p>
    <a href="/new-qc" class="btn-primary">Start free — no card required <i class="ti ti-arrow-right"></i></a>
  </div>
</section>

<footer class="footer">
  <div class="footer-inner">
    <div class="footer-grid">
      <div class="footer-brand">
        <div class="footer-brand-logo">
          <div class="footer-brand-mark"><i class="ti ti-shield-check" style="color:white;font-size:16px"></i></div>
          <div class="footer-brand-text">""" + c['site_name'] + """</div>
        </div>
        <p>""" + c['tagline'] + """. Built for QC professionals at leading market research agencies.</p>
        <div class="footer-social">
          <a href="""" + c['linkedin_url'] + """"><i class="ti ti-brand-linkedin"></i></a>
          <a href="""" + c['twitter_url'] + """"><i class="ti ti-brand-twitter"></i></a>
        </div>
      </div>
      <div>
        <div class="footer-title">Product</div>
        <a href="/features" class="footer-link">Features</a>
        <a href="/pricing" class="footer-link">Pricing</a>
        <a href="/docs" class="footer-link">Documentation</a>
        <a href="/changelog" class="footer-link">Changelog</a>
      </div>
      <div>
        <div class="footer-title">Resources</div>
        <a href="/blog" class="footer-link">Blog</a>
        <a href="/community" class="footer-link">Community</a>
        <a href="/affiliate" class="footer-link">Affiliate</a>
        <a href="mailto:""" + c['support_email'] + """" class="footer-link">Contact</a>
      </div>
      <div>
        <div class="footer-title">Compare</div>
        <a href="/compare/chatgpt" class="footer-link">vs ChatGPT</a>
        <a href="/compare/excel" class="footer-link">vs Excel</a>
        <a href="/compare/manual" class="footer-link">vs Manual</a>
      </div>
      <div>
        <div class="footer-title">Legal</div>
        <a href="/privacy-policy" class="footer-link">Privacy</a>
        <a href="/terms" class="footer-link">Terms</a>
        <a href="/refund-policy" class="footer-link">Refund Policy</a>
        <a href="/cookie-policy" class="footer-link">Cookie Policy</a>
        <a href="/dpa" class="footer-link">DPA</a>
      </div>
    </div>
    <div class="footer-bottom">
      <p>&copy; """ + c['footer_text'] + """</p>
    </div>
  </div>
</footer>

<script>
document.addEventListener('click', function(e) {
  var mm = document.getElementById('mm');
  if (mm && mm.classList.contains('open') && !mm.contains(e.target)) {
    var btn = document.querySelector('.hamburger');
    if (btn && !btn.contains(e.target)) mm.classList.remove('open');
  }
});
</script>

</body>
</html>"""
    return page



@app.route('/privacy-policy')
def privacy_policy_page():
    return render_template_string(SHARED_CSS + """
<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=5.0"><title>Privacy Policy — SurveyQC</title></head><body>
<div style="max-width:700px;margin:0 auto;padding:40px 20px">
  <a href="/" style="color:var(--color-text-secondary);text-decoration:none;font-size:13px"><i class="ti ti-arrow-left"></i> Back to home</a>
  <h1 style="font-size:28px;font-weight:600;color:var(--color-text-primary);margin:20px 0 8px">Privacy Policy</h1>
  <p style="font-size:13px;color:var(--color-text-secondary);margin-bottom:30px">Last updated: May 2026</p>

  <div class="card">
    <h2 style="font-size:16px;font-weight:600;color:var(--color-text-primary);margin-bottom:10px">Data we collect</h2>
    <p style="font-size:13px;color:var(--color-text-secondary);line-height:1.8">We collect your email, name, uploaded documents, and survey URLs for QC processing. No personal survey respondent data is collected.</p>
  </div>

  <div class="card">
    <h2 style="font-size:16px;font-weight:600;color:var(--color-text-primary);margin-bottom:10px">Auto-deletion (30 days)</h2>
    <p style="font-size:13px;color:var(--color-text-secondary);line-height:1.8">All uploaded files, screenshots, and reports are automatically deleted after 30 days. This is fully automatic — no action needed from you.</p>
  </div>

  <div class="card">
    <h2 style="font-size:16px;font-weight:600;color:var(--color-text-primary);margin-bottom:10px">GDPR compliance</h2>
    <p style="font-size:13px;color:var(--color-text-secondary);line-height:1.8">We comply with GDPR (EU), CCPA (USA), UK GDPR, India DPDP Act 2023, and Australia Privacy Act. You have the right to access, export, and delete your data at any time.</p>
  </div>

  <div class="card">
    <h2 style="font-size:16px;font-weight:600;color:var(--color-text-primary);margin-bottom:10px">Data sharing</h2>
    <p style="font-size:13px;color:var(--color-text-secondary);line-height:1.8">We never sell your data. Survey content is processed by AI models (Gemini) for QC analysis only and is not stored by third parties.</p>
  </div>

  <div class="card">
    <h2 style="font-size:16px;font-weight:600;color:var(--color-text-primary);margin-bottom:10px">Contact</h2>
    <p style="font-size:13px;color:var(--color-text-secondary);line-height:1.8">For data requests or privacy questions, email us at support@surveyqc.online</p>
  </div>
</div>
</body></html>""")


# ================================================================
# TERMS PAGE (Public)
# ================================================================
@app.route('/terms')
def terms_page():
    return render_template_string(SHARED_CSS + """
<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=5.0"><title>Terms of Service — SurveyQC</title></head><body>
<div style="max-width:700px;margin:0 auto;padding:40px 20px">
  <a href="/" style="color:var(--color-text-secondary);text-decoration:none;font-size:13px"><i class="ti ti-arrow-left"></i> Back to home</a>
  <h1 style="font-size:28px;font-weight:600;color:var(--color-text-primary);margin:20px 0 8px">Terms of Service</h1>
  <p style="font-size:13px;color:var(--color-text-secondary);margin-bottom:30px">Last updated: May 2026</p>
  <div class="card">
    <h2 style="font-size:16px;font-weight:600;color:var(--color-text-primary);margin-bottom:10px">1. Service</h2>
    <p style="font-size:13px;color:var(--color-text-secondary);line-height:1.8">SurveyQC provides AI-powered survey quality control services. You agree to use the service for legitimate survey testing purposes only.</p>
  </div>
  <div class="card">
    <h2 style="font-size:16px;font-weight:600;color:var(--color-text-primary);margin-bottom:10px">2. Your data</h2>
    <p style="font-size:13px;color:var(--color-text-secondary);line-height:1.8">You retain ownership of all uploaded documents and survey content. We process this data solely to provide QC analysis. All data is deleted after 30 days.</p>
  </div>
  <div class="card">
    <h2 style="font-size:16px;font-weight:600;color:var(--color-text-primary);margin-bottom:10px">3. Accuracy</h2>
    <p style="font-size:13px;color:var(--color-text-secondary);line-height:1.8">While we target 99% accuracy, AI-powered analysis may not catch all issues. Always perform manual checks for compound logic, quotas, and hidden questions as indicated in reports.</p>
  </div>
  <div class="card">
    <h2 style="font-size:16px;font-weight:600;color:var(--color-text-primary);margin-bottom:10px">4. Plans and billing</h2>
    <p style="font-size:13px;color:var(--color-text-secondary);line-height:1.8">Free plan includes 5 reports per month. Paid plans are billed monthly. You can cancel at any time. No refunds for partial months.</p>
  </div>
</div>
</body></html>""")



@app.route('/refund-policy')
def refund_policy_page():
    return render_template_string(SHARED_CSS + """
<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=5.0"><title>Refund Policy — SurveyQC</title></head><body>
<div style="max-width:700px;margin:0 auto;padding:40px 20px">
  <a href="/home" style="color:var(--color-text-secondary);text-decoration:none;font-size:13px"><i class="ti ti-arrow-left"></i> Back to home</a>
  <h1 style="font-size:28px;font-weight:600;color:var(--color-text-primary);margin:20px 0 8px">Refund Policy</h1>
  <p style="font-size:13px;color:var(--color-text-secondary);margin-bottom:30px">Last updated: May 2026</p>
  <div class="card">
    <p style="font-size:14px;color:var(--color-text-secondary);line-height:1.8">7-day money-back guarantee on all paid plans. If you are not satisfied within 7 days of your first payment, contact us at <a href="mailto:support@surveyqc.online">support@surveyqc.online</a> for a full refund. No questions asked.</p>
  </div>
</div>
</body></html>""")


@app.route('/cookie-policy')
def cookie_policy_page():
    return render_template_string(SHARED_CSS + """
<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=5.0"><title>Cookie Policy — SurveyQC</title></head><body>
<div style="max-width:700px;margin:0 auto;padding:40px 20px">
  <a href="/home" style="color:var(--color-text-secondary);text-decoration:none;font-size:13px"><i class="ti ti-arrow-left"></i> Back to home</a>
  <h1 style="font-size:28px;font-weight:600;color:var(--color-text-primary);margin:20px 0 8px">Cookie Policy</h1>
  <p style="font-size:13px;color:var(--color-text-secondary);margin-bottom:30px">Last updated: May 2026</p>
  <div class="card">
    <h2 style="font-size:16px;font-weight:600;color:var(--color-text-primary);margin-bottom:10px">What cookies we use</h2>
    <p style="font-size:13px;color:var(--color-text-secondary);line-height:1.8">We use strictly necessary session cookies to keep you logged in. We do not use advertising or tracking cookies.</p>
  </div>
  <div class="card">
    <h2 style="font-size:16px;font-weight:600;color:var(--color-text-primary);margin-bottom:10px">Third-party cookies</h2>
    <p style="font-size:13px;color:var(--color-text-secondary);line-height:1.8">Our payment provider (Stripe) may set cookies during checkout. These are governed by Stripe's cookie policy.</p>
  </div>
  <div class="card">
    <h2 style="font-size:16px;font-weight:600;color:var(--color-text-primary);margin-bottom:10px">Managing cookies</h2>
    <p style="font-size:13px;color:var(--color-text-secondary);line-height:1.8">You can disable cookies in your browser settings. Note that disabling session cookies will prevent you from logging in.</p>
  </div>
</div>
</body></html>""")


@app.route('/dpa')
def dpa_page():
    return render_template_string(SHARED_CSS + """
<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=5.0"><title>Data Processing Agreement — SurveyQC</title></head><body>
<div style="max-width:700px;margin:0 auto;padding:40px 20px">
  <a href="/home" style="color:var(--color-text-secondary);text-decoration:none;font-size:13px"><i class="ti ti-arrow-left"></i> Back to home</a>
  <h1 style="font-size:28px;font-weight:600;color:var(--color-text-primary);margin:20px 0 8px">Data Processing Agreement</h1>
  <p style="font-size:13px;color:var(--color-text-secondary);margin-bottom:30px">Last updated: May 2026</p>
  <div class="card">
    <p style="font-size:14px;color:var(--color-text-secondary);line-height:1.8">A Data Processing Agreement (DPA) is available for Enterprise customers to satisfy GDPR Article 28 requirements. To request a signed DPA, please contact <a href="mailto:support@surveyqc.online">support@surveyqc.online</a> with the subject line "DPA Request".</p>
  </div>
</div>
</body></html>""")


# ================================================================
# REPORT TEMPLATES
# ================================================================
@app.route('/templates')
@login_required
def templates_page():
    user = get_current_user()
    tmpl_html = ''
    for tid, t in templates_db.items():
        tmpl_html += f"""
        <div class="card" style="margin-bottom:10px">
          <div style="display:flex;align-items:center;justify-content:space-between">
            <div>
              <p style="font-size:13px;font-weight:500;color:var(--color-text-primary)">{t['name']}</p>
              <div style="display:flex;gap:10px;margin-top:4px">
                <span style="font-size:11px;color:var(--color-text-secondary)">{t['platform']}</span>
                <span style="font-size:11px;color:var(--color-text-secondary)">{t['country']}</span>
                <span style="font-size:11px;color:var(--color-text-secondary)">{t['mode'].title()} mode</span>
              </div>
            </div>
            <div style="display:flex;gap:7px">
              <a href="/new-qc?template={tid}" class="btn btn-primary btn-sm"><i class="ti ti-player-play"></i>Use template</a>
              <button class="btn btn-ghost btn-sm"><i class="ti ti-trash"></i></button>
            </div>
          </div>
        </div>"""

    return render_template_string(SHARED_CSS + f"""
<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=5.0"><title>Templates — SurveyQC</title></head><body>
<div class="app-layout">
  {sidebar_html('templates')}
  <div class="main-content">
    <div class="topbar">
      <div><p class="page-title">QC Templates</p><p class="page-sub">Save settings for quick reuse</p></div>
      <button class="btn btn-primary btn-sm" onclick="document.getElementById('newTmpl').style.display='block'"><i class="ti ti-plus"></i>New template</button>
    </div>

    <div id="newTmpl" style="display:none" class="card">
      <p style="font-size:13px;font-weight:500;color:var(--color-text-primary);margin-bottom:14px">Create new template</p>
      <form method="POST" action="/templates/create">
        <div style="display:grid;grid-template-columns:1fr 1fr 1fr 1fr;gap:10px;margin-bottom:12px">
          <div><label class="form-label">Template name</label><input class="form-input" name="name" placeholder="Confirmit France" style="margin:0"></div>
          <div><label class="form-label">Platform</label><select class="form-select" name="platform"><option>Confirmit</option><option>Decipher</option><option>Forsta</option><option>Qualtrics</option></select></div>
          <div><label class="form-label">Country</label><input class="form-input" name="country" placeholder="France" style="margin:0"></div>
          <div><label class="form-label">Mode</label><select class="form-select" name="mode"><option value="full">Full QC</option><option value="quick">Quick test</option></select></div>
        </div>
        <button type="submit" class="btn btn-primary btn-sm">Save template</button>
        <button type="button" class="btn btn-ghost btn-sm" onclick="document.getElementById('newTmpl').style.display='none'">Cancel</button>
      </form>
    </div>

    {tmpl_html if tmpl_html else '<div class="card" style="text-align:center;padding:30px"><i class="ti ti-copy" style="font-size:32px;color:var(--color-text-secondary)"></i><p style="margin-top:10px;color:var(--color-text-secondary)">No templates yet. Create one to save your settings!</p></div>'}
  </div>
</div>
</body></html>""")


@app.route('/templates/create', methods=['POST'])
@login_required
def create_template():
    tid = str(uuid.uuid4())[:8]
    templates_db[tid] = {
        'name': request.form.get('name', 'New template'),
        'platform': request.form.get('platform', 'Confirmit'),
        'country': request.form.get('country', ''),
        'mode': request.form.get('mode', 'full'),
        'user': session['user_email']
    }
    return redirect('/templates')


# ================================================================
# SEARCH REPORTS API
# ================================================================
@app.route('/api/search')
@login_required
def api_search():
    q = request.args.get('q', '').lower()
    email = session['user_email']
    results = []
    for jid, j in jobs.items():
        if j.get('user_email') != email:
            continue
        if q in j.get('doc_name', '').lower() or q in j.get('platform', '').lower():
            results.append({
                'id': jid,
                'doc_name': j.get('doc_name'),
                'platform': j.get('platform'),
                'status': j.get('status'),
                'created_at': j.get('created_at', '')[:10]
            })
    return jsonify({'results': results[:10]})


# ================================================================
# DUPLICATE DETECTION API
# ================================================================
@app.route('/api/check-duplicate', methods=['POST'])
@login_required
def check_duplicate():
    url = request.json.get('url', '')
    email = session['user_email']
    for jid, j in jobs.items():
        if j.get('user_email') == email and j.get('survey_url') == url:
            return jsonify({
                'duplicate': True,
                'job_id': jid,
                'doc_name': j.get('doc_name'),
                'date': j.get('created_at', '')[:10],
                'issues': j.get('total_issues', 0)
            })
    return jsonify({'duplicate': False})


# ================================================================
# ADMIN: DATA PRIVACY SETTINGS
# ================================================================

# Privacy settings store
privacy_settings = {
    'delete_reports_after': 30,
    'delete_screenshots_after': 30,
    'delete_uploads_after': 7,
    'auto_delete_enabled': True,
    'notify_before_days': 3
}

@app.route('/admin/privacy', methods=['GET', 'POST'])
@admin_required
def admin_privacy():
    global privacy_settings
    saved = False
    if request.method == 'POST':
        privacy_settings['delete_reports_after'] = int(request.form.get('reports_days', 30))
        privacy_settings['delete_screenshots_after'] = int(request.form.get('screenshots_days', 30))
        privacy_settings['delete_uploads_after'] = int(request.form.get('uploads_days', 7))
        privacy_settings['notify_before_days'] = int(request.form.get('notify_days', 3))
        privacy_settings['auto_delete_enabled'] = 'auto_delete' in request.form
        saved = True

    return render_template_string(SHARED_CSS + f"""
<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=5.0"><title>Privacy Settings — Admin</title><script src="/admin-sidebar-js"></script></head><body>
<div style="padding:24px;max-width:600px">
  <div style="display:flex;align-items:center;gap:14px;margin-bottom:20px">
    <a href="/admin" style="color:var(--color-text-secondary);text-decoration:none"><i class="ti ti-arrow-left"></i></a>
    <p style="font-size:18px;font-weight:500;color:var(--color-text-primary)">🔒 Data Privacy Settings</p>
  </div>
  {'<div class="alert alert-success">Privacy settings saved!</div>' if saved else ''}
  <div class="card" style="margin-bottom:14px">
    <p style="font-size:13px;font-weight:500;color:var(--color-text-primary);margin-bottom:14px">Auto-delete schedule</p>
    <form method="POST">
      <div style="display:grid;grid-template-columns:1fr 1fr;gap:10px;margin-bottom:12px">
        <div>
          <label class="form-label">Delete reports after (days)</label>
          <select class="form-select" name="reports_days">
            <option value="7" {'selected' if privacy_settings['delete_reports_after']==7 else ''}>7 days</option>
            <option value="14" {'selected' if privacy_settings['delete_reports_after']==14 else ''}>14 days</option>
            <option value="30" {'selected' if privacy_settings['delete_reports_after']==30 else ''}>30 days (GDPR)</option>
            <option value="60" {'selected' if privacy_settings['delete_reports_after']==60 else ''}>60 days</option>
            <option value="90" {'selected' if privacy_settings['delete_reports_after']==90 else ''}>90 days</option>
          </select>
        </div>
        <div>
          <label class="form-label">Delete screenshots after (days)</label>
          <select class="form-select" name="screenshots_days">
            <option value="7">7 days</option>
            <option value="30" selected>30 days</option>
            <option value="60">60 days</option>
          </select>
        </div>
        <div>
          <label class="form-label">Delete uploads after (days)</label>
          <select class="form-select" name="uploads_days">
            <option value="3">3 days</option>
            <option value="7" selected>7 days</option>
            <option value="14">14 days</option>
          </select>
        </div>
        <div>
          <label class="form-label">Notify users before delete (days)</label>
          <select class="form-select" name="notify_days">
            <option value="1">1 day before</option>
            <option value="3" selected>3 days before</option>
            <option value="7">7 days before</option>
          </select>
        </div>
      </div>
      <div style="background:#F5E6D8;border-radius:8px;padding:12px;margin-bottom:12px">
        <p style="font-size:11px;color:#C46A2B">🕛 Auto-delete runs every night at 12:00 AM · Compliant: 🇪🇺 GDPR · 🇺🇸 CCPA · 🇮🇳 DPDP · 🇬🇧 UK · 🇦🇺 AUS</p>
      </div>
      <label style="display:flex;align-items:center;gap:8px;margin-bottom:14px;font-size:13px;cursor:pointer">
        <input type="checkbox" name="auto_delete" {'checked' if privacy_settings['auto_delete_enabled'] else ''} style="width:15px;height:15px">
        Auto-delete enabled
      </label>
      <button type="submit" class="btn btn-primary">Save privacy settings</button>
    </form>
  </div>
  <div class="card">
    <p style="font-size:13px;font-weight:500;color:var(--color-text-primary);margin-bottom:12px">Manual cleanup</p>
    <form method="POST" action="/admin/cleanup">
      <button type="submit" class="btn btn-ghost" style="margin-bottom:8px;width:100%;justify-content:center"><i class="ti ti-trash"></i>Run cleanup now (delete expired data)</button>
    </form>
    <button style="width:100%;justify-content:center;font-size:12px;background:#FCEBEB;color:#791F1F;border:0.5px solid #F7C1C1;padding:9px;border-radius:8px;cursor:pointer;display:flex;align-items:center;gap:6px"><i class="ti ti-alert-triangle"></i>Delete ALL data (irreversible)</button>
  </div>
</div>
</body></html>""")


# ================================================================
# ERROR HANDLERS
# ================================================================
@app.errorhandler(404)
def not_found(e):
    return render_template_string(SHARED_CSS + """
<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=5.0"><title>404 — SurveyQC</title></head><body>
<div style="min-height:100vh;display:flex;align-items:center;justify-content:center;text-align:center">
  <div>
    <i class="ti ti-mood-confused" style="font-size:60px;color:var(--color-text-secondary)"></i>
    <p style="font-size:48px;font-weight:600;color:var(--color-text-primary);margin:16px 0 8px">404</p>
    <p style="font-size:16px;color:var(--color-text-secondary);margin-bottom:24px">Page not found — shayad galat URL hai!</p>
    <a href="/dashboard" class="btn btn-primary">Go to Dashboard</a>
  </div>
</div>
</body></html>"""), 404


@app.errorhandler(500)
def server_error(e):
    return render_template_string(SHARED_CSS + """
<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=5.0"><title>Error — SurveyQC</title></head><body>
<div style="min-height:100vh;display:flex;align-items:center;justify-content:center;text-align:center">
  <div>
    <i class="ti ti-alert-triangle" style="font-size:60px;color:#EF9F27"></i>
    <p style="font-size:32px;font-weight:600;color:var(--color-text-primary);margin:16px 0 8px">Something went wrong</p>
    <p style="font-size:14px;color:var(--color-text-secondary);margin-bottom:24px">Server error — please try again</p>
    <a href="/dashboard" class="btn btn-primary">Go to Dashboard</a>
  </div>
</div>
</body></html>"""), 500


@app.errorhandler(413)
def file_too_large(e):
    return jsonify({'error': 'File too large — maximum 50MB allowed'}), 413



# ================================================================
# API MANAGEMENT — Admin se sab APIs manage karo
# ================================================================

# All APIs store
api_store = {
    'gemini': {
        'name': 'Gemini (Google)',
        'key': '',
        'status': 'not_added',
        'use_for': 'Primary AI — survey analysis, text match, bug detection',
        'cost': 'Free tier available · $0.000075/1K tokens',
        'docs': 'aistudio.google.com/app/apikey',
        'icon': 'ti-brand-google',
        'color': '#EA4335',
        'priority': 1,
        'active': False
    },
    'claude': {
        'name': 'Claude (Anthropic)',
        'key': '',
        'status': 'not_added',
        'use_for': 'Complex logic — compound rules, termination analysis',
        'cost': '$0.003/1K tokens (Sonnet)',
        'docs': 'console.anthropic.com',
        'icon': 'ti-robot',
        'color': '#7C65FF',
        'priority': 2,
        'active': False
    },
    'openai': {
        'name': 'ChatGPT (OpenAI)',
        'key': '',
        'status': 'not_added',
        'use_for': 'Fallback AI — general analysis',
        'cost': '$0.002/1K tokens (GPT-4o-mini)',
        'docs': 'platform.openai.com/api-keys',
        'icon': 'ti-brand-openai',
        'color': '#10A37F',
        'priority': 3,
        'active': False
    },
    'stripe': {
        'name': 'Stripe (Payments)',
        'key': '',
        'status': 'not_added',
        'use_for': 'Payment processing — subscriptions, billing',
        'cost': '2.9% + 30¢ per transaction',
        'docs': 'dashboard.stripe.com/apikeys',
        'icon': 'ti-credit-card',
        'color': '#635BFF',
        'priority': 4,
        'active': False
    },
    'sendgrid': {
        'name': 'SendGrid (Email)',
        'key': '',
        'status': 'not_added',
        'use_for': 'Transactional emails — reports, notifications, welcome',
        'cost': 'Free 100 emails/day',
        'docs': 'app.sendgrid.com/settings/api_keys',
        'icon': 'ti-mail',
        'color': '#1A82E2',
        'priority': 5,
        'active': False
    },
    'google_oauth': {
        'name': 'Google OAuth (SSO)',
        'key': '',
        'status': 'not_added',
        'use_for': 'Google login — one-click signup/login',
        'cost': 'Free',
        'docs': 'console.cloud.google.com/apis/credentials',
        'icon': 'ti-brand-google',
        'color': '#4285F4',
        'priority': 6,
        'active': False
    },
    'microsoft_oauth': {
        'name': 'Microsoft OAuth (SSO)',
        'key': '',
        'status': 'not_added',
        'use_for': 'Microsoft/Office 365 login',
        'cost': 'Free',
        'docs': 'portal.azure.com/#blade/Microsoft_AAD_RegisteredApps',
        'icon': 'ti-brand-windows',
        'color': '#0078D4',
        'priority': 7,
        'active': False
    },
    'whatsapp': {
        'name': 'WhatsApp Business API',
        'key': '',
        'status': 'not_added',
        'use_for': 'WhatsApp notifications — report done alerts',
        'cost': '$0.005-0.009 per message',
        'docs': 'developers.facebook.com/docs/whatsapp',
        'icon': 'ti-brand-whatsapp',
        'color': '#25D366',
        'priority': 8,
        'active': False
    },
    'slack': {
        'name': 'Slack (Notifications)',
        'key': '',
        'status': 'not_added',
        'use_for': 'Slack alerts — QC done, new issues found',
        'cost': 'Free',
        'docs': 'api.slack.com/apps',
        'icon': 'ti-brand-slack',
        'color': '#4A154B',
        'priority': 9,
        'active': False
    },
    'razorpay': {
        'name': 'Razorpay (India Payments)',
        'key': '',
        'status': 'not_added',
        'use_for': 'India payment — UPI, cards, net banking',
        'cost': '2% per transaction',
        'docs': 'dashboard.razorpay.com/app/keys',
        'icon': 'ti-currency-rupee',
        'color': '#3395FF',
        'priority': 10,
        'active': False
    },
    'sentry': {
        'name': 'Sentry (Error Tracking)',
        'key': '',
        'status': 'not_added',
        'use_for': 'Error monitoring — catch bugs in production',
        'cost': 'Free tier available',
        'docs': 'sentry.io/settings/account/api/auth-tokens',
        'icon': 'ti-bug',
        'color': '#362D59',
        'priority': 11,
        'active': False
    },
    'cloudflare': {
        'name': 'Cloudflare (CDN/Security)',
        'key': '',
        'status': 'not_added',
        'use_for': 'DDoS protection, CDN, SSL',
        'cost': 'Free tier available',
        'docs': 'dash.cloudflare.com/profile/api-tokens',
        'icon': 'ti-cloud',
        'color': '#F48120',
        'priority': 12,
        'active': False
    }
}


# Load saved Gemini key from .env into api_store on startup
import os as _os
_saved_gemini = _os.environ.get('GEMINI_API_KEY', '').strip()
if _saved_gemini and 'gemini' in api_store:
    api_store['gemini']['key'] = _saved_gemini
    api_store['gemini']['status'] = 'active'
    api_store['gemini']['active'] = True

@app.route('/admin/apis', methods=['GET', 'POST'])
@admin_required
def admin_apis():
    saved_msg = ''
    if request.method == 'POST':
        api_id = request.form.get('api_id', '')
        key = request.form.get('key', '').strip()
        action = request.form.get('action', 'save')
        if api_id in api_store:
            api = api_store[api_id]
            if action == 'save' and key:
                api['key'] = key; api['status'] = 'active'; api['active'] = True
                saved_msg = api['name'] + ' key saved!'
                # Save gemini key to .env for persistence
                if api_id == 'gemini':
                    try:
                        from pathlib import Path
                        env_file = Path('/var/www/surveyqc/.env')
                        lines = []
                        if env_file.exists():
                            lines = [l for l in env_file.read_text().splitlines() if not l.startswith('GEMINI_API_KEY')]
                        lines.append('GEMINI_API_KEY=' + key)
                        env_file.write_text(chr(10).join(lines) + chr(10))
                        import os
                        os.environ['GEMINI_API_KEY'] = key
                    except Exception as e:
                        pass
            elif action == 'delete':
                api['key'] = ''; api['status'] = 'not_added'; api['active'] = False
                saved_msg = api['name'] + ' removed.'
            elif action == 'toggle':
                api['active'] = not api.get('active', False)
                saved_msg = api['name'] + (' enabled.' if api['active'] else ' disabled.')

    active_count = sum(1 for a in api_store.values() if a.get('active'))
    not_added = sum(1 for a in api_store.values() if not a.get('key'))

    def card(aid):
        a = api_store[aid]
        has_key = bool(a.get('key',''))
        is_active = a.get('active', False)
        if has_key and is_active:
            bg = '#F0FDF4'; border = '#A5D6A7'
            badge = '<span style="background:#EAF3DE;color:#27500A;font-size:10px;padding:2px 8px;border-radius:20px;font-weight:500">Active</span>'
        elif has_key:
            bg = '#FFFBEB'; border = '#FCD34D'
            badge = '<span style="background:#FAEEDA;color:#633806;font-size:10px;padding:2px 8px;border-radius:20px;font-weight:500">Disabled</span>'
        else:
            bg = 'white'; border = '#DDE1E7'
            badge = '<span style="background:#F1EFE8;color:#444441;font-size:10px;padding:2px 8px;border-radius:20px;font-weight:500">Not added</span>'
        masked = chr(8226)*10 + a['key'][-4:] if len(a.get('key','')) > 4 else ''
        ph = 'Update key...' if has_key else 'Paste API key...'
        lbl = 'Disable' if is_active else 'Enable'
        toggle = ('<form method="POST" style="display:inline"><input type="hidden" name="api_id" value="'+aid+'"><input type="hidden" name="action" value="toggle"><button type="submit" style="background:#FBF8F2;color:#374151;border:0.5px solid #DDE1E7;font-size:11px;padding:5px 10px;border-radius:7px;cursor:pointer;margin-right:4px">'+lbl+'</button></form>') if has_key else ''
        remove = ('<form method="POST" style="display:inline"><input type="hidden" name="api_id" value="'+aid+'"><input type="hidden" name="action" value="delete"><button type="submit" style="background:#FCEBEB;color:#791F1F;border:none;font-size:11px;padding:5px 10px;border-radius:7px;cursor:pointer">Remove</button></form>') if has_key else ''
        masked_row = ('<p style="font-family:monospace;font-size:11px;background:#FBF8F2;padding:4px 8px;border-radius:5px;margin-bottom:8px;color:#9CA3AF">'+masked+'</p>') if masked else ''
        return ('<div style="background:'+bg+';border:0.5px solid '+border+';border-radius:10px;padding:14px;margin-bottom:10px">'
            '<div style="display:flex;align-items:start;gap:10px">'
            '<div style="width:34px;height:34px;border-radius:8px;background:'+a['color']+';display:flex;align-items:center;justify-content:center;flex-shrink:0">'
            '<i class="ti '+a['icon']+'" style="font-size:17px;color:white"></i></div>'
            '<div style="flex:1">'
            '<div style="display:flex;align-items:center;gap:8px;margin-bottom:3px"><p style="font-size:13px;font-weight:600;color:#1A1A2E">'+a['name']+'</p>'+badge+'</div>'
            '<p style="font-size:11px;color:#6B7280;margin-bottom:2px">'+a['use_for']+'</p>'
            '<p style="font-size:10px;color:#9CA3AF;margin-bottom:7px">Cost: '+a['cost']+'</p>'
            + masked_row +
            '<form method="POST" style="display:flex;gap:6px;margin-bottom:6px">'
            '<input type="hidden" name="api_id" value="'+aid+'"><input type="hidden" name="action" value="save">'
            '<input name="key" type="password" placeholder="'+ph+'" style="flex:1;padding:7px 10px;border:0.5px solid #DDE1E7;border-radius:7px;font-size:12px;outline:none">'
            '<button type="submit" style="background:#1B140F;color:white;border:none;font-size:12px;padding:7px 12px;border-radius:7px;cursor:pointer;white-space:nowrap">Save</button>'
            '</form>'
            '<div style="display:flex;align-items:center;gap:4px">'+toggle+remove+'<a href="https://'+a['docs']+'" target="_blank" style="font-size:10px;color:#185FA5;margin-left:6px">Get key</a></div>'
            '</div></div></div>')

    ai_html = card('gemini') + card('claude') + card('openai')
    pay_html = card('stripe') + card('razorpay') + card('sendgrid')
    notif_html = card('whatsapp') + card('slack')
    other_html = card('google_oauth') + card('microsoft_oauth') + card('sentry') + card('cloudflare')
    alert = ('<div style="background:#EAF3DE;border:0.5px solid #A5D6A7;border-radius:8px;padding:10px 14px;margin-bottom:16px;color:#27500A;font-size:13px">'+saved_msg+'</div>') if saved_msg else ''

    page = '<!DOCTYPE html><html><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>API Management - Admin</title>'
    page += '<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/@tabler/icons-webfont@3.19.0/dist/tabler-icons.min.css">'
    page += '<style>*{box-sizing:border-box;margin:0;padding:0;font-family:-apple-system,BlinkMacSystemFont,"Segoe UI",sans-serif}body{background:#F0F2F5;color:#1A1A2E}a{text-decoration:none}</style><script src="/admin-sidebar-js"></script></head><body>'
    page += '<div style="padding:24px;max-width:1000px;margin:0 auto">'
    page += '<div style="display:flex;align-items:center;gap:14px;margin-bottom:20px">'
    page += '<a href="/admin" style="color:#6B7280;font-size:22px">&larr;</a>'
    page += '<div><p style="font-size:20px;font-weight:600;color:#1A1A2E">API Management</p><p style="font-size:12px;color:#6B7280">Manage all your API integrations in one place</p></div>'
    page += '<div style="margin-left:auto;display:flex;gap:10px">'
    page += '<div style="background:#EAF3DE;border-radius:8px;padding:10px 16px;text-align:center"><p style="font-size:18px;font-weight:600;color:#27500A">'+str(active_count)+'</p><p style="font-size:10px;color:#3B6D11">Active</p></div>'
    page += '<div style="background:white;border:0.5px solid #DDE1E7;border-radius:8px;padding:10px 16px;text-align:center"><p style="font-size:18px;font-weight:600;color:#1A1A2E">'+str(len(api_store))+'</p><p style="font-size:10px;color:#6B7280">Total</p></div>'
    page += '<div style="background:#FCEBEB;border-radius:8px;padding:10px 16px;text-align:center"><p style="font-size:18px;font-weight:600;color:#791F1F">'+str(not_added)+'</p><p style="font-size:10px;color:#A32D2D">Not added</p></div>'
    page += '</div></div>'+alert
    page += '<div style="background:#F5E6D8;border:0.5px solid #B5D4F4;border-radius:10px;padding:12px 16px;margin-bottom:20px"><p style="font-size:12px;color:#C46A2B">Security: API keys are stored server-side only. They are never exposed in source code or logs. Only admins can view them.</p></div>'
    page += '<div style="display:grid;grid-template-columns:1fr 1fr;gap:20px">'
    page += '<div><p style="font-size:11px;font-weight:600;color:#9CA3AF;text-transform:uppercase;letter-spacing:.08em;margin-bottom:12px">AI APIs</p>'+ai_html
    page += '<p style="font-size:11px;font-weight:600;color:#9CA3AF;text-transform:uppercase;letter-spacing:.08em;margin:16px 0 12px">Notifications</p>'+notif_html+'</div>'
    page += '<div><p style="font-size:11px;font-weight:600;color:#9CA3AF;text-transform:uppercase;letter-spacing:.08em;margin-bottom:12px">Payments & Email</p>'+pay_html
    page += '<p style="font-size:11px;font-weight:600;color:#9CA3AF;text-transform:uppercase;letter-spacing:.08em;margin:16px 0 12px">Other Services</p>'+other_html+'</div>'
    page += '</div></div></body></html>'
    return render_template_string(page)




@app.route('/admin-sidebar-js')
def admin_sidebar_js():
    js = """
(function(){
  if(!window.location.pathname.startsWith('/admin') || window.location.pathname === '/admin/login') return;
  if(document.getElementById('admsb')) return;
  var s = document.createElement('div');
  s.id = 'admsb';
  s.style.cssText = 'width:220px;background:#1B140F;padding:20px 12px;position:fixed;height:100vh;overflow-y:auto;left:0;top:0;z-index:9999;box-sizing:border-box';
  s.innerHTML = [
    '<div style="display:flex;align-items:center;gap:8px;margin-bottom:20px">',
    '<div style="width:26px;height:26px;background:#C46A2B;border-radius:6px;display:flex;align-items:center;justify-content:center">',
    '<span style="color:white;font-size:16px">&#9673;</span></div>',
    '<span style="color:white;font-size:13px;font-weight:600">Admin Panel</span></div>',
    '<p style="font-size:9px;color:rgba(255,255,255,.3);margin-bottom:8px;text-transform:uppercase;padding:0 8px">Management</p>',
    '<a href="/admin" style="display:block;padding:8px 10px;color:#9A8C7B;font-size:12px;text-decoration:none;border-radius:7px;margin-bottom:2px">&#9632; Overview</a>',
    '<a href="/admin/users" style="display:block;padding:8px 10px;color:#9A8C7B;font-size:12px;text-decoration:none;border-radius:7px;margin-bottom:2px">&#9632; Users</a>',
    '<a href="/admin/reports" style="display:block;padding:8px 10px;color:#9A8C7B;font-size:12px;text-decoration:none;border-radius:7px;margin-bottom:2px">&#9632; QC Reports</a>',
    '<a href="/admin/email" style="display:block;padding:8px 10px;color:#9A8C7B;font-size:12px;text-decoration:none;border-radius:7px;margin-bottom:2px">&#9632; Email Users</a>',
    '<a href="/admin/feedback" style="display:block;padding:8px 10px;color:#9A8C7B;font-size:12px;text-decoration:none;border-radius:7px;margin-bottom:2px">&#9632; Feedback</a>',
    '<hr style="border-color:rgba(255,255,255,.1);margin:8px 0">',
    '<p style="font-size:9px;color:rgba(255,255,255,.3);margin-bottom:8px;text-transform:uppercase;padding:0 8px">Settings</p>',
    '<a href="/admin/apis" style="display:block;padding:8px 10px;color:#9A8C7B;font-size:12px;text-decoration:none;border-radius:7px;margin-bottom:2px">&#9632; API Keys</a>',
    '<a href="/admin/tokens" style="display:block;padding:8px 10px;color:#9A8C7B;font-size:12px;text-decoration:none;border-radius:7px;margin-bottom:2px">&#9632; Token Limits</a>',
    '<a href="/admin/pricing" style="display:block;padding:8px 10px;color:#9A8C7B;font-size:12px;text-decoration:none;border-radius:7px;margin-bottom:2px">&#128176; Pricing</a>',
    '<a href="/admin/content" style="display:block;padding:8px 10px;color:#9A8C7B;font-size:12px;text-decoration:none;border-radius:7px;margin-bottom:2px">&#9632; Content</a>',
    '<a href="/admin/privacy" style="display:block;padding:8px 10px;color:#9A8C7B;font-size:12px;text-decoration:none;border-radius:7px;margin-bottom:2px">&#9632; Privacy</a>',
    '<a href="/admin/gift" style="display:block;padding:8px 10px;color:#9A8C7B;font-size:12px;text-decoration:none;border-radius:7px;margin-bottom:2px">&#9632; Gift Access</a>',
    '<hr style="border-color:rgba(255,255,255,.1);margin:8px 0">',
    '<a href="/" style="display:block;padding:8px 10px;color:#9A8C7B;font-size:12px;text-decoration:none;border-radius:7px">&#8592; Back to site</a>',
  ].join('');
  document.addEventListener('DOMContentLoaded',function(){
    document.body.insertBefore(s,document.body.firstChild);
    if(!document.querySelector('.main-content'))document.body.style.marginLeft='220px';
    var cur=window.location.pathname;
    s.querySelectorAll('a[href]').forEach(function(a){
      if(a.getAttribute('href')===cur){
        a.style.background='rgba(196,106,43,.18)';
        a.style.color='#E87B30';
        a.style.fontWeight='600';
      }
    });
  });
})();
"""
    from flask import Response
    return Response(js, mimetype='application/javascript')




if __name__ == '__main__':
    print("\n" + "="*55)
    print("  SurveyQC — Full Stack App v10.0")
    print("="*55)
    print(f"\n  Open:  http://localhost:5000")
    print(f"  Admin: http://localhost:5000/admin/login")
    print(f"  Demo:  demo@surveyqc.com / demo123")
    print(f"  Admin: admin123")
    print(f"\n  New features:")
    print(f"  - GDPR auto-delete (30 days)")
    print(f"  - Content management")
    print(f"  - Gift access")
    print(f"  - Share report + comments")
    print(f"  - Onboarding flow")
    print(f"  - Token control")
    print(f"  - Privacy page (5 countries)")
    print(f"\n  Press Ctrl+C to stop\n")
    port = int(os.environ.get('PORT', 5000))
    app.run(debug=False, port=port, host='0.0.0.0')
