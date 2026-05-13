"""
================================================================
  doc_analyzer.py — Document Parser (Multi-language)
================================================================
Yeh module .docx file se questions aur LOGIC tables nikalta hai.
Works with ANY language ke document.
"""

import re
from docx import Document
from docx.oxml.ns import qn


# QID detection — supports R/S/Q + bis/ter/Info/Ex suffixes
QID_PATTERN = re.compile(
    r'\[?\s*([RSQ]\d+(?:bis|ter|Info|info|Ex)?)\s*[\.\-–—\s\]]',
    re.IGNORECASE
)

# Junk patterns to skip
SKIP_PATTERNS = [
    r'^\s*PROGRAMMING\s+TABLE',
    r'^\s*\[NEXT SCREEN\]',
    r'^\s*\[SAME SCREEN',
    r'^\s*\[ALL COUNTRIES',
    r'^---+$', r'^\|\s*---', r'^\s*\|\s*$',
]
JUNK_RE = re.compile('|'.join(SKIP_PATTERNS), re.IGNORECASE)

# Termination keywords in MANY languages
TERMINATION_KEYWORDS = [
    # English
    'close', 'terminate', 'thank', 'thanks', 'screen out', 'disqualify',
    # Italian
    'grazie', 'chiudi', 'termina',
    # Spanish
    'gracias', 'cierre', 'terminar',
    # French
    'merci', 'fermer', 'fin',
    # German
    'vielen dank', 'schließen', 'beenden',
    # Generic
    'qualify', 'continue', 'go to',
]


def parse_questions(doc_path):
    """
    Parse all questions from the document.
    Returns: dict of {qid: {text, options, is_mandatory, ...}}
    """
    doc = Document(doc_path)
    questions = {}
    current_qid = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text or JUNK_RE.search(text):
            continue

        # Match QID
        match = QID_PATTERN.match(text)
        if match:
            qid = match.group(1)
            current_qid = qid
            if qid not in questions:
                questions[qid] = {
                    "text": "",
                    "options": [],
                    "is_mandatory": False,
                    "has_piping": False,
                }
            # Get rest of line after QID
            rest = text[match.end():].strip()
            rest = re.sub(r'^[\-–—\s]+[A-Z][A-Za-z\s"\'\-–—,]+\]', '', rest).strip()
            rest = re.sub(r'^\]', '', rest).strip()
            if rest and not JUNK_RE.search(rest):
                questions[qid]["text"] += " " + rest
            continue

        # If we have a current question, check for options
        if current_qid:
            # Option pattern: "1. Text" or "1) Text"
            opt_match = re.match(r'^(\d+)[\.\)]\s+(.+)', text)
            if opt_match:
                questions[current_qid]["options"].append({
                    "code": opt_match.group(1),
                    "text": opt_match.group(2).strip()
                })
            elif len(text) > 5:
                questions[current_qid]["text"] += " " + text
                if 'mandatory' in text.lower() or 'obligatoire' in text.lower():
                    questions[current_qid]["is_mandatory"] = True

    # Clean up whitespace in question texts
    for qid in questions:
        questions[qid]["text"] = re.sub(r'\s+', ' ', questions[qid]["text"]).strip()

    return questions


def extract_logic_tables(doc_path):
    """
    Extract all LOGIC/ROUTING tables from document, with host QID context.

    Returns: list of {host_qid, rows, flat_text}
    """
    doc = Document(doc_path)
    logic_tables = []
    current_qid = None

    # Walk through document body in order
    for child in doc.element.body.iterchildren():
        if child.tag == qn('w:p'):
            # It's a paragraph — find QID context
            for para in doc.paragraphs:
                if para._element is child:
                    text = para.text.strip()
                    m = QID_PATTERN.search(text)
                    if m:
                        current_qid = m.group(1)
                    break

        elif child.tag == qn('w:tbl'):
            # It's a table
            for tbl in doc.tables:
                if tbl._element is child:
                    rows_text = []
                    for row in tbl.rows:
                        row_text = []
                        for cell in row.cells:
                            ct = '\n'.join(
                                p.text.strip()
                                for p in cell.paragraphs
                                if p.text.strip()
                            )
                            row_text.append(ct)
                        rows_text.append(row_text)

                    flat = ' | '.join(' / '.join(r) for r in rows_text)

                    # Detect LOGIC/ROUTING table with termination keyword
                    has_logic_marker = bool(re.search(
                        r'\b(LOGIC|ROUTING|ROUTINE|TYPE)\b',
                        flat, re.IGNORECASE
                    ))
                    has_term_word = any(
                        kw.lower() in flat.lower()
                        for kw in TERMINATION_KEYWORDS
                    )

                    if has_logic_marker and has_term_word:
                        logic_tables.append({
                            'host_qid': current_qid,
                            'rows': rows_text,
                            'flat_text': flat
                        })
                    break

    return logic_tables


def analyze_document(doc_path):
    """
    Full document analysis.
    Returns: dict with 'questions' and 'logic_tables'
    """
    questions = parse_questions(doc_path)
    logic_tables = extract_logic_tables(doc_path)

    return {
        'questions': questions,
        'logic_tables': logic_tables,
        'stats': {
            'total_questions': len(questions),
            'questions_with_options': sum(1 for q in questions.values() if q['options']),
            'logic_tables_found': len(logic_tables),
        }
    }
